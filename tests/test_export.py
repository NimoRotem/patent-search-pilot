"""Export tests: /export produces a valid PDF and a valid DOCX with the expected content.
Uses the cached grabo report + locally-cached drawings (no paid APIs)."""
import os, tempfile
import pytest

SEL = ["US-11207792-B2", "US-9457478-B2", "US-4557659-A"]   # cached refs with local drawings


@pytest.fixture()
def model():
    import export_data
    return export_data.assemble("grabo_gripper_novelty", SEL)


def test_export_model_shape(model):
    assert model["mode"] == "novelty"
    assert len(model["references"]) == 3
    assert model["claim_chart"]["columns"], "claim chart must have reference columns"
    assert model["elements"], "elements present"
    # at least one selected reference has a locally embedded drawing + a quoted passage
    assert any(r["drawing_path"] and os.path.exists(r["drawing_path"]) for r in model["references"])
    assert any(r["quoted"] and r["quoted"]["text"] for r in model["references"])


def test_pdf_is_valid_and_complete(model, tmp_path):
    import export_pdf
    out = tmp_path / "r.pdf"
    export_pdf.render(model, out)
    data = out.read_bytes()
    assert data[:5] == b"%PDF-", "must be a real PDF"
    from pypdf import PdfReader
    reader = PdfReader(str(out))
    assert len(reader.pages) >= 5, f"expected a multi-page report, got {len(reader.pages)}"
    text = "\n".join((p.extract_text() or "") for p in reader.pages)
    #  ACCEPTANCE TEST for the export-disclosure fix. An exported document has to stand alone, so
    #  each of these was a real defect: the doc was titled like a search opinion, the grid was
    #  titled "claim chart", cells were shaded by retrieval score, and there was no scope
    #  disclosure of any kind.
    assert "Prior-Art Search Report" not in text, "the old authoritative title must be gone"
    assert "Prior-art retrieval report" in text
    assert "Element" in text and "retrieval map" in text
    assert "drafting aid" in text.lower()
    assert "not a verified claim chart" in text.lower()
    assert "Absence of results is not evidence of absence" in text
    assert "Measured recall" in text
    assert "CPC classes" in text
    assert SEL[0] in text, "a selected reference must appear"


def test_docx_is_valid_and_complete(model, tmp_path):
    import export_docx
    out = tmp_path / "r.docx"
    export_docx.render(model, out)
    from docx import Document
    doc = Document(str(out))
    paras = [p.text for p in doc.paragraphs]
    joined = "\n".join(paras)
    assert "Prior-Art Search Report" not in joined, "the old authoritative title must be gone"
    assert "Prior-art retrieval report" in joined
    assert any("retrieval map" in p.lower() for p in paras)
    assert any("drafting aid" in p.lower() for p in paras)
    assert any("Absence of results is not evidence of absence" in p for p in paras)
    assert any("Measured recall" in p for p in paras)
    assert len(doc.tables) >= 2, "retrieval map + appendix tables"
    assert len(doc.inline_shapes) >= 1, "at least one embedded drawing"
    # a selected reference appears somewhere (heading or appendix table)
    all_text = joined + "\n" + "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert any(s in all_text for s in SEL)


def test_export_chart_cells_carry_a_verification_verdict(model):
    """The root cause of the export-accuracy defect, pinned.

    export_data.assemble() builds its own view via webview.build_view(), and build_view does NOT
    verify anything -- verification was applied only to the cached view the browser reads. So every
    exported cell arrived with no `verify` key and the exporters shaded it green by retrieval
    score. If this assertion ever fails again, the exports have silently gone back to publishing
    unverified cells as coverage.
    """
    covered = [c for row in model["claim_chart"]["rows"] for c in row["cells"] if c.get("covered")]
    assert covered, "the gold report must have covered cells to check"
    missing = [c for c in covered if not c.get("verify")]
    assert not missing, f"{len(missing)} covered cells reached the export with no verdict"
    assert model["claim_chart"].get("verification"), "legend counts must reach the exporters"


def test_exports_never_shade_by_retrieval_score():
    """Cell colour must encode the verification verdict, never the fused retrieval score.

    The old behaviour made the most confidently-retrieved cell the greenest thing on the page --
    including cells the verifier had judged `unrelated`.
    """
    import inspect
    import export_docx, export_pdf

    def code_only(fn):
        """Strip comments: both functions EXPLAIN in a comment why they no longer use `intensity`,
        and that explanation must not itself trip the check."""
        out = []
        for line in inspect.getsource(fn).splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                continue
            out.append(line.split("#", 1)[0])
        return "\n".join(out)

    assert not hasattr(export_docx, "_green_shade"), "score-proportional green shading is back"
    assert "intensity" not in code_only(export_pdf._claim_chart_table), \
        "PDF chart is shading by retrieval score again"
    assert "intensity" not in code_only(export_docx._claim_chart), \
        "DOCX chart is shading by retrieval score again"


def test_disclosure_reports_the_independent_finding_not_the_self_grade():
    """'0 of 5 rejected' was the verifier grading its own work and read as reassurance the
    evidence does not support. The disclosure must lead with the independent re-read."""
    import corpus_facts, disclosure
    f = corpus_facts.facts()
    assert f["chart_indep_checked"] == 5
    assert f["chart_indep_overclaim"] == 2
    assert f["chart_indep_pct"] == 40
    warn = disclosure.chart_warning(f)
    assert "independent" in warn.lower()
    assert "overclaim" in warn.lower()
    assert "lower bound" in warn.lower(), "the shared-model-family caveat must survive"


# ── XLSX + Markdown: the two shapes added because the document formats could not carry them ──
#
# The spreadsheet exists to be the RESULT LIST (sortable, one sketch per row); the Markdown exists
# to carry the full claims and description that the page hides behind a tab and that no other
# export includes. Each test below pins the property that distinguishes its format — if a future
# change makes .xlsx a picture-free table or .md a text-free summary, the format has lost its
# reason to exist and these fail.

def test_xlsx_is_a_result_list_with_one_sketch_per_row(model, tmp_path):
    import export_xlsx
    from openpyxl import load_workbook
    out = tmp_path / "r.xlsx"
    export_xlsx.render(model, out)
    wb = load_workbook(str(out))
    assert {"Results", "Element x reference", "Search scope"} <= set(wb.sheetnames)

    ws = wb["Results"]
    headers = [c.value for c in ws[6]]
    # The fields a reader sees on the card, in the sheet. Each of these was web-only before.
    for h in ("Sketch", "Relevancy", "Legal status", "Family", "Found via", "Why relevant",
              "Publication", "Assignee", "CPC"):
        assert h in headers, f"result sheet lost the {h!r} column"
    assert ws.max_row >= 6 + len(model["references"])
    #  A row per selected reference, and a picture in at least one of them. Embedded images live
    #  on the worksheet, not in a cell value, so ws._images is the only place to see them.
    assert len(ws._images) >= 1, "no sketch was embedded — this format's whole point"
    assert len(ws._images) <= len(model["references"]), "one sketch per reference, not a gallery"

    scope = "\n".join(str(c.value) for row in wb["Search scope"].iter_rows() for c in row if c.value)
    assert "Measured recall" in scope
    assert "Absence of results is not evidence of absence" in scope
    assert model["query"][:40] in scope, "the sheet must say what was searched"


def test_xlsx_chart_colours_by_verification_state_not_by_score(model, tmp_path):
    """Same rule as the PDF/DOCX: a spreadsheet is the most authoritative-looking surface of the
    four, so a cell may only look like coverage when the verifier confirmed it."""
    import export_xlsx, disclosure, inspect
    src = inspect.getsource(export_xlsx._chart_sheet)
    assert "disclosure.cell_state" in src, "chart fill must come from the verification verdict"
    assert "intensity" not in src, "XLSX chart is shading by retrieval score"

    from openpyxl import load_workbook
    out = tmp_path / "r.xlsx"
    export_xlsx.render(model, out)
    ws = load_workbook(str(out))["Element x reference"]
    # openpyxl returns the ARGB string as written, so compare case-insensitively.
    fills = {str(c.fill.fgColor.rgb).lower() for row in ws.iter_rows(min_row=6) for c in row
             if c.fill and c.fill.fgColor and c.fill.fgColor.rgb}
    allowed = {("00" + f).lower() for _m, _l, f, _c in disclosure.CELL_STATES.values()}
    allowed |= {"00000000", "ffffffff"}                      # unfilled / default
    assert fills, "no cells were filled at all"
    assert fills <= allowed, f"unexpected chart fill: {fills - allowed}"


def test_markdown_carries_the_full_patent_text_the_other_formats_drop(tmp_path):
    import export_data, export_md
    m = export_data.assemble("grabo_gripper_novelty", SEL, include_text=True, include_drawings=False)
    out = tmp_path / "r.md"
    export_md.render(m, out)
    text = out.read_text(encoding="utf-8")

    #  The distinguishing property: whole claim sets and whole descriptions, not a single quoted
    #  passage. The web card caps description at 60 paragraphs and the PDF/DOCX quote one passage.
    assert any((r.get("text") or {}).get("claims") for r in m["references"]), \
        "include_text=True produced no claims"
    assert "#### Claims (" in text and "#### Description (" in text
    n_claims = sum(len((r.get("text") or {}).get("claims") or []) for r in m["references"])
    assert n_claims >= 10 and len(text) > 20_000, "the full text did not make it into the file"

    #  Deliberately image-free and link-free: that is what keeps it small and diffable.
    assert "![" not in text, "markdown export must not embed images"
    assert "](http" not in text, "markdown export must not carry hyperlinks"
    #  ...but never at the cost of the disclosure, which travels on every surface.
    assert "Measured recall" in text
    assert "Absence of results is not evidence of absence" in text
    assert "not a prior-art search opinion" in text


def test_assemble_carries_the_card_fields_every_export_needs(model):
    """The relevancy score, its written opinion, the legal-status tag and the family summary are
    read off the screen by whoever asks for an export; before this they existed only in the web
    view and every exported file quietly dropped them."""
    keys = ("relevancy", "relevancy_opinion", "status_label", "family_summary",
            "found_via", "n_images", "pdf_url")
    for r in model["references"]:
        for k in keys:
            assert k in r, f"reference is missing {k!r}"
    assert any(r.get("status_label") for r in model["references"])
    assert any(r.get("found_via") for r in model["references"])
    # report-wide facts behind the two collapsed panels on the page
    for k in ("rounds", "cpc_branches", "channels_used", "n_cards", "source_tags"):
        assert k in model, f"model is missing {k!r}"


def test_export_route_accepts_every_declared_format(app_client, monkeypatch):
    """/export used to hard-code ('pdf','docx'); the format table is now the single source of
    truth for the route, the export bar and these tests.

    Asserted against the table itself rather than a hardcoded set, so adding a format (the IDS
    citation listing was the fifth) does not require editing an assertion that was only ever
    restating the table."""
    import webapp
    assert set(webapp.EXPORT_FORMATS) >= {"pdf", "docx", "xlsx", "md", "ids"}
    r = app_client.post("/export", data={"slug": "grabo_gripper_novelty",
                                         "pubs": ",".join(SEL), "format": "tex"})
    assert r.status_code == 400
    body = r.get_json()
    for f in webapp.EXPORT_FORMATS:
        assert f in body["error"]

    r = app_client.post("/export", data={"slug": "grabo_gripper_novelty",
                                         "pubs": ",".join(SEL), "format": "md"})
    assert r.status_code == 200
    assert r.headers["Content-Type"].startswith("text/markdown")
    assert b"Prior-art retrieval report" in r.data
