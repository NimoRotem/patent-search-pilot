"""The drafting conversation: the checks, the workspace, the validator and the filing package.

Every test here is anchored on a defect that a machine-written patent application actually
produces — a numeral used and never defined, a claim depending on a claim that does not exist, a
citation to a publication that resolves to nothing, an abstract over the 150-word cap.  A check
that has never been shown to bite on a real failure is a check nobody should trust, so each one is
defect-injected: the good draft passes, and the same draft with one thing broken fails on exactly
that thing and nothing else.
"""
import hashlib
import json
import re
import threading
from contextlib import contextmanager
from pathlib import Path
from unittest.mock import Mock

import pytest

import draft_agent
import draft_cite
import draft_export
import draft_figures
import draft_qa
import draft_studio
import draft_uspto
import draft_workspace
import drafting


# =============================================================================================
# A small, internally consistent application to break in one place at a time
# =============================================================================================
GOOD = {
    "title": "Vacuum Lifting Tool With Interchangeable Sealing Ring",
    "cross_reference": "This application claims no priority.",
    "government_support": "Not applicable.",
    "field": "The disclosure relates to portable vacuum lifting tools.",
    "background": "Handheld vacuum lifters are known [REF:US-11223344-B2]. Such tools use a "
                  "single fixed seal, which limits the surfaces they can grip.",
    "summary": "A vacuum lifting tool has a body carrying a pump and a sealing ring that is "
               "removable from the body without a tool.",
    "drawing_descriptions": "FIG. 1 is a side elevation of the vacuum lifting tool.\n\n"
                            "FIG. 2 is an exploded view of the sealing ring and the body.",
    "detailed_description": "Referring to FIG. 1, a vacuum lifting tool 10 has a body 12 that "
                            "carries a pump 14. A sealing ring 16 is received in a groove 18 in "
                            "the body 12. As shown in FIG. 2, the sealing ring 16 is removable "
                            "from the groove 18 by hand, so that a ring suited to a rough "
                            "surface may replace one suited to glass. The pump 14 draws air "
                            "through a passage 20 in the body 12.",
    "claims": "1. A vacuum lifting tool comprising a body, a pump carried by the body, a groove "
              "in the body, and a sealing ring received in the groove and removable from the "
              "groove without a tool.\n\n"
              "2. The vacuum lifting tool of claim 1, wherein the body defines a passage between "
              "the pump and the groove.\n\n"
              "3. The vacuum lifting tool of claim 2, wherein the pump is battery powered.",
    "abstract": "A vacuum lifting tool has a body carrying a pump and a sealing ring received in "
                "a groove in the body. The sealing ring is removable by hand so a ring suited to "
                "one surface can replace a ring suited to another.",
}
NUMERALS = [
    {"numeral": "10", "part": "vacuum lifting tool"}, {"numeral": "12", "part": "body"},
    {"numeral": "14", "part": "pump"}, {"numeral": "16", "part": "sealing ring"},
    {"numeral": "18", "part": "groove"}, {"numeral": "20", "part": "passage"},
]
FIGURES = [
    {"label": "FIG. 1", "caption": "side elevation", "numerals": ["10 vacuum lifting tool",
                                                                  "12 body", "14 pump"]},
    {"label": "FIG. 2", "caption": "exploded view", "numerals": ["16 sealing ring", "18 groove",
                                                                    "20 passage"]},
]
ALLOWED = ["US-11223344-B2"]


@pytest.fixture(autouse=True)
def clean_source_review(monkeypatch, request):
    """The source-fidelity gate is its own text gate now.

    It used to be reached through ``_ensure_figures``, so every test that stubbed the drawing pass
    stubbed this along with it. Now that a drafting turn does not draw, it has to be stubbed on its
    own or these tests would try to run a real reviewing agent. A test that is ABOUT the gate
    overrides this after the runner is built.
    """
    if request.node.get_closest_marker("real_source_review"):
        return
    monkeypatch.setattr(
        draft_studio.TurnRunner, "_review_sources",
        lambda self, **_kwargs: {
            "status": "complete", "verdict": "pass", "summary": "clean",
            "checks": [], "findings": [], "counts": {}, "cost_usd": 0.0,
            "duration_ms": 0, "model_name": "", "last_error": ""})


@pytest.fixture(autouse=True)
def no_corpus(monkeypatch):
    """Citations resolve against a stub, so these tests never touch Postgres or a paid API."""
    known = {"US-11223344-B2": {"found": True, "publication_number": "US-11223344-B2",
                                "title": "Handheld vacuum lifter", "source": "corpus",
                                "publication_date": "2021-01-05", "kind_code": "B2",
                                "assignee": "Example Co", "url": "https://example.test/1"}}
    monkeypatch.setattr(draft_cite, "resolve",
                        lambda pub, **k: dict(known.get(draft_cite.normalize(pub) or "",
                                                        {"found": False,
                                                         "publication_number": str(pub),
                                                         "reason": "not in the local corpus"})))


def checks_for(sections=None, numerals=None, figures=None, allowed=ALLOWED):
    return {c["name"]: c for c in draft_qa.run_checks(
        sections=sections or GOOD, numerals=NUMERALS if numerals is None else numerals,
        figures=FIGURES if figures is None else figures, allowed_references=allowed,
        allow_remote=False)}


# =============================================================================================
# The clean draft passes
# =============================================================================================
def test_a_consistent_draft_passes_every_mechanical_check():
    checks = checks_for()
    failed = [name for name, check in checks.items() if check["status"] == "fail"]
    assert failed == [], f"a consistent draft should not fail anything: {failed}"
    assert draft_qa.verdict_for(list(checks.values()), []) in ("pass", "warn")


def test_reference_numeral_leader_cannot_end_in_an_arrowhead():
    figures = [
        {
            **FIGURES[0],
            "caption": (
                "The cord 10 is identified by a leader coming down from above and ending in an "
                "arrowhead whose tip touches the cord."
            ),
        },
        FIGURES[1],
    ]
    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]
    assert check["status"] == "fail"
    assert any("terminal dot" in item for item in check["items"])

    section_arrow = [
        {
            **FIGURES[0],
            "caption": (
                "A cutting-plane line crosses the view and carries an arrowhead at each end "
                "to indicate the viewing direction."
            ),
        },
        FIGURES[1],
    ]
    assert checks_for(figures=section_arrow)[
        "Drawing briefs are concise and drawable"]["status"] == "pass"


def test_flowchart_briefs_cannot_request_verbal_text_inside_shapes():
    figures = [{
        **FIGURES[0],
        "caption": (
            "A process flow diagram. The process starts at a block labeled "
            "\"Monitor branch current\" and proceeds to a decision block that asks "
            "\"Is current above the limit?\"."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "verbal drawing text" in " ".join(check["items"]).lower()
    assert "reference numerals" in " ".join(check["items"]).lower()


def test_flowchart_briefs_must_number_their_process_and_decision_shapes():
    figures = [{
        "label": "FIG. 1",
        "caption": (
            "A process flow diagram with four empty process rectangles connected in order "
            "and a return path from the last rectangle to the first rectangle."
        ),
        "numerals": [],
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "process-flow drawing has no numbered steps" in " ".join(check["items"]).lower()


def test_flowchart_briefs_with_numbered_steps_remain_renderable():
    figures = [{
        **FIGURES[0],
        "caption": (
            "A process flow diagram with a monitoring step 10 and a response step 12. "
            "A downward flow arrow connects the monitoring step 10 to the response step 12. "
            "Flowchart nodes: 10=process, 12=process, END=terminator. "
            "Flowchart directed edges: 10->12, 12->END."
        ),
        "numerals": ["10 vacuum lifting tool", "12 body"],
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]
    assert check["status"] == "pass", check["items"]


def test_flowchart_briefs_require_an_exact_machine_readable_topology():
    figures = [{
        **FIGURES[0],
        "caption": (
            "A process flow diagram with a monitoring step 10 and a response step 12. "
            "A downward flow arrow connects the monitoring step 10 to the response step 12."
        ),
        "numerals": ["10 vacuum lifting tool", "12 body"],
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    text = " ".join(check["items"]).lower()
    assert "flowchart nodes" in text and "flowchart directed edges" in text


def test_drafting_contract_uses_blank_start_and_end_connectors_for_split_flows():
    contract = re.sub(r"\s+", " ", draft_studio.DRAFT_SYSTEM.lower())

    assert "end=connector" in contract and "start=connector" in contract
    assert "never label a continuation connector" in contract


def test_an_empty_figure_brief_is_refused_before_drawing():
    figures = [{**FIGURES[0], "caption": " \n\t "}, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "empty drawing brief" in " ".join(check["items"]).lower()


def test_verdict_never_fails_on_an_advisory_check_alone():
    """A heuristic must not be able to condemn a draft — see the calibration note in draft_qa."""
    advisory = [{"name": "Antecedent basis", "status": "fail", "severity": "advisory",
                 "detail": "", "items": []}]
    assert draft_qa.verdict_for(advisory, []) == "warn"
    proven = [{"name": "Numerals", "status": "fail", "severity": "error", "detail": "",
               "items": []}]
    assert draft_qa.verdict_for(proven, []) == "fail"


# =============================================================================================
# Reference numerals
# =============================================================================================
def test_a_numeral_used_but_never_defined_is_caught():
    broken = dict(GOOD)
    broken["detailed_description"] += " A trigger 22 is mounted on the body 12."
    check = checks_for(broken)["Every numeral in the text is defined"]
    assert check["status"] == "fail"
    assert any(item.startswith("22") for item in check["items"])


def test_first_use_numeral_check_ignores_figure_labels():
    check = draft_qa._first_use_introduces(
        "FIG. 10 is a detail view of the reset mechanism. A container lid 10 supports the body.",
        {"10": "container lid"})

    assert check["status"] == "pass"


def test_first_use_numeral_check_ignores_claim_numbers():
    check = draft_qa._first_use_introduces(
        "11. A spring closes a valve through an outlet screen.\n\n"
        "12. The pressure relief cartridge 12 includes a retaining ring.",
        {"12": "pressure relief cartridge"})

    assert check["status"] == "pass"


def test_first_use_numeral_check_ignores_claim_references():
    check = draft_qa._first_use_introduces(
        "12. The method of claim 16, wherein pressure is released. "
        "A retaining ring 16 secures the cartridge.",
        {"16": "retaining ring"})

    assert check["status"] == "pass"


def test_two_numerals_for_the_same_part_are_caught():
    numerals = NUMERALS + [{"numeral": "24", "part": "sealing ring"}]
    broken = dict(GOOD)
    broken["detailed_description"] += " A second sealing ring 24 may be fitted."
    check = checks_for(broken, numerals=numerals)["One numeral per part"]
    assert check["status"] == "fail" and "sealing ring" in check["items"][0]


def test_every_numeral_entry_has_one_valid_number_and_one_named_part():
    broken = NUMERALS + [{"numeral": "22", "part": ""},
                         {"numeral": "12", "part": "different body"},
                         {"numeral": "part-x", "part": "invalid"}]
    check = checks_for(numerals=broken)["Every numeral-table row is complete and unique"]
    assert check["status"] == "fail"
    assert any("22" in item and "no part" in item for item in check["items"])
    assert any("12" in item and "more than once" in item for item in check["items"])
    assert any("part-x" in item and "invalid" in item for item in check["items"])


def test_a_numeral_on_a_drawing_that_the_table_does_not_define_is_caught():
    figures = [dict(FIGURES[0]), {"label": "FIG. 2", "caption": "exploded",
                                  "numerals": ["44 mystery bracket"]}]
    check = checks_for(figures=figures)["Numerals on the drawings are defined"]
    assert check["status"] == "fail" and "44" in check["items"]


def test_a_numeral_visible_on_a_drawing_but_absent_from_the_text_is_caught():
    numerals = NUMERALS + [{"numeral": "22", "part": "trigger"}]
    figures = FIGURES + [{"label": "FIG. 3", "caption": "trigger detail",
                          "numerals": ["22 trigger"]}]
    check = checks_for(numerals=numerals, figures=figures)[
        "Every drawing numeral appears in the specification"]
    assert check["status"] == "fail" and "22" in check["items"]


def test_a_text_numeral_missing_from_every_drawing_is_caught():
    figures = [{**FIGURES[0]}, {**FIGURES[1], "numerals": ["16 sealing ring", "18 groove"]}]
    check = checks_for(figures=figures)["Every specification numeral appears in a drawing"]
    assert check["status"] == "fail" and "20" in check["items"]
    assert "Do not remove a disclosed part" in check["detail"]


def test_a_reference_numeral_printed_twice_is_caught():
    figures = [{**FIGURES[0], "numerals": FIGURES[0]["numerals"] + ["10"]}, FIGURES[1]]
    check = checks_for(figures=figures)["Each drawing numeral appears once"]
    assert check["status"] == "fail" and check["items"] == ["FIG. 1: 10"]


def test_an_overcrowded_drawing_sheet_is_a_filing_blocker():
    numerals = [
        {"numeral": str(value), "part": f"part {value}"}
        for value in range(10, 28, 2)
    ]
    figures = [{
        "label": "FIG. 1", "caption": "assembly view",
        "numerals": [item["numeral"] for item in numerals],
    }]

    check = checks_for(numerals=numerals, figures=figures)[
        "Drawing sheets are not overcrowded"]

    assert check["status"] == "fail" and "9 numerals" in check["items"][0]


def test_an_overlong_drawing_brief_is_refused_before_image_generation():
    figures = [
        {**FIGURES[0], "caption": "plain rectangular body " * 160},
        FIGURES[1],
    ]
    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]
    assert check["status"] == "fail"
    assert "FIG. 1" in check["items"][0]

    with pytest.raises(
            draft_studio.FilingPreflightError,
            match="Drawing briefs are concise and drawable") as caught:
        draft_studio.validate_snapshot(
            {"sections": GOOD, "numerals": NUMERALS, "figures": figures}, ALLOWED)
    assert caught.value.category == "figures_and_numerals"


def test_a_legacy_figure_label_cut_off_mid_word_is_refused_before_drawing():
    figures = [{
        **FIGURES[0],
        "label": "FIG. 2 - Side elevation in vertical section, showing the cha",
        "caption": "The chamber 22 is bounded by the perimeter member 24.",
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "cut off mid-word" in check["items"][0]


def test_a_complete_word_at_character_sixty_is_not_treated_as_a_cutoff():
    prefix = "FIG. 1 - "
    label = prefix + ("x" * (60 - len(prefix) - len("view"))) + "view"
    figures = [{
        **FIGURES[0],
        "label": label,
        "caption": "The view shows the body 12 and pump 14.",
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "pass"


def test_a_self_contradictory_endpoint_target_is_refused_before_drawing():
    figures = [{
        **FIGURES[0],
        "caption": (
            "The pump is a rectangular upper block. The air-extraction mechanism 20 is "
            "identified on its flat right-hand face at mid-height, below that face."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "FIG. 1" in check["items"][0]
    assert "contradictory" in check["items"][0].lower()
    with pytest.raises(
            draft_studio.FilingPreflightError,
            match="Drawing briefs are concise and drawable"):
        draft_studio.validate_snapshot(
            {"sections": GOOD, "numerals": NUMERALS, "figures": figures}, ALLOWED)


def test_an_axial_hollow_cylinder_cannot_be_specified_as_an_annulus():
    figures = [{
        **FIGURES[0],
        "caption": (
            "A cross-sectional view taken on line 8-8 of FIG. 2. A single drill bushing 14 is "
            "shown in cross-section inside a carriage. The drill bushing is cylindrical, so "
            "its cross-section is a hatched annulus. The drill bushing has a vertical "
            "cylindrical bore passing completely through it. A threaded shank descends "
            "through the carriage beside the bushing."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "axial section" in check["items"][0]
    assert "two opposed sectioned walls" in check["items"][0]


def test_live_axial_bushing_word_order_cannot_bypass_annulus_preflight():
    figures = [{
        **FIGURES[0],
        "caption": (
            "A cross-sectional view taken on line A-A of FIG. 2. The first guide carriage 50 "
            "is shown in cross-section. A single, cylindrical drill bushing 54 is seated in a "
            "vertical bore within the first guide carriage 50. In this cross-sectional view, "
            "the drill bushing 54 appears as a hollow annulus, with its solid wall shown in "
            "section with hatching. The drill bushing 54 has a central, vertical bore passing "
            "completely through it. A threaded shank extends vertically downward from a clamp "
            "knob and passes through the first guide carriage 50."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "axial section" in check["items"][0]
    assert "two opposed sectioned walls" in check["items"][0]


def test_vertical_bore_axis_cannot_be_aligned_with_longitudinal_slot_axis():
    figures = [{
        **FIGURES[0],
        "caption": (
            "A cross-sectional view taken on line A-A of FIG. 2. A longitudinal slot 16 "
            "extends along the rail 10. A drill bushing 54 has a central vertical bore that "
            "passes completely through the bushing. The central axis of the bore of the drill "
            "bushing 54 is vertically aligned with the central axis of the longitudinal slot "
            "16 of the rail 10."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "vertical bore axis" in check["items"][0]
    assert "longitudinal slot axis" in check["items"][0]
    assert "intersects the open slot" in check["items"][0]


def test_a_figure_numeral_cannot_target_a_different_grouping_shape():
    figures = [{
        **FIGURES[0],
        "caption": (
            "A large rectangle, the vacuum lifting tool 10, surrounds the other parts."
        ),
        "numerals": [
            "10 vacuum lifting tool: leader ends on the vertical spine of the large square "
            "bracket.",
            "12 body",
            "14 pump",
        ],
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "numeral 10" in check["items"][0]
    assert "rectangle" in check["items"][0]
    assert "square bracket" in check["items"][0]
    with pytest.raises(
            draft_studio.FilingPreflightError,
            match="Drawing briefs are concise and drawable"):
        draft_studio.validate_snapshot(
            {"sections": GOOD, "numerals": NUMERALS, "figures": figures}, ALLOWED)

    consistent = [{
        **figures[0],
        "caption": (
            "A large square bracket, the vacuum lifting tool 10, groups the other parts."
        ),
    }, FIGURES[1]]
    assert checks_for(figures=consistent)[
        "Drawing briefs are concise and drawable"]["status"] == "pass"


def test_a_remote_part_numeral_cannot_label_its_offsheet_connection_stub():
    figures = [{
        **FIGURES[0],
        "caption": (
            "A large rectangle is the body 12. A short vertical line extends upward from the "
            "body 12, representing a connection to the pump 14."
        ),
        "numerals": [
            "10 vacuum lifting tool",
            "12 body: leader ends inside the large rectangle.",
            "14 pump: leader ends on the short vertical line extending upward from the body 12.",
        ],
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "numeral 14" in check["items"][0]
    assert "off-sheet connection" in check["items"][0]
    drawing_problems = []
    draft_studio.validate_snapshot(
        {"sections": GOOD, "numerals": NUMERALS, "figures": figures}, ALLOWED,
        drawing_problems)
    assert any("off-sheet connection" in problem for problem in drawing_problems)

    actual_bus = draft_qa._offsheet_connection_target(
        "108 isolated local bus: leader ends on the bus line between the controller and branch.",
        "A solid horizontal line, the isolated local bus, extends below the controller. A "
        "vertical line connects the controller to the left end of the isolated local bus.")
    assert actual_bus is None

    corrected = [{
        **figures[0],
        "caption": (
            "A large rectangle is the body 12. A smaller rectangle above it is the pump 14. "
            "A short vertical line connects the body 12 to the pump 14."
        ),
        "numerals": [
            "10 vacuum lifting tool",
            "12 body: leader ends inside the large rectangle.",
            "14 pump: leader ends inside the smaller rectangle.",
        ],
    }, FIGURES[1]]
    assert checks_for(figures=corrected)[
        "Drawing briefs are concise and drawable"]["status"] == "pass"


def test_a_drawn_tile_cannot_coexist_with_a_no_other_panel_constraint():
    figures = [{
        **FIGURES[0],
        "caption": (
            "A large plain tile fills the lower part of the sheet. The base 12 is the lowest "
            "slab of the assembly, the one slab on the sheet; no other slab, plate or panel is "
            "drawn."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "contradictory sheet exclusivity" in check["items"][0].lower()
    with pytest.raises(
            draft_studio.FilingPreflightError,
            match="Drawing briefs are concise and drawable"):
        draft_studio.validate_snapshot(
            {"sections": GOOD, "numerals": NUMERALS, "figures": figures}, ALLOWED)


def test_blanket_shape_background_and_stroke_controls_are_refused_before_drawing():
    figures = [{
        **FIGURES[0],
        "caption": (
            "No circle, ring, disc, hole or ellipse appears anywhere on the sheet. "
            "A plain tile has no joint line and no other tile is shown. "
            "The cord is a strip bounded by two long lines."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    issues = " ".join(check["items"]).lower()
    assert "blanket shape exclusion" in issues
    assert "background exclusion" in issues
    assert "exact stroke count" in issues
    with pytest.raises(
            draft_studio.FilingPreflightError,
            match="Drawing briefs are concise and drawable"):
        draft_studio.validate_snapshot(
            {"sections": GOOD, "numerals": NUMERALS, "figures": figures}, ALLOWED)


def test_an_arbitrary_blank_drawing_area_is_refused_before_drawing():
    figures = [{
        **FIGURES[0],
        "caption": (
            "A tray side wall stands beside a downward guide duct. The drawing area to the "
            "left of the tray side wall is blank."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "background exclusion" in check["items"][0]
    assert "drawing area" in check["items"][0]


def test_generic_negative_boundary_controls_are_refused_before_drawing():
    figures = [{
        **FIGURES[0],
        "caption": (
            "A slab stands on a separate band and carries two closed housings. "
            "No face has a rim, ledge, chamfer or second boundary drawn inside its edges."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "generic negative linework control" in " ".join(check["items"]).lower()
    with pytest.raises(
            draft_studio.FilingPreflightError,
            match="Drawing briefs are concise and drawable"):
        draft_studio.validate_snapshot(
            {"sections": GOOD, "numerals": NUMERALS, "figures": figures}, ALLOWED)


@pytest.mark.parametrize("control", [
    "Every outline is one thin unbroken line.",
    "Where two faces meet, they meet along one shared edge drawn once and serving both.",
])
def test_generic_face_linework_controls_are_refused_before_drawing(control):
    figures = [{
        **FIGURES[0],
        "caption": "A slab stands on a separate band and carries two closed housings. " + control,
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "generic face-linework control" in " ".join(check["items"]).lower()


def test_an_exact_separator_line_count_is_refused_before_drawing():
    figures = [{
        **FIGURES[0],
        "caption": "Three stacked hatched bands are separated by one horizontal line.",
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "exact stroke count" in " ".join(check["items"]).lower()


@pytest.mark.parametrize("caption", [
    (
        "The electrical supply cord is a slender strip of even width with plain white paper "
        "along its interior. Identified well inside that strip."
    ),
    (
        "The pulling element cable runs away from the device in one sweep. It is drawn as a "
        "broad strip of even width with plain white paper along its interior."
    ),
])
def test_a_white_interior_cord_strip_is_refused_before_drawing(caption):
    figures = [{
        **FIGURES[0],
        "caption": caption,
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "ambiguous multi-stroke cord" in " ".join(check["items"]).lower()


@pytest.mark.parametrize("caption", [
    "Each body is large, with open white paper between neighbours.",
    "The duct is broad, with open paper between it and the motor housing.",
])
def test_renderer_only_open_paper_between_solid_bodies_is_refused(caption):
    figures = [{**FIGURES[0], "caption": caption}, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "open-paper spacing" in " ".join(check["items"]).lower()


@pytest.mark.parametrize("caption", [
    "The cord runs away from the body and leaves the sheet at its left edge.",
    "The column runs from the top edge of the sheet to the lower band.",
    "The exposed face runs unbroken across the sheet.",
    "The covering element fills the lower part of the sheet.",
])
def test_figure_linework_cannot_be_directed_to_a_physical_sheet_edge(caption):
    figures = [{**FIGURES[0], "caption": caption}, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "physical sheet edge" in " ".join(check["items"]).lower()


def test_a_clear_sheet_margin_instruction_remains_renderable():
    figures = [{
        **FIGURES[0],
        "caption": "The whole of the drawing stands clear of the edges of the sheet.",
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "pass"


def test_every_open_paper_spacing_on_one_sheet_is_reported_together():
    figures = [{
        **FIGURES[0],
        "caption": (
            "Each body is large, with open white paper between neighbours. "
            "The duct is broad, with open paper between it and the motor housing."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]
    issues = [item for item in check["items"] if "open-paper spacing" in item]

    assert len(issues) == 2


def test_an_endpoint_deliberately_disconnected_from_its_named_part_is_refused():
    figures = [{
        **FIGURES[0],
        "caption": (
            "The first side 14 is the straight upper edge line of the slab. Identified in the "
            "open white paper directly above that upper edge line, touching no line."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "disconnected endpoint" in check["items"][0].lower()
    with pytest.raises(
            draft_studio.FilingPreflightError,
            match="Drawing briefs are concise and drawable"):
        draft_studio.validate_snapshot(
            {"sections": GOOD, "numerals": NUMERALS, "figures": figures}, ALLOWED)


@pytest.mark.parametrize("target", [
    "Identified at the centre of its flat front face.",
    "Identified at mid-height on its right side face.",
    "Identified in the lower-left quarter of the surface.",
    "Identified at the topmost point of the ring.",
    "Identified on the upright line near the left end.",
    "Identified halfway along the boundary toward its right end.",
    "Identified well inside the left-hand part of the chamber.",
    "Identified in the band along the right side of the sheet.",
    "Identified in the margin along the bottom of the sheet.",
    "Identified on the second rectangle.",
])
def test_an_arbitrary_exact_numeral_target_is_refused_before_drawing(target):
    figures = [{
        **FIGURES[0],
        "caption": "The pump 20 is a visible rectangular body. " + target,
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "fail"
    assert "arbitrary exact endpoint" in check["items"][0].lower()
    with pytest.raises(
            draft_studio.FilingPreflightError,
            match="Drawing briefs are concise and drawable"):
        draft_studio.validate_snapshot(
            {"sections": GOOD, "numerals": NUMERALS, "figures": figures}, ALLOWED)


def test_every_arbitrary_target_on_one_sheet_is_reported_together():
    figures = [{
        **FIGURES[0],
        "caption": (
            "The pump 14 is visible. Identified at the centre of its face. "
            "The ring 16 is visible. Identified at its topmost point."
        ),
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]
    issues = [item for item in check["items"] if item.startswith("FIG. 1")]

    assert len(issues) == 2
    assert "centre" in issues[0]
    assert "topmost" in issues[1]


@pytest.mark.parametrize("target", [
    "Identified well inside its flat front face.",
    "Identified at any point along the boundary line.",
    "Identified within the outermost rectangle.",
    "Identified anywhere within the middle band.",
])
def test_a_broad_stable_numeral_target_remains_renderable(target):
    figures = [{
        **FIGURES[0],
        "caption": "The pump 20 is a visible rectangular body. " + target,
    }, FIGURES[1]]

    check = checks_for(figures=figures)["Drawing briefs are concise and drawable"]

    assert check["status"] == "pass"


def test_preflight_reports_every_over_specific_figure_before_drawing():
    figures = [
        {**FIGURES[0], "caption": "The pump 14 is visible. Identified at its centre."},
        {**FIGURES[1], "caption": "The ring 16 is visible. Identified at its topmost point."},
    ]

    with pytest.raises(draft_studio.FilingPreflightError) as caught:
        draft_studio.validate_snapshot(
            {"sections": GOOD, "numerals": NUMERALS, "figures": figures}, ALLOWED)

    assert "FIG. 1" in str(caught.value)
    assert "FIG. 2" in str(caught.value)


def test_an_explicit_figure_numeral_declaration_must_match_the_sheet_list():
    figures = [
        {
            **FIGURES[0],
            "caption": (
                "A side elevation of the body and pump.\n\n"
                "**Numerals appearing on this figure:** 10, 12"
            ),
        },
        FIGURES[1],
    ]
    check = checks_for(figures=figures)[
        "Figure brief numeral declarations match sheet lists"]

    assert check["status"] == "fail"
    assert "FIG. 1" in check["items"][0]
    assert "14" in check["items"][0]

    with pytest.raises(
            draft_studio.FilingPreflightError,
            match="Figure brief numeral declarations match sheet lists") as caught:
        draft_studio.validate_snapshot(
            {"sections": GOOD, "numerals": NUMERALS, "figures": figures}, ALLOWED)
    assert caught.value.category == "figures_and_numerals"


def test_letter_qualified_reference_numerals_are_compared_exactly():
    version = {**GOOD, "detailed_description":
               GOOD["detailed_description"] + " A secondary spacer 10a is beside a lug A12."}
    numerals = NUMERALS + [{"numeral": "10A", "part": "secondary spacer"},
                           {"numeral": "A12", "part": "lug"}]
    figures = [{**FIGURES[0], "numerals": FIGURES[0]["numerals"] + ["10a", "A12"]},
               FIGURES[1]]
    checks = checks_for(version, numerals=numerals, figures=figures)
    assert checks["Numerals on the drawings are defined"]["status"] == "pass"
    assert checks["Every drawing numeral appears in the specification"]["status"] == "pass"
    assert checks["Every specification numeral appears in a drawing"]["status"] == "pass"


def test_the_review_never_reads_drawing_pixels():
    """Nothing generates a sheet any more, so nothing inspects one.

    ``figures_for_qa`` used to join every figure specification to the vision audits of the sheet
    the product had drawn for it. With no generation there is no audit, and a check keyed on a
    ``drawn`` or ``numeral_audit`` field would report every sheet the applicant has not uploaded
    yet as a FAILED pixel review - the worst possible reading of "you have not sent us your
    drawings". The specifications go to the checks exactly as the agent wrote them.
    """
    assert not hasattr(draft_studio, "figures_for_qa")
    names = set(checks_for(figures=FIGURES))
    assert "Drawing pixels were inspected" not in names
    assert "Drawing content matches its specification" not in names
    assert "Drawing leaders identify the named features" not in names
    #  The checks that are about the TEXT of the drawings still run, because a figure described
    #  and never listed, or listed and never described, is a real defect in the application.
    assert "Every figure used is described" in names
    assert "Every drawing sheet is described" in names


def test_a_missing_drawing_does_not_block_a_filing_ready_version():
    """A sheet the applicant has not supplied is their calendar, not our gate.

    While the product drew its own figures, a drawing defect was ours and blocking on it was
    right. It uploads them now, so refusing to call the TEXT filing-ready because FIG. 3 has not
    arrived would be the product holding a draft hostage over a file it cannot produce itself.
    The finding still appears in Review; it just does not veto the version.
    """
    report = {
        "status": "complete",
        "checks": [
            {"name": "Every drawing sheet is described", "status": "fail",
             "category": "figures_and_numerals", "detail": "FIG. 3 has no sheet."},
            {"name": "Every section is written", "status": "fail",
             "detail": "The Abstract is empty."},
        ],
        "findings": [],
    }
    blockers = draft_studio.filing_blockers(report)
    assert any("Every section is written" in item for item in blockers)
    assert not any("Every drawing sheet is described" in item for item in blockers)
    #  And the drawing side is still classified, so Review can show it under its own heading.
    assert draft_studio.drawing_blockers(report)


def test_measurements_and_years_are_not_read_as_reference_numerals():
    """The generous failure here would report every dimension as an undefined part."""
    text = ("The body 12 is 45 mm long, weighs 2.5 kg, holds 90 percent vacuum at 25 degrees, "
            "and was first described in 2019. See FIG. 3 and claim 7.")
    assert set(draft_qa.numerals_used(text)) == {"12"}


def test_figure_series_are_not_read_as_reference_numerals():
    text = (
        "The body 12 includes regions enlarged in FIGS. 3, 4 and 5. "
        "Other regions appear in FIGS. 6, 7, and 8 and FIG. 9."
    )

    assert set(draft_qa.numerals_used(text)) == {"12"}


def test_alphanumeric_cutting_line_marks_are_not_read_as_reference_numerals():
    text = (
        "FIG. 2A is taken on line 2A-2A of FIG. 1 and shows the body 12 in section."
    )

    assert set(draft_qa.numerals_used(text)) == {"12"}


# =============================================================================================
# Figures
# =============================================================================================
def test_a_figure_used_in_the_description_but_never_described_is_caught():
    broken = dict(GOOD)
    broken["detailed_description"] += " FIG. 4 shows an alternative pump."
    check = checks_for(broken)["Every figure used is described"]
    assert check["status"] == "fail" and "FIG. 4" in check["items"]


def test_a_figure_range_counts_every_figure_in_it():
    assert draft_qa.figures_mentioned("As shown in FIGS. 1-3 and FIG. 7") == {"1", "2", "3", "7"}


# =============================================================================================
# Claims
# =============================================================================================
def test_a_claim_depending_on_a_claim_that_does_not_exist_is_caught():
    broken = dict(GOOD)
    broken["claims"] += "\n\n4. The vacuum lifting tool of claim 9, wherein the body is aluminium."
    check = checks_for(broken)["Claim dependencies are valid"]
    assert check["status"] == "fail" and "claim 9" in " ".join(check["items"])


def test_a_claim_depending_on_a_later_claim_is_caught():
    broken = dict(GOOD)
    broken["claims"] = broken["claims"].replace(
        "2. The vacuum lifting tool of claim 1,", "2. The vacuum lifting tool of claim 3,")
    assert checks_for(broken)["Claim dependencies are valid"]["status"] == "fail"


def test_claims_must_be_numbered_consecutively_from_one():
    broken = dict(GOOD)
    broken["claims"] = broken["claims"].replace("3. The vacuum", "5. The vacuum")
    assert checks_for(broken)["Claims are numbered consecutively"]["status"] == "fail"


def test_claims_stay_within_the_standard_uspto_count_before_excess_fees():
    broken = dict(GOOD)
    claims = [f"{number}. A distinct apparatus comprising a body and a pump."
              for number in range(1, 5)]
    claims.extend(
        f"{number}. The apparatus of claim 1, wherein the body includes a groove."
        for number in range(5, 22))
    broken["claims"] = "\n\n".join(claims)

    check = checks_for(broken)["Standard USPTO claim count"]

    assert check["status"] == "fail"
    assert any("21 total claims" in item for item in check["items"])
    assert any("4 independent claims" in item for item in check["items"])


def test_twenty_total_and_three_independent_claims_need_no_excess_claim_fee():
    clean = dict(GOOD)
    claims = [f"{number}. A distinct apparatus comprising a body and a pump."
              for number in range(1, 4)]
    claims.extend(
        f"{number}. The apparatus of claim 1, wherein the body includes a groove."
        for number in range(4, 21))
    clean["claims"] = "\n\n".join(claims)

    assert checks_for(clean)["Standard USPTO claim count"]["status"] == "pass"


def test_method_claim_rejects_a_gerund_mixed_into_coordinated_base_verbs():
    broken = dict(GOOD)
    broken["claims"] = """
    1. A method of controlling charging, the method comprising:
    at each allocation interval, subtract a measured load from a branch limit, assign a
    sustaining current to each connector, and distributing a remainder among the connectors.
    """

    check = checks_for(broken)["Method claim steps use parallel verb forms"]

    assert check["status"] == "fail"
    assert any("claim 1" in item and "distributing" in item for item in check["items"])


def test_dependent_method_claim_rejects_mixed_verbs_without_serial_comma():
    broken = dict(GOOD)
    broken["claims"] = """
    1. A method for managing charging, the method comprising:
    detecting an overcurrent condition; and
    opening a contactor in response to the overcurrent condition.

    2. The method of claim 1, further comprising:
    determining that a welded-contactor condition exists; and
    in response to determining the welded-contactor condition, inhibiting operation of other
    connector assemblies and activate a fault indicator.
    """

    check = checks_for(broken)["Method claim steps use parallel verb forms"]

    assert check["status"] == "fail"
    assert any(
        "claim 2" in item and "inhibiting" in item and "activate" in item
        for item in check["items"])


def test_filing_prose_rejects_third_person_and_base_verb_coordination():
    broken = dict(GOOD)
    broken["detailed_description"] += (
        " In response, the controller inhibits operation of the other connector assemblies "
        "and activate a fault indicator."
    )

    check = checks_for(broken)["Filing prose uses parallel coordinated verbs"]

    assert check["status"] == "fail"
    assert any(
        "inhibits" in item and "activate" in item for item in check["items"])
    clean = dict(broken)
    clean["detailed_description"] = clean["detailed_description"].replace(
        "and activate a fault indicator", "and activates a fault indicator")
    assert checks_for(clean)["Filing prose uses parallel coordinated verbs"]["status"] == "pass"


def test_method_claim_accepts_coordinated_base_verbs_with_matching_forms():
    clean = dict(GOOD)
    clean["claims"] = """
    1. A method of controlling charging, the method comprising:
    at each allocation interval, subtract a measured load from a branch limit, assign a
    sustaining current to each connector, and distribute a remainder among the connectors;
    sending the assigned limit through a control-pilot interface, and verifying current with
    a sensor.
    """

    check = checks_for(clean)["Method claim steps use parallel verb forms"]

    assert check["status"] == "pass"


def test_a_multiple_dependent_claim_on_another_is_a_rule_violation():
    broken = dict(GOOD)
    broken["claims"] += ("\n\n4. The vacuum lifting tool of any one of claims 1 or 2, wherein the "
                         "groove is annular."
                         "\n\n5. The vacuum lifting tool of any one of claims 3 or 4, wherein the "
                         "body is aluminium.")
    checks = checks_for(broken)
    assert checks["No multiple dependent claim depends on another"]["status"] == "fail"


def test_claim_dependency_parsing_handles_the_forms_attorneys_write():
    assert draft_qa.claim_dependencies("The tool of claim 1, wherein") == [1]
    assert draft_qa.claim_dependencies("The tool according to claim 2") == [2]
    assert draft_qa.claim_dependencies("The tool as set forth in claim 3") == [3]
    assert draft_qa.claim_dependencies("The tool of any one of claims 1 to 3") == [1, 2, 3]
    assert draft_qa.claim_dependencies("The tool of claims 1 or 4") == [1, 4]
    assert draft_qa.claim_dependencies("A tool comprising a body") == []


def test_antecedent_basis_is_advisory_and_finds_a_real_miss():
    broken = dict(GOOD)
    broken["claims"] = ("1. A vacuum lifting tool comprising a body and a pump carried by the "
                        "body, wherein the sealing ring is received in the body.")
    check = checks_for(broken)["Antecedent basis in the claims"]
    assert check["severity"] == "advisory"
    assert check["status"] == "warn" and any("sealing ring" in i for i in check["items"])


def test_a_claim_term_absent_from_the_description_is_flagged_but_cannot_fail_the_draft():
    broken = dict(GOOD)
    broken["claims"] += "\n\n4. The vacuum lifting tool of claim 1, wherein the body is titanium."
    check = checks_for(broken)["Claim terms appear in the description"]
    assert check["severity"] == "advisory" and check["status"] == "warn"
    assert any("titanium" in item for item in check["items"])


def test_claim_support_ignores_operable_and_further_as_boilerplate():
    broken = dict(GOOD)
    broken["claims"] += ("\n\n4. The vacuum lifting tool of claim 1, wherein the pump is operable "
                         "to draw air, and further comprising the passage.")
    check = checks_for(broken)["Claim terms appear in the description"]
    reported = " ".join(check.get("items") or [])
    assert "operable" not in reported
    assert "further" not in reported


def test_claim_support_ignores_dependency_vocabulary():
    check = checks_for()["Claim terms appear in the description"]
    reported = " ".join(check.get("items") or [])
    assert "“claim”" not in reported
    assert "“claims”" not in reported


# =============================================================================================
# Citations
# =============================================================================================
def test_a_citation_that_resolves_to_nothing_is_caught():
    broken = dict(GOOD)
    broken["background"] += " A second lifter is described in [REF:US-9999999-B9]."
    check = checks_for(broken, allowed=ALLOWED + ["US-9999999-B9"])[
        "Every citation resolves to a real publication"]
    assert check["status"] == "fail" and "US-9999999-B9" in check["items"][0]


def test_a_citation_to_a_reference_the_project_was_never_given_is_caught():
    broken = dict(GOOD)
    broken["background"] += " See [REF:US-11223344-B2] and [REF:EP-1234567-A1]."
    check = checks_for(broken)["Citations are to supplied references"]
    assert check["status"] == "fail" and "EP-1234567-A1" in check["items"]


def test_a_citation_in_the_claims_is_a_drafting_error():
    broken = dict(GOOD)
    broken["claims"] = broken["claims"].replace(
        "1. A vacuum", "1. A vacuum lifter improving on [REF:US-11223344-B2]. A vacuum")
    check = checks_for(broken)["Citations sit where they belong"]
    assert check["status"] == "fail" and "Claims" in check["items"][0]


def test_a_malformed_citation_token_is_caught_rather_than_normalised():
    broken = dict(GOOD)
    broken["background"] += " See [REF:the Smith patent]."
    assert checks_for(broken)["Citation tokens are well formed"]["status"] == "fail"


def test_a_publication_number_written_without_a_token_is_surfaced():
    broken = dict(GOOD)
    broken["background"] += " US 7,654,321 B1 describes a related tool."
    check = checks_for(broken)["Publication numbers use citation tokens"]
    assert check["status"] == "warn" and "US-7654321-B1" in check["items"][0]


def test_supplied_art_that_is_never_cited_is_a_filing_error():
    check = checks_for(allowed=ALLOWED + ["US-8888888-B2"])["Supplied art is addressed"]
    assert check["status"] == "fail" and check["severity"] == "error"
    assert "US-8888888-B2" in check["items"]


def test_uncited_supplied_art_is_refused_before_any_image_call():
    snapshot = {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}

    with pytest.raises(
            draft_studio.FilingPreflightError,
            match="Supplied art is addressed") as caught:
        draft_studio.validate_snapshot(snapshot, ALLOWED + ["US-8888888-B2"])
    assert caught.value.category == "internal_logic"


def test_a_supplied_art_set_with_no_citations_is_a_filing_error():
    sections = {**GOOD, "background": "Known handheld lifting tools use fixed seals."}

    check = checks_for(sections=sections)["Prior art is cited"]

    assert check["status"] == "fail" and check["severity"] == "error"


def test_citation_token_extraction():
    assert draft_cite.citations_in("a [REF:US-1-A] b [REF:EP-2222222-A1]") == ["US-1-A",
                                                                              "EP-2222222-A1"]
    assert draft_cite.malformed_citations_in("see [REF:the thing]") == ["the thing"]


# =============================================================================================
# Formalities
# =============================================================================================
def test_an_abstract_over_the_word_cap_fails():
    broken = dict(GOOD)
    broken["abstract"] = " ".join(["word"] * 200)
    check = checks_for(broken)["Abstract is in filing form"]
    assert check["status"] == "fail" and "200 words" in check["detail"]


def test_abstract_word_cap_counts_hyphenated_compounds_conservatively():
    broken = dict(GOOD)
    broken["abstract"] = " ".join(["word"] * 148 + ["air-extraction", "low-friction"])

    check = checks_for(broken)["Abstract is in filing form"]

    assert check["status"] == "fail"
    assert "152 words" in check["detail"]


def test_an_abstract_in_two_paragraphs_is_a_warning_not_a_failure():
    broken = dict(GOOD)
    broken["abstract"] = "First paragraph of the abstract.\n\nSecond paragraph."
    check = checks_for(broken)["Abstract is in filing form"]
    assert check["status"] == "warn" and "single paragraph" in check["detail"]


def test_a_puffed_title_is_objected_to():
    broken = dict(GOOD)
    broken["title"] = "New and Improved Vacuum Lifting Tool"
    assert checks_for(broken)["Title is in filing form"]["status"] == "warn"


def test_an_empty_section_fails():
    broken = dict(GOOD)
    broken["summary"] = ""
    assert checks_for(broken)["Every section is written"]["status"] == "fail"


def test_missing_government_support_fails_the_filing_preflight():
    broken = dict(GOOD)
    del broken["government_support"]

    with pytest.raises(drafting.DraftingValidationError,
                       match="Statement Regarding Federally Sponsored Research or Development"):
        draft_studio.validate_snapshot(
            {"sections": broken, "numerals": NUMERALS, "figures": FIGURES}, ALLOWED)


def test_open_drafting_notes_are_a_filing_blocker():
    broken = dict(GOOD)
    broken["detailed_description"] += " [DRAFTING NOTE: confirm the ring material.]"
    check = checks_for(broken)["No unresolved drafting notes"]
    assert check["status"] == "fail" and check["severity"] == "error"
    assert "ring material" in check["items"][0]


@pytest.mark.parametrize("placeholder", [
    "[TODO: add dimensions]", "TBD", "TO BE PROVIDED", "<INSERT MATERIAL>",
    "{{applicant_name}}", "_______", "Part names are for the draftsperson only.",
    "[VERIFY: confirm the material]", "Confirm with the inventor before filing.",
    "The applicant should provide the missing dimension.",
    "Manually add the connector label.", "Human intervention is required.",
])
def test_every_placeholder_form_is_a_filing_blocker(placeholder):
    broken = dict(GOOD)
    broken["summary"] += " " + placeholder
    check = checks_for(broken)["No unresolved drafting notes"]
    assert check["status"] == "fail" and check["items"]


# =============================================================================================
# Workspace
# =============================================================================================
def test_sections_survive_a_write_and_read_round_trip(tmp_path):
    draft_workspace.write_sections(tmp_path, GOOD)
    assert draft_workspace.read_sections(tmp_path) == GOOD


def test_workspace_refresh_removes_noncanonical_section_aliases(tmp_path):
    draft_workspace.write_sections(tmp_path, GOOD)
    draft_workspace.write_numerals(tmp_path, NUMERALS)
    draft = tmp_path / "draft"
    (draft / "08-claims.md").write_text("Wrong claim file", encoding="utf-8")
    (draft / "10-end.md").write_text("Wrong terminal file", encoding="utf-8")

    draft_workspace.write_sections(tmp_path, GOOD)

    assert not (draft / "08-claims.md").exists()
    assert not (draft / "10-end.md").exists()
    assert (draft / "09-claims.md").read_text(encoding="utf-8").strip() == GOOD["claims"]
    assert (draft / "numerals.md").exists()


def test_snapshot_rejects_a_noncanonical_section_file_created_during_the_turn(tmp_path):
    draft_workspace.write_sections(tmp_path, GOOD)
    draft_workspace.write_numerals(tmp_path, NUMERALS)
    draft_workspace.write_figures(tmp_path, FIGURES)
    alias = tmp_path / "draft" / "08-claims.md"
    alias.write_text(
        "Claims written under the wrong filename", encoding="utf-8")

    with pytest.raises(drafting.DraftingValidationError, match="08-claims.md.*09-claims.md"):
        draft_workspace.snapshot(tmp_path)
    assert not alias.exists()


def test_government_support_has_a_standalone_workspace_file_in_filing_order(tmp_path):
    draft_workspace.write_sections(tmp_path, GOOD)

    keys = [key for key, _name, _heading in draft_workspace.SECTION_FILES]
    assert keys.index("cross_reference") < keys.index("government_support") < keys.index("field")
    assert [name for _key, name, _heading in draft_workspace.SECTION_FILES] == [
        "01-title.md", "02-cross-reference.md", "03-government-support.md", "04-field.md",
        "05-background.md", "06-summary.md", "07-drawings.md",
        "08-detailed-description.md", "09-claims.md", "10-abstract.md",
    ]
    path = tmp_path / "draft" / "03-government-support.md"
    assert path.read_text(encoding="utf-8") == "Not applicable.\n"

    path.write_text(
        "## Statement Regarding Federally Sponsored Research or Development\n\n"
        "Not applicable.\n",
        encoding="utf-8")
    assert draft_workspace.read_sections(tmp_path)["government_support"] == "Not applicable."

    path.write_text(
        "STATEMENT REGARDING FEDERALLY SPONSORED RESEARCH OR DEVELOPMENT\n\n"
        "Not applicable.\n",
        encoding="utf-8")
    assert draft_workspace.read_sections(tmp_path)["government_support"] == "Not applicable."


def test_legacy_workspace_section_files_are_migrated_without_mixing_bodies(tmp_path):
    draft = tmp_path / "draft"
    draft.mkdir()
    legacy = {
        "01-title.md": "Legacy title",
        "02-cross-reference.md": "Legacy cross reference",
        "03-field.md": "Legacy field",
        "04-background.md": "Legacy background",
        "05-summary.md": "Legacy summary",
        "06-drawings.md": "Legacy drawings",
        "07-detailed-description.md": "Legacy detail",
        "08-claims.md": "Legacy claims",
        "09-abstract.md": "Legacy abstract",
        "10-government-support.md": "Legacy government support",
    }
    for name, body in legacy.items():
        (draft / name).write_text(body + "\n", encoding="utf-8")

    sections = draft_workspace.read_sections(tmp_path)

    assert sections == {
        "title": "Legacy title",
        "cross_reference": "Legacy cross reference",
        "government_support": "Legacy government support",
        "field": "Legacy field",
        "background": "Legacy background",
        "summary": "Legacy summary",
        "drawing_descriptions": "Legacy drawings",
        "detailed_description": "Legacy detail",
        "claims": "Legacy claims",
        "abstract": "Legacy abstract",
    }
    assert (draft / "03-government-support.md").read_text(encoding="utf-8") == \
        "Legacy government support\n"
    assert (draft / "10-abstract.md").read_text(encoding="utf-8") == "Legacy abstract\n"
    assert not (draft / "03-field.md").exists()
    assert not (draft / "10-government-support.md").exists()

def test_a_heading_the_agent_added_back_is_dropped_but_real_headings_survive(tmp_path):
    draft_workspace.write_sections(tmp_path, GOOD)
    path = tmp_path / "draft" / "05-background.md"
    path.write_text("## Background\n\nHandheld vacuum lifters are known.\n", encoding="utf-8")
    detail = tmp_path / "draft" / "08-detailed-description.md"
    detail.write_text("## The pump\n\nThe pump 14 draws air.\n", encoding="utf-8")
    out = draft_workspace.read_sections(tmp_path)
    assert out["background"] == "Handheld vacuum lifters are known."
    assert out["detailed_description"].startswith("## The pump")


def test_the_numeral_table_round_trips(tmp_path):
    draft_workspace.write_numerals(tmp_path, NUMERALS)
    assert draft_workspace.read_numerals(tmp_path) == NUMERALS


def test_the_numeral_reader_accepts_plain_agent_list_formats(tmp_path):
    path = tmp_path / "draft" / "numerals.md"
    path.parent.mkdir(parents=True)
    path.write_text(
        "# Reference numerals\n\n"
        "10, cartridge body\n"
        "12: valve seat\n"
        "- 14 - compression spring\n",
        encoding="utf-8")

    assert draft_workspace.read_numerals(tmp_path) == [
        {"numeral": "10", "part": "cartridge body"},
        {"numeral": "12", "part": "valve seat"},
        {"numeral": "14", "part": "compression spring"},
    ]


def test_workspace_draft_files_contain_no_process_scaffolding(tmp_path):
    draft_workspace.write_sections(tmp_path, GOOD)
    draft_workspace.write_numerals(tmp_path, NUMERALS)

    text = "\n".join(
        path.read_text(encoding="utf-8")
        for path in sorted((tmp_path / "draft").glob("*.md")))

    assert "<!--" not in text
    assert "checked mechanically" not in text
    assert "Keep this table" not in text


def test_a_valid_numeral_table_may_carry_extra_audit_columns(tmp_path):
    directory = tmp_path / "draft"
    directory.mkdir()
    (directory / "numerals.md").write_text(
        "# Reference numerals\n\n"
        "| Numeral | Part | First introduced |\n"
        "| --- | --- | --- |\n"
        "| 10 | vacuum lifting tool | FIG. 1 |\n"
        "| 12 | body | paragraph 12 |\n",
        encoding="utf-8")
    assert draft_workspace.read_numerals(tmp_path) == [
        {"numeral": "10", "part": "vacuum lifting tool"},
        {"numeral": "12", "part": "body"},
    ]


def test_figures_round_trip(tmp_path):
    figures = tmp_path / "figures"
    figures.mkdir()
    (figures / "agent-created.svg").write_text("<svg/>", encoding="utf-8")
    (figures / "rendered-stale.png").write_bytes(b"stale")
    (figures / "scratch").mkdir()
    (figures / "scratch" / "notes.md").write_text("not a figure", encoding="utf-8")
    draft_workspace.write_figures(tmp_path, FIGURES)
    out = draft_workspace.read_figures(tmp_path)
    assert [f["label"] for f in out] == ["FIG. 1", "FIG. 2"]
    assert out[1]["numerals"] == ["16 sealing ring", "18 groove", "20 passage"]
    assert not (figures / "agent-created.svg").exists()
    assert not (figures / "rendered-stale.png").exists()
    assert not (figures / "scratch").exists()


def test_figures_round_trip_keeps_numeric_order_after_figure_nine(tmp_path):
    figures = [
        {"label": f"FIG. {number}", "caption": f"View {number}.", "numerals": []}
        for number in range(1, 11)
    ]

    draft_workspace.write_figures(tmp_path, figures)

    assert [item["label"] for item in draft_workspace.read_figures(tmp_path)] == [
        f"FIG. {number}" for number in range(1, 11)
    ]


def test_figure_snapshot_removes_and_rejects_a_filename_alias_from_the_current_turn(tmp_path):
    draft_workspace.write_figures(tmp_path, FIGURES)
    alias = tmp_path / "figures" / "FIG-1-SHORT.md"
    alias.write_text("# FIG. 1\n\nAn aliased duplicate brief.\n", encoding="utf-8")

    with pytest.raises(drafting.DraftingValidationError, match="FIG-1-SHORT.md.*FIG-1.md"):
        draft_workspace.read_figures(tmp_path)

    assert not alias.exists()
    assert (tmp_path / "figures" / "FIG-1.md").exists()


def test_figure_snapshot_allows_rendered_review_evidence(tmp_path):
    draft_workspace.write_figures(tmp_path, FIGURES)
    evidence = tmp_path / "figures" / "rendered-FIG-1.png"
    evidence.write_bytes(b"checked pixels")

    assert len(draft_workspace.read_figures(tmp_path)) == 2
    assert evidence.exists()


def test_long_figure_labels_round_trip_without_mid_word_truncation(tmp_path):
    label = (
        "FIG. 2 - Side elevation in vertical section showing the chamber and both "
        "air paths"
    )
    draft_workspace.write_figures(tmp_path, [{
        "label": label,
        "caption": "A complete drawing brief.",
        "numerals": ["10 body"],
    }])

    assert draft_workspace.read_figures(tmp_path)[0]["label"] == label


def test_an_existing_draft_is_split_on_its_headings():
    document = ("Vacuum Lifting Tool\n\nBACKGROUND OF THE INVENTION\n\nLifters are known.\n\n"
                "SUMMARY OF THE INVENTION\n\nA better lifter.\n\n"
                "DETAILED DESCRIPTION\n\nThe tool 10 has a body 12.\n\n"
                "WHAT IS CLAIMED IS:\n\n1. A tool.\n\nABSTRACT\n\nA tool.")
    out = draft_workspace.seed_sections_from_document(document)
    assert out["title"] == "Vacuum Lifting Tool"
    assert out["background"] == "Lifters are known."
    assert out["claims"] == "1. A tool."
    assert "body 12" in out["detailed_description"]


def test_an_existing_draft_preserves_a_government_support_statement():
    document = (
        "Vacuum Tool\n\n"
        "CROSS-REFERENCE TO RELATED APPLICATIONS\n\nNot applicable.\n\n"
        "STATEMENT REGARDING FEDERALLY SPONSORED RESEARCH OR DEVELOPMENT\n\n"
        "This invention was made with support under Award 123.\n\n"
        "FIELD OF THE INVENTION\n\nThe disclosure relates to tools."
    )

    out = draft_workspace.seed_sections_from_document(document)

    assert out["cross_reference"] == "Not applicable."
    assert out["government_support"] == "This invention was made with support under Award 123."
    assert out["field"] == "The disclosure relates to tools."


def test_a_document_with_no_recognisable_headings_is_kept_whole_not_dropped():
    out = draft_workspace.seed_sections_from_document("Just a paragraph about a tool.")
    assert out == {"detailed_description": "Just a paragraph about a tool."}


def test_numerals_are_harvested_from_an_existing_draft():
    found = {item["numeral"]: item["part"]
             for item in draft_workspace.numerals_from_sections(
                 {"detailed_description": "The vacuum lifting tool 10 has a body 12."})}
    assert found["10"].endswith("tool") and found["12"] == "body"


def test_the_workspace_lives_outside_the_home_directory_or_says_why():
    #  Claude Code collects CLAUDE.md files walking up from its working directory; a workspace
    #  under the operator's home would inherit that box's own operating instructions.
    assert draft_workspace.DEFAULT_ROOT.startswith("/srv") or \
        "DRAFT_WORKSPACE_ROOT" in Path(draft_workspace.__file__).read_text()


def test_the_lookup_tool_is_written_with_a_resolvable_source_path(tmp_path):
    draft_workspace.install_tools(tmp_path, src_dir="/opt/app/src")
    body = (tmp_path / "tools" / "patent_lookup.py").read_text()
    assert '"/opt/app/src"' in body and "draft_cite" in body
    compile(body, "patent_lookup.py", "exec")


# =============================================================================================
# The validator that decides whether a turn becomes a version
# =============================================================================================
def test_a_good_draft_validates():
    assert draft_studio.validate_sections(GOOD, ALLOWED)["title"] == GOOD["title"]


def test_an_empty_section_is_refused():
    with pytest.raises(Exception) as caught:
        draft_studio.validate_sections({**GOOD, "claims": ""}, ALLOWED)
    assert "Claims" in str(caught.value)


def test_a_citation_to_an_unsupplied_reference_is_refused():
    broken = {**GOOD, "background": GOOD["background"] + " [REF:US-4444444-A]"}
    with pytest.raises(Exception) as caught:
        draft_studio.validate_sections(broken, ALLOWED)
    assert "US-4444444-A" in str(caught.value)


def test_a_citation_is_refused_when_the_project_has_no_supplied_references():
    broken = {**GOOD, "background": GOOD["background"] + " [REF:US-11223344-B2]"}
    with pytest.raises(Exception) as caught:
        draft_studio.validate_sections(broken, [])
    assert "not among this project's sources" in str(caught.value)


def test_a_legal_conclusion_is_refused():
    broken = {**GOOD, "summary": GOOD["summary"] + " The claimed tool is clearly non-obvious."}
    with pytest.raises(Exception) as caught:
        draft_studio.validate_sections(broken, ALLOWED)
    assert "legal conclusion" in str(caught.value)


def test_a_placeholder_is_refused_before_a_version_can_be_saved():
    broken = {**GOOD, "cross_reference": "[DRAFTING NOTE: confirm priority.]"}
    with pytest.raises(Exception) as caught:
        draft_studio.validate_sections(broken, ALLOWED)
    assert "placeholder" in str(caught.value).lower()


@pytest.mark.parametrize("field,value", [
    ("numerals", [{"numeral": "10", "part": "TBD component"}]),
    ("figures", [{"label": "FIG. 1", "caption": "<INSERT VIEW>", "numerals": ["10"]}]),
    ("figures", [{"label": "FIG. 1",
                  "caption": "Part names are for the draftsperson only.",
                  "numerals": ["10"]}]),
])
def test_placeholders_outside_the_sections_are_refused_before_save(field, value):
    snapshot = {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}
    snapshot[field] = value
    with pytest.raises(draft_studio.FilingPreflightError, match="placeholder") as caught:
        draft_studio.validate_snapshot(snapshot, ALLOWED)
    assert caught.value.category == "figures_and_numerals"


def test_overcrowded_sheet_is_refused_before_any_image_call():
    snapshot = {
        "sections": GOOD,
        "numerals": [
            {"numeral": str(value), "part": f"part {value}"}
            for value in range(10, 28, 2)
        ],
        "figures": [{
            "label": "FIG. 1", "caption": "assembly view",
            "numerals": [str(value) for value in range(10, 28, 2)],
        }],
    }

    with pytest.raises(draft_studio.FilingPreflightError,
                       match="more than 8 numerals") as caught:
        draft_studio.validate_snapshot(snapshot, ALLOWED)
    assert caught.value.category == "figures_and_numerals"


def test_redundant_leaf_labels_are_removed_from_an_overcrowded_overview():
    numerals = [
        {"numeral": "10", "part": "rail"},
        {"numeral": "16", "part": "longitudinal slot"},
        {"numeral": "18", "part": "distance marks"},
        {"numeral": "30", "part": "fence"},
        {"numeral": "50", "part": "first guide carriage"},
        {"numeral": "54", "part": "drill bushing of the first guide carriage"},
        {"numeral": "58", "part": "clamp knob of the first guide carriage"},
        {"numeral": "70", "part": "second guide carriage"},
        {"numeral": "74", "part": "drill bushing of the second guide carriage"},
        {"numeral": "78", "part": "clamp knob of the second guide carriage"},
    ]
    overview = {
        "label": "FIG. 2",
        "caption": (
            "A top view with a drill bushing 54 and clamp knob 58 on the first guide "
            "carriage 50, and a drill bushing 74 and clamp knob 78 on the second guide "
            "carriage 70. A cutting-plane line 5-5 establishes FIG. 5, and line 8-8 "
            "establishes FIG. 8. The disclosed calibration target contains 54 apertures."
        ),
        "numerals": [item["numeral"] for item in numerals],
    }
    focused_first = {
        "label": "FIG. 5", "caption": "A first-carriage sectional view.",
        "numerals": ["10", "16", "50", "54", "58"],
    }
    focused_second = {
        "label": "FIG. 8", "caption": "A second-carriage sectional view.",
        "numerals": ["10", "16", "70", "74", "78"],
    }

    repaired, changes = draft_studio.normalize_overcrowded_figure_plans({
        "sections": GOOD,
        "numerals": numerals,
        "figures": [overview, focused_first, focused_second],
    })

    repaired_overview = repaired["figures"][0]
    assert repaired_overview["numerals"] == [
        "10", "16", "18", "30", "50", "58", "70", "78",
    ]
    assert "bushing 54" not in repaired_overview["caption"]
    assert "bushing 74" not in repaired_overview["caption"]
    assert "a drill bushing and clamp knob 58" in repaired_overview["caption"]
    assert "a drill bushing and clamp knob 78" in repaired_overview["caption"]
    assert "line 5-5 establishes FIG. 5" in repaired_overview["caption"]
    assert "line 8-8 establishes FIG. 8" in repaired_overview["caption"]
    assert "contains 54 apertures" in repaired_overview["caption"]
    assert changes == ["FIG. 2: moved redundant labels 54, 74 to focused sheets"]
    assert {
        draft_qa._drawing_numeral(value)
        for figure in repaired["figures"] for value in figure["numerals"]
    } == {item["numeral"] for item in numerals}


def test_overcrowded_plan_is_left_for_agent_repair_without_a_focused_label_view():
    numerals = [
        {"numeral": str(value), "part": f"assembly part {value}"}
        for value in range(10, 28, 2)
    ]
    entries = [item["numeral"] for item in numerals]
    snapshot = {
        "sections": GOOD,
        "numerals": numerals,
        "figures": [
            {"label": "FIG. 1", "caption": "A top plan view.", "numerals": entries},
            {"label": "FIG. 2", "caption": "A side elevation view.", "numerals": entries},
        ],
    }

    repaired, changes = draft_studio.normalize_overcrowded_figure_plans(snapshot)

    assert repaired["figures"] == snapshot["figures"]
    assert changes == []


def test_overcrowded_plan_never_drops_one_unpaired_label_by_heuristic():
    numerals = [
        {"numeral": str(value), "part": f"assembly part {value}"}
        for value in range(10, 28, 2)
    ]
    entries = [item["numeral"] for item in numerals]
    snapshot = {
        "sections": GOOD,
        "numerals": numerals,
        "figures": [
            {"label": "FIG. 1", "caption": "A top plan view.", "numerals": entries},
            {"label": "FIG. 2", "caption": "A focused detail view.", "numerals": ["10"]},
        ],
    }

    repaired, changes = draft_studio.normalize_overcrowded_figure_plans(snapshot)

    assert repaired["figures"] == snapshot["figures"]
    assert changes == []


def test_a_text_numeral_missing_from_every_figure_is_refused_before_any_image_call():
    figures = [dict(FIGURES[0]), {**FIGURES[1], "numerals": ["16", "18"]}]
    snapshot = {"sections": GOOD, "numerals": NUMERALS, "figures": figures}

    with pytest.raises(
            draft_studio.FilingPreflightError,
            match="Every specification numeral appears in a drawing") as caught:
        draft_studio.validate_snapshot(snapshot, ALLOWED)
    assert caught.value.category == "figures_and_numerals"


def test_agent_filing_artifacts_are_normalized_for_human_facing_text():
    numerals = [dict(item) for item in NUMERALS]
    numerals[0]["part"] = "vacuum lifting tool\u2014assembly"
    figures = [{**item, "numerals": list(item["numerals"])} for item in FIGURES]
    figures[0]["caption"] = "side\u2014elevation"
    snapshot = {
        "sections": {**GOOD, "summary": GOOD["summary"] + " One view\u2014shown below."},
        "numerals": numerals,
        "figures": figures,
    }
    clean = draft_studio.validate_snapshot(snapshot, ALLOWED)
    assert "\u2014" not in json.dumps(clean, ensure_ascii=False)
    assert "One view - shown below." in clean["sections"]["summary"]
    assert clean["numerals"][0]["part"] == "vacuum lifting tool - assembly"
    assert clean["figures"][0]["caption"] == "side - elevation"


def test_a_citation_in_the_detailed_description_is_allowed_here():
    """Incorporation by reference is real practice; the reviewer judges it, the validator does not
    throw fifteen minutes of drafting away over it."""
    ok = {**GOOD, "detailed_description":
          GOOD["detailed_description"] + " US-11223344-B2 is incorporated by reference "
          "[REF:US-11223344-B2]."}
    assert draft_studio.validate_sections(ok, ALLOWED)


def test_citations_are_collected_in_first_use_order():
    sections = {**GOOD, "detailed_description":
                GOOD["detailed_description"] + " [REF:US-11223344-B2]"}
    assert draft_studio.citations_of(sections) == ["US-11223344-B2"]


def test_the_rendered_markdown_keeps_the_application_order():
    rendered = draft_studio.render_markdown(GOOD)
    positions = [rendered.index(heading) for _k, _n, heading in draft_workspace.SECTION_FILES]
    assert positions == sorted(positions)


# =============================================================================================
# Filing
# =============================================================================================
def test_the_fee_profile_counts_a_multiple_dependent_claim_as_several():
    profile = draft_uspto.fee_profile(
        GOOD["claims"] + "\n\n4. The tool of any one of claims 1 to 3, wherein it is red.")
    assert profile["total"] == 4 and profile["independent"] == 1
    assert profile["multiple_dependent"] == 1 and profile["billable"] == 6
    assert any("multiple dependent" in s for s in profile["surcharges"])


def test_more_than_three_independent_claims_triggers_a_surcharge():
    claims = "\n\n".join(f"{n}. An apparatus comprising a part {n}." for n in range(1, 6))
    assert any("independent" in s for s in draft_uspto.fee_profile(claims)["surcharges"])


def checked_figures(*labels):
    labels = labels or tuple(spec["label"] for spec in FIGURES)
    out = []
    for sheet_index, label in enumerate(labels, 1):
        spec = next(item for item in FIGURES
                    if draft_figures.figure_key(item["label"]) == draft_figures.figure_key(label))
        expected = draft_figures.expected_entries(spec, NUMERALS)
        digest = draft_figures.specification_hash(spec["label"], spec["caption"], expected)
        out.append({"figure_label": label, "active_version": 1, "versions": [{
            "version_no": 1, "numeral_audit": {
                "ok": True, "inspected": True,
                "prompt_version": draft_figures.OCR_PROMPT_VERSION,
                "correct_figure_label": True,
                "expected_sheet_number": f"{sheet_index}/{len(labels)}",
                "detected_sheet_numbers": [f"{sheet_index}/{len(labels)}"],
                "correct_sheet_number": True,
                "expected_section_designations": [],
                "detected_section_designations": [],
                "correct_section_designations": True,
            },
            "leader_audit": {
                "ok": True, "inspected": True, "specification_hash": digest,
                "model_name": draft_figures.vision_model(),
                "prompt_version": draft_figures.LEADER_PROMPT_VERSION,
                "review_count": draft_figures.LEADER_REVIEW_COUNT,
                "section_mark_anchor_audit":
                    draft_figures._section_mark_anchor_audit([], []),
                "flowchart_topology_audit": {
                    "ok": True, "inspected": False, "required": False,
                    "model_name": "deterministic-parser",
                    "specification_hash": digest,
                    "prompt_version": draft_figures.FLOWCHART_TOPOLOGY_PROMPT_VERSION,
                    "review_count": 0, "expected": [], "observed": [],
                    "missing": [], "unexpected": [], "duplicates": [], "errors": [],
                },
            },
            "semantic_audit": {
                "ok": True, "inspected": True, "specification_hash": digest,
                "model_name": draft_figures.vision_model(),
                "prompt_version": draft_figures.SEMANTIC_PROMPT_VERSION,
                "review_count": draft_figures.SEMANTIC_REVIEW_COUNT,
                "pixel_anchor_audit": {
                    "ok": True, "inspected": True,
                    "version": draft_figures.PIXEL_ANCHOR_VERSION,
                },
                "topology_audit": {
                    "ok": True, "inspected": False, "required": False,
                    "version": draft_figures.CLOSED_REGION_AUDIT_VERSION,
                },
                "section_mark_audit": {
                    "ok": True, "inspected": False, "required": False,
                    "expected": [], "marks": [], "errors": [], "review_count": 0,
                    "model_name": "deterministic-parser",
                    "prompt_version": draft_figures.SECTION_MARK_PROMPT_VERSION,
                },
                "marked_anchor_audit": {
                    "ok": True, "inspected": True, "specification_hash": digest,
                    "model_name": draft_figures.vision_model(),
                    "prompt_version": draft_figures.MARKED_ANCHOR_PROMPT_VERSION,
                    "review_count": draft_figures.MARKED_ANCHOR_REVIEW_COUNT,
                },
            }}]})
    return out


def clean_version(version_no=1, sections=None):
    return {"version_no": version_no, "sections": sections or GOOD,
            "citations": ["US-11223344-B2"], "numerals": NUMERALS,
            "figure_specs": FIGURES}


def clean_qa(version_no=1):
    return {"version_no": version_no, "status": "complete", "verdict": "pass",
            "checks": [], "findings": []}


def test_readiness_blocks_on_an_open_drafting_note_and_on_a_failed_check():
    version = clean_version(sections={**GOOD, "summary":
                                    GOOD["summary"] + " [DRAFTING NOTE: get the ring material.]"})
    qa = {"version_no": 1, "status": "complete", "verdict": "fail",
          "checks": [{"name": "Every numeral in the text is defined",
                                         "status": "fail", "severity": "error",
                                         "detail": "22 is undefined", "items": ["22"]}],
          "findings": []}
    report = draft_uspto.readiness(project={"inventors": "Dana Drafter", "applicant": "Example"},
                                   version=version, qa=qa, figures=checked_figures())
    titles = [b["title"] for b in report["blockers"]]
    assert not report["ready"]
    assert any("drafting note" in t for t in titles)
    assert "Every numeral in the text is defined" in titles


def test_readiness_blocks_when_no_inventor_is_named():
    report = draft_uspto.readiness(project={"inventors": "", "applicant": ""},
                                   version=clean_version(),
                                   qa=clean_qa(), figures=checked_figures())
    assert any("inventor" in b["title"].lower() for b in report["blockers"])


def test_an_unreviewed_draft_is_never_reported_as_ready():
    report = draft_uspto.readiness(project={"inventors": "Dana", "applicant": "Example"},
                                   version=clean_version(),
                                   qa=None, figures=[{"id": 1}])
    assert not report["ready"]


def test_a_clean_draft_reports_no_blockers_but_still_lists_what_a_person_must_do():
    report = draft_uspto.readiness(project={"inventors": "Dana", "applicant": "Example"},
                                   version=clean_version(),
                                   qa=clean_qa(), figures=checked_figures())
    assert report["ready"] and len(report["remaining"]) >= 5
    assert any("oath or declaration" in item for item in report["remaining"])


def test_readiness_rechecks_the_active_drawing_instead_of_trusting_old_qa():
    figures = [{"figure_label": "FIG. 1", "active_version": 2, "versions": [{
        "version_no": 2, "numeral_audit": {"ok": True},
        "semantic_audit": {"ok": False}}]}]
    report = draft_uspto.readiness(
        project={"inventors": "Dana", "applicant": "Example"},
        version=clean_version(),
        qa=clean_qa(), figures=figures)
    assert not report["ready"]
    assert any("active drawings" in item["title"] for item in report["blockers"])


def test_readiness_rejects_pixels_that_are_not_bound_to_the_current_exact_renderer(
        monkeypatch):
    monkeypatch.setattr(
        draft_figures, "current_geometry_binding",
        lambda *_args, **_kwargs: False, raising=False)

    report = draft_uspto.readiness(
        project={"user_id": 7, "inventors": "Dana", "applicant": "Example"},
        version=clean_version(), qa=clean_qa(), figures=checked_figures())

    assert not report["ready"]
    assert any(
        "current deterministic geometry" in item["items"]
        for item in report["blockers"])


def test_readiness_rejects_a_wrong_or_stale_drawing_sheet_number():
    figures = checked_figures()
    figures[0]["versions"][0]["numeral_audit"]["expected_sheet_number"] = "1/3"
    figures[0]["versions"][0]["numeral_audit"]["detected_sheet_numbers"] = ["1/3"]

    report = draft_uspto.readiness(
        project={"inventors": "Dana", "applicant": "Example"},
        version=clean_version(), qa=clean_qa(), figures=figures)

    assert not report["ready"]
    assert any("sheet 1/2" in item["items"] for item in report["blockers"])


def test_readiness_blocks_a_sheet_without_final_leader_placement_approval():
    figures = checked_figures()
    figures[0]["versions"][0]["leader_audit"] = {
        "inspected": True, "ok": False, "errors": ["12 points to the body"]}
    report = draft_uspto.readiness(
        project={"inventors": "Dana", "applicant": "Example"},
        version=clean_version(), qa=clean_qa(), figures=figures)
    assert not report["ready"]
    assert any("leader placement" in item["items"] for item in report["blockers"])


def test_readiness_rejects_a_leader_review_from_an_older_gate():
    figures = checked_figures()
    figures[0]["versions"][0]["leader_audit"]["prompt_version"] = "old"
    report = draft_uspto.readiness(
        project={"inventors": "Dana", "applicant": "Example"},
        version=clean_version(), qa=clean_qa(), figures=figures)
    assert not report["ready"]
    assert any("leader" in item["items"] for item in report["blockers"])


def test_readiness_rejects_a_semantic_review_from_an_older_gate():
    figures = checked_figures()
    figures[0]["versions"][0]["semantic_audit"]["prompt_version"] = "old"
    report = draft_uspto.readiness(
        project={"inventors": "Dana", "applicant": "Example"},
        version=clean_version(), qa=clean_qa(), figures=figures)
    assert not report["ready"]
    assert any("semantic" in item["items"] for item in report["blockers"])


def test_readiness_requires_review_for_the_exact_exported_version():
    report = draft_uspto.readiness(
        project={"inventors": "Dana", "applicant": "Example"},
        version=clean_version(version_no=2),
        qa=clean_qa(version_no=1), figures=checked_figures())
    assert not report["ready"]
    assert any("exact version" in item["title"].lower() for item in report["blockers"])


def test_readiness_blocks_every_review_finding_and_nonpassing_check():
    qa = clean_qa() | {
        "verdict": "warn",
        "checks": [{"name": "Ambiguous term", "status": "warn", "severity": "advisory",
                    "detail": "The term may be unclear.", "items": []}],
        "findings": [{"severity": "minor", "title": "Description drift",
                      "detail": "One phrase differs.", "where": "summary"}],
    }
    report = draft_uspto.readiness(
        project={"inventors": "Dana", "applicant": "Example"},
        version=clean_version(),
        qa=qa, figures=checked_figures())
    titles = [item["title"] for item in report["blockers"]]
    assert not report["ready"]
    assert "Ambiguous term" in titles and "Description drift" in titles


def test_readiness_rejects_todos_and_missing_or_mismatched_drawing_audits():
    version = clean_version(sections={**GOOD, "summary": GOOD["summary"] + " TODO"})
    report = draft_uspto.readiness(
        project={"inventors": "Dana", "applicant": "Example"}, version=version,
        qa=clean_qa(), figures=[{"figure_label": "FIG. 2", "active_version": 0}])
    titles = [item["title"].lower() for item in report["blockers"]]
    assert not report["ready"]
    assert any("unfinished" in title for title in titles)
    assert any("drawing" in title for title in titles)


@pytest.mark.parametrize("field,value", [
    ("numerals", [{"numeral": "10", "part": "TBD component"}] + NUMERALS[1:]),
    ("figure_specs", [{**FIGURES[0], "caption": "[VERIFY: complete the view]"}] + FIGURES[1:]),
])
def test_readiness_rescans_versioned_drawing_sources_for_editorial_markers(field, value):
    version = {**clean_version(), field: value}
    report = draft_uspto.readiness(
        project={"inventors": "Dana", "applicant": "Example"},
        version=version, qa=clean_qa(), figures=checked_figures())
    assert not report["ready"]
    assert any("unfinished marker" in item["title"].lower()
               for item in report["blockers"])


def test_readiness_rejects_a_previous_sheet_for_a_changed_figure_specification():
    figures = checked_figures()
    figures[0]["versions"][0]["semantic_audit"]["specification_hash"] = "old-specification"
    report = draft_uspto.readiness(
        project={"inventors": "Dana", "applicant": "Example"},
        version=clean_version(), qa=clean_qa(), figures=figures)
    assert not report["ready"]
    assert any("different drawing specification" in item["items"]
               for item in report["blockers"])


def test_readiness_rejects_duplicate_active_sheet_numbers():
    figures = checked_figures()
    figures.append(dict(figures[0]))
    report = draft_uspto.readiness(
        project={"inventors": "Dana", "applicant": "Example"},
        version=clean_version(), qa=clean_qa(), figures=figures)
    assert not report["ready"]
    assert any("duplicate" in item["title"].lower() for item in report["blockers"])


def test_the_filing_text_numbers_its_paragraphs_and_orders_the_parts():
    text = draft_uspto.filing_text({"title": "x"}, {"sections": GOOD})
    assert "[0001]" in text
    assert text.index("BACKGROUND OF THE INVENTION") < text.index("CLAIMS")
    assert text.index("CLAIMS") < text.index("ABSTRACT OF THE DISCLOSURE")


def test_government_support_is_exported_once_in_required_filing_order():
    from docx import Document

    heading = "STATEMENT REGARDING FEDERALLY SPONSORED RESEARCH OR DEVELOPMENT"
    text = draft_uspto.filing_text({"title": "x"}, {"sections": GOOD})
    document = Document(draft_uspto.render_filing_docx(
        {"title": GOOD["title"]}, {"sections": GOOD},
        readiness_report={"blockers": []}))
    word_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    for exported in (text, word_text):
        assert exported.count(heading) == 1
        cross = exported.index("CROSS-REFERENCE TO RELATED APPLICATIONS")
        support = exported.index(heading)
        statement = exported.index("Not applicable.")
        field = exported.index("FIELD OF THE INVENTION")
        assert cross < support < statement < field


def test_legacy_candidate_missing_government_support_is_preserved_for_automatic_repair():
    legacy = {key: value for key, value in GOOD.items() if key != "government_support"}

    repaired = draft_studio.candidate_snapshot_for_repair({
        "sections": legacy, "numerals": NUMERALS, "figures": FIGURES,
    })

    assert repaired is not None
    assert repaired["sections"]["government_support"] == ""


def test_source_reviewer_reads_the_standalone_government_support_file():
    assert "input/brief.md" in draft_qa.SOURCE_REVIEW_PROMPT
    assert "authority only for filing formalities" in draft_qa.SOURCE_REVIEW_SYSTEM
    assert "draft/03-government-support.md" in draft_qa.SOURCE_REVIEW_PROMPT


def test_drafting_agent_has_a_filing_clean_default_for_no_government_support():
    assert "If no government support was supplied" in draft_studio.DRAFT_SYSTEM
    assert "government-support section" in draft_studio.DRAFT_SYSTEM


def test_deterministic_source_gate_rejects_an_unsupported_close_fit(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir(parents=True)
    (input_dir / "disclosure.md").write_text(
        "A key passes into a longitudinal slot and prevents rotation while allowing sliding.",
        encoding="utf-8",
    )
    (input_dir / "conversation.md").write_text(
        "The key is received in the slot.", encoding="utf-8")
    draft_workspace.write_figures(tmp_path, [{
        "label": "FIG. 5",
        "caption": "The key 52 has a width that fits closely within the slot 16.",
        "numerals": ["52 key", "16 slot"],
    }])

    findings = draft_qa.deterministic_source_fidelity_findings(tmp_path)

    assert len(findings) == 1
    assert findings[0]["where"].startswith("figures/FIG-5.md:")
    assert "close-fit" in findings[0]["title"].lower()
    assert "fits closely within" in findings[0]["evidence"]
    enforced = draft_qa.enforce_deterministic_source_fidelity({
        "status": "complete", "verdict": "pass", "summary": "Model review passed.",
        "checks": [], "findings": [], "counts": {},
    }, tmp_path)
    assert enforced["verdict"] == "fail"
    assert enforced["checks"][0]["status"] == "fail"
    assert enforced["findings"] == findings

    (input_dir / "disclosure.md").write_text(
        "A drill bushing has a press fit within a carriage.", encoding="utf-8")
    assert len(draft_qa.deterministic_source_fidelity_findings(tmp_path)) == 1

    (input_dir / "disclosure.md").write_text(
        "The key fits closely within the longitudinal slot.", encoding="utf-8")
    assert draft_qa.deterministic_source_fidelity_findings(tmp_path) == []

    draft_workspace.write_figures(tmp_path, [{
        "label": "FIG. 5",
        "caption": "The key 52 has a width that fits within the slot 16.",
        "numerals": ["52 key", "16 slot"],
    }])
    assert draft_qa.deterministic_source_fidelity_findings(tmp_path) == []


def test_deterministic_source_gate_rejects_a_review_that_misquotes_inventor_text(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir(parents=True)
    (input_dir / "disclosure.md").write_text(
        "The poppet returns to the seat after pressure falls below a lower closing pressure.",
        encoding="utf-8",
    )
    (input_dir / "conversation.md").write_text("", encoding="utf-8")
    report = {
        "status": "complete",
        "verdict": "fail",
        "summary": "The qualifier is an unsupported narrowing.",
        "checks": [{
            "name": "Source fidelity is clean before rendering",
            "status": "fail",
            "severity": "error",
            "category": "disclosure_fidelity",
            "detail": "The qualifier is an unsupported narrowing.",
            "items": ["Unsupported narrowing of disclosed pressure condition"],
        }],
        "findings": [{
            "severity": "critical",
            "category": "disclosure_fidelity",
            "title": "Unsupported narrowing of disclosed pressure condition",
            "where": "draft/09-claims.md",
            "detail": "The inventor did not disclose a lower closing pressure.",
            "evidence": "The source says only a closing pressure.",
            "fix": (
                "<edit><search>a lower closing pressure</search>"
                "<replace>a closing pressure</replace></edit>"
            ),
        }],
        "counts": {},
    }

    enforced = draft_qa.enforce_deterministic_source_fidelity(report, tmp_path)

    assert enforced["verdict"] == "pass"
    assert enforced["checks"][0]["status"] == "pass"
    assert enforced["findings"] == []
    assert len(enforced["reconciled_findings"]) == 1


def test_source_quote_reconciliation_keeps_a_nonverbatim_unsupported_relationship(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir(parents=True)
    (input_dir / "disclosure.md").write_text(
        "The poppet returns to the seat after pressure falls below a lower closing pressure.",
        encoding="utf-8",
    )
    finding = {
        "severity": "critical",
        "category": "disclosure_fidelity",
        "title": "Unsupported pressure relationship",
        "detail": "The comparison is unsupported.",
        "fix": (
            "<edit><search>a closing pressure lower than the opening pressure</search>"
            "<replace>a lower closing pressure</replace></edit>"
        ),
    }

    kept, reconciled = draft_qa.reconcile_explicit_source_support_findings(
        tmp_path, [finding])

    assert kept == [finding]
    assert reconciled == []


def test_standard_exports_contain_only_clean_application_text():
    from docx import Document
    version = {"version_no": 1, "sections": GOOD, "status": "approved"}
    markdown = draft_export.render_markdown({"title": GOOD["title"]}, version, [])
    assert "WORKING DRAFT" not in markdown
    assert "DRAFTING SOURCE TRACE" not in markdown
    document = Document(draft_export.render_docx({"title": GOOD["title"]}, version, []))
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs] +
        [paragraph.text for section in document.sections
         for paragraph in list(section.header.paragraphs) + list(section.footer.paragraphs)])
    assert "WORKING DRAFT" not in text
    assert "attorney review" not in text.lower()
    assert "DRAFTING SOURCE TRACE" not in text
    assert "1. A vacuum lifting tool comprising" in text
    assert "3. The vacuum lifting tool of claim 2" in text


def test_every_application_export_replaces_internal_citation_tokens():
    from docx import Document
    from pypdf import PdfReader

    project = {"title": GOOD["title"]}
    version = {"version_no": 1, "sections": GOOD, "citations": ALLOWED}
    expected = "U.S. Patent No. 11,223,344"

    markdown = draft_export.render_markdown(project, version, [])
    word = Document(draft_export.render_docx(project, version, []))
    word_text = "\n".join(paragraph.text for paragraph in word.paragraphs)
    pdf = PdfReader(draft_export.render_pdf(project, version, []))
    pdf_text = "\n".join((page.extract_text() or "") for page in pdf.pages)
    filing_text = draft_uspto.filing_text(project, version)
    filing_word = Document(draft_uspto.render_filing_docx(
        project, version,
        readiness_report={"blockers": [], "formalities": [], "remaining": [],
                          "fees": {"total": 3, "independent": 1,
                                   "multiple_dependent": 0, "billable": 3}}))
    filing_word_text = "\n".join(paragraph.text for paragraph in filing_word.paragraphs)

    for exported in (markdown, word_text, pdf_text, filing_text, filing_word_text):
        assert "[REF:" not in exported
        assert expected in exported


def test_filing_citation_display_handles_us_application_publications_and_rejects_bad_tokens():
    rendered = draft_cite.filing_citations(
        "A related device is described in [REF:US-2023103821-A1].")
    assert rendered == ("A related device is described in U.S. Patent Application Publication "
                        "No. US 2023/0103821 A1.")
    with pytest.raises(ValueError, match="Malformed internal citation token"):
        draft_cite.filing_citations("See [REF:the related patent].")


def test_filing_docx_uses_clean_formal_drawing_pages_with_required_margins():
    import io
    from PIL import Image
    from docx import Document
    image = Image.new("RGB", (640, 420), "white")
    png = io.BytesIO()
    image.save(png, format="PNG")
    version = {"version_no": 1, "sections": GOOD, "citations": []}
    output = draft_uspto.render_filing_docx(
        {"title": GOOD["title"]}, version,
        readiness_report={"blockers": [], "formalities": [], "remaining": [],
                          "fees": {"total": 3, "independent": 1,
                                   "multiple_dependent": 0, "billable": 3}},
        figure_images=[{"label": "FIG. 1", "png": png.getvalue()}])
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert len(document.inline_shapes) == 1
    assert len(document.sections) == 2
    drawing_section = document.sections[-1]
    assert drawing_section.top_margin.inches == pytest.approx(1.0)
    assert drawing_section.left_margin.inches == pytest.approx(1.0)
    assert drawing_section.right_margin.inches == pytest.approx(0.625)
    assert drawing_section.bottom_margin.inches == pytest.approx(0.375)
    shape = document.inline_shapes[0]
    assert shape.width.inches <= 6.875
    assert shape.height.inches <= 9.625
    assert "DRAWING SHEETS" not in text
    drawing_paragraph = next(
        paragraph for paragraph in document.paragraphs
        if paragraph._p.xpath(".//w:drawing"))
    assert drawing_paragraph.text == ""
    for forbidden in ("(not supplied)", "STILL REQUIRED", "NOT READY", "legal advice"):
        assert forbidden.lower() not in text.lower()


def test_filing_docx_refuses_to_export_a_blocked_version():
    with pytest.raises(drafting.DraftingValidationError, match="filing gate"):
        draft_uspto.render_filing_docx(
            {"title": GOOD["title"]}, {"version_no": 1, "sections": GOOD},
            readiness_report={"blockers": [{"title": "Drawing mismatch"}]})


def test_paragraph_numbering_is_continuous_across_sections():
    first, cursor = draft_uspto.numbered_paragraphs("one\n\ntwo", 1)
    second, _ = draft_uspto.numbered_paragraphs("three", cursor)
    assert first[0].startswith("[0001]") and second[0].startswith("[0003]")


def test_no_dollar_amount_is_printed_anywhere_in_the_filing_module():
    """Fee amounts change by rulemaking; a number baked in here would be quietly wrong."""
    body = Path(draft_uspto.__file__).read_text()
    body = re.sub(r'""".*?"""', "", body, flags=re.S)          # docstrings may discuss the policy
    assert not re.search(r"\$\s?\d", body)


def test_the_ads_fields_never_invent_a_value():
    fields = draft_uspto.ads_fields({"inventors": "", "applicant": ""},
                                    {"sections": GOOD})
    inventor = next(f for f in fields if f["field"] == "Inventor(s)")
    assert inventor["value"] == "(not supplied)"


# =============================================================================================
# The agent bridge
# =============================================================================================
def test_a_structured_result_is_parsed_from_the_json_string_the_cli_returns():
    parsed = draft_agent._parse_result('{"action":"revised","summary":"did a thing"}')
    assert parsed["action"] == "revised"


def test_a_fenced_result_is_repaired_rather_than_thrown_away():
    assert draft_agent._parse_result('```json\n{"a":1}\n```')["a"] == 1


def test_a_result_that_is_not_an_object_is_a_failure_not_an_empty_dict():
    assert draft_agent._parse_result("no json here") is None
    assert draft_agent._parse_result("") is None


def test_tool_calls_are_summarised_to_workspace_relative_paths():
    steps = []
    draft_agent._summarize_event({
        "type": "assistant",
        "message": {"content": [
            {"type": "tool_use", "name": "Edit",
             "input": {"file_path": "/srv/patent-drafts/p7/draft/09-claims.md"}},
            {"type": "text", "text": "narrowed claim 1"},
            {"type": "tool_use", "name": "StructuredOutput", "input": {"summary": "x"}}]}}, steps)
    assert steps[0] == {"kind": "tool", "tool": "Edit", "detail": "draft/09-claims.md"}
    assert steps[1]["kind"] == "say"
    assert len(steps) == 2, "the structured answer is the result, not a step"


def test_availability_reports_the_reason_it_cannot_run(monkeypatch):
    monkeypatch.setattr(draft_agent, "binary", lambda: "")
    state = draft_agent.availability()
    assert state["ok"] is False and "not installed" in state["reason"]


def test_strings_bounds_and_cleans_a_model_list():
    assert draft_agent.strings(["a", "", "b"]) == ["a", "b"]
    assert draft_agent.strings("single") == ["single"]
    assert len(draft_agent.strings(["x"] * 100)) == 40


# =============================================================================================
# The reviewer's own output contract
# =============================================================================================
def test_a_finding_without_evidence_is_dropped():
    """An unfalsifiable warning in a review panel is the one users learn to scroll past."""
    findings = draft_qa.normalize_findings([
        {"severity": "critical", "title": "Claim 1 unsupported", "evidence": "", "detail": "d",
         "where": "claims", "category": "claim_support", "fix": ""},
        {"severity": "major", "title": "Numeral 16 drifts", "evidence": "the ring 16 …",
         "detail": "d", "where": "detailed description", "category": "figures_and_numerals",
         "fix": "f"}])
    assert [f["title"] for f in findings] == ["Numeral 16 drifts"]


def test_non_actionable_source_gap_is_not_returned_as_a_repair_finding():
    findings = draft_qa.normalize_findings([
        {
            "severity": "minor", "title": "Undisclosed fluid route", "evidence": "quote",
            "detail": "The source gives the relationship but does not specify the route.",
            "where": "claims", "category": "enablement",
            "fix": "No text change is available. Leave as filed or ask the inventor.",
        },
        {
            "severity": "minor", "title": "Prior-art inference", "evidence": "quote",
            "detail": "The sentence adds an inference.", "where": "background",
            "category": "citations", "fix": "Delete the trailing inference.",
        },
    ])

    assert [finding["title"] for finding in findings] == ["Prior-art inference"]


def test_source_review_does_not_require_every_supported_claim_relationship_in_a_figure():
    omission = {
        "severity": "major",
        "category": "formalities",
        "title": (
            "Controller coupling is recited in claim 1 and the description but is shown in no "
            "figure"
        ),
        "where": "figures/FIG-2.md; draft/09-claims.md claim 1",
        "detail": (
            "No figure depicts the coupling, leaving a claim element undepicted. This is a "
            "depiction gap only; the coupling is fully supported by the disclosure."
        ),
        "evidence": (
            "Claim 1 recites a controller coupled to a sensor. The disclosure says that the "
            "controller verifies current with the corresponding sensor."
        ),
        "fix": "Depict the disclosed coupling by adding a dashed connection to FIG. 2.",
    }
    linked = dict(
        omission,
        title="FIG. 2 omits the coupling it says that it shows",
        detail=(
            "The detailed description states that FIG. 2 shows the controller coupled to the "
            "sensor, but the figure brief omits that relationship."
        ),
    )

    kept, reconciled = draft_qa.reconcile_source_drawing_omission_findings(
        [omission, linked])

    assert [finding["title"] for finding in kept] == [linked["title"]]
    assert [finding["title"] for finding in reconciled] == [omission["title"]]
    assert "need not depict every claim limitation" in reconciled[0]["reconciliation"]


def test_source_review_cannot_invent_a_flow_arrow_to_repair_an_unpromised_omission():
    finding = {
        "severity": "major",
        "category": "figures_and_numerals",
        "title": "Incomplete Process Flow in FIG. 5",
        "where": "figures/FIG-5.md",
        "detail": (
            "The reclosure path is separate and has no defined entry point. The disclosure "
            "states prerequisites for reclosure but does not state that shedding triggers it."
        ),
        "evidence": (
            "The brief says that a separate path is shown for reclosure. The inventor says a "
            "contactor is not reclosed until its sensor reports zero and a delay expires."
        ),
        "fix": (
            "Add a dashed line with an arrowhead from the shedding step 304 to the reclosure "
            "check step 310."
        ),
    }

    kept, reconciled = draft_qa.reconcile_source_drawing_omission_findings([finding])

    assert kept == []
    assert [item["title"] for item in reconciled] == [finding["title"]]
    assert "must not invent a connection" in reconciled[0]["reconciliation"]


def test_source_review_cannot_turn_a_general_open_command_into_a_shedding_trigger():
    finding = {
        "severity": "critical",
        "category": "internal_logic",
        "title": "Figure flow diagram omits a trigger condition described in the text",
        "where": "figures/FIG-5.md",
        "detail": (
            "The flow diagram shows the welded-contactor check as a separate process with no "
            "trigger. The source states that this check is performed in response to an open "
            "command. The shedding step is one such command, and the diagram should reflect "
            "this dependency."
        ),
        "evidence": (
            "The brief has no connecting line between the two flows. The source says a welded "
            "condition is recorded when current remains above a threshold after an open command."
        ),
        "fix": (
            "In the process-flow list, add the following item:\n\n- A line leaves the right side of "
            "the shedding step 304 and enters the left vertex of the welded-contactor check "
            "step 306."
        ),
    }

    kept, reconciled = draft_qa.reconcile_source_drawing_omission_findings([finding])

    assert kept == []
    assert [item["title"] for item in reconciled] == [finding["title"]]
    assert "must not invent a connection" in reconciled[0]["reconciliation"]


def test_source_review_allows_arrows_to_depict_an_exact_disclosed_flow_path():
    finding = {
        "severity": "critical",
        "category": "figures_and_numerals",
        "title": "Unsupported flow arrows in figure",
        "where": "figures/FIG-2.md",
        "detail": (
            "The inventor described the circulation path in text but did not ask for it to be "
            "depicted with arrows in a drawing."
        ),
        "evidence": (
            "The source passage describes the path from the central inlet, along the cassette "
            "faces, down the guide ducts, through the tray perforations, and back through the "
            "return passage, but does not mention or request that arrows be drawn."
        ),
        "fix": "Delete the flow-arrow instruction from figures/FIG-2.md.",
    }

    kept, reconciled = draft_qa.reconcile_source_depiction_convention_findings([finding])

    assert kept == []
    assert [item["title"] for item in reconciled] == [finding["title"]]
    assert "need not prescribe patent-drawing notation" in reconciled[0]["reconciliation"]


def test_source_review_keeps_an_arrow_finding_when_the_notation_changes_direction():
    finding = {
        "severity": "critical",
        "category": "figures_and_numerals",
        "title": "Flow arrows contradict the disclosed direction",
        "where": "figures/FIG-2.md",
        "detail": "The arrows point in the wrong direction and do not match the source path.",
        "evidence": (
            "The source passage describes the clockwise flow path but does not request arrows. "
            "The brief instead specifies counterclockwise arrows."
        ),
        "fix": "Remove the incorrect flow-arrow instruction from figures/FIG-2.md.",
    }

    kept, reconciled = draft_qa.reconcile_source_depiction_convention_findings([finding])

    assert [item["title"] for item in kept] == [finding["title"]]
    assert reconciled == []


def test_source_preflight_reconciles_a_claim_only_drawing_omission(monkeypatch):
    monkeypatch.setattr(draft_agent, "run", lambda **_kwargs: draft_agent.AgentRun(
        ok=True,
        result={
            "summary": (
                "Every claim limitation, numeral, numbered part, figure brief, drawing "
                "description, and affirmative inventor source was checked and traced in full "
                "without sampling any section."
            ),
            "findings": [{
                "severity": "major",
                "category": "formalities",
                "title": "A supported claim relationship is shown in no figure",
                "where": "figures/FIG-2.md; draft/09-claims.md claim 1",
                "detail": (
                    "No figure depicts the relationship, but the relationship is fully "
                    "supported by the inventor disclosure."
                ),
                "evidence": (
                    "Claim 1 recites the relationship and the disclosure affirmatively "
                    "describes it."
                ),
                "fix": "Depict the relationship by adding a dashed connection to FIG. 2.",
            }],
        },
    ))

    outcome = draft_qa.review_sources(Path("/tmp"))

    assert outcome["ok"] is True, outcome
    assert outcome["findings"] == []
    assert len(outcome["reconciled_findings"]) == 1
    assert "No unresolved source-fidelity findings remain" in outcome["summary"]


def test_findings_are_ordered_by_severity():
    findings = draft_qa.normalize_findings([
        {"severity": "minor", "title": "m", "evidence": "e", "detail": "", "where": "",
         "category": "terminology", "fix": ""},
        {"severity": "critical", "title": "c", "evidence": "e", "detail": "", "where": "",
         "category": "citations", "fix": ""}])
    assert [f["severity"] for f in findings] == ["critical", "minor"]


def test_the_review_schema_is_valid_json_schema_the_cli_will_accept():
    encoded = json.dumps(draft_qa.REVIEW_SCHEMA)
    assert '"additionalProperties": false' in encoded
    assert draft_qa.REVIEW_SCHEMA["required"] == ["summary", "findings"]
    assert json.dumps(draft_studio.TURN_SCHEMA)


def test_the_reviewer_never_resumes_the_drafting_session(monkeypatch):
    """The whole value of the second pass is that it has not heard the drafter's argument."""
    seen = {}

    def fake_run(**kwargs):
        seen.update(kwargs)
        return draft_agent.AgentRun(ok=True, result={"summary": "s", "findings": []})

    monkeypatch.setattr(draft_agent, "run", fake_run)
    cancel = threading.Event()
    draft_qa.review(Path("/tmp"), checks=[], cancel=cancel)
    assert seen["resume"] is False
    assert "Bash" in seen["tools"] and "Write" not in seen["tools"]
    assert seen["cancel"] is cancel


def test_source_preflight_is_independent_read_only_and_ignores_pixels(monkeypatch):
    seen = {}
    monkeypatch.setattr(draft_agent, "run", lambda **kwargs: seen.update(kwargs) or
                        draft_agent.AgentRun(
                            ok=True, model="review-model",
                            result={
                                "summary": (
                                    "Every claim limitation and numbered part was traced to the "
                                    "inventor disclosure. Every numeral and figure brief is "
                                    "consistent with those sources, and no drafting note, "
                                    "placeholder, question, or instruction remains in the draft."
                                ),
                                "findings": [],
                            }))

    cancel = threading.Event()
    outcome = draft_qa.review_sources(Path("/tmp"), cancel=cancel)

    assert outcome["ok"] is True
    assert seen["resume"] is False
    assert seen["tools"] == "Read,Glob,Grep"
    assert "Ignore rendered image files" in seen["prompt"]
    assert draft_qa.SOURCE_REVIEW_VERSION in seen["prompt"]
    assert "both required root properties" in seen["prompt"]
    assert '"findings": []' in seen["prompt"]
    assert seen["cancel"] is cancel


@pytest.mark.real_source_review
def test_cancelled_source_review_is_a_turn_interruption_not_a_filing_finding(
        tmp_path, monkeypatch):
    stop = threading.Event()
    stop.set()
    qa = Mock()
    qa.SOURCE_REVIEW_VERSION = draft_qa.SOURCE_REVIEW_VERSION
    qa.review_sources.return_value = {
        "ok": False, "cancelled": True, "error": "Stopped at your request.",
        "summary": "", "findings": [], "cost_usd": 0.0,
        "duration_ms": 1, "model": "review-model",
    }
    runner = draft_studio.TurnRunner(Mock(), Mock(), qa=qa, stop_event=stop)
    render = Mock(return_value={"ok": True})
    monkeypatch.setattr(draft_figures, "ensure_project_figures", render)

    with pytest.raises(drafting.DraftingConflict):
        runner._ensure_figures(
            turn_id=3, lease="lease", project_id=7, user_id=91,
            sections=GOOD, numerals=NUMERALS, figures=FIGURES,
            disclosure="inventor source", workspace=tmp_path)

    render.assert_not_called()
    assert qa.review_sources.call_args.kwargs["cancel"].is_set() is True


def test_cancelled_independent_review_is_a_turn_interruption(tmp_path, monkeypatch):
    stop = threading.Event()
    stop.set()
    qa = Mock()
    qa.run_checks.return_value = []
    qa.review.return_value = {
        "ok": False, "cancelled": True, "error": "Stopped at your request.",
        "summary": "", "findings": [], "cost_usd": 0.0,
        "duration_ms": 1, "model": "review-model",
    }
    runner = draft_studio.TurnRunner(Mock(), Mock(), qa=qa, stop_event=stop)
    monkeypatch.setattr(
        runner, "_load", lambda _project_id: {"project": {"user_id": 91}})

    with pytest.raises(drafting.DraftingConflict):
        runner.evaluate(
            7, version_no=1, workspace=tmp_path, allowed=[], sections=GOOD,
            numerals=NUMERALS, figures=FIGURES, turn_id=3, lease="lease")

    assert qa.review.call_args.kwargs["cancel"].is_set() is True


def test_source_preflight_retries_structured_output_exhaustion_in_a_fresh_session(monkeypatch):
    results = iter([
        draft_agent.AgentRun(
            ok=False, model="review-model",
            error="Failed to provide valid structured output after 5 attempts"),
        draft_agent.AgentRun(ok=True, model="review-model", result={
            "summary": (
                "Every claim limitation and numbered part was traced to affirmative inventor "
                "disclosure. The numeral table and every figure brief use those supported parts "
                "consistently, and no drafting note, placeholder, question, or instruction "
                "remains anywhere in the candidate."
            ),
            "findings": [],
        }),
    ])
    sessions = []

    def fake_run(**kwargs):
        sessions.append(kwargs["session_id"])
        return next(results)

    monkeypatch.setattr(draft_agent, "run", fake_run)

    outcome = draft_qa.review_sources(Path("/tmp"))

    assert outcome["ok"] is True
    assert len(sessions) == 2
    assert sessions[0] != sessions[1]


def test_source_preflight_remains_fail_closed_after_two_schema_exhaustions(monkeypatch):
    calls = 0

    def fake_run(**_kwargs):
        nonlocal calls
        calls += 1
        return draft_agent.AgentRun(
            ok=False, model="review-model",
            error="Failed to provide valid structured output after 5 attempts")

    monkeypatch.setattr(draft_agent, "run", fake_run)

    outcome = draft_qa.review_sources(Path("/tmp"))

    assert outcome["ok"] is False
    assert "structured output" in outcome["error"]
    assert calls == 2


def test_source_preflight_retries_a_non_substantive_structured_result(monkeypatch):
    results = iter([
        draft_agent.AgentRun(ok=True, model="review-model", result={
            "summary": "Test",
            "findings": [{
                "severity": "minor", "category": "terminology", "title": "t",
                "where": "w", "detail": "d", "evidence": "e", "fix": "f",
            }],
        }),
        draft_agent.AgentRun(ok=True, model="review-model", result={
            "summary": (
                "Every claim limitation and numbered part was traced to affirmative inventor "
                "disclosure. The numeral table and every figure brief use those supported parts "
                "consistently, and no drafting note, placeholder, question, or instruction "
                "remains anywhere in the candidate."
            ),
            "findings": [],
        }),
    ])
    sessions = []

    def fake_run(**kwargs):
        sessions.append(kwargs["session_id"])
        return next(results)

    monkeypatch.setattr(draft_agent, "run", fake_run)

    outcome = draft_qa.review_sources(Path("/tmp"))

    assert outcome["ok"] is True
    assert len(sessions) == 2
    assert sessions[0] != sessions[1]


def test_vertex_source_preflight_must_read_exact_case_figure_manifest(monkeypatch, tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    for name in ("disclosure.md", "conversation.md", "brief.md"):
        (input_dir / name).write_text("affirmative inventor source", encoding="utf-8")
    draft_workspace.write_sections(tmp_path, GOOD)
    draft_workspace.write_numerals(tmp_path, NUMERALS)
    draft_workspace.write_figures(tmp_path, FIGURES)
    required = [
        "input/disclosure.md", "input/conversation.md", "input/brief.md",
        *[f"draft/{name}" for _key, name, _heading in draft_workspace.SECTION_FILES],
        f"draft/{draft_workspace.NUMERALS_FILE}",
        "figures/FIG-1.md", "figures/FIG-2.md",
    ]
    prompts = []

    def fake_run(**kwargs):
        prompts.append(kwargs["prompt"])
        read_paths = required if len(prompts) == 2 else required[:-2]
        return draft_agent.AgentRun(
            ok=True,
            model="vertex/gemini-2.5-pro",
            steps=[{"kind": "tool", "tool": "read_file", "detail": path}
                   for path in read_paths],
            result={
                "summary": (
                    "Every claim limitation and numbered part was traced to affirmative inventor "
                    "source disclosure. Every numeral and every figure brief was checked in full, "
                    "and the complete claims, drawings, and inventor sources contain no "
                    "unsupported technical assertion."
                ),
                "findings": [],
            },
        )

    monkeypatch.setattr(draft_agent, "run", fake_run)

    outcome = draft_qa.review_sources(tmp_path)

    assert outcome["ok"] is True
    assert len(prompts) == 2
    assert "figures/FIG-1.md" in prompts[0]
    assert "did not read every required file" in prompts[1]
    assert "figures/FIG-2.md" in prompts[1]


def test_source_preflight_accepts_a_precise_file_path_as_finding_location(monkeypatch):
    calls = 0

    def fake_run(**_kwargs):
        nonlocal calls
        calls += 1
        return draft_agent.AgentRun(ok=True, model="review-model", result={
            "summary": (
                "Every claim limitation and numbered part was traced to affirmative inventor "
                "disclosure. The numeral table and every figure brief use those supported parts "
                "consistently. One formalities finding remains in the abstract, and no drafting "
                "note, placeholder, question, or instruction remains elsewhere in the draft."
            ),
            "findings": [{
                "severity": "minor",
                "category": "formalities",
                "title": "Abstract exceeds the conservative word limit",
                "where": "draft/10-abstract.md",
                "detail": (
                    "The conservative count treats hyphenated compounds as separate words and "
                    "places the abstract above the filing limit."
                ),
                "evidence": (
                    "The abstract contains the phrases air-extraction mechanism and "
                    "low-friction perimeter member."
                ),
                "fix": (
                    "Replace the longer phrase with supported shorter wording while preserving "
                    "the disclosed technical substance."
                ),
            }],
        })

    monkeypatch.setattr(draft_agent, "run", fake_run)

    outcome = draft_qa.review_sources(Path("/tmp"))

    assert outcome["ok"] is True
    assert [item["where"] for item in outcome["findings"]] == ["draft/10-abstract.md"]
    assert calls == 1


def test_source_preflight_fails_closed_after_repeated_non_substantive_results(monkeypatch):
    calls = 0

    def fake_run(**_kwargs):
        nonlocal calls
        calls += 1
        return draft_agent.AgentRun(ok=True, model="review-model", result={
            "summary": "Test",
            "findings": [{
                "severity": "minor", "category": "terminology", "title": "t",
                "where": "w", "detail": "d", "evidence": "e", "fix": "f",
            }],
        })

    monkeypatch.setattr(draft_agent, "run", fake_run)

    outcome = draft_qa.review_sources(Path("/tmp"))

    assert outcome["ok"] is False
    assert "non-substantive" in outcome["error"]
    assert calls == 2


def test_source_preflight_retry_explains_the_rejected_quality_gate(monkeypatch):
    prompts = []

    def fake_run(**kwargs):
        prompts.append(kwargs["prompt"])
        if len(prompts) == 1:
            return draft_agent.AgentRun(ok=True, model="review-model", result={
                "summary": "The complete source ledger is clean.",
                "findings": [],
            })
        assert "previous source-review output was rejected" in kwargs["prompt"].lower()
        assert "claims, numerals, figures, and inventor sources" in kwargs["prompt"].lower()
        return draft_agent.AgentRun(ok=True, model="review-model", result={
            "summary": (
                "Every claim limitation and numbered part was traced to affirmative inventor "
                "source disclosure. Every numeral is used consistently, and every figure brief "
                "depicts only those supported structures. The complete claims, numerals, figures, "
                "and inventor sources were reviewed without an unsupported technical assertion."
            ),
            "findings": [],
        })

    monkeypatch.setattr(draft_agent, "run", fake_run)

    outcome = draft_qa.review_sources(Path("/tmp"))

    assert outcome["ok"] is True
    assert len(prompts) == 2


def test_source_preflight_fails_closed_on_a_malformed_finding(monkeypatch):
    monkeypatch.setattr(draft_agent, "run", lambda **_kwargs: draft_agent.AgentRun(
        ok=True, model="review-model", result={
            "summary": "A source gap exists.",
            "findings": [{
                "severity": "critical", "category": "disclosure_fidelity",
                "title": "Unsupported duct", "where": "draft/numerals.md",
                "detail": "The duct is not in the disclosure.", "evidence": "",
                "fix": "Remove it.",
            }],
        }))

    outcome = draft_qa.review_sources(Path("/tmp"))

    assert outcome["ok"] is False
    assert "malformed finding" in outcome["error"]


def test_the_reviewer_is_told_which_checks_already_ran(monkeypatch):
    seen = {}
    monkeypatch.setattr(draft_agent, "run", lambda **k: seen.update(k) or draft_agent.AgentRun(
        ok=True, result={"summary": "", "findings": []}))
    draft_qa.review(Path("/tmp"), checks=[
        {"name": "Numerals", "status": "fail", "detail": "22 undefined", "items": ["22"]},
        {"name": "Claims", "status": "pass", "detail": "fine", "items": []}])
    assert "22 undefined" in seen["prompt"] and "Claims" not in seen["prompt"]
    assert "rendered-*.png" in draft_qa.REVIEW_PROMPT
    assert "review/figure-audit-evidence.json" in draft_qa.REVIEW_PROMPT
    normalized = " ".join(draft_qa.REVIEW_SYSTEM.split())
    assert "reconcile that disagreement" in normalized
    assert "raw image coordinates" in normalized
    assert "byte-exact section-hatch certificate" in normalized


def test_review_reconciles_only_a_hatch_claim_disproved_by_an_exact_render_certificate(
        monkeypatch, tmp_path):
    review_dir = tmp_path / "review"
    review_dir.mkdir()
    figures_dir = tmp_path / "figures"
    figures_dir.mkdir()
    rendered = b"exact checked sheet"
    (figures_dir / "rendered-FIG-2.png").write_bytes(rendered)
    (review_dir / "figure-audit-evidence.json").write_text(json.dumps({
        "schema_version": 1,
        "figures": [{
            "figure_label": "FIG. 2",
            "rendered_file": "rendered-FIG-2.png",
            "rendered_sha256": hashlib.sha256(rendered).hexdigest(),
            "geometry": {"ok": True},
            "deterministic_section_hatching": {
                "ok": True,
                "version": draft_figures.DETERMINISTIC_SECTION_HATCH_CERTIFICATE_VERSION,
                "exact_renderer_match": True,
                "renderer": "chamber_section",
                "coordinate_space": "raw_pixels_origin_upper_left_y_down",
                "raw_png_sha256": "a" * 64,
                "components": [
                    {"component": "base slab", "angle_degrees": -45,
                     "direction": "rises_to_right"},
                    {"component": "left perimeter leg", "angle_degrees": 45,
                     "direction": "falls_to_right"},
                    {"component": "covering-element band", "angle_degrees": -75,
                     "direction": "rises_to_right"},
                ],
            },
        }],
    }), encoding="utf-8")
    monkeypatch.setattr(draft_agent, "run", lambda **_kwargs: draft_agent.AgentRun(
        ok=True, model="review-model", result={
            "summary": "FIG. 2 appears to use identical hatching.",
            "findings": [{
                "severity": "major", "category": "figures_and_numerals",
                "title": "Rendered FIG. 2 uses the same hatch angle for three bodies",
                "where": "figures/rendered-FIG-2.png",
                "detail": "The visible hatch directions appear parallel rather than distinct.",
                "evidence": "The base, legs and bottom band appear to rise to the right.",
                "fix": "Render the three bodies again with distinct hatch angles.",
            }],
        }))

    outcome = draft_qa.review(tmp_path, checks=[])

    assert outcome["ok"] is True
    assert outcome["findings"] == []
    assert len(outcome["reconciled_findings"]) == 1
    assert outcome["reconciled_findings"][0]["figure_label"] == "FIG. 2"
    assert "exact-image reconciliation" in outcome["summary"]

    (figures_dir / "rendered-FIG-2.png").write_bytes(b"changed after review")
    kept, reconciled = draft_qa.reconcile_exact_section_hatch_findings(
        tmp_path, outcome["reconciled_findings"])
    assert len(kept) == 1
    assert reconciled == []

    (figures_dir / "rendered-FIG-2.png").write_bytes(rendered)
    text_finding = dict(outcome["reconciled_findings"][0], category="internal_logic")
    kept, reconciled = draft_qa.reconcile_exact_section_hatch_findings(
        tmp_path, [text_finding])
    assert len(kept) == 1
    assert reconciled == []


def test_independent_review_cannot_reintroduce_an_unpromised_flow_arrow(
        tmp_path, monkeypatch):
    monkeypatch.setattr(draft_agent, "run", lambda **_kwargs: draft_agent.AgentRun(
        ok=True, model="review-model", result={
            "summary": "FIG. 5 omits the trigger for a separate safety sequence.",
            "findings": [{
                "severity": "major", "category": "internal_logic",
                "title": "FIG. 5 flowchart omits the trigger for the welded contactor check",
                "where": "draft/08-detailed-description.md; figures/FIG-5.md",
                "detail": (
                    "The drawing shows two separate unconnected process flows even though an "
                    "open command is a condition for the welded contactor check."
                ),
                "evidence": (
                    "The brief says that the welded contactor check is a separate sequence and "
                    "that its trigger is described in the specification but not shown with a "
                    "separate entry arrow."
                ),
                "fix": (
                    "Add a dashed-line arrow from the shedding step 304 to the welded contactor "
                    "check step 306."
                ),
            }],
        }))

    outcome = draft_qa.review(tmp_path, checks=[])

    assert outcome["ok"] is True
    assert outcome["findings"] == []
    assert len(outcome["reconciled_findings"]) == 1
    assert "must not invent a connection" in (
        outcome["reconciled_findings"][0]["reconciliation"])


def test_independent_review_does_not_label_an_offsheet_connection_as_the_remote_part(
        tmp_path, monkeypatch):
    figures_dir = tmp_path / "figures"
    figures_dir.mkdir()
    draft_dir = tmp_path / "draft"
    draft_dir.mkdir()
    (figures_dir / "FIG-3.md").write_text(
        "# FIG. 3\n\n"
        "A short solid line extends upward from the controller. This line is the connection to "
        "the branch current sensor. A short solid line extends downward from the controller. "
        "This line is the connection to the isolated local bus.\n\n"
        "## Numerals shown on this figure\n\n"
        "- 106 edge controller\n- 110 network interface\n- 112 service input\n",
        encoding="utf-8",
    )
    (draft_dir / "numerals.md").write_text(
        "# Reference numerals\n\n| Numeral | Part |\n| --- | --- |\n"
        "| 104 | branch current sensor |\n| 106 | edge controller |\n"
        "| 108 | isolated local bus |\n| 110 | network interface |\n"
        "| 112 | service input |\n",
        encoding="utf-8",
    )
    monkeypatch.setattr(draft_agent, "run", lambda **_kwargs: draft_agent.AgentRun(
        ok=True, model="review-model", result={
            "summary": "FIG. 3 omits two remote-part numerals.",
            "findings": [{
                "severity": "minor", "category": "figures_and_numerals",
                "title": "FIG. 3 omits labels for two offsheet connections",
                "where": "figures/FIG-3.md; figures/rendered-FIG-3.png",
                "detail": (
                    "The lines extending to the branch current sensor and isolated local bus are "
                    "not labelled with the numerals of those remote parts."
                ),
                "evidence": (
                    "The brief calls the lines connections to the branch current sensor and the "
                    "isolated local bus."
                ),
                "fix": (
                    "Add leaders and numerals for the branch current sensor (104) and isolated "
                    "local bus (108), pointing to the respective connection lines."
                ),
            }],
        }))

    outcome = draft_qa.review(tmp_path, checks=[])

    assert outcome["ok"] is True
    assert outcome["findings"] == []
    assert len(outcome["reconciled_findings"]) == 1
    assert "offsheet connection" in (
        outcome["reconciled_findings"][0]["reconciliation"])


def test_independent_reviewer_requires_a_source_supported_automatic_fix():
    prompt = draft_qa.REVIEW_SYSTEM

    assert "Do not report a gap in the inventor's disclosure" in prompt
    assert "Do not recommend asking the inventor" in prompt
    assert "need not depict every claim limitation" in prompt
    assert "generic depiction convention" in prompt
    assert "cutting-plane line" in prompt
    assert "section designation" in prompt


def test_a_broken_reviewer_is_a_finding_not_an_exception(monkeypatch):
    def explode(**_kwargs):
        raise draft_agent.AgentUnavailable("no CLI here")
    monkeypatch.setattr(draft_agent, "run", explode)
    out = draft_qa.review(Path("/tmp"), checks=[])
    assert out["ok"] is False and "no CLI" in out["error"]


def test_the_drafting_prompt_states_the_rules_it_must_not_break():
    system = draft_studio.DRAFT_SYSTEM
    for phrase in ("prior_art/INDEX.md", "[REF:KEY]", "numerals.md",
                   "never state or imply", "Never invent", "Not applicable"):
        assert phrase in system, f"the drafting prompt no longer says: {phrase}"
    normalized = " ".join(system.split())
    assert "Operational requests to resume, preserve, repair, inspect, or audit" in normalized
    assert "do not disclose or affirm its technical content" in normalized
    assert "[DRAFTING NOTE" not in system
    assert "Return `questions` as an empty array" in system
    assert "No placeholder" in system
    assert "at least one figure" in system
    assert "Normally use two to four figures" in system
    assert "Do not list more than eight numerals on one sheet" in system
    assert str(draft_qa.MAX_FIGURE_BRIEF_CHARS) in system
    assert "arbitrary exact counts" in system
    assert "shown schematically" in normalized
    assert "least-specific schematic outline" in system
    assert "which end is deeper" in system
    assert "runout direction" in normalized
    assert "open paper between solid bodies" in system
    assert "keep all line work inside the drawing area" in system
    assert "white-interior strip" in system
    assert "generic negative bans on linework" in system
    assert "shared face edge to be drawn once" in system
    assert "Numeral endpoint instructions identify the part" in system
    assert "cutting-plane line" in system
    assert "same repeated designation" in normalized
    assert "not a reference numeral" in normalized
    assert "axial section through a hollow cylindrical part" in normalized
    assert "two opposed sectioned walls" in normalized
    assert "transverse section" in normalized
    assert "vertical bore axis" in normalized
    assert "longitudinal slot axis" in normalized
    assert "actual intersection or center-plane relationship" in normalized
    assert "broad interior target" in draft_studio.FINALIZE_PROMPT
    assert "generic negative linework controls" in draft_studio.FINALIZE_PROMPT
    assert "generic face-linework controls" in draft_studio.FINALIZE_PROMPT
    assert "Replace ordinal geometry references" in draft_studio.FINALIZE_PROMPT
    assert "Never address or mention a draftsperson" in " ".join(system.split())
    assert "broadest statement of the invention that the description fully supports" in system
    assert "Every reference listed in prior_art/INDEX.md must be addressed" in system
    assert "Never omit a listed reference solely because it is less relevant" in system


def test_drawing_repairs_can_never_change_the_invention_to_match_bad_pixels():
    prompt = draft_studio.FINALIZE_PROMPT
    normalized = " ".join(prompt.split())
    assert "Generated drawing pixels are evidence to inspect, never authority" \
        in draft_studio.DRAFT_SYSTEM
    assert "Generated pixels are never authority for the invention" in prompt
    assert "Never change the claims, description, numeral table, or disclosed embodiments" \
        in prompt
    assert "regenerate the sheet from the authoritative text" in normalized


def test_figure_plan_repairs_remove_stale_references_from_the_old_sheet():
    prompt = " ".join(draft_studio.DRAFT_SYSTEM.split())
    assert "remove every use of that numeral and its canonical part name from the old sheet" \
        in prompt
    assert "describe it generically as an unnumbered" in prompt
    assert "Never delete the focused sheet merely to fix stale references" in prompt


def test_the_independent_reviewer_checks_source_fidelity_before_internal_consistency():
    system = draft_qa.REVIEW_SYSTEM
    normalized = " ".join(system.split())
    assert "DISCLOSURE FIDELITY" in system
    assert "input/disclosure.md" in system
    assert "generated drawing artifact" in normalized
    assert "every independent and dependent claim limitation" in normalized
    assert "Common engineering knowledge is not source support" in normalized
    assert "corrective instruction that names a candidate detail only to reject" in normalized
    assert "Prior-art characterisations trace to prior_art/" in normalized
    assert "disclosure_fidelity" in json.dumps(draft_qa.REVIEW_SCHEMA)
    preflight = " ".join(draft_qa.SOURCE_REVIEW_SYSTEM.split())
    assert "Build a complete source ledger before returning" in preflight
    assert "every limitation in every claim" in preflight
    assert "corrective USER message that names a candidate detail only to reject" in preflight
    assert "passages under headings labeled USER are the authority" in preflight
    assert "YOU, REVIEWER, or SYSTEM are context, never inventor support" in preflight
    assert "Instructions merely to resume, preserve, repair, inspect, or audit a candidate" \
        in preflight
    assert "numeral or figure counts, labels, and filing gates do not affirm" in preflight
    assert "Prior-art characterizations in the Background do not require inventor support" \
        in preflight
    assert "simple generic outline" in preflight
    assert "depiction convention" in preflight
    assert "which end is deeper" in preflight
    assert "runout direction" in preflight
    assert "functional face, slot, joint, cam, ramp, seal, port, or flow boundary" in preflight
    assert "substantive even when the brief calls the geometry schematic" in preflight
    assert "must stay confined to the figure brief" in preflight
    assert "by way of example" in preflight
    assert "any closed outline" in preflight
    assert "Do not inspect or rely on rendered images" in preflight


def test_source_reviews_preserve_disclosed_features_in_both_directions():
    preflight = " ".join(draft_qa.SOURCE_REVIEW_SYSTEM.split())
    final_review = " ".join(draft_qa.REVIEW_SYSTEM.split())
    drafting_prompt = " ".join(draft_studio.DRAFT_SYSTEM.split())

    for prompt in (preflight, final_review):
        assert "Build the disclosure ledger in both directions" in prompt
        assert "installation or calibration procedure" in prompt
        assert "data-recording behavior" in prompt
        assert "commercially distinct embodiment" in prompt
        assert "Do not require every optional feature in an independent claim" in prompt
    assert "Never silently drop affirmative technical matter" in drafting_prompt
    for prompt in (preflight, final_review, drafting_prompt):
        assert "Description-only preservation is not claim coverage" in prompt
        assert "below 20 total claims" in prompt
        assert "technical safeguard against misconfiguration or failure" in prompt


def test_drafting_and_review_preserve_compound_source_conditions_and_verifiers():
    prompts = (
        " ".join(draft_studio.DRAFT_SYSTEM.split()),
        " ".join(draft_qa.SOURCE_REVIEW_SYSTEM.split()),
        " ".join(draft_qa.REVIEW_SYSTEM.split()),
    )

    for prompt in prompts:
        assert "conditional, temporal, negative, exception, threshold, actor, and verification" \
            in prompt
        assert "indivisible source constraint" in prompt
        assert "sensor-confirmed agreement with human confirmation" in prompt
        assert "unexpired-token condition with generic authorization" in prompt


def test_drafting_and_review_repairs_respect_standard_claim_count_limits():
    instruction = (
        "No automatic fix may leave more than 20 total claims or more than three "
        "independent claims"
    )

    assert instruction in " ".join(draft_studio.DRAFT_SYSTEM.split())
    assert instruction in " ".join(draft_qa.SOURCE_REVIEW_SYSTEM.split())
    assert instruction in " ".join(draft_qa.REVIEW_SYSTEM.split())


def test_review_fixes_respect_the_drawing_sheet_numeral_limit():
    instruction = (
        "Never propose an automatic fix that leaves more than eight reference numerals "
        "on one drawing sheet"
    )

    assert instruction in " ".join(draft_qa.SOURCE_REVIEW_SYSTEM.split())
    assert instruction in " ".join(draft_qa.REVIEW_SYSTEM.split())


def test_drafting_prompt_removes_superseded_figure_briefs():
    prompt = " ".join(draft_studio.DRAFT_SYSTEM.split())

    assert "Never leave two figure files with the same FIG. number" in prompt
    assert "delete the superseded file before returning" in prompt


@pytest.mark.real_source_review
def test_source_fidelity_preflight_blocks_rendering_unsupported_geometry(
        monkeypatch, tmp_path):
    qa = Mock()
    qa.review_sources.return_value = {
        "ok": True,
        "summary": "A numbered duct has no inventor source.",
        "findings": [{
            "severity": "critical", "category": "disclosure_fidelity",
            "title": "Unsupported duct", "where": "draft/numerals.md",
            "detail": "The disclosure never introduces a duct.",
            "evidence": "32 | duct", "fix": "Remove or generalize it.",
        }],
        "cost_usd": 0.1, "duration_ms": 100, "model": "review-model",
    }
    runner = draft_studio.TurnRunner(Mock(), Mock(), qa=qa)
    render = Mock(return_value={"ok": True})
    monkeypatch.setattr(draft_figures, "ensure_project_figures", render)
    monkeypatch.setattr(draft_figures, "checkpoint_project_figures", Mock())

    with pytest.raises(draft_studio.SourceFidelityInspectionError) as caught:
        runner._ensure_figures(
            turn_id=3, lease="lease", project_id=7, user_id=91,
            sections=GOOD, numerals=NUMERALS, figures=FIGURES,
            disclosure="the inventor disclosed a body and pump", workspace=tmp_path)

    render.assert_not_called()
    assert caught.value.report["findings"][0]["title"] == "Unsupported duct"
    assert caught.value.report["verdict"] == "fail"


@pytest.mark.real_source_review
def test_source_reviewer_outage_retries_saved_candidate_without_draft_repair(
        monkeypatch, tmp_path):
    qa = Mock()
    qa.review_sources.return_value = {
        "ok": False,
        "error": "API Error: 529 Overloaded. This is a server-side issue.",
        "summary": "", "findings": [], "cost_usd": 0.0,
        "duration_ms": 100, "model": "review-model",
    }
    runner = draft_studio.TurnRunner(Mock(), Mock(), qa=qa)
    render = Mock(return_value={"ok": True})
    monkeypatch.setattr(draft_figures, "ensure_project_figures", render)
    monkeypatch.setattr(draft_figures, "checkpoint_project_figures", Mock())

    with pytest.raises(draft_studio.SourceReviewUnavailable) as caught:
        runner._ensure_figures(
            turn_id=3, lease="lease", project_id=7, user_id=91,
            sections=GOOD, numerals=NUMERALS, figures=FIGURES,
            disclosure="the inventor disclosed a body and pump", workspace=tmp_path)

    render.assert_not_called()
    assert caught.value.retry_without_repair is True
    assert "529 Overloaded" in str(caught.value)


@pytest.mark.real_source_review
def test_clean_source_preflight_is_cached_before_repeated_rendering(monkeypatch, tmp_path):
    qa = Mock()
    qa.review_sources.return_value = {
        "ok": True, "summary": "Every candidate detail has affirmative support.",
        "findings": [], "cost_usd": 0.1, "duration_ms": 100,
        "model": "review-model",
    }
    runner = draft_studio.TurnRunner(Mock(), Mock(), qa=qa)
    render = Mock(return_value={"ok": True})
    monkeypatch.setattr(draft_figures, "ensure_project_figures", render)
    monkeypatch.setattr(draft_figures, "checkpoint_project_figures", Mock())
    monkeypatch.setattr(draft_figures, "materialize_review_images", Mock(return_value=[]))

    values = dict(
        turn_id=3, lease="lease", project_id=7, user_id=91,
        sections=GOOD, numerals=NUMERALS, figures=FIGURES,
        disclosure="the inventor disclosed a body and pump", workspace=tmp_path)
    runner._ensure_figures(**values)
    runner._ensure_figures(**values)

    qa.review_sources.assert_called_once()
    assert render.call_count == 2


@pytest.mark.real_source_review
def test_clean_source_preflight_cache_survives_a_new_worker_runner(monkeypatch, tmp_path):
    qa = Mock()
    qa.review_sources.return_value = {
        "ok": True, "summary": "Every candidate detail has affirmative support.",
        "findings": [], "cost_usd": 0.1, "duration_ms": 100,
        "model": "review-model",
    }
    durable = {}
    repository = Mock()
    repository.source_review_cache.side_effect = durable.get
    repository.save_source_review_cache.side_effect = (
        lambda source_hash, report: durable.__setitem__(source_hash, dict(report)))
    render = Mock(return_value={"ok": True})
    monkeypatch.setattr(draft_figures, "ensure_project_figures", render)
    monkeypatch.setattr(draft_figures, "checkpoint_project_figures", Mock())
    monkeypatch.setattr(draft_figures, "materialize_review_images", Mock(return_value=[]))
    values = dict(
        turn_id=3, lease="lease", project_id=7, user_id=91,
        sections=GOOD, numerals=NUMERALS, figures=FIGURES,
        disclosure="the inventor disclosed a body and pump", workspace=tmp_path)

    draft_studio.TurnRunner(repository, Mock(), qa=qa)._ensure_figures(**values)
    draft_studio.TurnRunner(repository, Mock(), qa=qa)._ensure_figures(**values)

    qa.review_sources.assert_called_once()
    assert render.call_count == 2
    assert len(durable) == 1


@pytest.mark.real_source_review
def test_durable_source_preflight_cache_changes_with_the_review_model(monkeypatch, tmp_path):
    qa = Mock()
    qa.review_sources.return_value = {
        "ok": True, "summary": "Every candidate detail has affirmative support.",
        "findings": [], "cost_usd": 0.1, "duration_ms": 100,
        "model": "review-model",
    }
    durable = {}
    repository = Mock()
    repository.source_review_cache.side_effect = durable.get
    repository.save_source_review_cache.side_effect = (
        lambda source_hash, report: durable.__setitem__(source_hash, dict(report)))
    monkeypatch.setattr(
        draft_figures, "ensure_project_figures", Mock(return_value={"ok": True}))
    monkeypatch.setattr(draft_figures, "checkpoint_project_figures", Mock())
    monkeypatch.setattr(draft_figures, "materialize_review_images", Mock(return_value=[]))
    values = dict(
        turn_id=3, lease="lease", project_id=7, user_id=91,
        sections=GOOD, numerals=NUMERALS, figures=FIGURES,
        disclosure="the inventor disclosed a body and pump", workspace=tmp_path)

    first = draft_studio.TurnRunner(repository, Mock(), qa=qa)
    first._review_sources(**{key: values[key] for key in (
        "turn_id", "lease", "sections", "numerals", "figures", "disclosure", "workspace")},
        model="review-model-a")
    second = draft_studio.TurnRunner(repository, Mock(), qa=qa)
    second._review_sources(**{key: values[key] for key in (
        "turn_id", "lease", "sections", "numerals", "figures", "disclosure", "workspace")},
        model="review-model-b")

    assert qa.review_sources.call_count == 2
    assert len(durable) == 2


@pytest.mark.real_source_review
def test_source_preflight_cache_changes_when_filing_brief_changes(monkeypatch, tmp_path):
    qa = Mock()
    qa.review_sources.return_value = {
        "ok": True, "summary": "Every candidate detail has affirmative support.",
        "findings": [], "cost_usd": 0.1, "duration_ms": 100,
        "model": "review-model",
    }
    runner = draft_studio.TurnRunner(Mock(), Mock(), qa=qa)
    monkeypatch.setattr(draft_figures, "ensure_project_figures", Mock(return_value={"ok": True}))
    monkeypatch.setattr(draft_figures, "checkpoint_project_figures", Mock())
    monkeypatch.setattr(draft_figures, "materialize_review_images", Mock(return_value=[]))
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    brief = input_dir / "brief.md"
    brief.write_text("Applicant: First Applicant\n", encoding="utf-8")
    values = dict(
        turn_id=3, lease="lease", project_id=7, user_id=91,
        sections=GOOD, numerals=NUMERALS, figures=FIGURES,
        disclosure="the inventor disclosed a body and pump", workspace=tmp_path)

    runner._ensure_figures(**values)
    brief.write_text("Applicant: Corrected Applicant\n", encoding="utf-8")
    runner._ensure_figures(**values)

    assert qa.review_sources.call_count == 2


def test_drawing_only_repairs_cannot_mutate_filing_sources_or_figure_membership(tmp_path):
    draft_workspace.write_sections(tmp_path, GOOD)
    draft_workspace.write_numerals(tmp_path, NUMERALS)
    changed_figures = [
        {**FIGURES[0], "caption": "A corrected geometry brief.",
         "numerals": [*FIGURES[0]["numerals"], "99 pixel artifact"]},
        {**FIGURES[1], "numerals": []},
        {"label": "FIG. 9", "caption": "A pixel-derived extra sheet.", "numerals": ["99"]},
    ]
    draft_workspace.write_figures(tmp_path, changed_figures)
    baseline = {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}
    draft_workspace.write_sections(tmp_path, {**GOOD, "summary": "Pixel-derived embodiment."})
    draft_workspace.write_numerals(tmp_path, [*NUMERALS, {"numeral": "99", "part": "artifact"}])
    report = {
        "checks": [{
            "name": "Every drawing sheet passes geometry, leader, and OCR inspection",
            "status": "fail",
        }],
        "findings": [{"category": "figures_and_numerals", "title": "Floating leader"}],
    }

    assert draft_studio.restore_text_after_drawing_only_review(tmp_path, baseline, report) is True
    assert draft_workspace.read_sections(tmp_path) == GOOD
    assert draft_workspace.read_numerals(tmp_path) == NUMERALS
    restored_figures = draft_workspace.read_figures(tmp_path)
    assert [item["label"] for item in restored_figures] == ["FIG. 1", "FIG. 2"]
    assert restored_figures[0]["caption"] == "A corrected geometry brief."
    assert restored_figures[0]["numerals"] == FIGURES[0]["numerals"]
    assert restored_figures[1]["numerals"] == FIGURES[1]["numerals"]


def test_drawing_only_repair_preserves_unrelated_figure_briefs(tmp_path):
    baseline = {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}
    revised_figures = [
        {**FIGURES[0], "caption": "An unrelated rewrite of the first sheet."},
        {**FIGURES[1], "caption": "A focused correction to the second sheet."},
    ]
    draft_workspace.write_sections(tmp_path, GOOD)
    draft_workspace.write_numerals(tmp_path, NUMERALS)
    draft_workspace.write_figures(tmp_path, revised_figures)
    report = {
        "checks": [],
        "findings": [{
            "category": "figures_and_numerals",
            "title": "FIG. 2 contains surplus geometry",
            "detail": "Regenerate FIG. 2 without the extra body.",
        }],
    }

    assert draft_studio.restore_text_after_drawing_only_review(
        tmp_path, baseline, report) is True

    restored = draft_workspace.read_figures(tmp_path)
    assert restored[0]["caption"] == FIGURES[0]["caption"]
    assert restored[1]["caption"] == "A focused correction to the second sheet."


def test_figure_plan_repairs_keep_authoritative_text_and_numerals(tmp_path):
    baseline = {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}
    revised_sections = {
        **GOOD,
        "drawing_descriptions": (
            "FIG. 1 is a focused body view.\n\n"
            "FIG. 2 is a ring detail.\n\n"
            "FIG. 3 is a focused passage view."
        ),
        "detailed_description": "Removed disclosed structure to make the picture easier.",
        "claims": "1. A different invention.",
    }
    revised_figures = [
        {**FIGURES[0], "numerals": ["10", "12"]},
        {**FIGURES[1], "numerals": ["14", "16", "18"]},
        {"label": "FIG. 3", "caption": "focused passage view", "numerals": ["20"]},
    ]
    draft_workspace.write_sections(tmp_path, revised_sections)
    draft_workspace.write_numerals(tmp_path, NUMERALS[:-1])
    draft_workspace.write_figures(tmp_path, revised_figures)
    report = {
        "checks": [{
            "name": "Saved candidate passes the current filing preflight",
            "status": "fail", "category": "figures_and_numerals",
        }],
        "findings": [],
    }

    assert draft_studio.restore_sources_after_figure_plan_review(
        tmp_path, baseline, report) is True

    sections = draft_workspace.read_sections(tmp_path)
    assert sections["drawing_descriptions"] == revised_sections["drawing_descriptions"]
    assert sections["detailed_description"] == GOOD["detailed_description"]
    assert sections["claims"] == GOOD["claims"]
    assert draft_workspace.read_numerals(tmp_path) == NUMERALS
    assert [{key: item[key] for key in ("label", "caption", "numerals")}
            for item in draft_workspace.read_figures(tmp_path)] == revised_figures


def test_source_locks_never_restore_an_empty_or_incomplete_baseline(tmp_path):
    draft_workspace.write_sections(tmp_path, GOOD)
    draft_workspace.write_numerals(tmp_path, NUMERALS)
    draft_workspace.write_figures(tmp_path, FIGURES)
    before = draft_workspace.snapshot(tmp_path)
    figure_plan_report = {
        "checks": [{
            "name": "Drawing briefs are concise and drawable",
            "status": "fail", "category": "figures_and_numerals",
        }],
        "findings": [],
    }
    drawing_report = {
        "checks": [{
            "name": "Every drawing sheet passes geometry, leader, and OCR inspection",
            "status": "fail", "category": "figures_and_numerals",
        }],
        "findings": [],
    }

    assert not draft_studio.restore_sources_after_figure_plan_review(
        tmp_path, {}, figure_plan_report)
    assert draft_workspace.snapshot(tmp_path) == before
    assert not draft_studio.restore_text_after_drawing_only_review(
        tmp_path, {"sections": GOOD}, drawing_report)
    assert draft_workspace.snapshot(tmp_path) == before


def test_mixed_review_does_not_lock_a_legitimate_source_repair(tmp_path):
    revised = {**GOOD, "summary": "A source-faithful correction."}
    draft_workspace.write_sections(tmp_path, revised)
    draft_workspace.write_numerals(tmp_path, NUMERALS)
    draft_workspace.write_figures(tmp_path, FIGURES)
    report = {
        "checks": [{
            "name": "Saved candidate passes the current filing preflight",
            "status": "fail", "category": "figures_and_numerals",
        }],
        "findings": [{
            "title": "Unsupported relationship", "category": "disclosure_fidelity",
        }],
    }

    assert draft_studio.restore_sources_after_figure_plan_review(
        tmp_path, {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES},
        report) is False
    assert draft_workspace.read_sections(tmp_path) == revised


def test_a_non_drawing_review_may_repair_the_patent_text(tmp_path):
    draft_workspace.write_sections(tmp_path, {**GOOD, "summary": "Needs source repair."})
    baseline = {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}
    report = {
        "checks": [],
        "findings": [{"category": "disclosure_fidelity", "title": "Unsupported structure"}],
    }

    assert draft_studio.restore_text_after_drawing_only_review(tmp_path, baseline, report) is False
    assert draft_workspace.read_sections(tmp_path)["summary"] == "Needs source repair."


# =============================================================================================
# The workspace is a cache, not the record
# =============================================================================================
def test_the_figure_directory_mirrors_the_stored_version(tmp_path, monkeypatch):
    """A drawing removed from the draft must not survive on disk and reappear in the next review."""
    monkeypatch.setattr(draft_workspace, "root", lambda: tmp_path)
    project = {"id": 1, "title": "t", "disclosure_text": "d" * 60, "input_kind": "description"}
    draft_workspace.build(project=project, figures=FIGURES)
    assert len(draft_workspace.read_figures(draft_workspace.for_project(1))) == 2
    draft_workspace.build(project=project, figures=[FIGURES[0]])
    assert [f["label"] for f in draft_workspace.read_figures(draft_workspace.for_project(1))] == \
        ["FIG. 1"]


def test_prior_art_files_and_the_index_carry_the_citation_keys(tmp_path, monkeypatch):
    monkeypatch.setattr(draft_workspace, "root", lambda: tmp_path)
    workspace = draft_workspace.build(
        project={"id": 2, "title": "t", "disclosure_text": "d" * 60},
        references=[{"publication_number": "US-11223344-B2", "title": "Handheld lifter",
                     "origin": "report", "relevance_summary": "close art",
                     "snapshot": {"abstract": "A lifter.", "claims": "1. A lifter."}}],
        documents=[{"kind": "prior_art", "filename": "brochure.pdf", "title": "A brochure",
                    "body": "Some product text.", "note": "found at a trade show"}])
    index = (workspace / "prior_art" / "INDEX.md").read_text()
    assert "[REF:PUBLICATION]" in index and "US-11223344-B2" in index and "UPLOAD-01" in index
    reference = (workspace / "prior_art" / "US-11223344-B2.md").read_text()
    assert "`[REF:US-11223344-B2]`" in reference and "1. A lifter." in reference
    upload = (workspace / "prior_art" / "UPLOAD-01.md").read_text()
    assert "has NOT been ranked" in upload and "trade show" in upload


def test_a_project_with_no_prior_art_is_told_so_rather_than_left_to_guess(tmp_path, monkeypatch):
    monkeypatch.setattr(draft_workspace, "root", lambda: tmp_path)
    workspace = draft_workspace.build(
        project={"id": 3, "title": "t", "disclosure_text": "d" * 60, "search_slug": ""})
    brief = (workspace / "input" / "brief.md").read_text()
    assert "no search was run" in brief.lower() or "none was run" in brief.lower()
    assert "may be incomplete" in brief


def test_legacy_intake_notes_cannot_request_placeholders_in_a_workspace(tmp_path, monkeypatch):
    monkeypatch.setattr(draft_workspace, "root", lambda: tmp_path)
    workspace = draft_workspace.build(project={
        "id": 4,
        "title": "t",
        "disclosure_text": "d" * 60,
        "inventor_notes": (
            "Filing and drafting instructions:\n"
            "Priority status is not confirmed; leave a drafting note requesting it.\n"
            "Government support status is not confirmed; leave a drafting note requesting it."
        ),
    })

    brief = (workspace / "input" / "brief.md").read_text()
    assert brief.count("Not applicable.") == 2
    assert not re.search(r"drafting note|not confirmed|placeholder", brief, re.IGNORECASE)


def test_the_previous_review_reaches_the_next_iteration(tmp_path, monkeypatch):
    monkeypatch.setattr(draft_workspace, "root", lambda: tmp_path)
    workspace = draft_workspace.build(
        project={"id": 5, "title": "t", "disclosure_text": "d" * 60},
        qa_report={"verdict": "fail", "summary": "numeral 22 is undefined",
                   "checks": [{"name": "Every numeral in the text is defined", "status": "fail",
                               "detail": "22 is not in the table",
                               "items": ["FIG. 2: wrong fastener axis",
                                         "FIG. 7: missing process arrow"]}],
                   "findings": [{"severity": "critical", "title": "Claim 1 unsupported",
                                 "where": "claims", "detail": "no support", "fix": "add support"}]})
    review = (workspace / "review" / "previous-qa.md").read_text()
    assert "numeral 22 is undefined" in review and "Claim 1 unsupported" in review
    assert "add support" in review
    assert "FIG. 2: wrong fastener axis" in review
    assert "FIG. 7: missing process arrow" in review


def test_figure_labels_match_across_spellings():
    import draft_studio_service
    assert draft_studio_service._figure_key("FIG. 1") == draft_studio_service._figure_key("Fig 1")
    assert draft_studio_service._figure_key("FIGURE 2") == draft_studio_service._figure_key("FIG.2")
    assert draft_studio_service._figure_key("FIG. 1") != draft_studio_service._figure_key("FIG. 2")
    verbose = "FIG. 2: side elevation through the pressure chamber and all conduits"
    assert draft_studio_service._figure_key(verbose) == \
        draft_studio_service._figure_key(verbose[:28])


def test_filing_gate_requires_every_check_and_independent_review_to_be_clean():
    clean = {"status": "complete", "checks": [
        {"name": "Claims", "status": "pass"}], "findings": []}
    assert draft_studio.filing_blockers(clean) == []
    assert "Claims" in draft_studio.filing_blockers({
        **clean, "checks": [{"name": "Claims", "status": "warn"}]})[0]
    assert "Unsupported limitation" in draft_studio.filing_blockers({
        **clean, "findings": [{"title": "Unsupported limitation"}]})[0]
    assert "review" in draft_studio.filing_blockers({**clean, "status": "failed"})[0].lower()


@pytest.mark.parametrize("check_name", [
    "Drawing pixels were inspected",
    "Section views have matching source-view cutting lines",
    "Drawing content matches its specification",
    "Drawing leaders identify the named features",
])
def test_drawing_evidence_failures_stay_out_of_the_text_repair_lane(check_name):
    report = {
        "status": "complete",
        "checks": [{"name": check_name, "status": "fail"}],
        "findings": [],
    }

    assert draft_studio.text_blockers(report) == []
    assert check_name in draft_studio.drawing_blockers(report)[0]


def test_default_finalization_budget_allows_drawing_and_text_repair_rounds():
    assert draft_studio.MAX_FINALIZATION_ROUNDS == 6


def test_valid_candidate_is_checkpointed_before_the_long_review(monkeypatch, tmp_path):
    repository = Mock()
    agent = Mock()
    agent.DRAFT_MODEL = "draft-model"
    agent.DRAFT_TIMEOUT = 60
    agent.new_session_id.return_value = "new-session"
    agent.run.return_value = draft_agent.AgentRun(
        ok=True, session_id="draft-session", model="draft-model",
        cost_usd=0.5, duration_ms=1000,
        result={"action": "revised", "summary": "complete candidate",
                "reasoning": [], "changes": [], "questions": [],
                "prior_art_strategy": "", "answer": ""})
    workspace = Mock()
    workspace.snapshot.return_value = {
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, qa=draft_qa, workspace=workspace)
    gate_resume = {
        "session_id": "draft-session", "model": "draft-model", "cost_usd": 0.5,
        "duration_ms": 1000, "num_turns": 1, "steps": [],
        "result": agent.run.return_value.result,
    }
    monkeypatch.setattr(runner, "prepare", lambda _turn: {
        "workspace": tmp_path,
        "project": {"user_id": 91, "agent_session_id": "draft-session",
                    "latest_version_no": 1,
                    "disclosure_text": "disclosure"},
        "references": [{"publication_number": ALLOWED[0]}], "documents": [],
        "seeded": False, "had_version": True, "resuming_candidate": True,
        "resuming_candidate_turn_id": 2,
        "prepared_snapshot": {"sections": GOOD, "numerals": NUMERALS,
                              "figures": FIGURES},
        "prepared_qa": {"_gate_resume": gate_resume},
        "previous_sections": GOOD,
    })
    monkeypatch.setattr(runner, "_reconcile_drawings", Mock(return_value=[]))
    monkeypatch.setattr(
        runner, "evaluate",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(KeyboardInterrupt()))

    with pytest.raises(KeyboardInterrupt):
        runner.run({"id": 3, "lease_token": "lease", "project_id": 7,
                    "turn_no": 2, "kind": "gate_resume", "attempts": 1,
                    "idempotency_key": "auto-filing-repair-2-1"})

    checkpoint = repository.save_retry_candidate.call_args.kwargs
    assert checkpoint["snapshot"]["sections"] == GOOD
    assert checkpoint["report"]["_gate_resume"]["session_id"] == "draft-session"
    assert checkpoint["report"]["_gate_resume"]["result"]["summary"] == \
        "complete candidate"


def test_worker_shutdown_reaches_the_inflight_agent_cancel_event(tmp_path):
    stop = threading.Event()
    stop.set()
    repository = Mock()
    agent = Mock()
    agent.DRAFT_MODEL = "draft-model"
    agent.DRAFT_TIMEOUT = 60
    agent.run.return_value = draft_agent.AgentRun(
        ok=False, cancelled=True, error="worker restart")
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, stop_event=stop)

    with pytest.raises(drafting.DraftingConflict):
        runner._run_agent(
            turn_id=3, lease="lease", workspace=tmp_path, prompt="draft",
            session_id="session", resume=False, transcript=tmp_path / "turn.jsonl",
            stage="drafting")

    assert agent.run.call_args.kwargs["cancel"].is_set() is True


def test_worker_shutdown_stops_between_drawing_sheets(monkeypatch, tmp_path):
    stop = threading.Event()
    stop.set()
    repository = Mock()
    runner = draft_studio.TurnRunner(
        repository, object(), qa=Mock(), stop_event=stop)
    monkeypatch.setattr(
        runner, "_review_sources",
        lambda **_kwargs: {"ok": True, "findings": []})
    monkeypatch.setattr(draft_studio, "filing_blockers", lambda _report: [])
    monkeypatch.setattr(draft_figures, "checkpoint_project_figures", Mock())
    monkeypatch.setattr(draft_figures, "materialize_review_images", Mock(return_value=0))
    observed = {}

    def ensure_project_figures(*_args, check_cancel, **_kwargs):
        observed["continue"] = check_cancel()
        return {"ok": False, "budget_spent": True, "errors": []}

    monkeypatch.setattr(
        draft_figures, "ensure_project_figures", ensure_project_figures)

    runner._ensure_figures(
        turn_id=3, lease="lease", project_id=7, user_id=91,
        sections=GOOD, numerals=NUMERALS, figures=FIGURES,
        disclosure="inventor source", workspace=tmp_path)

    assert observed["continue"] is False
    repository.heartbeat.assert_called_once_with(
        3, "lease", stage="drawing and inspecting figures")


def test_transient_drawing_capacity_retries_the_durable_drawing_turn(
        monkeypatch, tmp_path):
    runner = draft_studio.TurnRunner(Mock(), object(), agent=Mock(), workspace=Mock())
    monkeypatch.setattr(runner, "_drawings_already_match", lambda *_a, **_k: False)
    monkeypatch.setattr(
        runner, "_ensure_figures",
        lambda **_kwargs: (_ for _ in ()).throw(
            draft_figures.FigureTransientError("429 RESOURCE_EXHAUSTED")))

    with pytest.raises(draft_figures.FigureTransientError, match="RESOURCE_EXHAUSTED"):
        runner._reconcile_drawings(
            turn_id=3, lease="lease", project_id=7, user_id=91,
            sections=GOOD, numerals=NUMERALS, figures=FIGURES,
            disclosure="disclosure", workspace=tmp_path)


def test_paid_image_generation_is_serialized_across_worker_slots(monkeypatch):
    active = 0
    maximum_active = 0
    lock = threading.Lock()
    rendezvous = threading.Barrier(2)
    outputs = []

    def unavailable_cursor(*_args, **_kwargs):
        raise RuntimeError("cache database intentionally unavailable")

    def generated(prompt, previous_png=None):
        nonlocal active, maximum_active
        with lock:
            active += 1
            maximum_active = max(maximum_active, active)
        try:
            try:
                rendezvous.wait(timeout=0.2)
            except threading.BrokenBarrierError:
                pass
            return (prompt + str(bool(previous_png))).encode()
        finally:
            with lock:
                active -= 1

    monkeypatch.setattr(draft_figures.db, "cursor", unavailable_cursor)
    monkeypatch.setattr(draft_figures, "generate_png", generated)
    threads = [threading.Thread(
        target=lambda value=value: outputs.append(draft_figures._cached_generate(value)))
        for value in ("first", "second")]
    for thread in threads:
        thread.start()
    for thread in threads:
        thread.join(timeout=2)

    assert len(outputs) == 2
    assert maximum_active == 1


@pytest.mark.parametrize(("configured", "expected"), [
    (None, 1), ("", 1), ("garbage", 1), ("0", 1), ("2", 2), ("99", 4),
])
def test_image_generation_slot_limit_is_bounded_and_invalid_values_fail_safe(
        configured, expected):
    assert draft_figures._image_generation_slot_limit(configured) == expected


def test_default_drawing_window_covers_a_full_multi_sheet_review():
    assert draft_studio.DEFAULT_DRAWING_BUDGET_SECONDS == 3600
    assert draft_studio.DRAWING_BUDGET_SECONDS >= 3600


@pytest.mark.parametrize(("limits", "spent"), [
    ({"max_agent_runs": 14}, {
        "agent_runs": 14, "spend_usd": 0, "tokens_total": 2_414_288,
    }),
    ({"max_spend_usd": 12}, {
        "agent_runs": 8, "spend_usd": 12, "tokens_total": 800_000,
    }),
])
def test_turn_ceiling_message_never_asks_for_manual_intervention(limits, spent):
    runner = draft_studio.TurnRunner(Mock(), object(), qa=Mock(), workspace=Mock())
    runner._budget = limits

    with pytest.raises(draft_studio.TurnBudgetSpent) as failure:
        runner._check_budget(3, spent)

    message = str(failure.value)
    assert "continue automatically" in message
    assert "No manual ceiling change is required" in message
    assert "Raise the ceiling" not in message
    assert "smaller change" not in message


def test_drawing_budget_exhaustion_never_becomes_a_publishable_fault_list(
        monkeypatch, tmp_path):
    runner = draft_studio.TurnRunner(Mock(), object(), qa=Mock(), workspace=Mock())
    monkeypatch.setattr(runner, "_drawings_already_match", lambda *_a, **_k: False)
    monkeypatch.setattr(
        runner, "_ensure_figures",
        lambda **_kwargs: (_ for _ in ()).throw(
            draft_studio.DrawingBudgetSpent("drawing work is incomplete")))

    with pytest.raises(draft_studio.DrawingBudgetSpent, match="incomplete"):
        runner._reconcile_drawings(
            turn_id=3, lease="lease", project_id=7, user_id=91,
            sections=GOOD, numerals=NUMERALS, figures=FIGURES,
            disclosure="disclosure", workspace=tmp_path)


def test_drawing_budget_returns_collected_sheet_faults_before_retrying(
        monkeypatch, tmp_path):
    runner = draft_studio.TurnRunner(Mock(), object(), qa=Mock(), workspace=Mock())
    monkeypatch.setattr(runner, "_drawings_already_match", lambda *_a, **_k: False)
    monkeypatch.setattr(runner, "_ensure_figures", lambda **_kwargs: {
        "ok": False,
        "budget_spent": True,
        "errors": ["FIG. 1: the disclosed linkage is missing"],
    })

    assert runner._reconcile_drawings(
        turn_id=3, lease="lease", project_id=7, user_id=91,
        sections=GOOD, numerals=NUMERALS, figures=FIGURES,
        disclosure="disclosure", workspace=tmp_path,
    ) == ["FIG. 1: the disclosed linkage is missing"]


def test_drawing_budget_without_a_sheet_fault_retries_the_saved_candidate(
        monkeypatch, tmp_path):
    runner = draft_studio.TurnRunner(Mock(), object(), qa=Mock(), workspace=Mock())
    monkeypatch.setattr(runner, "_drawings_already_match", lambda *_a, **_k: False)
    monkeypatch.setattr(runner, "_ensure_figures", lambda **_kwargs: {
        "ok": False, "budget_spent": True, "errors": [],
    })

    with pytest.raises(draft_studio.DrawingBudgetSpent, match="time budget"):
        runner._reconcile_drawings(
            turn_id=3, lease="lease", project_id=7, user_id=91,
            sections=GOOD, numerals=NUMERALS, figures=FIGURES,
            disclosure="disclosure", workspace=tmp_path)


@pytest.mark.parametrize("error", [
    "DrawingBudgetSpent: the bounded drawing pass stopped between sheets",
    "This turn reached its ceiling of $12.00. The saved draft was not published.",
])
def test_drawing_continuation_retry_preserves_completed_sheets(error):
    import draft_studio_service

    runner = Mock()
    runner.repository.fail_turn.return_value = {"status": "queued"}
    result = draft_studio_service._fail(
        runner,
        {"id": 4, "project_id": 7, "requested_by_user_id": 91,
         "project_revision": 1, "lease_token": "lease"},
        error,
        retryable=True,
    )

    assert result["status"] == "queued"
    runner.restore_figures.assert_not_called()


def test_stale_drawing_approval_cannot_skip_the_current_gates(monkeypatch):
    spec = FIGURES[0]
    expected_hash = draft_figures.specification_hash(
        spec["label"], spec["caption"], draft_figures.expected_entries(spec, NUMERALS))
    stored = [{
        "id": 11, "figure_label": spec["label"], "active_version": 2,
        "versions": [{
            "version_no": 2,
            "semantic_audit": {"ok": True, "specification_hash": expected_hash},
            "leader_audit": {"ok": True, "specification_hash": expected_hash},
            "numeral_audit": {"ok": True, "inspected": True},
        }],
    }]
    monkeypatch.setattr(draft_figures, "listing", lambda *_args: stored)
    monkeypatch.setattr(draft_figures, "current_ocr_audit", lambda *_a, **_k: True)
    monkeypatch.setattr(draft_figures, "current_leader_audit", lambda *_a, **_k: True)
    monkeypatch.setattr(draft_figures, "current_semantic_audit", lambda *_a, **_k: False)
    runner = draft_studio.TurnRunner(Mock(), object(), agent=Mock(), workspace=Mock())

    assert runner._drawings_already_match(7, 91, NUMERALS, [spec]) is False

    monkeypatch.setattr(draft_figures, "current_semantic_audit", lambda *_a, **_k: True)
    assert runner._drawings_already_match(7, 91, NUMERALS, [spec]) is True


def test_generated_pixels_cannot_skip_a_new_current_exact_renderer(monkeypatch):
    spec = FIGURES[0]
    expected_hash = draft_figures.specification_hash(
        spec["label"], spec["caption"], draft_figures.expected_entries(spec, NUMERALS))
    stored = [{
        "id": 11, "figure_label": spec["label"], "active_version": 2,
        "versions": [{
            "version_no": 2,
            "semantic_audit": {"ok": True, "specification_hash": expected_hash},
            "leader_audit": {"ok": True, "specification_hash": expected_hash},
            "numeral_audit": {"ok": True, "inspected": True},
        }],
    }]
    monkeypatch.setattr(draft_figures, "listing", lambda *_args: stored)
    monkeypatch.setattr(draft_figures, "current_ocr_audit", lambda *_a, **_k: True)
    monkeypatch.setattr(draft_figures, "current_leader_audit", lambda *_a, **_k: True)
    monkeypatch.setattr(draft_figures, "current_semantic_audit", lambda *_a, **_k: True)
    binding = Mock(return_value=False)
    monkeypatch.setattr(draft_figures, "current_geometry_binding", binding)
    runner = draft_studio.TurnRunner(Mock(), object(), agent=Mock(), workspace=Mock())

    assert runner._drawings_already_match(7, 91, NUMERALS, [spec]) is False
    binding.assert_called_once_with(stored[0], 91, stored[0]["versions"][0], spec["caption"])

    binding.return_value = True
    assert runner._drawings_already_match(7, 91, NUMERALS, [spec]) is True


def test_current_drawings_are_materialized_before_the_independent_review(monkeypatch, tmp_path):
    runner = draft_studio.TurnRunner(Mock(), object(), agent=Mock(), workspace=Mock())
    monkeypatch.setattr(runner, "_drawings_already_match", lambda *_args: True)
    materialize = Mock(return_value=len(FIGURES))
    monkeypatch.setattr(draft_figures, "materialize_review_images", materialize)
    ensure = Mock()
    monkeypatch.setattr(runner, "_ensure_figures", ensure)

    faults = runner._reconcile_drawings(
        turn_id=3, lease="lease", project_id=7, user_id=91,
        sections=GOOD, numerals=NUMERALS, figures=FIGURES,
        disclosure="disclosure", workspace=tmp_path)

    assert faults == []
    materialize.assert_called_once_with(7, 91, tmp_path)
    ensure.assert_not_called()


def test_missing_review_image_bytes_block_the_final_review(monkeypatch, tmp_path):
    runner = draft_studio.TurnRunner(Mock(), object(), agent=Mock(), workspace=Mock())
    monkeypatch.setattr(runner, "_drawings_already_match", lambda *_args: True)
    monkeypatch.setattr(draft_figures, "materialize_review_images", Mock(return_value=1))
    monkeypatch.setattr(runner, "_ensure_figures", Mock())

    faults = runner._reconcile_drawings(
        turn_id=3, lease="lease", project_id=7, user_id=91,
        sections=GOOD, numerals=NUMERALS, figures=FIGURES,
        disclosure="disclosure", workspace=tmp_path)

    assert len(faults) == 1
    assert "1 of 2" in faults[0]
    assert "independent review" in faults[0]


def test_newly_checked_drawings_need_every_review_image(monkeypatch, tmp_path):
    runner = draft_studio.TurnRunner(Mock(), object(), agent=Mock(), workspace=Mock())
    monkeypatch.setattr(runner, "_drawings_already_match", lambda *_args: False)
    monkeypatch.setattr(
        runner, "_ensure_figures", Mock(return_value={"ok": True, "review_images": 1}))

    faults = runner._reconcile_drawings(
        turn_id=3, lease="lease", project_id=7, user_id=91,
        sections=GOOD, numerals=NUMERALS, figures=FIGURES,
        disclosure="disclosure", workspace=tmp_path)

    assert len(faults) == 1
    assert "1 of 2" in faults[0]
    assert "independent review" in faults[0]


def test_restart_resumes_a_checkpointed_candidate_without_rerunning_the_agent(
        monkeypatch, tmp_path):
    monkeypatch.setattr(draft_figures, "discard_project_figure_checkpoint",
                        lambda _turn_id: False)
    repository = Mock()
    repository.save_version.return_value = {"version_no": 2}
    repository.save_qa.return_value = {
        "id": 5, "verdict": "pass", "checks": [], "findings": [], "counts": {}}
    repository.complete_turn.return_value = {"status": "complete"}
    agent = Mock()
    agent.DRAFT_MODEL = "draft-model"
    agent.DRAFT_TIMEOUT = 60
    agent.strings.side_effect = lambda value, **_kwargs: list(value or [])
    workspace = Mock()
    workspace.snapshot.return_value = {
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, qa=draft_qa, workspace=workspace)
    gate_resume = {
        "session_id": "draft-session", "model": "draft-model", "cost_usd": 0.5,
        "duration_ms": 1000, "num_turns": 1, "steps": [],
        "result": {"action": "revised", "summary": "complete candidate",
                   "reasoning": [], "changes": [], "questions": [],
                   "prior_art_strategy": "", "answer": ""},
    }
    monkeypatch.setattr(runner, "prepare", lambda _turn: {
        "workspace": tmp_path,
        "project": {"user_id": 91, "agent_session_id": "draft-session",
                    "latest_version_no": 1, "disclosure_text": "disclosure"},
        "references": [{"publication_number": ALLOWED[0]}], "documents": [],
        "seeded": False, "had_version": True, "resuming_candidate": True,
        "resuming_candidate_turn_id": 3,
        "prepared_snapshot": {"sections": GOOD, "numerals": NUMERALS,
                              "figures": FIGURES},
        "prepared_qa": {"_gate_resume": gate_resume},
        "previous_sections": {},
    })
    ensure_figures = Mock(return_value={"ok": True, "review_images": len(FIGURES)})
    monkeypatch.setattr(runner, "_ensure_figures", ensure_figures)
    monkeypatch.setattr(runner, "evaluate", lambda *args, **kwargs: {
        "status": "complete", "verdict": "pass", "summary": "ready",
        "checks": [], "findings": [], "counts": {}, "cost_usd": 0,
        "duration_ms": 1, "model_name": "review"})

    runner.run({"id": 4, "lease_token": "lease", "project_id": 7,
                "turn_no": 2, "kind": "gate_resume", "attempts": 1,
                "idempotency_key": "auto-filing-repair-3-1"})

    agent.run.assert_not_called()
    ensure_figures.assert_called_once()
    assert repository.save_version.call_count == 1
    assert repository.complete_turn.call_args.kwargs["session_id"] == "draft-session"


def test_resumed_figure_plan_is_compared_with_the_published_version():
    repaired_figures = [
        {**FIGURES[0], "caption": FIGURES[0]["caption"] + " The outlet is shown separately."},
        FIGURES[1],
    ]
    context = {
        "published_snapshot": {
            "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES},
        "prepared_snapshot": {
            "sections": GOOD, "numerals": NUMERALS, "figures": repaired_figures},
        "previous_sections": GOOD,
    }

    assert draft_studio._candidate_differs_from_published(context, {
        "sections": GOOD, "numerals": NUMERALS, "figures": repaired_figures,
    }) is True
    assert draft_studio._candidate_differs_from_published(context, {
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES,
    }) is False


def test_drawing_faults_are_repaired_before_the_independent_review(monkeypatch, tmp_path):
    repository = Mock()
    agent = Mock()
    agent.DRAFT_MODEL = "draft-model"
    agent.DRAFT_TIMEOUT = 60
    agent.strings.side_effect = lambda value, **_kwargs: list(value or [])
    workspace = Mock()
    workspace.snapshot.return_value = {
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, qa=draft_qa, workspace=workspace)
    gate_resume = {
        "session_id": "draft-session", "model": "draft-model", "cost_usd": 0.5,
        "duration_ms": 1000, "num_turns": 1, "steps": [],
        "result": {"action": "revised", "summary": "complete candidate",
                   "reasoning": [], "changes": [], "questions": [],
                   "prior_art_strategy": "", "answer": ""},
    }
    monkeypatch.setattr(runner, "prepare", lambda _turn: {
        "workspace": tmp_path,
        "project": {"user_id": 91, "agent_session_id": "draft-session",
                    "latest_version_no": 1, "disclosure_text": "disclosure"},
        "references": [{"publication_number": ALLOWED[0]}], "documents": [],
        "seeded": False, "had_version": True, "resuming_candidate": True,
        "resuming_candidate_turn_id": 3,
        "prepared_snapshot": {"sections": GOOD, "numerals": NUMERALS,
                              "figures": FIGURES},
        "prepared_qa": {"_gate_resume": gate_resume},
        "previous_sections": {},
    })
    monkeypatch.setattr(
        runner, "_reconcile_drawings",
        Mock(return_value=["FIG. 1 failed final-pixel endpoint inspection"]))
    independent_review = Mock(return_value={
        "status": "complete", "verdict": "pass", "summary": "ready",
        "checks": [], "findings": [], "counts": {}, "cost_usd": 0,
        "duration_ms": 1, "model_name": "review"})
    monkeypatch.setattr(runner, "evaluate", independent_review)
    monkeypatch.setattr(
        runner, "_run_agent",
        Mock(side_effect=RuntimeError("stop after the first drawing repair report")))

    with pytest.raises(RuntimeError, match="first drawing repair report"):
        runner.run({"id": 4, "lease_token": "lease", "project_id": 7,
                    "turn_no": 2, "kind": "gate_resume", "attempts": 1,
                    "idempotency_key": "auto-filing-repair-3-1"})

    independent_review.assert_not_called()
    saved_report = repository.save_retry_candidate.call_args.kwargs["report"]
    assert saved_report["findings"] == []
    assert saved_report["checks"][0]["category"] == "figures_and_numerals"
    assert saved_report["summary"] == "1 drawing issue requires automatic repair."
    assert saved_report["checks"][0]["detail"].startswith("1 drawing issue failed.")
    assert saved_report["checks"][0]["items"] == [
        "FIG. 1 failed final-pixel endpoint inspection"]

    error = draft_studio.DrawingInspectionError([
        "FIG. 1 is overcrowded", "FIG. 2 is overcrowded", "FIG. 2 is overlong"])
    assert str(error) == "3 drawing issues did not pass inspection."
    assert str(draft_studio.DrawingInspectionError(["FIG. 1 is overcrowded"])) == (
        "1 drawing issue did not pass inspection."
    )


def test_plan_preflight_fault_is_repaired_before_any_image_call(monkeypatch, tmp_path):
    repository = Mock()
    agent = Mock()
    agent.DRAFT_MODEL = "draft-model"
    agent.DRAFT_TIMEOUT = 60
    agent.strings.side_effect = lambda value, **_kwargs: list(value or [])
    figures = [
        {
            **FIGURES[0],
            "caption": (
                "The cord 10 is identified by a leader ending in an arrowhead whose tip "
                "touches the cord."
            ),
        },
        FIGURES[1],
    ]
    workspace = Mock()
    workspace.snapshot.return_value = {
        "sections": GOOD, "numerals": NUMERALS, "figures": figures}
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, qa=draft_qa, workspace=workspace)
    gate_resume = {
        "session_id": "draft-session", "model": "draft-model", "cost_usd": 0.5,
        "duration_ms": 1000, "num_turns": 1, "steps": [],
        "result": {"action": "revised", "summary": "complete candidate",
                   "reasoning": [], "changes": [], "questions": [],
                   "prior_art_strategy": "", "answer": ""},
    }
    monkeypatch.setattr(runner, "prepare", lambda _turn: {
        "workspace": tmp_path,
        "project": {"user_id": 91, "agent_session_id": "draft-session",
                    "latest_version_no": 1, "disclosure_text": "disclosure"},
        "references": [{"publication_number": ALLOWED[0]}], "documents": [],
        "seeded": False, "had_version": True, "resuming_candidate": True,
        "resuming_candidate_turn_id": 3,
        "prepared_snapshot": {"sections": GOOD, "numerals": NUMERALS,
                              "figures": figures},
        "prepared_qa": {"_gate_resume": gate_resume},
        "previous_sections": {},
    })
    reconcile = Mock(side_effect=AssertionError("image lane invoked"))
    monkeypatch.setattr(runner, "_reconcile_drawings", reconcile)
    monkeypatch.setattr(
        runner, "_run_agent",
        Mock(side_effect=RuntimeError("stop after the plan repair report")))

    with pytest.raises(RuntimeError, match="plan repair report"):
        runner.run({"id": 4, "lease_token": "lease", "project_id": 7,
                    "turn_no": 2, "kind": "gate_resume", "attempts": 1,
                    "idempotency_key": "auto-filing-repair-3-1"})

    reconcile.assert_not_called()
    saved_report = repository.save_retry_candidate.call_args.kwargs["report"]
    assert saved_report["checks"][0]["name"] == "Drawing plans pass deterministic preflight"
    assert "terminal dot" in saved_report["checks"][0]["items"][0]


def test_automatic_continuation_reuses_a_prior_turn_gate_checkpoint():
    marker = {
        "session_id": "draft-session", "model": "draft-model", "cost_usd": 0.5,
        "duration_ms": 1000, "num_turns": 1, "steps": [],
        "result": {"action": "revised", "summary": "complete candidate",
                   "reasoning": [], "changes": [], "questions": [],
                   "prior_art_strategy": "", "answer": ""},
    }
    context = {
        "resuming_candidate_turn_id": 60,
        "prepared_qa": {"_gate_resume": marker},
    }

    run = draft_studio._gate_resume_run(context, {
        "id": 61, "kind": "gate_resume",
        "idempotency_key": "auto-filing-repair-60-1",
    })

    assert run is not None
    assert run.session_id == "draft-session"
    assert run.result["summary"] == "complete candidate"


def test_internal_gate_resume_never_redrafts_when_its_retry_key_was_rewritten():
    marker = {
        "session_id": "draft-session", "model": "draft-model", "cost_usd": 0.5,
        "duration_ms": 1000, "num_turns": 1, "steps": [],
        "result": {"action": "revised", "summary": "complete candidate",
                   "reasoning": [], "changes": [], "questions": [],
                   "prior_art_strategy": "", "answer": ""},
    }
    context = {
        "resuming_candidate_turn_id": 60,
        "prepared_qa": {"_gate_resume": marker},
    }

    run = draft_studio._gate_resume_run(context, {
        "id": 61, "kind": "gate_resume",
        "idempotency_key": "operator-recovery-60-1",
    })

    assert run is not None
    assert run.session_id == "draft-session"


def test_internal_gate_resume_uses_a_legacy_candidate_without_a_result_marker():
    run = draft_studio._gate_resume_run({
        "resuming_candidate_turn_id": 60,
        "prepared_qa": {"verdict": "fail", "summary": "legacy checkpoint"},
        "project": {"agent_session_id": ""},
    }, {
        "id": 61, "kind": "gate_resume",
        "idempotency_key": "auto-filing-repair-60-1",
    })

    assert run is not None
    assert run.ok
    assert run.session_id == ""
    assert run.result["action"] == "revised"


def test_repeated_gate_resume_reuses_a_saved_checkpoint_with_no_provider_session():
    marker = {
        "session_id": "", "model": "saved-candidate", "cost_usd": 0.0,
        "duration_ms": 0, "num_turns": 0, "steps": [],
        "result": {"action": "revised", "summary": "restored candidate",
                   "reasoning": [], "changes": [], "questions": [],
                   "prior_art_strategy": "", "answer": ""},
    }
    run = draft_studio._gate_resume_run({
        "resuming_candidate_turn_id": 60,
        "prepared_qa": {"_gate_resume": marker},
    }, {
        "id": 61, "kind": "gate_resume",
        "idempotency_key": "auto-filing-repair-60-1",
    })

    assert run is not None
    assert run.session_id == ""
    assert run.model == "saved-candidate"
    assert run.result["summary"] == "restored candidate"


def test_client_qa_fix_cannot_reuse_a_prior_turn_gate_checkpoint():
    context = {
        "resuming_candidate_turn_id": 60,
        "prepared_qa": {"_gate_resume": {
            "session_id": "draft-session",
            "result": {"action": "revised", "summary": "complete candidate"},
        }},
    }

    assert draft_studio._gate_resume_run(context, {
        "id": 61, "kind": "qa_fix",
        "idempotency_key": "auto-filing-repair-60-1",
    }) is None


def test_turn_runner_publishes_text_and_queues_no_drawing_gate(
        monkeypatch, tmp_path):
    """A finished turn is finished. There is no second, automatic drawing phase behind it.

    The runner used to publish the text and then queue itself a `gate_resume` turn that generated
    every sheet, inspected the pixels, and only then let the project be called filing-ready. That
    lane is gone with the image generation it drove, so what has to be true now is the opposite of
    what this test used to assert: no continuation, and no candidate held back waiting for one.
    """
    class Repository:
        def __init__(self):
            self.saved_versions = []
            self.saved_reports = []
            self.retry_candidates = []
            self.messages = []

        def heartbeat(self, *_args, **_kwargs):
            pass

        def save_version(self, *_args, **kwargs):
            self.saved_versions.append(kwargs)
            return {"version_no": 1}

        def save_qa(self, _project_id, **kwargs):
            self.saved_reports.append(kwargs)
            return {"id": 4, **kwargs["report"]}

        def save_retry_candidate(self, _turn_id, _lease, *, snapshot, report):
            self.retry_candidates.append((snapshot, report))

        def add_message(self, _project_id, role, body, **kwargs):
            self.messages.append((role, body, kwargs))

        def complete_turn(self, *_args, **kwargs):
            return {"status": "complete", **kwargs}

    class Agent:
        DRAFT_MODEL = "draft-model"
        DRAFT_TIMEOUT = 60

        def __init__(self, repository):
            self.calls = []
            self.repository = repository

        @staticmethod
        def new_session_id():
            return "new-session"

        @staticmethod
        def strings(value, **_kwargs):
            return list(value or [])

        def run(self, **kwargs):
            self.calls.append(kwargs)
            if len(self.calls) == 2:
                assert self.repository.saved_versions == [], \
                    "a failing intermediate draft must never be published"
            return draft_agent.AgentRun(
                ok=True, session_id="draft-session", model="draft-model",
                cost_usd=0.5, duration_ms=1000,
                result={"action": "revised", "summary": f"round {len(self.calls)}",
                        "reasoning": [], "changes": [], "questions": [],
                        "prior_art_strategy": "", "answer": ""})

    class Workspace:
        def __init__(self):
            self.review_reports = []

        def snapshot(self, _path):
            return {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}

        def _write_review(self, _path, report):
            self.review_reports.append(report)

    repository = Repository()
    agent = Agent(repository)
    workspace = Workspace()
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, qa=object(), workspace=workspace)
    monkeypatch.setattr(runner, "prepare", lambda _turn: {
        "workspace": tmp_path, "project": {"user_id": 91, "agent_session_id": "",
                                             "latest_version_no": 0,
                                             "disclosure_text": "disclosure"},
        "references": [{"publication_number": ALLOWED[0]}], "documents": [],
        "seeded": False, "had_version": False,
        "previous_sections": {},
    })
    text_report = {
        "status": "complete", "verdict": "pass", "summary": "text gates passed",
        "checks": [{"name": "Claims", "status": "pass"}], "findings": [],
        "counts": {}, "cost_usd": 0, "duration_ms": 1,
        "model_name": "deterministic checks"}
    mechanical = Mock(return_value=text_report)
    final_review = Mock()
    monkeypatch.setattr(runner, "mechanical_report", mechanical)
    monkeypatch.setattr(runner, "evaluate", final_review)

    out = runner.run({"id": 3, "lease_token": "lease", "project_id": 7,
                      "turn_no": 1, "kind": "initial"})

    assert len(agent.calls) == 1
    assert len(repository.saved_versions) == 1
    assert len(repository.saved_reports) == 1
    assert repository.saved_reports[0]["report"]["verdict"] == "pass"
    assert repository.retry_candidates[0][0]["sections"] == GOOD
    mechanical.assert_called_once()
    final_review.assert_not_called()
    assert out["turn"]["continuation"] is None
    assert out["turn"]["discard_candidates"] is True
    assert out["version"]["version_no"] == 1


def test_retry_preparation_uses_the_durable_checked_candidate_instead_of_published_text(
        monkeypatch, tmp_path):
    candidate_sections = {**GOOD, "summary": "Repaired candidate summary."}
    candidate_report = {
        "status": "complete", "verdict": "fail", "summary": "Fix claim wording.",
        "checks": [{"name": "Claims", "status": "fail", "items": ["claim 1"]}],
        "findings": [],
    }
    repository = Mock()
    repository.documents.return_value = []
    repository.messages.return_value = []
    repository.latest_qa.return_value = {"summary": "stale published review"}
    repository.retry_candidate.return_value = {
        "snapshot": {"sections": candidate_sections, "numerals": NUMERALS,
                     "figures": FIGURES},
        "qa_report": candidate_report,
    }
    repository.latest_retry_candidate.return_value = None
    workspace = Mock()
    workspace.build.return_value = tmp_path
    runner = draft_studio.TurnRunner(repository, Mock(), workspace=workspace)
    monkeypatch.setattr(runner, "_load", lambda _project_id: {
        "project": {"id": 7, "user_id": 91, "input_kind": "description"},
        "references": [{"publication_number": ALLOWED[0]}],
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES,
    })

    context = runner.prepare({"id": 33, "project_id": 7, "attempts": 1,
                              "user_message": "Finish automatically."})

    values = workspace.build.call_args.kwargs
    assert values["sections"] == candidate_sections
    assert values["qa_report"] == candidate_report
    assert context["previous_sections"] == GOOD
    assert context["published_snapshot"] == {
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}


def test_turn_preparation_keeps_the_complete_bounded_inventor_history(monkeypatch, tmp_path):
    repository = Mock()
    repository.documents.return_value = []
    repository.messages.return_value = [
        {"id": index, "role": "user", "body": f"inventor amendment {index}"}
        for index in range(1, 41)
    ]
    repository.latest_qa.return_value = None
    repository.retry_candidate.return_value = None
    repository.latest_retry_candidate.return_value = None
    workspace = Mock()
    workspace.build.return_value = tmp_path
    runner = draft_studio.TurnRunner(repository, Mock(), workspace=workspace)
    monkeypatch.setattr(runner, "_load", lambda _project_id: {
        "project": {"id": 7, "user_id": 91, "input_kind": "description"},
        "references": [], "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES,
    })

    runner.prepare({"id": 33, "project_id": 7, "attempts": 1,
                    "user_message": "Finish automatically."})

    history = workspace.build.call_args.kwargs["conversation"]
    assert [item["id"] for item in history] == list(range(1, 41))
    repository.messages.assert_called_once_with(7, limit=400)


def test_new_turn_preparation_uses_the_latest_failed_turn_candidate(
        monkeypatch, tmp_path):
    candidate_sections = {**GOOD, "summary": "Best unpublished candidate."}
    candidate_report = {
        "status": "complete", "verdict": "fail", "summary": "Finish claim repair.",
        "checks": [{"name": "Claims", "status": "fail", "items": ["claim 1"]}],
        "findings": [],
    }
    repository = Mock()
    repository.documents.return_value = []
    repository.messages.return_value = []
    repository.latest_qa.return_value = {"summary": "stale published review"}
    repository.retry_candidate.return_value = None
    repository.latest_retry_candidate.return_value = {
        "turn_id": 32,
        "snapshot": {"sections": candidate_sections, "numerals": NUMERALS,
                     "figures": FIGURES},
        "qa_report": candidate_report,
    }
    workspace = Mock()
    workspace.build.return_value = tmp_path
    runner = draft_studio.TurnRunner(repository, Mock(), workspace=workspace)
    monkeypatch.setattr(runner, "_load", lambda _project_id: {
        "project": {"id": 7, "user_id": 91, "input_kind": "description"},
        "references": [{"publication_number": ALLOWED[0]}],
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES,
    })

    context = runner.prepare({"id": 33, "project_id": 7, "attempts": 1,
                              "user_message": "Finish automatically."})

    values = workspace.build.call_args.kwargs
    assert values["sections"] == candidate_sections
    assert values["qa_report"] == candidate_report
    repository.latest_retry_candidate.assert_called_once_with(7, before_turn_id=33)
    assert context["resuming_candidate"] is True
    assert context["previous_sections"] == GOOD


def test_a_candidate_blocked_by_a_new_preflight_gate_is_retained_for_repair(
        monkeypatch, tmp_path):
    candidate_figures = [{
        "label": "FIG. 1", "caption": "assembly view",
        "numerals": [str(value) for value in range(10, 28, 2)],
    }]
    candidate_snapshot = {
        "sections": GOOD,
        "numerals": [
            {"numeral": str(value), "part": f"part {value}"}
            for value in range(10, 28, 2)
        ],
        "figures": candidate_figures,
    }
    repository = Mock()
    repository.documents.return_value = []
    repository.messages.return_value = []
    repository.latest_qa.return_value = {"summary": "stale published review"}
    repository.retry_candidate.return_value = None
    repository.latest_retry_candidate.return_value = {
        "turn_id": 32,
        "snapshot": candidate_snapshot,
        "qa_report": {
            "status": "complete", "verdict": "fail", "summary": "Drawing repair needed.",
            "checks": [{"name": "Drafting run completed", "status": "fail"}],
            "findings": [],
        },
    }
    workspace = Mock()
    workspace.build.return_value = tmp_path
    runner = draft_studio.TurnRunner(repository, Mock(), workspace=workspace)
    monkeypatch.setattr(runner, "_load", lambda _project_id: {
        "project": {"id": 7, "user_id": 91, "input_kind": "description"},
        "references": [{"publication_number": ALLOWED[0]}],
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES,
    })

    context = runner.prepare({"id": 33, "project_id": 7, "attempts": 1,
                              "user_message": "Finish automatically."})

    values = workspace.build.call_args.kwargs
    assert values["figures"] == candidate_figures
    assert values["qa_report"]["verdict"] == "fail"
    assert any("current filing preflight" in item["name"].lower()
               for item in values["qa_report"]["checks"])
    current_gate = next(item for item in values["qa_report"]["checks"]
                        if "current filing preflight" in item["name"].lower())
    assert current_gate["category"] == "figures_and_numerals"
    assert all(item["name"] != "Drafting run completed"
               for item in values["qa_report"]["checks"])
    assert "more than 8 numerals" in json.dumps(values["qa_report"])
    assert context["resuming_candidate"] is True
    assert context["prepared_snapshot"] == candidate_snapshot
    assert context["prepared_qa"] == values["qa_report"]
    repository.discard_retry_candidate.assert_not_called()


def test_a_structurally_corrupt_candidate_is_discarded_instead_of_repaired(
        monkeypatch, tmp_path):
    repository = Mock()
    repository.documents.return_value = []
    repository.messages.return_value = []
    repository.latest_qa.return_value = {"summary": "published review"}
    repository.retry_candidate.return_value = None
    repository.latest_retry_candidate.return_value = {
        "turn_id": 32,
        "snapshot": {"sections": {"title": "partial"}, "numerals": "bad", "figures": []},
        "qa_report": {"verdict": "fail"},
    }
    workspace = Mock()
    workspace.build.return_value = tmp_path
    runner = draft_studio.TurnRunner(repository, Mock(), workspace=workspace)
    monkeypatch.setattr(runner, "_load", lambda _project_id: {
        "project": {"id": 7, "user_id": 91, "input_kind": "description"},
        "references": [{"publication_number": ALLOWED[0]}],
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES,
    })

    context = runner.prepare({"id": 33, "project_id": 7, "attempts": 1,
                              "user_message": "Finish automatically."})

    values = workspace.build.call_args.kwargs
    assert values["sections"] == GOOD and values["figures"] == FIGURES
    assert context["resuming_candidate"] is False
    repository.discard_retry_candidate.assert_called_once_with(32)


def test_an_agent_budget_stop_checkpoints_the_valid_workspace_for_retry(monkeypatch, tmp_path):
    repository = Mock()
    repository.documents.return_value = []
    agent = Mock()
    agent.DRAFT_MODEL = "draft-model"
    agent.DRAFT_TIMEOUT = 60
    agent.new_session_id.return_value = "new-session"
    agent.run.return_value = draft_agent.AgentRun(
        ok=False, session_id="session", model="draft-model",
        error="Reached maximum budget ($12)")
    workspace = Mock()
    workspace.snapshot.return_value = {
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, qa=draft_qa, workspace=workspace)
    monkeypatch.setattr(runner, "prepare", lambda _turn: {
        "workspace": tmp_path, "project": {"user_id": 91, "agent_session_id": "",
                                             "latest_version_no": 0,
                                             "disclosure_text": "disclosure"},
        "references": [{"publication_number": ALLOWED[0]}], "documents": [],
        "seeded": False, "had_version": False,
        "previous_sections": {},
    })

    with pytest.raises(draft_studio.StudioError, match="maximum budget"):
        runner.run({"id": 3, "lease_token": "lease", "project_id": 7,
                    "turn_no": 1, "kind": "initial"})

    snapshot = repository.save_retry_candidate.call_args.kwargs["snapshot"]
    report = repository.save_retry_candidate.call_args.kwargs["report"]
    assert snapshot["sections"] == GOOD
    assert report["verdict"] == "fail"
    assert "continue" in report["summary"].lower()


def test_an_interrupted_repair_keeps_the_existing_filing_gate_findings(tmp_path):
    prior_report = {
        "status": "failed",
        "verdict": "fail",
        "summary": "Two drawing endpoints require automatic repair.",
        "checks": [{
            "name": "Every drawing sheet passes inspection",
            "status": "fail",
            "severity": "error",
            "items": ["FIG. 2: numeral 26 misses the bearing face"],
        }],
        "findings": [{"title": "FIG. 2 endpoint mismatch", "severity": "major"}],
        "counts": {"checks": 1, "checks_failed": 1, "findings": 1, "major": 1},
    }
    repository = Mock()
    repository.retry_candidate.return_value = {
        "snapshot": {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES},
        "qa_report": prior_report,
    }
    workspace = Mock()
    workspace.snapshot.return_value = {
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}
    runner = draft_studio.TurnRunner(
        repository, object(), agent=Mock(), qa=draft_qa, workspace=workspace)

    runner._checkpoint_interrupted_agent(
        turn_id=3, lease="lease", workspace=tmp_path, allowed=ALLOWED,
        error=draft_studio.StudioError("Reached maximum repair budget"))

    saved = repository.save_retry_candidate.call_args.kwargs["report"]
    assert saved["checks"] == prior_report["checks"]
    assert saved["findings"] == prior_report["findings"]
    assert saved["counts"] == prior_report["counts"]
    assert saved["summary"].startswith(prior_report["summary"])
    assert "maximum repair budget" in saved["last_error"]


def test_terminal_failure_retains_candidate_until_a_project_completes(monkeypatch):
    queries = []

    class Cursor:
        def execute(self, query, params=()):
            queries.append((query, params))

        def fetchone(self):
            return {
                "id": 33, "project_id": 7, "turn_no": 4,
                "attempts": 3, "max_attempts": 3, "status": "failed",
            }

    @contextmanager
    def cursor_factory(**_kwargs):
        yield Cursor()

    repository = draft_studio.StudioRepository(cursor_factory, migrate=False)
    turn = {
        "id": 33, "project_id": 7, "turn_no": 4,
        "attempts": 3, "max_attempts": 3, "status": "running",
    }
    monkeypatch.setattr(repository, "_verify", lambda *_args: turn)

    result = repository.fail_turn(33, "lease", "budget stopped", retryable=True)

    assert result["status"] == "failed"
    assert not any("DELETE FROM app_draft_turn_candidates" in query for query, _ in queries)

    queries.clear()
    repository.complete_turn(
        33, "lease", result={}, session_id="session", cost_usd=1,
        duration_ms=1000, model_name="model")
    cleanup = next(query for query, _ in queries
                   if "DELETE FROM app_draft_turn_candidates" in query)
    assert "USING app_draft_turns" in cleanup and "t.project_id" in cleanup

    queries.clear()
    repository.complete_turn(
        33, "lease", result={}, session_id="session", cost_usd=1,
        duration_ms=1000, model_name="model", discard_candidates=False)
    assert not any("DELETE FROM app_draft_turn_candidates" in query for query, _ in queries)


def test_cancelling_an_ordinary_turn_restores_the_published_version_state():
    queries = []

    class Cursor:
        def execute(self, query, params=()):
            queries.append((query, params))

        def fetchone(self):
            return {"kind": "revise"}

    @contextmanager
    def cursor_factory(**_kwargs):
        yield Cursor()

    repository = draft_studio.StudioRepository(cursor_factory, migrate=False)
    repository.cancel_turn(7, 33)

    assert any("THEN 'ready'" in query for query, _params in queries)
    assert any("DELETE FROM app_draft_turn_candidates" in query for query, _params in queries)


def test_complete_turn_atomically_queues_its_drawing_continuation(monkeypatch):
    queries = []
    rows = iter([
        {"id": 33, "project_id": 7, "turn_no": 4, "requested_by_user_id": 91,
         "project_revision": 2, "status": "complete"},
        {"n": 5},
        {"id": 44},
    ])

    class Cursor:
        def execute(self, query, params=()):
            queries.append((query, params))

        def fetchone(self):
            return next(rows)

    @contextmanager
    def cursor_factory(**_kwargs):
        yield Cursor()

    repository = draft_studio.StudioRepository(cursor_factory, migrate=False)
    monkeypatch.setattr(repository, "_verify", lambda *_args: {
        "id": 33, "project_id": 7, "turn_no": 4, "requested_by_user_id": 91,
        "project_revision": 2, "status": "running",
    })

    result = repository.complete_turn(
        33, "lease", result={}, session_id="session", cost_usd=1,
        duration_ms=1000, model_name="model", discard_candidates=False,
        continuation={
            "kind": "gate_resume", "idempotency_key": "auto-filing-repair-33-1",
            "user_message": "Finish all drawings automatically.",
        })

    assert result["continuation_turn_id"] == 44
    insert = next((query, params) for query, params in queries
                  if "INSERT INTO app_draft_turns" in query)
    assert "gate_resume" in insert[1]
    assert "auto-filing-repair-33-1" in insert[1]
    assert any("SET status='queued'" in query for query, _params in queries)
    assert not any("DELETE FROM app_draft_turn_candidates" in query for query, _ in queries)


def test_complete_turn_refuses_ready_when_checked_drawing_rows_disappear(monkeypatch):
    queries = []

    class Cursor:
        last_query = ""

        def execute(self, query, params=()):
            self.last_query = query
            queries.append((query, params))

        def fetchone(self):
            if "SELECT user_id FROM app_drafting_projects" in self.last_query:
                return {"user_id": 91}
            if "AS figure_count" in self.last_query:
                return {"figure_count": 0, "active_png_count": 0}
            raise AssertionError(self.last_query)

    @contextmanager
    def cursor_factory(**_kwargs):
        yield Cursor()

    repository = draft_studio.StudioRepository(cursor_factory, migrate=False)
    monkeypatch.setattr(repository, "_verify", lambda *_args: {
        "id": 33, "project_id": 7, "turn_no": 4, "requested_by_user_id": 91,
        "project_revision": 2, "status": "running",
    })

    with pytest.raises(drafting.DraftingValidationError, match="drawing set changed"):
        repository.complete_turn(
            33, "lease", result={}, session_id="session", cost_usd=1,
            duration_ms=1000, model_name="model", required_figure_count=2)

    assert not any("SET status='complete'" in query for query, _params in queries)
    assert not any("SET status='ready'" in query for query, _params in queries)


def test_invalid_workspace_is_automatic_repair_input_not_a_failed_turn(monkeypatch, tmp_path):
    repository = Mock()
    repository.save_version.return_value = {"version_no": 1}
    repository.save_qa.return_value = {"id": 5, "verdict": "pass", "checks": [],
                                       "findings": [], "counts": {}}
    repository.complete_turn.return_value = {"status": "complete"}
    agent = Mock()
    agent.DRAFT_MODEL = "draft-model"
    agent.DRAFT_TIMEOUT = 60
    agent.new_session_id.return_value = "new-session"
    agent.strings.side_effect = lambda value, **_kwargs: list(value or [])
    agent.run.side_effect = [
        draft_agent.AgentRun(ok=True, session_id="session", model="draft-model",
                             result={"action": "revised", "summary": "initial",
                                     "reasoning": [], "changes": [], "questions": [],
                                     "prior_art_strategy": "", "answer": ""}),
        draft_agent.AgentRun(ok=True, session_id="session", model="draft-model",
                             result={"action": "revised", "summary": "repaired",
                                     "reasoning": [], "changes": [], "questions": [],
                                     "prior_art_strategy": "", "answer": ""}),
    ]
    workspace = Mock()
    workspace.snapshot.side_effect = [
        {"sections": {**GOOD, "cross_reference": ""}, "numerals": NUMERALS,
         "figures": FIGURES},
        {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES},
    ]
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, qa=draft_qa, workspace=workspace)
    monkeypatch.setattr(runner, "prepare", lambda _turn: {
        "workspace": tmp_path, "project": {"user_id": 91, "agent_session_id": "",
                                             "latest_version_no": 0,
                                             "disclosure_text": "disclosure"},
        "references": [{"publication_number": ALLOWED[0]}], "documents": [],
        "seeded": False, "had_version": False,
        "previous_sections": {},
    })
    monkeypatch.setattr(runner, "_ensure_figures", lambda **_kwargs: {"ok": True})
    text_report = {
        "status": "complete", "verdict": "pass", "summary": "ready",
        "checks": [], "findings": [], "counts": {}, "cost_usd": 0,
        "duration_ms": 1, "model_name": "review"}
    monkeypatch.setattr(runner, "mechanical_report", Mock(return_value=text_report))
    final_review = Mock()
    monkeypatch.setattr(runner, "evaluate", final_review)

    runner.run({"id": 3, "lease_token": "lease", "project_id": 7,
                "turn_no": 1, "kind": "initial"})

    assert agent.run.call_count == 2
    assert repository.save_version.call_count == 1
    assert workspace._write_review.call_count >= 1
    assert "missing Cross-Reference" in workspace._write_review.call_args_list[0].args[1][
        "summary"]
    final_review.assert_not_called()


def test_initial_turn_cannot_finish_as_an_answer_without_a_filing_candidate(monkeypatch, tmp_path):
    repository = Mock()
    repository.save_version.return_value = {"version_no": 1}
    repository.save_qa.return_value = {"id": 5, "verdict": "pass", "checks": [],
                                       "findings": [], "counts": {}}
    repository.complete_turn.return_value = {"status": "complete"}
    agent = Mock()
    agent.DRAFT_MODEL = "draft-model"
    agent.DRAFT_TIMEOUT = 60
    agent.new_session_id.return_value = "new-session"
    agent.strings.side_effect = lambda value, **_kwargs: list(value or [])
    agent.run.side_effect = [
        draft_agent.AgentRun(ok=True, session_id="session", model="draft-model",
                             result={"action": "answered", "summary": "I need more detail",
                                     "reasoning": [], "changes": [], "questions": [],
                                     "prior_art_strategy": "", "answer": "Please clarify."}),
        draft_agent.AgentRun(ok=True, session_id="session", model="draft-model",
                             result={"action": "revised", "summary": "complete application",
                                     "reasoning": [], "changes": [], "questions": [],
                                     "prior_art_strategy": "", "answer": ""}),
    ]
    workspace = Mock()
    workspace.snapshot.return_value = {
        "sections": GOOD, "numerals": NUMERALS, "figures": FIGURES}
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, qa=draft_qa, workspace=workspace)
    monkeypatch.setattr(runner, "prepare", lambda _turn: {
        "workspace": tmp_path, "project": {"user_id": 91, "agent_session_id": "",
                                             "latest_version_no": 0,
                                             "disclosure_text": "disclosure"},
        "references": [{"publication_number": ALLOWED[0]}], "documents": [],
        "seeded": False, "had_version": False,
        "previous_sections": {},
    })
    monkeypatch.setattr(runner, "_ensure_figures", lambda **_kwargs: {"ok": True})
    text_report = {
        "status": "complete", "verdict": "pass", "summary": "ready",
        "checks": [], "findings": [], "counts": {}, "cost_usd": 0,
        "duration_ms": 1, "model_name": "review"}
    monkeypatch.setattr(runner, "mechanical_report", Mock(return_value=text_report))
    final_review = Mock()
    monkeypatch.setattr(runner, "evaluate", final_review)

    runner.run({"id": 3, "lease_token": "lease", "project_id": 7,
                "turn_no": 1, "kind": "initial"})

    assert agent.run.call_count == 2
    assert repository.save_version.call_count == 1
    first_report = workspace._write_review.call_args_list[0].args[1]
    assert "filing candidate" in first_report["summary"].lower()
    final_review.assert_not_called()


def test_resumed_figure_plan_repair_is_source_locked_before_validation(monkeypatch, tmp_path):
    monkeypatch.setattr(draft_figures, "discard_project_figure_checkpoint",
                        lambda _turn_id: False)
    repository = Mock()
    repository.save_version.return_value = {"version_no": 2}
    repository.save_qa.return_value = {
        "id": 5, "verdict": "pass", "checks": [], "findings": [], "counts": {}}
    repository.complete_turn.return_value = {"status": "complete"}
    agent = Mock()
    agent.DRAFT_MODEL = "draft-model"
    agent.DRAFT_TIMEOUT = 60
    agent.new_session_id.return_value = "new-session"
    agent.strings.side_effect = lambda value, **_kwargs: list(value or [])
    revised_figures = [
        {**FIGURES[0], "numerals": ["10", "12"]},
        {**FIGURES[1], "numerals": ["14", "16", "18", "20"]},
    ]

    def mutate_sources(**_kwargs):
        draft_workspace.write_sections(tmp_path, {
            **GOOD,
            "drawing_descriptions": "FIG. 1 is a body view.\n\nFIG. 2 is a ring view.",
            "detailed_description": "Deleted part 20 to simplify the drawings.",
            "claims": "1. A different invention.",
        })
        draft_workspace.write_numerals(tmp_path, NUMERALS[:-1])
        draft_workspace.write_figures(tmp_path, revised_figures)
        return draft_agent.AgentRun(
            ok=True, session_id="session", model="draft-model",
            result={"action": "revised", "summary": "rebalanced figures",
                    "reasoning": [], "changes": [], "questions": [],
                    "prior_art_strategy": "", "answer": ""})

    agent.run.side_effect = mutate_sources
    draft_workspace.write_sections(tmp_path, GOOD)
    draft_workspace.write_numerals(tmp_path, NUMERALS)
    draft_workspace.write_figures(tmp_path, FIGURES)
    workspace = Mock()
    workspace.snapshot.side_effect = lambda _path: draft_workspace.snapshot(tmp_path)
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, qa=draft_qa, workspace=workspace)
    monkeypatch.setattr(runner, "prepare", lambda _turn: {
        "workspace": tmp_path,
        "project": {"user_id": 91, "agent_session_id": "", "latest_version_no": 1,
                    "disclosure_text": "disclosure"},
        "references": [{"publication_number": ALLOWED[0]}], "documents": [],
        "seeded": False, "had_version": True, "resuming_candidate": True,
        "prepared_snapshot": {"sections": GOOD, "numerals": NUMERALS,
                              "figures": FIGURES},
        "prepared_qa": {"checks": [{
            "name": "Saved candidate passes the current filing preflight",
            "status": "fail", "category": "figures_and_numerals",
        }], "findings": []},
        "previous_sections": {},
    })
    monkeypatch.setattr(runner, "_ensure_figures", lambda **_kwargs: {"ok": True})
    text_report = {
        "status": "complete", "verdict": "pass", "summary": "ready",
        "checks": [], "findings": [], "counts": {}, "cost_usd": 0,
        "duration_ms": 1, "model_name": "review"}
    monkeypatch.setattr(runner, "mechanical_report", Mock(return_value=text_report))
    final_review = Mock()
    monkeypatch.setattr(runner, "evaluate", final_review)

    runner.run({"id": 3, "lease_token": "lease", "project_id": 7,
                "turn_no": 2, "kind": "revise"})

    saved = repository.save_version.call_args.kwargs
    assert saved["sections"]["detailed_description"] == GOOD["detailed_description"]
    assert saved["sections"]["claims"] == GOOD["claims"]
    assert saved["numerals"] == NUMERALS
    assert [{key: item[key] for key in ("label", "caption", "numerals")}
            for item in saved["figures"]] == revised_figures
    final_review.assert_not_called()


def test_figure_plan_defect_publishes_the_text_and_reports_the_defect(
        monkeypatch, tmp_path):
    """An overcrowded drawing brief is a finding, not a second turn.

    It used to be both: the text published, the defect was reported, and a `gate_resume` turn was
    queued to redraw the sheet and re-inspect it. Nothing redraws now, so the finding stands on
    its own in Review for the agent or the applicant to act on.
    """
    monkeypatch.setattr(draft_figures, "discard_project_figure_checkpoint",
                        lambda _turn_id: False)
    repository = Mock()
    repository.save_version.return_value = {"version_no": 2}
    repository.save_qa.return_value = {
        "id": 5, "verdict": "pass", "checks": [], "findings": [], "counts": {}}
    repository.complete_turn.return_value = {"status": "complete"}
    agent = Mock()
    agent.DRAFT_MODEL = "draft-model"
    agent.DRAFT_TIMEOUT = 60
    agent.new_session_id.return_value = "new-session"
    agent.strings.side_effect = lambda value, **_kwargs: list(value or [])
    overlong_figures = [
        {**FIGURES[0], "caption": "plain rectangular body " * 160},
        FIGURES[1],
    ]
    repaired_figures = [
        {**FIGURES[0], "caption": "plain rectangular body"},
        FIGURES[1],
    ]

    def run_agent(**_kwargs):
        if agent.run.call_count == 1:
            return draft_agent.AgentRun(
                ok=True, session_id="session", model="draft-model",
                result={"action": "revised", "summary": "candidate",
                        "reasoning": [], "changes": [], "questions": [],
                        "prior_art_strategy": "", "answer": ""})
        draft_workspace.write_sections(tmp_path, {
            key: ("FIG. 1 is an assembly view. FIG. 2 is a ring view."
                  if key == "drawing_descriptions" else "")
            for key, _name, _heading in draft_workspace.SECTION_FILES
        })
        draft_workspace.write_numerals(tmp_path, [])
        draft_workspace.write_figures(tmp_path, repaired_figures)
        return draft_agent.AgentRun(
            ok=True, session_id="session", model="draft-model",
            result={"action": "revised", "summary": "shortened drawing brief",
                    "reasoning": [], "changes": [], "questions": [],
                    "prior_art_strategy": "", "answer": ""})

    agent.run.side_effect = run_agent
    draft_workspace.write_sections(tmp_path, GOOD)
    draft_workspace.write_numerals(tmp_path, NUMERALS)
    draft_workspace.write_figures(tmp_path, overlong_figures)
    workspace = Mock()
    workspace.snapshot.side_effect = lambda _path: draft_workspace.snapshot(tmp_path)
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, qa=draft_qa, workspace=workspace)
    monkeypatch.setattr(runner, "prepare", lambda _turn: {
        "workspace": tmp_path,
        "project": {"user_id": 91, "agent_session_id": "", "latest_version_no": 1,
                    "disclosure_text": "disclosure"},
        "references": [{"publication_number": ALLOWED[0]}], "documents": [],
        "seeded": False, "had_version": True, "resuming_candidate": False,
        "previous_sections": {},
    })
    monkeypatch.setattr(runner, "_ensure_figures", lambda **_kwargs: {"ok": True})
    text_report = {
        "status": "complete", "verdict": "pass", "summary": "ready",
        "checks": [], "findings": [], "counts": {}, "cost_usd": 0,
        "duration_ms": 1, "model_name": "review"}
    monkeypatch.setattr(runner, "mechanical_report", Mock(return_value=text_report))
    final_review = Mock()
    monkeypatch.setattr(runner, "evaluate", final_review)

    runner.run({"id": 3, "lease_token": "lease", "project_id": 7,
                "turn_no": 2, "kind": "revise"})

    assert agent.run.call_count == 1
    saved = repository.save_version.call_args.kwargs
    assert saved["sections"]["claims"] == GOOD["claims"]
    assert saved["sections"]["detailed_description"] == GOOD["detailed_description"]
    assert saved["numerals"] == NUMERALS
    assert [item["label"] for item in saved["figures"]] == [
        item["label"] for item in overlong_figures]
    assert saved["figures"][0]["caption"] == overlong_figures[0]["caption"].strip()
    report = repository.save_qa.call_args.kwargs["report"]
    assert any(check.get("category") == "figures_and_numerals"
               for check in report["checks"] if check.get("status") != "pass")
    completion = repository.complete_turn.call_args.kwargs
    assert completion["discard_candidates"] is True
    assert completion["continuation"] is None
    final_review.assert_not_called()


def test_answering_a_question_does_not_discard_an_unpublished_candidate(monkeypatch, tmp_path):
    repository = Mock()
    repository.complete_turn.return_value = {"status": "complete"}
    agent = Mock()
    agent.DRAFT_MODEL = "draft-model"
    agent.DRAFT_TIMEOUT = 60
    agent.new_session_id.return_value = "new-session"
    agent.run.return_value = draft_agent.AgentRun(
        ok=True, session_id="session", model="draft-model",
        result={"action": "answered", "summary": "answered",
                "reasoning": [], "changes": [], "questions": [],
                "prior_art_strategy": "", "answer": "The answer."})
    runner = draft_studio.TurnRunner(
        repository, object(), agent=agent, qa=draft_qa, workspace=Mock())
    monkeypatch.setattr(runner, "prepare", lambda _turn: {
        "workspace": tmp_path, "project": {"user_id": 91, "agent_session_id": "prior",
                                             "latest_version_no": 1,
                                             "disclosure_text": "disclosure"},
        "references": [], "documents": [], "seeded": False, "had_version": True,
        "resuming_candidate": True, "previous_sections": GOOD,
    })

    runner.run({"id": 3, "lease_token": "lease", "project_id": 7,
                "turn_no": 2, "kind": "question"})

    assert repository.complete_turn.call_args.kwargs["discard_candidates"] is False


def test_a_drawing_prompt_gets_only_the_numerals_for_that_figure():
    import draft_studio_service
    version = {"numerals": NUMERALS, "figure_specs": FIGURES}
    assert draft_studio_service._expected_numerals(version, "Figure 1") == [
        "10 = vacuum lifting tool", "12 = body", "14 = pump"]
    assert "20 = passage" in draft_studio_service._expected_numerals(version, "FIG. 2")


def test_uploading_a_sheet_lands_on_the_figure_the_draft_is_missing(monkeypatch):
    """An upload with no label goes where the specification says it belongs.

    The alternative - numbering it after however many sheets already exist - makes the first
    upload against a two-figure draft "FIG. 3", an orphan the reviewer then reports and the
    applicant has to explain to the agent. The draft already says which sheets it wants.
    """
    import draft_figures
    import draft_studio_service

    class DraftingService:
        def get_project(self, _principal, project_id, include_versions=True):
            assert project_id == 7 and include_versions is True
            return {"id": 7, "user_id": 91, "latest_version_no": 1,
                    "versions": [{"version_no": 1, "sections": GOOD, "numerals": NUMERALS,
                                  "figure_specs": FIGURES}]}

    made = {}
    monkeypatch.setattr(draft_figures, "normalize_source_image",
                        lambda data, content_type="": b"PNG:" + data)
    #  FIG. 1 is already supplied; FIG. 2 is the sheet the draft is still waiting for.
    monkeypatch.setattr(draft_figures, "listing", lambda *_a: [
        {"id": 4, "figure_label": FIGURES[0]["label"], "sort_order": 1}])
    monkeypatch.setattr(draft_figures, "create_figure",
                        lambda project_id, user_id, label, caption="", sort_order=0:
                        made.setdefault("figure", {"id": 5, "label": label,
                                                   "caption": caption}))
    monkeypatch.setattr(draft_figures, "add_version",
                        lambda figure_id, **kwargs: made.update(
                            {"version": {"figure_id": figure_id, **kwargs}}))

    class Repository:
        def add_message(self, *_args, **_kwargs):
            made["announced"] = True

    service = draft_studio_service.StudioService(DraftingService(), repository=Repository())
    out = service.upload_figure(object(), 7, image=b"raw", content_type="image/png")

    assert out["label"] == FIGURES[1]["label"]
    assert made["version"]["figure_id"] == 5
    assert made["version"]["png"] == b"PNG:raw"
    assert made["version"]["source_kind"] == "uploaded"
    assert made["announced"] is True


def test_rejected_studio_message_does_not_pollute_the_conversation(monkeypatch):
    import draft_studio_service

    drafting_service = Mock()
    drafting_service.repository.get_project.return_value = {
        "id": 7, "status": "generating", "revision": 4,
    }
    repository = Mock()
    repository.enqueue_turn_safely.side_effect = drafting.DraftingConflict(
        "The drafting agent is still working on the previous message.")
    service = draft_studio_service.StudioService(
        drafting_service, repository=repository)
    monkeypatch.setattr(
        draft_studio_service.draft_agent, "availability", lambda: {"ok": True})

    with pytest.raises(drafting.DraftingConflict):
        service.start_turn(
            Mock(user_id=91), 7, message="Apply the source-fidelity repair.")

    repository.add_message.assert_not_called()


# =============================================================================================
# False positives, each measured on a real 20-claim draft before it was fixed
# =============================================================================================
def test_a_claims_own_number_is_not_a_reference_numeral():
    """Before this, every claim in a 20-claim set was reported as an undefined part."""
    check = checks_for()["Numerals in the claims are defined"] if \
        "Numerals in the claims are defined" in checks_for() else None
    assert check is None, "a clean claim set must raise nothing about numerals in the claims"


def test_a_numeral_genuinely_recited_in_a_claim_is_still_caught():
    broken = dict(GOOD)
    broken["claims"] += ("\n\n4. The vacuum lifting tool of claim 1, wherein the retaining "
                         "shoulder 44 is annular.")
    check = checks_for(broken)["Numerals in the claims are defined"]
    assert check["status"] == "warn" and "44" in check["items"]


def test_a_figure_caption_that_cross_references_another_figure_still_matches_its_sheet():
    """"FIG. 3 — Enlarged detail III of FIG. 2" is figure 3, not figure 32."""
    assert draft_qa.figure_number("FIG. 3 — Enlarged detail III of FIG. 2") == "3"
    assert draft_qa.figure_number("FIG. 1 — Perspective view") == "1"
    figures = [{"label": "FIG. 1 — Perspective view", "caption": "", "numerals": []},
               {"label": "FIG. 2 — Section on 2—2 of FIG. 1", "caption": "", "numerals": []}]
    checks = checks_for(figures=figures)
    assert "Each described figure has a drawing sheet" not in checks


def test_figure_specifications_must_have_unique_contiguous_numbers():
    figures = [
        {"label": "FIG. 1: perspective", "caption": "view", "numerals": ["10"]},
        {"label": "FIG. 1: duplicate", "caption": "other", "numerals": ["12"]},
        {"label": "FIG. 3", "caption": "gap", "numerals": ["14"]},
    ]
    check = checks_for(figures=figures)["Figure-sheet numbering is unique and contiguous"]
    assert check["status"] == "fail"
    assert any("duplicate" in item for item in check["items"])
    assert any("expected" in item for item in check["items"])


def test_figure_specifications_must_be_in_numeric_filing_order():
    labels = ["FIG. 1", "FIG. 10"] + [f"FIG. {number}" for number in range(2, 10)]
    figures = [
        {"label": label, "caption": "view", "numerals": []}
        for label in labels
    ]

    check = checks_for(figures=figures)["Figure-sheet numbering is unique and contiguous"]

    assert check["status"] == "fail"
    assert any("filing order" in item for item in check["items"])


def test_a_participle_after_the_noun_does_not_break_antecedent_basis():
    """"an evacuable chamber" … "the evacuable chamber displaces" is the same chamber."""
    broken = dict(GOOD)
    broken["claims"] = ("1. A tool comprising a body, a pump, and an evacuable chamber bounded by "
                        "the body, wherein the evacuable chamber displaces a sealing ring, and "
                        "wherein the body carries the pump.")
    check = checks_for(broken)["Antecedent basis in the claims"]
    assert check["status"] == "pass", check["items"]


def test_antecedent_basis_still_catches_a_term_that_was_never_introduced():
    broken = dict(GOOD)
    broken["claims"] = "1. A tool comprising a body, wherein the retaining shoulder is annular."
    check = checks_for(broken)["Antecedent basis in the claims"]
    assert check["status"] == "warn" and any("retaining shoulder" in i for i in check["items"])


def test_antecedent_basis_still_catches_an_unintroduced_plural_term():
    broken = dict(GOOD)
    broken["claims"] = "1. A tool comprising a body, wherein the retaining shoulders are annular."
    check = checks_for(broken)["Antecedent basis in the claims"]
    assert check["status"] == "warn"
    assert any("retaining shoulders" in item for item in check["items"])


def test_antecedent_basis_recognises_a_method_step_gerund_in_the_parent_claim():
    broken = dict(GOOD)
    broken["claims"] = (
        "1. A method comprising: translating a device across a surface.\n\n"
        "2. The method of claim 1, wherein the translating is performed by vibration."
    )
    check = checks_for(broken)["Antecedent basis in the claims"]
    assert check["status"] == "pass", check["items"]


def test_antecedent_basis_recognises_quantified_plural_parts():
    broken = dict(GOOD)
    broken["claims"] = (
        "1. A carrier comprising two opposed phase-change cassettes and two downward guide "
        "ducts, wherein the guide ducts route air along the phase-change cassettes."
    )
    check = checks_for(broken)["Antecedent basis in the claims"]
    assert check["status"] == "pass", check["items"]


def test_antecedent_basis_recognises_bare_plural_parts_introduced_before_reference():
    broken = dict(GOOD)
    broken["claims"] = (
        "1. A carrier comprising a frame that carries resilient feet and has ledges beneath "
        "the resilient feet, wherein the ledges support the resilient feet."
    )
    check = checks_for(broken)["Antecedent basis in the claims"]
    assert check["status"] == "pass", check["items"]


def test_antecedent_basis_recognises_method_goods_and_thereafter_step():
    broken = dict(GOOD)
    broken["claims"] = (
        "1. A method of transporting goods, the method comprising: placing the goods in a "
        "carrier; and thereafter transporting the goods.\n\n"
        "2. The method of claim 1, wherein the transporting is performed without a fan."
    )
    check = checks_for(broken)["Antecedent basis in the claims"]
    assert check["status"] == "pass", check["items"]


def test_antecedent_basis_recognises_bare_mass_nouns_and_action_nouns():
    broken = dict(GOOD)
    broken["claims"] = (
        "1. An equipment assembly comprising a boss; equipment supported on the boss, the "
        "equipment being a sensor.\n\n"
        "2. A method comprising observing displacement of an indicator, the displacement "
        "indicating a compression range."
    )
    check = checks_for(broken)["Antecedent basis in the claims"]
    assert check["status"] == "pass", check["items"]


def test_antecedent_basis_does_not_treat_anaphors_or_superlatives_as_components():
    broken = dict(GOOD)
    broken["claims"] = (
        "1. A clamp comprising three jaws, one jaw advancing no farther than the others, and "
        "three sensor values, a controller selecting the greatest of the sensor values."
    )
    check = checks_for(broken)["Antecedent basis in the claims"]
    assert check["status"] == "pass", check["items"]


def test_morphological_variants_are_not_reported_as_unsupported_claim_terms():
    broken = dict(GOOD)
    broken["detailed_description"] += (" A controller energises the pump 14 and connects it to a "
                                       "battery, driving the diaphragm, and each ring is fitted "
                                       "and received in the groove 18 through a cycle.")
    broken["claims"] += ("\n\n4. The tool of claim 1, wherein a controller energising the pump "
                         "is connecting a battery, driving the diaphragm, the ring being fittable "
                         "and receivable, cycles being repeated.")
    check = checks_for(broken)["Claim terms appear in the description"]
    reported = " ".join(check.get("items") or [])
    for variant in ("energising", "connecting", "driving", "fittable", "receivable", "cycles"):
        assert variant not in reported, f"{variant} is the same word as one in the description"


def test_claim_support_recognises_live_derivational_variants():
    broken = dict(GOOD)
    broken["detailed_description"] += (
        " The lid carries the rib. The controller verifies the current and later reclosed the "
        "contactor. A command is withheld until the delay expires."
    )
    broken["claims"] += (
        "\n\n4. The tool of claim 1, further comprising a lid carrying a rib and a controller "
        "configured for verifying current before reclosure while withholding a command."
    )
    check = checks_for(broken)["Claim terms appear in the description"]
    reported = " ".join(check.get("items") or [])
    for variant in ("carrying", "verifying", "reclosure", "withholding"):
        assert variant not in reported, f"{variant} has a supported derivational form"


def test_claim_support_ignores_noncomponent_drafting_vocabulary():
    broken = dict(GOOD)
    broken["claims"] += (
        "\n\n4. A method comprising continuing to move the body, thereby observing a pin "
        "indicating motion, thereafter placing a wick lying under a frame sized for use, and "
        "declining another request for a respective device."
    )
    check = checks_for(broken)["Claim terms appear in the description"]
    reported = " ".join(check.get("items") or [])
    for word in ("continuing", "thereby", "observing", "indicating", "thereafter", "under",
                 "lying", "sized", "declining", "respective"):
        assert word not in reported, f"{word} is claim grammar, not a component term"


def test_a_genuinely_different_word_is_still_reported():
    broken = dict(GOOD)
    broken["claims"] += "\n\n4. The tool of claim 1, wherein the housing is titanium."
    check = checks_for(broken)["Claim terms appear in the description"]
    reported = " ".join(check.get("items") or [])
    assert "housing" in reported and "titanium" in reported


def test_the_stemmer_is_consistent_across_the_forms_that_bit():
    pairs = [("energising", "energises"), ("driving", "drive"), ("connecting", "connect"),
             ("fittable", "fitted"), ("receivable", "received"), ("cycles", "cycle"),
             ("carrying", "carries"), ("verifying", "verified"),
             ("reclosure", "reclosed"), ("withholding", "withheld")]
    for left, right in pairs:
        assert draft_qa._stem(left) == draft_qa._stem(right), (left, right)
    assert draft_qa._stem("housing") != draft_qa._stem("body")


def test_the_first_version_names_the_project():
    """Until a draft exists the title is a placeholder — usually the first line of a paste."""
    assert draft_studio.project_title_from(1, GOOD) == GOOD["title"]
    assert draft_studio.project_title_from(1, {"title": "A Title\nand a stray line"}) == "A Title"
    assert draft_studio.project_title_from(1, {"title": "   "}) == ""


def test_a_later_version_never_renames_the_project():
    """A user who renamed the project must not have it overwritten by the next iteration."""
    assert draft_studio.project_title_from(2, GOOD) == ""
    assert draft_studio.project_title_from(7, GOOD) == ""


def test_recovery_leaves_a_turn_whose_worker_is_still_alive():
    """Expiring a live lease would put two agents on one workspace, not just two rows."""
    import os
    import draft_studio_service as service
    assert service._worker_is_alive(f"draft-turn-{os.getpid()}-140234") is True
    assert service._worker_is_alive("draft-turn-999999-1") is False
    assert service._worker_is_alive("") is False
    assert service._worker_is_alive(None) is False


def test_restart_recovery_does_not_spend_a_drafting_attempt(monkeypatch):
    import db
    import draft_studio_service as service

    queries = []

    class Cursor:
        def execute(self, query, params=()):
            queries.append((query, params))

        def fetchall(self):
            return [{"id": 33, "claimed_by": "draft-turn-999999-1"}]

    @contextmanager
    def cursor_factory(**_kwargs):
        yield Cursor()

    monkeypatch.setattr(db, "cursor", cursor_factory)

    assert service.recover_interrupted_turns() == 1
    update = next(query for query, _params in queries
                  if "stage='resuming after a restart'" in query)
    assert "attempts=greatest(0,attempts-1)" in update


@pytest.mark.parametrize("failure", [
    drafting.DraftingValidationError("candidate failed"),
    drafting.DraftingConflict("turn superseded"),
])
def test_a_failed_or_superseded_turn_restores_the_published_drawing_set(monkeypatch, failure):
    import draft_studio_service as service

    class Repository:
        def claim_turn(self, worker_id):
            return {"id": 31, "lease_token": "lease", "project_id": 7}

        def fail_turn(self, turn_id, lease, error, retryable=True):
            return {"status": "failed"}

        def add_message(self, *args, **kwargs):
            pass

    class Runner:
        repository = Repository()

        def __init__(self):
            self.restored = []

        def run(self, claimed):
            raise failure

        def restore_figures(self, turn_id):
            self.restored.append(turn_id)
            return True

    runner = Runner()
    monkeypatch.setattr(service, "_RUNNER_FACTORY", lambda: runner)

    service.process_one()

    assert runner.restored == [31]


def test_terminal_filing_gate_failure_continues_from_saved_candidate_without_user_input():
    import draft_studio_service as service

    repository = Mock()
    repository.fail_turn.return_value = {
        "id": 31,
        "project_id": 7,
        "requested_by_user_id": 91,
        "project_revision": 4,
        "idempotency_key": None,
        "status": "failed",
    }
    repository.enqueue_turn_safely.return_value = {"id": 32, "status": "queued"}
    runner = Mock(repository=repository)
    claimed = {
        "id": 31,
        "project_id": 7,
        "requested_by_user_id": 91,
        "project_revision": 4,
        "idempotency_key": None,
        "lease_token": "lease",
    }

    result = service._fail(
        runner, claimed,
        "The automatic filing gate could not clear: FIG. 2 endpoint inspection failed",
        retryable=True)

    assert result["status"] == "failed"
    queued = repository.enqueue_turn_safely.call_args
    assert queued.args == (7, 91)
    assert queued.kwargs["kind"] == "qa_fix"
    assert queued.kwargs["project_revision"] == 4
    assert queued.kwargs["idempotency_key"] == "auto-filing-repair-31-1"
    assert "not new invention disclosure" in queued.kwargs["user_message"]
    message = repository.add_message.call_args.args[2]
    assert "No action is required" in message
    assert "Try again" not in message


@pytest.mark.parametrize("error", [
    "StudioError: API Error: Connection lost mid-response. The response may be incomplete.",
    "StudioError: No conversation found with session ID: stopped-review-session",
    ("StudioError: The Vertex drafting fallback finished without returning the required "
     "structured answer."),
    "SourceReviewUnavailable: The drafting agent produced no result (exit code 143).",
    "SourceReviewUnavailable: Failed to provide valid structured output after 5 attempts",
    ("FigureTransientError: Cross-provider geometry inspection failed: Anthropic geometry "
     "audit did not return complete JSON."),
    ("FigureTransientError: the image model could not draw this figure: the image model "
     "returned no response parts (IMAGE_RECITATION)"),
    "DrawingBudgetSpent: the bounded drawing caller stopped before every sheet passed",
    ("This turn reached its ceiling of $12.00 (14 agent runs, 4,451,679 tokens). "
     "The draft it had reached is saved; nothing was published."),
])
def test_terminal_provider_disconnect_continues_a_saved_candidate_without_user_input(error):
    import draft_studio_service as service

    repository = Mock()
    repository.fail_turn.return_value = {
        "id": 31,
        "project_id": 7,
        "requested_by_user_id": 91,
        "project_revision": 4,
        "idempotency_key": None,
        "status": "failed",
    }
    repository.retry_candidate.return_value = {
        "turn_id": 31,
        "snapshot": {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES},
        "qa_report": {"verdict": "fail"},
    }
    repository.enqueue_turn_safely.return_value = {"id": 32, "status": "queued"}
    runner = Mock(repository=repository)
    claimed = {
        "id": 31,
        "project_id": 7,
        "requested_by_user_id": 91,
        "project_revision": 4,
        "idempotency_key": None,
        "lease_token": "lease",
    }

    result = service._fail(
        runner, claimed, error, retryable=True)

    assert result["status"] == "failed"
    queued = repository.enqueue_turn_safely.call_args
    assert queued.args == (7, 91)
    assert queued.kwargs["kind"] == "gate_resume"
    assert queued.kwargs["idempotency_key"] == "auto-filing-repair-31-1"
    assert "No action is required" in repository.add_message.call_args.args[2]


def test_automatic_filing_repair_chain_continues_past_three_bounded_drawing_turns():
    import draft_studio_service as service

    repository = Mock()
    repository.fail_turn.return_value = {
        "id": 35,
        "project_id": 7,
        "requested_by_user_id": 91,
        "project_revision": 4,
        "idempotency_key": "auto-filing-repair-31-3",
        "status": "failed",
    }
    runner = Mock(repository=repository)
    claimed = {
        "id": 35,
        "project_id": 7,
        "requested_by_user_id": 91,
        "project_revision": 4,
        "idempotency_key": "auto-filing-repair-31-3",
        "lease_token": "lease",
    }

    service._fail(
        runner, claimed,
        "The automatic filing gate could not clear: source fidelity review failed",
        retryable=True)

    queued = repository.enqueue_turn_safely.call_args
    assert queued.kwargs["idempotency_key"] == "auto-filing-repair-31-4"
    assert "No action is required" in repository.add_message.call_args.args[2]


def test_automatic_filing_repair_chain_stops_at_its_durable_safety_limit():
    import draft_studio_service as service

    repository = Mock()
    repository.fail_turn.return_value = {
        "id": 38,
        "project_id": 7,
        "requested_by_user_id": 91,
        "project_revision": 4,
        "idempotency_key": "auto-filing-repair-31-6",
        "status": "failed",
    }
    runner = Mock(repository=repository)
    claimed = {
        "id": 38,
        "project_id": 7,
        "requested_by_user_id": 91,
        "project_revision": 4,
        "idempotency_key": "auto-filing-repair-31-6",
        "lease_token": "lease",
    }

    service._fail(
        runner, claimed,
        "The automatic filing gate could not clear: source fidelity review failed",
        retryable=True)

    repository.enqueue_turn_safely.assert_not_called()
    assert "safety limit" in repository.add_message.call_args.args[2]


@pytest.mark.parametrize("error", [
    "DrawingBudgetSpent: the drawing pass reached its time budget.",
    ("This turn reached its ceiling of 14 agent runs ($7.06 spent, 559,213 tokens). "
     "The draft it had reached is saved; nothing was published."),
])
def test_technical_candidate_continuation_starts_a_new_chain_at_the_repair_limit(error):
    import draft_studio_service as service

    repository = Mock()
    repository.fail_turn.return_value = {
        "id": 38,
        "project_id": 7,
        "requested_by_user_id": 91,
        "project_revision": 4,
        "idempotency_key": "auto-filing-repair-31-6",
        "status": "failed",
    }
    repository.retry_candidate.return_value = {
        "turn_id": 38,
        "snapshot": {"sections": GOOD, "numerals": NUMERALS, "figures": FIGURES},
        "qa_report": {"verdict": "fail"},
    }
    repository.enqueue_turn_safely.return_value = {"id": 39, "status": "queued"}
    runner = Mock(repository=repository)
    claimed = {
        "id": 38,
        "project_id": 7,
        "requested_by_user_id": 91,
        "project_revision": 4,
        "idempotency_key": "auto-filing-repair-31-6",
        "lease_token": "lease",
    }

    service._fail(runner, claimed, error, retryable=False)

    queued = repository.enqueue_turn_safely.call_args
    assert queued.kwargs["kind"] == "gate_resume"
    assert queued.kwargs["idempotency_key"] == "auto-filing-repair-38-1"
    assert "No action is required" in repository.add_message.call_args.args[2]


def test_a_publication_number_in_prose_is_not_three_reference_numerals():
    """Measured live: the turn that finally cited a reference properly FAILED the numeral check,
    because "US 9,108,319 B2" beside its token read as numerals 9, 108 and 319."""
    text = ("US 9,108,319 B2 [REF:US-9108319-B2] describes a suction cup assembly having a "
            "housing 12. The tool 10 lifts 60 kg with a 200 mm ring.")
    assert sorted(draft_qa.numerals_used(text), key=int) == ["10", "12"]


def test_a_cited_reference_does_not_fail_the_numeral_check():
    broken = dict(GOOD)
    broken["background"] += (" US 9,108,319 B2 [REF:US-11223344-B2] describes a suction cup "
                             "assembly.")
    check = checks_for(broken)["Every numeral in the text is defined"]
    assert check["status"] == "pass", check["items"]


def test_a_figure_written_as_prose_still_yields_its_numerals(tmp_path):
    """The agent writes a drawing brief, not a bullet list; the numerals must be found anyway."""
    directory = tmp_path / "figures"
    directory.mkdir()
    (directory / "FIG-1-PERSPECTIVE-VIEW.md").write_text(
        "# FIG. 1 — Perspective view\n\n**View type:** isometric.\n\n"
        "The body 12 carries the handle 14; the sealing ring 16 is visible at the underside.\n",
        encoding="utf-8")
    figure = draft_workspace.read_figures(tmp_path)[0]
    assert figure["label"].startswith("FIG. 1")
    assert set(figure["numerals"]) == {"12", "14", "16"}


def test_an_explicit_numeral_list_still_wins_over_the_prose_scan(tmp_path):
    directory = tmp_path / "figures"
    directory.mkdir()
    (directory / "FIG-2.md").write_text(
        "# FIG. 2\n\nSection through the body 12.\n\n## Numerals shown on this figure\n\n"
        "- 12 body\n- 14 handle\n", encoding="utf-8")
    figure = draft_workspace.read_figures(tmp_path)[0]
    assert figure["numerals"] == ["12 body", "14 handle"]


def test_a_numbered_list_is_not_a_set_of_reference_numerals():
    """Measured: an ordered list inside a drawing brief reported 1, 2 and 3 as undefined parts."""
    text = ("**What is shown**\n\n1. The heel portion 38 is deformed.\n"
            "2. The lip portion 40 rolls.\n\nThe body 12 is unchanged.")
    assert sorted(draft_qa.numerals_used(text), key=int) == ["12", "38", "40"]


def test_a_numeral_at_the_end_of_a_sentence_still_counts():
    """The list-marker rule must not swallow a real numeral; one never OPENS a sentence."""
    assert sorted(draft_qa.numerals_used("The tool comprises a body 12."), key=int) == ["12"]


def test_a_misnamed_figure_file_is_renamed_rather_than_deleted(tmp_path):
    """The canonical name is in the file. Throwing the file away to enforce it is not a trade.

    Measured on the first real drafting agent to meet this convention: it wrote six figure briefs,
    ran publish, and had all six deleted for being on filenames nobody had told it about. The name
    is derived from each file's own heading, so it was recoverable the whole time and the work was
    not.
    """
    figures = tmp_path / "figures"
    figures.mkdir()
    (figures / "fig1.md").write_text(
        "# FIG. 1 - side view of the clamp\n\nThe clamp 10 in side view.\n", encoding="utf-8")
    (figures / "notes.txt").write_text("scratch", encoding="utf-8")

    with pytest.raises(drafting.DraftingValidationError) as caught:
        draft_workspace._reject_noncanonical_figure_entries(figures)

    #  The brief survived, on its canonical name.
    kept = figures / "FIG-1-SIDE-VIEW-OF-THE-CLAMP.md"
    assert kept.exists() and "clamp 10 in side view" in kept.read_text(encoding="utf-8")
    assert not (figures / "fig1.md").exists()
    #  Only the thing that could not be placed is removed, and only that is reported.
    assert not (figures / "notes.txt").exists()
    assert "notes.txt" in str(caught.value)
    assert "fig1.md" not in str(caught.value)


def test_a_renamed_figure_file_alone_does_not_refuse_the_publish(tmp_path):
    figures = tmp_path / "figures"
    figures.mkdir()
    (figures / "whatever.md").write_text("# FIG. 2 - end view\n\nThe body 12.\n", encoding="utf-8")
    draft_workspace._reject_noncanonical_figure_entries(figures)      # no exception
    assert (figures / "FIG-2-END-VIEW.md").exists()


def test_two_files_claiming_one_heading_is_a_conflict_not_a_rename(tmp_path):
    figures = tmp_path / "figures"
    figures.mkdir()
    (figures / "FIG-3-DETAIL.md").write_text("# FIG. 3 - detail\n\nKeep me.\n", encoding="utf-8")
    (figures / "copy.md").write_text("# FIG. 3 - detail\n\nDrop me.\n", encoding="utf-8")

    with pytest.raises(drafting.DraftingValidationError, match="duplicates"):
        draft_workspace._reject_noncanonical_figure_entries(figures)
    assert "Keep me." in (figures / "FIG-3-DETAIL.md").read_text(encoding="utf-8")
    assert not (figures / "copy.md").exists()


def test_cancelling_any_turn_leaves_a_published_project_ready():
    """There is no drawing continuation left to be pending on.

    A cancelled `gate_resume` used to park the project at 'active', because a text version was
    not filing-ready until its automatic drawing phase had generated and inspected every sheet.
    With that phase gone, the same branch would leave a project reading "generating" with nothing
    left to generate and nothing that could ever move it on.
    """
    statements = []

    class Cursor:
        def execute(self, sql, params=None):
            statements.append((" ".join(sql.split()), params))

        def fetchone(self):
            return {"kind": "gate_resume", "idempotency_key": "auto-filing-repair-3-1"}

        def __enter__(self):
            return self

        def __exit__(self, *_exc):
            return False

    repository = draft_studio.StudioRepository(cursor_factory=lambda **_kw: Cursor(),
                                               migrate=False)
    repository.cancel_turn(7, 3)

    project_update = next(sql for sql, _p in statements if "app_drafting_projects" in sql)
    assert "latest_version_no>0 THEN 'ready'" in project_update
    assert "SET status='active'" not in project_update
    #  ...and the candidate goes with it, rather than being kept for a phase that cannot run.
    assert any("DELETE FROM app_draft_turn_candidates" in sql for sql, _p in statements)
