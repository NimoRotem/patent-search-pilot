"""The filing package, and every check that decides whether it may be handed over.

DEFECT INJECTION IS THE POINT.  A test that builds a good package and asserts the audit is silent
proves nothing about the audit: an audit that always returns nothing passes it.  So each check
here is given a document with exactly one thing wrong and asserted to name that thing, and the
same document with the defect removed and asserted to go quiet.
"""
from __future__ import annotations

import io
import re
import zipfile

import pytest

import figure_facts
import filing_pack
import filing_profile
import filing_rules

TITLE = "Vibration device for bedding covering elements"

SECTIONS = {
    "title": TITLE,
    "cross_reference": "Not applicable.",
    "government_support": "Not applicable.",
    "field": "The present disclosure relates to devices for bedding covering elements.",
    "background": "Vibration assists bedding.\n\nA sliding base needs continuous operator force.",
    "summary": "A vibration device is disclosed.\n\nThe perimeter member is not airtight.",
    "drawing_descriptions": (
        "FIG. 1 is a perspective view of the vibration device.\n\n"
        "FIG. 2 is a sectional view of the vibration device of FIG. 1."),
    "detailed_description": (
        "Referring to FIG. 1, a vibration device 10 includes a base 12.\n\n"
        "As shown in FIG. 2, the chamber 22 is bounded by the second side 16."),
    "claims": (
        "1. A vibration device comprising: a rigid base having a first side and a second side; "
        "and a chamber at the second side.\n\n"
        "2. The vibration device of claim 1, wherein the base is metal.\n\n"
        "3. The vibration device of claim 1, wherein the chamber is annular.\n\n"
        "4. A method comprising placing a vibration device on a covering element."),
    "abstract": ("A vibration device for bedding a covering element includes a rigid base, a "
                 "vibration motor, a chamber and a rigid perimeter member terminating in a "
                 "sliding bearing face."),
}
NUMERALS = [{"numeral": "10", "part": "vibration device"},
            {"numeral": "12", "part": "base"},
            {"numeral": "16", "part": "second side of the base"},
            {"numeral": "22", "part": "chamber"}]
PROJECT = {"id": 1, "user_id": 1, "title": TITLE, "inventors": "Ada Lovelace", "applicant": ""}
VERSION = {"version_no": 3, "sections": SECTIONS}

COMPLETE_PROFILE = {
    "inventors": [{"given_name": "Ada", "middle_name": "", "family_name": "Lovelace",
                   "city": "Reno", "state": "NV", "country": "US",
                   "mailing_address": "1 Example Way", "mailing_city": "Reno",
                   "mailing_state": "NV", "mailing_postcode": "89501",
                   "mailing_country": "US"}],
    "correspondence_name": "Ada Lovelace", "correspondence_address": "1 Example Way",
    "correspondence_city": "Reno", "correspondence_state": "NV",
    "correspondence_postcode": "89501", "correspondence_country": "US",
    "correspondence_email": "ada@example.com",
}


# =============================================================================================
# Text hygiene: the reason a perfect DOCX came back from the validator
# =============================================================================================
def test_smart_punctuation_is_folded_to_ascii():
    folded = filing_rules.to_filing_text(
        "terms such as “upper,” “lower” – and the inventor’s own")
    assert filing_rules.non_latin_characters(folded) == []
    assert '"upper,"' in folded and "inventor's" in folded


def test_the_scan_names_the_character_it_objects_to():
    strays = filing_rules.non_latin_characters("a “quoted” word")
    assert [count for _c, count in strays] == [1, 1]
    assert "LEFT DOUBLE QUOTATION MARK" in filing_rules.describe_character("“")


def test_a_specification_carrying_curly_quotes_is_reported():
    sections = dict(SECTIONS, background="Terms such as “upper” are used loosely.")
    #  Straight past to_filing_text, to prove the CHECK works rather than the folding.
    blob = _docx_with(sections, fold=False)
    titles = [item["title"] for item in filing_rules.audit_specification_docx(blob)]
    assert any("non-Latin script" in title for title in titles)
    assert not any("non-Latin script" in item["title"] for item in
                   filing_rules.audit_specification_docx(_docx_with(sections)))


def _docx_with(sections, *, fold: bool = True) -> bytes:
    version = {"version_no": 1,
               "sections": {key: (filing_rules.to_filing_text(value) if fold else value)
                            for key, value in sections.items()}}
    if fold:
        return filing_pack.specification_docx(PROJECT, version)
    #  Build without the fold by putting the raw text back after the fold has run.
    from docx import Document
    blob = filing_pack.specification_docx(PROJECT, version)
    document = Document(io.BytesIO(blob))
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if "upper" in run.text:
                run.text = run.text.replace('"upper"', "“upper”")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


# =============================================================================================
# 37 CFR 1.52: the specification as a filing document
# =============================================================================================
@pytest.fixture(scope="module")
def specification() -> bytes:
    return filing_pack.specification_docx(PROJECT, VERSION)


def test_the_built_specification_passes_every_formal_check(specification):
    findings = filing_rules.audit_specification_docx(specification)
    assert [item for item in findings if item["severity"] == "blocker"] == []


def test_paragraphs_are_numbered_consecutively_from_0001(specification):
    from docx import Document
    numbers = [re.match(r"^\[(\d{4})\]", paragraph.text.strip())
               for paragraph in Document(io.BytesIO(specification)).paragraphs]
    found = [int(match.group(1)) for match in numbers if match]
    assert found == list(range(1, len(found) + 1))
    assert len(found) >= 8


def test_the_paragraph_number_is_bold_and_the_body_is_not(specification):
    from docx import Document
    for paragraph in Document(io.BytesIO(specification)).paragraphs:
        if paragraph.text.strip().startswith("[0001]"):
            assert paragraph.runs[0].bold is True
            assert not paragraph.runs[-1].bold
            return
    pytest.fail("no numbered paragraph was written")


def test_claims_and_abstract_each_commence_on_their_own_page(specification):
    from docx import Document
    document = Document(io.BytesIO(specification))
    texts = [paragraph.text.strip() for paragraph in document.paragraphs]
    breaks = filing_rules._page_break_indexes(document.paragraphs)
    for heading in ("CLAIMS", "ABSTRACT OF THE DISCLOSURE"):
        index = texts.index(heading)
        assert any(0 < index - value <= 2 for value in breaks), heading
    assert "What is claimed is:" in texts


def test_the_package_carries_no_part_the_validator_warns_about(specification):
    names = set(zipfile.ZipFile(io.BytesIO(specification)).namelist())
    assert "docProps/thumbnail.jpeg" not in names
    assert "word/stylesWithEffects.xml" not in names
    assert not [name for name in names if name.startswith("customXml/")]
    assert not [name for name in names if re.fullmatch(r"word/comments\d*\.xml", name)]


def test_stripping_those_parts_leaves_a_document_python_docx_can_still_open(specification):
    from docx import Document
    assert Document(io.BytesIO(specification)).paragraphs[0].text == "TITLE OF THE INVENTION"


def test_single_spacing_is_reported_as_a_blocker():
    from docx import Document
    from docx.shared import Pt
    document = Document(io.BytesIO(filing_pack.specification_docx(PROJECT, VERSION)))
    document.styles["Normal"].paragraph_format.line_spacing = 1.0
    document.styles["Normal"].font.size = Pt(12)
    output = io.BytesIO()
    document.save(output)
    findings = filing_rules.audit_specification_docx(output.getvalue())
    assert any("1.5 or double spaced" in item["title"] or "not 1.5" in item["title"]
               for item in findings)


def test_a_narrow_margin_is_reported():
    from docx import Document
    from docx.shared import Inches
    document = Document(io.BytesIO(filing_pack.specification_docx(PROJECT, VERSION)))
    document.sections[0].left_margin = Inches(0.5)
    output = io.BytesIO()
    document.save(output)
    assert any("left margin is below the minimum" in item["title"]
               for item in filing_rules.audit_specification_docx(output.getvalue()))


# =============================================================================================
# 37 CFR 1.63 and 1.76: the papers that did not exist before
# =============================================================================================
def test_the_declaration_carries_every_statement_the_rule_requires():
    blob = filing_pack.declaration_pdf(PROJECT, VERSION,
                                       filing_profile.resolve(COMPLETE_PROFILE, PROJECT))
    text = filing_pack._pdf_text(blob)
    assert filing_rules.audit_declaration_text(text) == []
    assert "was made or was authorized to be made" in text


def test_a_declaration_missing_1_63_a_4_is_reported():
    text = ("This declaration is directed to the attached application. I believe that I am the "
            "original inventor. I have reviewed and understand the contents. I acknowledge the "
            "duty to disclose. Any willful false statement is punishable. Signature /Ada/")
    titles = [item["title"] for item in filing_rules.audit_declaration_text(text)]
    assert any("authorized to be made" in title for title in titles)


def test_the_application_data_sheet_carries_every_section_heading():
    blob = filing_pack.application_data_sheet_pdf(
        PROJECT, VERSION, filing_profile.resolve(COMPLETE_PROFILE, PROJECT),
        drawing_sheets_count=2)
    assert filing_rules.audit_ads_text(filing_pack._pdf_text(blob)) == []


def test_an_ads_missing_a_heading_is_reported():
    titles = [item["title"] for item in filing_rules.audit_ads_text(
        "Application Data Sheet\nInventor Information\nSignature /Ada/")]
    assert any("section heading" in title for title in titles)


def test_every_font_in_every_paper_is_embedded_under_a_name_with_no_space():
    resolved = filing_profile.resolve(COMPLETE_PROFILE, PROJECT)
    for blob in (filing_pack.declaration_pdf(PROJECT, VERSION, resolved),
                 filing_pack.application_data_sheet_pdf(PROJECT, VERSION, resolved,
                                                        drawing_sheets_count=0),
                 filing_pack.fee_worksheet_pdf(
                     resolved, filing_pack.fee_profile(SECTIONS["claims"])),
                 filing_pack.information_disclosure_pdf([])):
        findings = filing_rules.audit_pdf(blob, where="paper.pdf")
        assert [item for item in findings if item["severity"] == "blocker"] == []


# =============================================================================================
# 37 CFR 1.16: the counts
# =============================================================================================
def test_the_claim_counts_drive_the_right_surcharges():
    profile = filing_pack.fee_profile(SECTIONS["claims"])
    assert profile["total"] == 4
    assert profile["independent"] == 2
    assert profile["multiple_dependent"] == 0
    assert profile["triggered"] == []


def test_a_multiple_dependent_claim_is_counted_as_the_claims_it_refers_to():
    claims = (SECTIONS["claims"] +
              "\n\n5. The vibration device of any one of claims 1, 2 or 3, wherein it is red.")
    profile = filing_pack.fee_profile(claims)
    assert profile["multiple_dependent"] == 1
    assert profile["billable"] > profile["total"]
    assert any(item["key"] == "multiple" for item in profile["triggered"])


def test_the_application_size_fee_starts_above_one_hundred_sheets():
    assert filing_pack.fee_profile(SECTIONS["claims"],
                                   specification_sheets=99,
                                   drawing_sheets_count=1)["size_blocks"] == 0
    profile = filing_pack.fee_profile(SECTIONS["claims"], specification_sheets=140,
                                      drawing_sheets_count=11)
    assert profile["size_blocks"] == 2


def test_no_dollar_figure_is_printed_on_the_fee_worksheet():
    text = filing_pack._pdf_text(filing_pack.fee_worksheet_pdf(
        filing_profile.resolve(COMPLETE_PROFILE, PROJECT),
        filing_pack.fee_profile(SECTIONS["claims"])))
    assert "$" not in text
    assert "1011" in text and "uspto.gov" in text


# =============================================================================================
# The profile: a missing field is a defect that names itself
# =============================================================================================
def test_every_missing_field_is_named_the_way_the_ads_names_it():
    fields = [gap["field"] for gap in filing_profile.gaps(
        filing_profile.resolve({}, {"inventors": "Ada Lovelace"}))]
    assert "Inventor 1: City of residence" in fields
    assert "Correspondence email" in fields
    assert filing_profile.gaps(filing_profile.resolve(COMPLETE_PROFILE, PROJECT)) == []


def test_a_seeded_inventor_keeps_the_given_and_family_name_apart():
    row = filing_profile.resolve({}, {"inventors": "Ada Byron Lovelace"})["inventors"][0]
    assert (row["given_name"], row["middle_name"], row["family_name"]) == \
        ("Ada", "Byron", "Lovelace")


def test_an_unusable_value_is_refused_rather_than_corrected():
    with pytest.raises(ValueError):
        filing_profile.clean({"correspondence_email": "not-an-address"})
    with pytest.raises(ValueError):
        filing_profile.clean({"entity_status": "cheap"})
    with pytest.raises(ValueError):
        filing_profile.clean({"nonsense": "x"})


# =============================================================================================
# Figures: what the sheet says against what the text says
# =============================================================================================
def _sheet(label, views, **extra):
    facts = {"label": label, "views": views, "unlabelled_views": [], "numerals": [],
             "text_labels": [], "numeral_key_table": False, "divider_rules": 0,
             "sheet_number_text": "",
             "smallest_reference_character_height_fraction": 0.03}
    facts.update(extra)
    if "numerals" not in extra:
        seen = {}
        for view in list(views) + list(facts["unlabelled_views"]):
            for item in view.get("numerals") or []:
                seen.setdefault(item["value"], dict(item))
        facts["numerals"] = list(seen.values())
    return facts


def _numeral(value, **extra):
    row = {"value": value, "lead_lines": 1, "points_at": "", "designates_one_part": True,
           "matches_declared_part": "yes"}
    row.update(extra)
    return row


def _view(legend, values, **extra):
    row = {"legend": legend, "printed_caption": "", "kind": "perspective", "bbox": [0, 0, 1, 1],
           "numerals": [_numeral(value) for value in values], "character_height": 0.03}
    row.update(extra)
    return row


def test_a_figure_citation_names_every_figure_in_the_list_or_the_range():
    assert figure_facts.figure_numbers("Referring to FIGS. 1 and 2, the device") == \
        ["FIG. 1", "FIG. 2"]
    assert figure_facts.figure_numbers("as in FIGS. 2-4") == ["FIG. 2", "FIG. 3", "FIG. 4"]
    assert figure_facts.figure_numbers("FIGS. 3A, 3B and 3C") == \
        ["FIG. 3A", "FIG. 3B", "FIG. 3C"]


def test_a_view_with_no_number_is_a_blocker():
    sheets = [_sheet("FIG. 2", [_view("FIG. 2", ["10", "22"])],
                     unlabelled_views=[{"description": "ENLARGED DETAIL", "bbox": [0, 0, 1, 1],
                                        "numerals": [_numeral("22")]}])]
    findings = figure_facts.reconcile(sheets=sheets, sections=SECTIONS, numerals=NUMERALS)
    assert any(item["rule"] == "37 CFR 1.84(u)" and item["severity"] == "blocker" and
               "no figure number" in item["title"] for item in findings)


def test_two_leaders_on_one_part_are_left_alone_and_two_on_two_parts_are_not():
    same = [_sheet("FIG. 1", [_view("FIG. 1", [])],
                   numerals=[_numeral("12", lead_lines=2, points_at="two feet of the base")])]
    assert not [item for item in figure_facts.reconcile(
        sheets=same, sections=SECTIONS, numerals=NUMERALS)
        if item["rule"] == "37 CFR 1.84(p)(2)"]
    different = [_sheet("FIG. 1", [_view("FIG. 1", [])],
                        numerals=[_numeral("12", lead_lines=2, designates_one_part=False,
                                           points_at="the base; the slab")])]
    assert any(item["rule"] == "37 CFR 1.84(p)(2)" and item["severity"] == "blocker"
               for item in figure_facts.reconcile(
                   sheets=different, sections=SECTIONS, numerals=NUMERALS))


def test_a_numeral_that_contradicts_the_table_is_a_blocker():
    sheets = [_sheet("FIG. 2", [_view("FIG. 2", [])],
                     numerals=[_numeral("16", matches_declared_part="no",
                                        points_at="the first side")])]
    findings = figure_facts.reconcile(sheets=sheets, sections=SECTIONS, numerals=NUMERALS)
    assert any("does not call it" in item["title"] for item in findings)


def test_a_stale_cross_reference_is_found_and_a_correct_one_is_not():
    good = [_sheet("FIG. 1", [_view("FIG. 1", ["10", "12"])]),
            _sheet("FIG. 2", [_view("FIG. 2", ["16", "22"])])]
    assert not [item for item in figure_facts.reconcile(
        sheets=good, sections=SECTIONS, numerals=NUMERALS)
        if "cross-reference" in item["title"]]
    stale = [_sheet("FIG. 1", [_view("FIG. 1", ["10", "12"])]),
             _sheet("FIG. 2", [_view("FIG. 2", ["10", "12"])])]
    assert any("cross-reference" in item["title"] for item in figure_facts.reconcile(
        sheets=stale, sections=SECTIONS, numerals=NUMERALS))


def test_a_drawing_description_that_hedges_is_a_blocker():
    sections = dict(SECTIONS, drawing_descriptions=(
        SECTIONS["drawing_descriptions"] +
        "; an enlarged portion of the same figure may illustrate the bearing face."))
    sheets = [_sheet("FIG. 1", [_view("FIG. 1", ["10", "12"])]),
              _sheet("FIG. 2", [_view("FIG. 2", ["16", "22"])])]
    findings = figure_facts.reconcile(sheets=sheets, sections=sections, numerals=NUMERALS)
    assert any("does not number" in item["title"] and item["severity"] == "blocker"
               for item in findings)


def test_a_described_view_with_no_sheet_and_a_drawn_view_with_no_description():
    sheets = [_sheet("FIG. 1", [_view("FIG. 1", ["10", "12"])]),
              _sheet("FIG. 7", [_view("FIG. 7", ["22"])])]
    titles = [item["title"] for item in figure_facts.reconcile(
        sheets=sheets, sections=SECTIONS, numerals=NUMERALS)]
    assert any("have no sheet" in title for title in titles)
    assert any("are not described" in title for title in titles)


def test_a_claimed_feature_that_no_sheet_shows_is_a_blocker():
    numerals = NUMERALS + [{"numeral": "28", "part": "port"}]
    sections = dict(SECTIONS, claims=SECTIONS["claims"].replace(
        "and a chamber at the second side.",
        "a chamber at the second side; and a port in fluid communication with the chamber."))
    sheets = [_sheet("FIG. 1", [_view("FIG. 1", ["10", "12"])]),
              _sheet("FIG. 2", [_view("FIG. 2", ["16", "22"])])]
    findings = figure_facts.reconcile(
        sheets=sheets, sections=sections, numerals=numerals,
        claim_terms=figure_facts.claim_terms_from(sections["claims"]))
    assert any(item["rule"] == "37 CFR 1.83(a)" and "28" in item["title"]
               for item in findings)


def test_words_on_a_sheet_that_duplicate_a_numeral_are_reported():
    sheets = [_sheet("FIG. 3", [_view("FIG. 3", ["10"])],
                     text_labels=[{"text": "FLEXIBLE PULLING ELEMENT",
                                   "duplicates_numeral": "46"}],
                     numeral_key_table=True, divider_rules=1)]
    titles = [item["title"] for item in figure_facts.reconcile(
        sheets=sheets, sections=SECTIONS, numerals=NUMERALS)]
    assert any("duplicate reference numeral 46" in title for title in titles)
    assert any("key is printed" in title for title in titles)
    assert any("rule is drawn" in title for title in titles)


# =============================================================================================
# Sheets: the measurement that decides 37 CFR 1.84(p)(3)
# =============================================================================================
def _artwork(width, height, digit_height, count=40):
    from PIL import Image, ImageDraw
    image = Image.new("L", (width, height), 255)
    draw = ImageDraw.Draw(image)
    for index in range(count):
        x = 20 + (index % 10) * (width // 12)
        y = 20 + (index // 10) * (height // 6)
        #  A filled box the size of a digit is close enough for a height measurement, and it is
        #  exactly what the component labeller sees.
        draw.rectangle([x, y, x + max(2, digit_height // 2), y + digit_height], fill=0)
    return image


def test_the_reference_character_height_is_measured_from_the_artwork():
    buffer = io.BytesIO()
    _artwork(1400, 1100, 34).save(buffer, format="PNG")
    measured = figure_facts.measure_character_height(buffer.getvalue())
    assert 30 <= measured["pixels"] <= 38
    assert measured["samples"] >= 12


def test_a_sheet_whose_characters_are_too_small_is_blocked_and_a_good_one_is_not():
    small = filing_rules.audit_drawing_geometry(
        sheet_count=1, described_figures=["FIG. 1"], sheet_labels=[["FIG. 1"]],
        character_points=4.0)
    assert any(item["rule"] == "37 CFR 1.84(p)(3)" and item["severity"] == "blocker"
               for item in small)
    fine = filing_rules.audit_drawing_geometry(
        sheet_count=1, described_figures=["FIG. 1"], sheet_labels=[["FIG. 1"]],
        character_points=12.0)
    assert not [item for item in fine if item["rule"] == "37 CFR 1.84(p)(3)"]


def test_a_sheet_with_no_view_number_is_blocked():
    findings = filing_rules.audit_drawing_geometry(
        sheet_count=2, described_figures=["FIG. 1"],
        sheet_labels=[["FIG. 1"], []], character_points=12.0)
    assert any("carries no view number" in item["title"] for item in findings)


def test_two_views_on_one_upload_become_two_sheets():
    from PIL import Image
    left = _artwork(600, 500, 20)
    right = _artwork(600, 500, 20)
    canvas = Image.new("L", (1300, 500), 255)
    canvas.paste(left, (0, 0))
    canvas.paste(right, (700, 0))
    buffer = io.BytesIO()
    canvas.save(buffer, format="PNG")
    png = buffer.getvalue()
    facts = _sheet("FIG. 2", [
        _view("FIG. 2", ["10"], bbox=[0.02, 0.02, 0.46, 0.98]),
        _view("FIG. 3", ["22"], bbox=[0.54, 0.02, 0.98, 0.98])])
    sheets = filing_pack.drawing_sheets(
        [{"label": "FIG. 2", "png": png}], {figure_facts.sheet_key(png): facts})
    assert [item["label"] for item in sheets] == ["FIG. 2", "FIG. 3"]
    assert all(item["cropped"] for item in sheets)
    for item in sheets:
        with Image.open(io.BytesIO(item["png"])) as cropped:
            assert cropped.size[0] < 800


def test_the_crop_grows_out_to_the_gutter_rather_than_clipping_a_lead_line():
    from PIL import Image, ImageDraw
    canvas = Image.new("L", (1000, 400), 255)
    draw = ImageDraw.Draw(canvas)
    draw.rectangle([300, 100, 700, 300], outline=0, width=3)      # the drawing
    draw.line([40, 200, 300, 200], fill=0, width=3)               # a lead line running off left
    buffer = io.BytesIO()
    canvas.save(buffer, format="PNG")
    ink = filing_pack._ink_profiles(canvas)
    box = filing_pack._snap_box([0.29, 0.2, 0.72, 0.8], [], ink=ink, width=1000, height=400)
    assert box[0] < 60, "the lead line at x=40 was cropped away"


def test_a_divider_rule_between_two_views_is_trimmed_off():
    from PIL import Image, ImageDraw
    canvas = Image.new("L", (800, 400), 255)
    draw = ImageDraw.Draw(canvas)
    draw.rectangle([100, 60, 700, 380], outline=0, width=3)
    draw.line([0, 6, 800, 6], fill=0, width=4)                    # the divider, full width
    trimmed = filing_pack._trim_rules(canvas, (0, 0, 800, 400))
    assert trimmed[1] > 6


# =============================================================================================
# The package as a whole
# =============================================================================================
def test_the_package_names_every_paper_and_audits_what_it_wrote():
    built = filing_pack.build(project=PROJECT, version=VERSION, profile=COMPLETE_PROFILE,
                              figures=[], citations=[])
    names = set(built["files"])
    for name in (filing_pack.README_NAME, filing_pack.SPEC_NAME, filing_pack.ADS_NAME,
                 filing_pack.WEB_ADS_NAME, filing_pack.DECLARATION_NAME,
                 filing_pack.FEES_NAME, filing_pack.IDS_NAME, filing_pack.AUDIT_NAME):
        assert name in names
    #  The only blocker with no drawings is the absence of drawings, which is the truth.
    blockers = [item["title"] for item in built["findings"] if item["severity"] == "blocker"]
    assert blockers == ["There are no drawing sheets"]
    assert "37 CFR 1.16(u)" in built["files"][filing_pack.README_NAME].decode("utf-8")


def test_a_missing_party_field_blocks_the_package_and_says_which_field():
    built = filing_pack.build(project=PROJECT, version=VERSION, profile={}, figures=[],
                              citations=[])
    titles = [item["title"] for item in built["findings"] if item["severity"] == "blocker"]
    assert any("Correspondence email is not supplied" in title for title in titles)
    assert not built["ready"]


def test_the_archive_holds_exactly_the_files_that_were_built():
    built = filing_pack.build(project=PROJECT, version=VERSION, profile=COMPLETE_PROFILE,
                              figures=[], citations=[])
    archive = zipfile.ZipFile(io.BytesIO(filing_pack.zip_bytes(built["files"])))
    assert sorted(archive.namelist()) == sorted(built["files"])


def test_the_audit_text_separates_blockers_from_formalities():
    text = filing_pack.audit_text([
        filing_rules.finding("37 CFR 1.84(u)", "blocker", "sheet 1", "A", "a"),
        filing_rules.finding("37 CFR 1.84(o)", "formality", "sheet 1", "B", "b")])
    assert text.index("BLOCKERS") < text.index("FORMALITIES")
    assert "37 CFR 1.84(u)" in text


# =============================================================================================
# The intake form: many inventors, and what it hands the filing profile
# =============================================================================================
def test_inventor_cards_are_read_by_index_rather_than_as_parallel_lists():
    """A browser posts parallel lists in document order, so one empty field in the middle shifts
    every value after it onto the wrong inventor. The index is in the field name for that reason."""
    from werkzeug.datastructures import MultiDict
    import webapp
    form = MultiDict([
        ("inventor_0_given_name", "Ada"), ("inventor_0_family_name", "Lovelace"),
        ("inventor_0_city", "Reno"),
        ("inventor_1_given_name", ""), ("inventor_1_family_name", ""),
        ("inventor_2_given_name", "Grace"), ("inventor_2_family_name", "Hopper"),
        ("inventor_2_country", "US"), ("title", "not an inventor field"),
        ("inventor_0_nonsense", "ignored"),
    ])
    rows = webapp._inventor_rows_from_form(form)
    assert [(row["given_name"], row["family_name"]) for row in rows] == \
        [("Ada", "Lovelace"), ("Grace", "Hopper")]
    assert rows[1]["country"] == "US"
    assert "nonsense" not in rows[0]


def test_an_inventor_card_holds_every_field_the_ads_and_the_declaration_need():
    keys = {key for key, _label, _required in filing_profile.INVENTOR_FIELDS}
    assert {"given_name", "family_name", "city", "country",
            "mailing_address", "mailing_postcode"} <= keys
    rows = [{"given_name": "Ada", "family_name": "Lovelace"}]
    profile = filing_profile.resolve({"inventors": rows})
    assert filing_profile.full_name(profile["inventors"][0]) == "Ada Lovelace"
    #  Naming an inventor is not the same as being able to file for them, and the gap list says
    #  which field is still missing rather than that the section is incomplete.
    missing = [gap["field"] for gap in filing_profile.gaps(profile)]
    assert "Inventor 1: City of residence" in missing
