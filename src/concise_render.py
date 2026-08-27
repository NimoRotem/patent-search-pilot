"""Render a concise description of relevance to the two formats the USPTO accepts on filing.

PDF is the filing artefact (Patent Center takes PDF for a 37 CFR 1.290 submission) and DOCX is the
editable copy the attorney marks up before filing. Both carry the same content from the same model,
so the paper that gets filed is the paper that was reviewed.

Layout follows the attorney's own examples: a running "Re: U.S. App No. ..." line, the statutory
heading, the application identification, the document's biblio block, a one-paragraph
characterisation, then the two-column table. Letter paper, 1 inch margins, 11pt serif — the
formalities a paper has to satisfy to be accepted rather than returned.
"""
from __future__ import annotations

import html as _html
import io
import re

import pdf_fonts

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (BaseDocTemplate, Frame, PageTemplate, Paragraph, Table,
                                TableStyle)

HEADING = ("CONCISE DESCRIPTION OF RELEVANCE: THIRD-PARTY SUBMISSION "
           "UNDER 37 CFR § 1.290")


def _esc(s):
    """Escape for reportlab's Paragraph markup, and keep every glyph drawable.

    The second half matters as much as the first. Asked for a character the current face has no
    glyph for, reportlab silently substitutes ZapfDingbats, and a Chinese name printed on a filed
    document list as two solid black squares. Runs outside the Latin face are wrapped in a span
    pointing at an embedded fallback, so the paper says what it means or the render fails loudly.

    Every caller feeds a Paragraph. The running head goes to canvas.drawString and does NOT come
    through here, which is why this can return markup.
    """
    return pdf_fonts.with_fallback(_html.escape(str(s or "")))


def is_latin(name):
    """True when every letter can be typeset by the filing font. U+2E80 starts the CJK blocks."""
    return all(ord(c) < 0x2E80 for c in str(name or ""))


def printable_party(biblio):
    """Who to name on the paper for this document, and what to call them. -> (label, value)

    1.290(e)(3) identifies a foreign document by "the applicant, patentee, or first named
    inventor", and that OR is the way out of a problem this hit twice. The filing font has no CJK
    glyphs, so a Chinese or Japanese personal name with no Latin form in the record printed as a
    row of solid black boxes on a paper filed at the USPTO: CN 216190291 U went out identifying its
    inventor as "■■".

    Preferring a romanisation is not always possible, because a romanisation is not always in the
    record: for that CN document all seven inventors are Chinese-only and nothing carries a Latin
    form. Transliterating is not an option either, since a name nobody verified does not belong on
    a filing. The applicant does carry a Latin name, and the rule accepts it, so the paper names
    the applicant and says so.

    Falls back to the unprintable name only when there is no Latin alternative at all, and the
    audit fails the packet in that case rather than filing boxes.
    """
    b = biblio or {}
    inventor = str(b.get("inventor") or "").strip()
    if inventor and is_latin(inventor):
        return "First Named Inventor", inventor
    applicant = str(b.get("assignee") or "").strip()
    #  An assignee field can hold several, comma-joined. The first is the one to name.
    first = applicant.split(",")[0].strip() if applicant else ""
    if first and is_latin(first):
        return "Applicant", first
    return "First Named Inventor", inventor


def running_head(subject):
    """The line at the head and foot of every page of the filed paper.

    "Re: U.S. Application No. 19/318,450 (Publication No. US 2026/0070232 A1)"

    THIS USED TO PRINT THE PUBLICATION NUMBER UNDER THE WORDS "App No." whenever the application
    number was blank, which is how a filed document came to say `Re: U.S. App No. US 2026/0070232
    A1`. A publication number is not an application number and a paper that says it is, is wrong on
    its face. With no application number the line names the publication AS a publication instead.
    """
    app = str(subject.get("app_no") or "").strip()
    pub = str(subject.get("pub_no") or "").strip()
    if app and pub:
        return "Re: U.S. Application No. %s (Publication No. %s)" % (app, pub)
    if app:
        return "Re: U.S. Application No. %s" % app
    if pub:
        return "Re: U.S. Publication No. %s" % pub
    return "Re: the above-identified application"


def subject_line(subject):
    """U.S. Application No. 18/915,337 (Publication No. US 2025/0033224 A1), "Title", Inventor."""
    bits = []
    if subject.get("app_no"):
        bits.append("U.S. Application No. %s" % subject["app_no"])
    if subject.get("pub_no"):
        bits.append("(Publication No. %s)" % subject["pub_no"] if bits
                    else "U.S. Publication No. %s" % subject["pub_no"])
    head = " ".join(bits)
    tail = []
    if subject.get("title"):
        tail.append("“%s”" % subject["title"])
    if subject.get("inventor"):
        tail.append(subject["inventor"])
    return ", ".join([x for x in [head] + tail if x])


def _left_cell(row, pub_no_label):
    if row["quote_claim"] and row["claim_text"]:
        return "Claim %s: “%s”" % (row["claim_no"], row["claim_text"])
    para = row["claim_paraphrase"] or row["claim_text"]
    return "Claim %s (%s)" % (row["claim_no"], para)


def _right_bits(row):
    """(prose, [citations]) — the quotation is only used where a passage survived grounding."""
    prose = row.get("disclosure") or row.get("note") or ""
    if row.get("quote_translated") and row.get("quote_original"):
        #  1.290(d)(3) wants the translation; an examiner checking it wants the original. Both.
        orig = " ".join(str(row["quote_original"]).split())
        prose = "%s\n[original] \u201c%s\u201d" % (prose, orig[:400]) if prose else orig[:400]
    if row.get("strong") and row.get("quote"):
        #  The stored passage is capped mid-word by the reader's character budget. Filed text may
        #  be elided, but it may not look like a transcription error, so cut back to the last whole
        #  word and mark the elision. 40 words is the cap the reader was asked to quote within.
        q = " ".join(row["quote"].split())
        words = q.split()
        elided = len(words) > 40
        if elided:
            words = words[:40]
        if words and not re.search(r"[.!?\"\u201d)]$", words[-1]):
            elided = True
        q = " ".join(words)
        if elided:
            q = q.rstrip(" ,;:-") + " …"
        prose = "%s\n“%s”" % (prose, q) if prose else "“%s”" % q
    return prose, list(row.get("cites") or [])


def filing_notes(doc_model):
    """The lines a practitioner must read before signing this paper.

    NEVER RENDERED INTO THE PDF OR THE DOCX. They used to be, under a heading that asked the reader
    to "delete this block before filing", and the obvious thing happened: they were filed. Two of
    them are actively harmful in front of an examiner. The translation note is addressed to the
    practitioner and says so in its own words; and

        "5 of 14 quotations could not be found in the stored text of this reference and were
         removed"

    is a statement, on the face of a paper filed at the Office, that the process which produced it
    failed its own source verification. A caveat that is safe in a review tool is not safe on a
    filing, and an instruction to delete something is not a control.

    They are shown on the build page instead, next to the download links, which is where the
    practitioner reviews the package. The paper carries only what is filed.
    """
    c = doc_model.get("compliance") or {}
    out = []
    q = c.get("qualify") or {}
    if q.get("note"):
        out.append(("Prior-art basis", q["note"]))
    sc = c.get("self_collision") or {}
    if sc.get("note"):
        out.append(("Common ownership", sc["note"]))
    #  AN EXCLUSION IS ABOUT ONE PUBLICATION AND A FAMILY IS MANY. Where the sweep found a member
    #  that published earlier, that is the most valuable line on this page: the disclosure is
    #  available even though this publication of it is not. See family_sweep.
    sib = c.get("sibling") or {}
    if sib.get("best"):
        out.append(("An earlier member of this family", sib.get("note") or ""))
    elif sib.get("checked") is False and sib.get("note"):
        out.append(("Family", sib["note"]))
    tr = c.get("translation") or {}
    if tr.get("translated"):
        out.append(("Translation", "%d relied-on passage%s machine-translated into English; the "
                                   "original-language text is retained in the record. %s"
                    % (tr["translated"], "" if tr["translated"] == 1 else "s",
                       "A verified human translation may be required before filing.")))
    elif tr.get("note"):
        out.append(("Translation", tr["note"]))
    qz = c.get("quotes") or {}
    if qz.get("note"):
        out.append(("Quotations", qz["note"]))
    if c.get("neutralised"):
        out.append(("Neutral language",
                    "Argumentative phrasing was removed so the description states disclosure only, "
                    "as 37 CFR 1.290 requires: %s." % "; ".join(c["neutralised"][:6])))
    return out


# --------------------------------------------------------------------------- PDF


def _styles():
    base = ParagraphStyle("cd", fontName=pdf_fonts.font(pdf_fonts.SERIF), fontSize=11, leading=13.5,
                          alignment=TA_LEFT, spaceAfter=0)
    return {
        "h": ParagraphStyle("h", parent=base, fontName=pdf_fonts.font(pdf_fonts.SERIF_BOLD), fontSize=11.5, leading=14,
                            spaceAfter=6),
        "app": ParagraphStyle("app", parent=base, fontName=pdf_fonts.font(pdf_fonts.SERIF_ITALIC), spaceAfter=8),
        "doc": ParagraphStyle("doc", parent=base, fontName=pdf_fonts.font(pdf_fonts.SERIF_BOLD), spaceAfter=3),
        "bib": ParagraphStyle("bib", parent=base, leftIndent=20, spaceAfter=1),
        "body": ParagraphStyle("body", parent=base, spaceBefore=6, spaceAfter=9),
        "th": ParagraphStyle("th", parent=base, fontName=pdf_fonts.font(pdf_fonts.SERIF_BOLD), fontSize=10.5, leading=13),
        "td": ParagraphStyle("td", parent=base, fontSize=10.5, leading=13),
        "cite": ParagraphStyle("cite", parent=base, fontSize=10.5, leading=13, leftIndent=10,
                               bulletIndent=2, spaceBefore=1),
        "run": ParagraphStyle("run", parent=base, fontSize=9.5, leading=11,
                              textColor=colors.HexColor("#333333")),
    }


def to_pdf(doc_model) -> bytes:
    st = _styles()
    subj = doc_model["subject"]
    running = running_head(subj)
    buf = io.BytesIO()

    def _page(canv, docobj):
        canv.saveState()
        canv.setFont(pdf_fonts.font(pdf_fonts.SERIF), 9.5)
        canv.setFillColor(colors.HexColor("#333333"))
        canv.drawString(inch, letter[1] - 0.62 * inch, running)
        canv.drawString(inch, 0.62 * inch, running)
        canv.restoreState()

    tmpl = BaseDocTemplate(buf, pagesize=letter, leftMargin=inch, rightMargin=inch,
                           topMargin=0.95 * inch, bottomMargin=0.95 * inch,
                           title="Concise Description of Relevance: %s" % doc_model["pub"],
                           author="Third-party submission under 37 CFR 1.290")
    frame = Frame(inch, 0.95 * inch, letter[0] - 2 * inch, letter[1] - 1.9 * inch, id="f")
    tmpl.addPageTemplates([PageTemplate(id="p", frames=[frame], onPage=_page)])

    story = [Paragraph(_esc(HEADING), st["h"]),
             Paragraph(_esc(subject_line(subj)), st["app"])]

    b = doc_model["biblio"]
    story.append(Paragraph("Document %s: %s (“Document %s”)" % (
        doc_model["n"], _esc(b["label"]), doc_model["n"]), st["doc"]))
    #  1.290(e) IDENTIFICATION ONLY. A U.S. publication is identified by its publication
    #  number, its first named inventor and its publication date. Assignee and earliest
    #  priority date were neither required nor safe to print: assignment changes hands over
    #  time and these values come from a cache that does not always match the current public
    #  record, so the paper asserted ownership facts it had no need to assert.
    for lbl, val in (printable_party(b),
                     ("Issue Date" if b.get("kind") == "patent" else "Publication Date",
                      b.get("issue_date_pretty")),
                     ("Title", b.get("title"))):
        if val:
            story.append(Paragraph("<i>%s:</i> %s" % (_esc(lbl), _esc(val)), st["bib"]))

    story.append(Paragraph(
        _esc(doc_model["summary"]) +
        " Its potential relevance to the claims of the above-identified application is set forth "
        "below.", st["body"]))

    pub_no = subj.get("pub_no") or subj.get("app_no") or ""
    head = [Paragraph("Claim language (%s)" % _esc(pub_no), st["th"]),
            Paragraph("Relevant disclosure of this document", st["th"])]
    data = [head]
    for row in doc_model["rows"]:
        prose, cites = _right_bits(row)
        right = [Paragraph(_esc(prose).replace("\n", "<br/>"), st["td"])]
        for c in cites:
            right.append(Paragraph("• %s" % _esc(c), st["cite"]))
        data.append([Paragraph(_esc(_left_cell(row, pub_no)), st["td"]), right])

    w = letter[0] - 2 * inch
    tbl = Table(data, colWidths=[w * 0.46, w * 0.54], repeatRows=1)
    tbl.setStyle(TableStyle([
        #  A Table whose cells are Paragraphs STILL emits its own default cell font, and
        #  that default is an unembedded Helvetica. Naming it here is what keeps a base-14
        #  resource out of a paper Patent Center validates.
        ("FONTNAME", (0, 0), (-1, -1), pdf_fonts.font(pdf_fonts.SERIF)),
        ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#444444")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEFEF")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(tbl)
    #  NO FILING NOTES ON THE PAPER. See filing_notes: they are for the practitioner and are shown
    #  on the build page. One of them announced that quotations had failed verification, on a
    #  document filed at the Office.
    tmpl.build(story)
    return buf.getvalue()


# --------------------------------------------------------------------------- DOCX


def to_docx(doc_model) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    subj = doc_model["subject"]
    d = Document()
    s = d.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(s, m, Inches(1))
    normal = d.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    running = running_head(subj)
    for holder in (s.header, s.footer):
        p = holder.paragraphs[0]
        p.text = running
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.size = Pt(9.5)

    h = d.add_paragraph()
    h.add_run(HEADING).bold = True
    ap = d.add_paragraph()
    ap.add_run(subject_line(subj)).italic = True

    b = doc_model["biblio"]
    dp = d.add_paragraph()
    dp.add_run("Document %s: %s (“Document %s”)" % (
        doc_model["n"], b["label"], doc_model["n"])).bold = True
    #  Identification fields only. See the note in to_pdf.
    for lbl, val in (printable_party(b),
                     ("Issue Date" if b.get("kind") == "patent" else "Publication Date",
                      b.get("issue_date_pretty")),
                     ("Title", b.get("title"))):
        if val:
            p = d.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.space_after = Pt(1)
            p.add_run("%s: " % lbl).italic = True
            p.add_run(str(val))

    d.add_paragraph(doc_model["summary"] +
                    " Its potential relevance to the claims of the above-identified application "
                    "is set forth below.")

    pub_no = subj.get("pub_no") or subj.get("app_no") or ""
    t = d.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for cell, text in zip(hdr, ("Claim language (%s)" % pub_no,
                                "Relevant disclosure of this document")):
        cell.paragraphs[0].add_run(text).bold = True
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "EFEFEF")
        cell._tc.get_or_add_tcPr().append(shd)

    for row in doc_model["rows"]:
        prose, cites = _right_bits(row)
        cells = t.add_row().cells
        cells[0].text = _left_cell(row, pub_no)
        first = True
        for chunk in (prose or "").split("\n"):
            p = cells[1].paragraphs[0] if first else cells[1].add_paragraph()
            p.text = chunk
            first = False
        for c in cites:
            p = cells[1].add_paragraph()
            p.paragraph_format.left_indent = Inches(0.14)
            p.text = "• %s" % c

    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10.5)

    #  NO FILING NOTES HERE EITHER. The DOCX is the copy the attorney marks up and then files, so
    #  anything left in it is a thing that gets filed. See filing_notes.
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def filename(doc_model, ext):
    pub = re.sub(r"[^A-Za-z0-9]+", "", doc_model["pub"] or "doc")
    return "ConciseDescription_Doc%s_%s.%s" % (doc_model["n"], pub, ext)
