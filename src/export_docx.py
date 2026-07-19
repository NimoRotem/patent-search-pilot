"""Render the export model to an editable DOCX (python-docx). Milestone 4 §1."""
from __future__ import annotations
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

import disclosure

NAVY = RGBColor(0x0b, 0x25, 0x45)
BLUE = RGBColor(0x00, 0x50, 0xA0)
GOOD = RGBColor(0x0a, 0x7d, 0x4d)
SECRET = RGBColor(0x7b, 0x2f, 0xbf)
WARN = RGBColor(0xb2, 0x5e, 0x00)
MUTED = RGBColor(0x5b, 0x6b, 0x82)


def _basis_label(b):
    return {"public_prior_art": "PUBLIC prior art", "secret_prior_art": "SECRET prior art (novelty only)",
            "priority_interval": "priority-interval art"}.get(b, "not dated")


def _shade(cell, hex_no_hash):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_no_hash)
    tcPr.append(shd)


#  _green_shade() DELETED ON PURPOSE.
#
#  It blended white -> green in proportion to `intensity`, which is the normalised RETRIEVAL score.
#  So in a document headed "Prior-Art Search Report", the strongest visual assertion of coverage
#  was made by whichever cell the retriever was most confident about -- including, in the
#  grabo_gripper_novelty report, US-9266686-B2 cl 6, a cell the verifier judged UNRELATED (it
#  describes a vent hole and air filter, not a rigid base element with sides). Cell appearance now
#  comes from disclosure.CELL_STATES, i.e. from the verification verdict. See disclosure.py.


def _run(p, text, bold=False, color=None, size=None, italic=False):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return r


def render(model, out_path):
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    # ---- cover ----
    h = doc.add_heading(disclosure.DOC_TITLE, level=0)
    for run in h.runs:
        run.font.color.rgb = NAVY
    sp = doc.add_paragraph()
    _run(sp, disclosure.DOC_SUBTITLE, bold=True, color=WARN, size=10)
    p = doc.add_paragraph()
    _run(p, "Search mode: ", bold=True); _run(p, model["mode"].replace("_", " ").title())
    _run(p, "     Subject: ", bold=True); _run(p, str(model.get("subject") or "—"))
    _run(p, "     Date: ", bold=True); _run(p, model["generated"])

    doc.add_heading("Search query / invention", level=2)
    doc.add_paragraph(model["query"][:2000])
    doc.add_heading("Method & corpus", level=2)
    mp = doc.add_paragraph(); _run(mp, model["corpus_note"], size=8.5, color=MUTED)

    p = doc.add_paragraph()
    _run(p, f"{model['n_covered']}/{model['n_elements']} elements disclosed", bold=True)
    _run(p, f"   ·   {len(model['references'])} references cited   ·   ")
    _run(p, f"{len(model['uncovered_elements'])} apparently novel", bold=True, color=WARN)

    # ---- scope + measured reliability ----
    # On the cover, ahead of any result. An exported document has to stand alone: the reader may
    # never have seen the web page, cannot hover a tooltip, and may be reading it months later.
    _scope_section(doc)

    # ---- executive summary ----
    doc.add_heading("Executive summary", level=1)
    if model["strongest"]:
        doc.add_paragraph("Strongest references:").runs[0].bold = True
        for s in model["strongest"]:
            pp = doc.add_paragraph(style="List Bullet")
            _run(pp, f"{s['pub']} — {s['title']} ", bold=True)
            _run(pp, f"(match {s['score']}; {_basis_label(s['basis'])})", color=MUTED, size=9)
    p = doc.add_paragraph()
    _run(p, "Elements disclosed by the cited art: ", bold=True)
    _run(p, ", ".join(model["covered_elements"]) or "—")
    p = doc.add_paragraph()
    _run(p, "Apparently novel (not found in corpus): ", bold=True)
    _run(p, ", ".join(model["uncovered_elements"]) or "none identified", color=WARN)

    # ---- retrieval map (formerly mis-titled "claim chart") ----
    doc.add_heading(disclosure.CHART_TITLE, level=1)
    tp = doc.add_paragraph()
    _run(tp, disclosure.CHART_TAG.upper(), bold=True, color=WARN, size=9)
    wp = doc.add_paragraph()
    _run(wp, disclosure.chart_warning(), size=8.5)
    _claim_chart(doc, model)

    # ---- combination ----
    cv = model["combination_view"]
    if cv.get("primary"):
        doc.add_heading("Combination / inventive-step analysis", level=1)
        p = doc.add_paragraph(); _run(p, "Primary reference: ", bold=True)
        _run(p, f"{cv['primary']} — discloses: {', '.join(cv.get('covers', [])) or '—'}")
        for sec in cv.get("secondaries", []):
            p = doc.add_paragraph(); _run(p, "+ Secondary: ", bold=True)
            _run(p, f"{sec['ref']} — supplies: {', '.join(sec['supplies'])}")
        if cv.get("uncovered_elements"):
            p = doc.add_paragraph(); _run(p, "Not disclosed by any reference: ", bold=True)
            _run(p, ", ".join(cv["uncovered_elements"]), color=WARN)

    # ---- per reference ----
    doc.add_page_break()
    doc.add_heading("Cited references", level=1)
    for i, r in enumerate(model["references"]):
        _reference(doc, r, i + 1)

    # ---- appendix ----
    doc.add_page_break()
    doc.add_heading("Appendix — full ranked reference list", level=1)
    tbl = doc.add_table(rows=1, cols=6); tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for j, t in enumerate(["#", "Publication", "Title", "Match", "Basis", "Channels"]):
        hdr[j].paragraphs[0].add_run(t).bold = True
    for a in model["appendix"]:
        c = tbl.add_row().cells
        c[0].text = str(a.get("rank") or ""); c[1].text = a["pub"]
        c[2].text = (a.get("title") or "")[:60]; c[3].text = str(a.get("score") or "")
        c[4].text = _basis_label(a.get("basis")).split(" (")[0]
        c[5].text = ", ".join(a.get("channels", []))

    doc.save(str(out_path))
    return str(out_path)


def _scope_section(doc):
    """The full scope + measured-reliability disclosure, printed, not linked."""
    doc.add_heading("Scope and measured reliability — read before relying on this document", level=1)
    for heading, body in disclosure.scope_paragraphs():
        p = doc.add_paragraph()
        _run(p, f"{heading}. ", bold=True, size=9.5)
        _run(p, body, size=9.5)
    p = doc.add_paragraph()
    _run(p, "Indexed CPC classes: ", bold=True, size=9)
    _run(p, "; ".join(disclosure.cpc_lines()), size=9)
    p = doc.add_paragraph()
    _run(p, disclosure.not_indexed(), size=9, color=MUTED)


def _claim_chart(doc, model):
    chart = model["claim_chart"]; cols = chart["columns"]
    tbl = doc.add_table(rows=1, cols=1 + len(cols)); tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    hdr[0].paragraphs[0].add_run("Invention element").bold = True
    for j, c in enumerate(cols):
        run = hdr[j + 1].paragraphs[0].add_run(c["pub"]); run.bold = True; run.font.size = Pt(7.5)
        hdr[j + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in chart["rows"]:
        cells = tbl.add_row().cells
        er = cells[0].paragraphs[0].add_run(row["element"]); er.font.size = Pt(8.5)
        for j, cell in enumerate(row["cells"]):
            para = cells[j + 1].paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cell.get("covered"):
                _mark, _label, fill, is_cov = disclosure.cell_state(cell)
                word = disclosure.cell_word(cell)
                # The MARK carries the meaning and leads the cell; the retrieval score follows it
                # in muted small type. Previously the score was the bold headline of the cell,
                # which is exactly backwards -- a high retrieval score is not evidence of
                # disclosure, and printing it in bold green said that it was.
                mr = para.add_run(word); mr.bold = True; mr.font.size = Pt(8)
                sr = para.add_run(f"\n{cell.get('score')}")
                sr.font.size = Pt(7); sr.font.color.rgb = MUTED
                if cell.get("coord"):
                    cr = para.add_run(f"\n{cell['coord']}")
                    cr.font.size = Pt(6.5); cr.font.color.rgb = MUTED
                _shade(cells[j + 1], fill)
            else:
                para.add_run("·").font.color.rgb = MUTED
    # Legend + this-report verification rate, directly under the grid where the marks are read.
    lp = doc.add_paragraph()
    _run(lp, "Legend: ", bold=True, size=8.5)
    _run(lp, "  ·  ".join(f"{n} {w} — {lbl}" for w, n, lbl in disclosure.legend_lines(chart))
             or "no cells to verify", size=8.5)
    summ = disclosure.verification_summary(chart)
    if summ:
        sp = doc.add_paragraph(); _run(sp, summ, size=8.5, bold=True)


def _reference(doc, r, n):
    doc.add_heading(f"{n}. {r['pub']} — {r['title'] or '(untitled)'}", level=2)
    # biblio + drawing side by side via a 2-col table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.columns[0].width = Inches(4.4); tbl.columns[1].width = Inches(2.2)
    left, right = tbl.rows[0].cells
    def line(label, val):
        p = left.add_paragraph(); _run(p, f"{label}: ", bold=True); _run(p, str(val or "—")); p.paragraph_format.space_after = Pt(1)
    line("Assignee", "; ".join(r["assignees"]))
    line("Inventors", ", ".join(r["inventors"]))
    line("Priority / Filing / Published", f"{r['priority_date'] or '—'} / {r['filing_date'] or '—'} / {r['publication_date'] or '—'}")
    line("Family", r["family_id"]); line("CPC", ", ".join(r["cpc"]))
    line("Legal status", r["legal_status"])
    bp = left.add_paragraph(); _run(bp, "Prior-art basis: ", bold=True)
    _run(bp, _basis_label(r["basis"]), color=(GOOD if r["basis"] == "public_prior_art" else SECRET if r["basis"] == "secret_prior_art" else MUTED))
    _run(bp, f"     Match: {r['match_score']}")
    if r.get("covers_elements"):
        cp = left.add_paragraph(); _run(cp, "Reads on: ", bold=True, size=9); _run(cp, ", ".join(r["covers_elements"]), size=9)
    if r.get("drawing_path"):
        try:
            with PILImage.open(r["drawing_path"]) as im:
                w, hh = im.size
            width = Inches(min(2.1, 2.1));
            right.paragraphs[0].add_run().add_picture(r["drawing_path"], width=Inches(2.0))
            cap = right.add_paragraph(); _run(cap, f"{r['pub']} · drawing", size=7, color=MUTED)
        except Exception:
            _run(right.paragraphs[0], "[facsimile not digitized]", size=8, color=WARN)
    else:
        _run(right.paragraphs[0], "[facsimile not digitized]", size=8, color=WARN)
    if r.get("why"):
        p = doc.add_paragraph(); _run(p, "Why relevant: ", bold=True); _run(p, r["why"])
    if r.get("quoted"):
        q = r["quoted"]
        p = doc.add_paragraph(); _run(p, f"Relevant passage ({q['kind']} {q['coord']}, similarity {q['score']}):", bold=True, size=9)
        qp = doc.add_paragraph(); qp.paragraph_format.left_indent = Inches(0.2)
        qr = qp.add_run(q["text"]); qr.italic = True; qr.font.size = Pt(9)
    doc.add_paragraph()
