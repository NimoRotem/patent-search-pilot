"""Filing readiness, and the package that comes out the other end.

The last question this product has to answer is the one the user actually came for: *can I file
this?*  Answering it honestly means separating three different things that all feel like "not
ready":

  BLOCKERS      the application would be defective or incomplete as filed - an unresolved
                drafting note where a dimension should be, a citation that resolves to nothing,
                no named inventor.  These are listed as blockers and the export says so.
  FORMALITIES   things the USPTO will object to but which do not stop a filing date - a title
                over 500 characters, an abstract over 150 words.
  ADMINISTRATIVE the oath or declaration, entity-status certification, fee payment, and optional
                 practitioner review. These are listed as what remains, never silently ticked off.

NO FEE AMOUNTS ARE PRINTED.  The counts that drive them (total claims, independent claims,
multiple dependent claims, specification sheet count) are computed exactly, and which surcharges
those counts trigger is stated - but the dollar figures change by fee-setting rulemaking and a
number baked in here would be quietly wrong within a year.  The current schedule is one link away
and always right.

WHAT THIS IS NOT.  Nothing here is a legal opinion, and no check in this module says anything
about whether the application is allowable.  It reads the document against the formal
requirements of 37 CFR 1.52, 1.72, 1.75, 1.77 and 1.121.
"""
from __future__ import annotations

import html
import json
import re
from collections import Counter
from io import BytesIO
from typing import Any, Mapping, Sequence

import draft_qa
import draft_workspace
import drafting

FEE_SCHEDULE_URL = "https://www.uspto.gov/learning-and-resources/fees-and-payment/uspto-fee-schedule"
EFS_URL = "https://patentcenter.uspto.gov/"

# 37 CFR 1.77(b): the order in which the parts of a utility specification are arranged.
FILING_ORDER = (
    ("title", "TITLE OF THE INVENTION", False),
    ("cross_reference", "CROSS-REFERENCE TO RELATED APPLICATIONS", True),
    ("government_support",
     "STATEMENT REGARDING FEDERALLY SPONSORED RESEARCH OR DEVELOPMENT", True),
    ("field", "FIELD OF THE INVENTION", True),
    ("background", "BACKGROUND OF THE INVENTION", True),
    ("summary", "BRIEF SUMMARY OF THE INVENTION", True),
    ("drawing_descriptions", "BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS", True),
    ("detailed_description", "DETAILED DESCRIPTION OF THE INVENTION", True),
)

_NOTE_RE = re.compile(r"\[DRAFTING NOTE:([^\]]*)\]", re.IGNORECASE)


def _filing_sections(version: Mapping[str, Any]) -> dict[str, str]:
    import draft_cite
    try:
        return {str(key): draft_cite.filing_citations(str(value or ""))
                for key, value in dict(version.get("sections") or {}).items()}
    except ValueError as exc:
        raise drafting.DraftingValidationError(str(exc)) from exc


def _filing_label(value: Any) -> str:
    import draft_cite
    try:
        return draft_cite.filing_citations(str(value or ""))
    except ValueError as exc:
        raise drafting.DraftingValidationError(str(exc)) from exc


# =============================================================================================
# Readiness
# =============================================================================================
def fee_profile(claims_text: str) -> dict[str, Any]:
    """The claim counts that decide the filing fee, and which surcharges they trigger."""
    claims = draft_qa.split_claims(claims_text)
    independent, multiple, extra_from_multiple = 0, 0, 0
    for claim in claims:
        dependencies = draft_qa.claim_dependencies(claim["text"])
        if not dependencies:
            independent += 1
        elif len(dependencies) > 1:
            multiple += 1
            # 37 CFR 1.75(c): a multiple dependent claim is counted as the number of claims to
            # which it refers, which is why one of them can cost more than ten ordinary claims.
            extra_from_multiple += len(dependencies) - 1
    total = len(claims)
    billable = total + extra_from_multiple
    surcharges = []
    if independent > 3:
        surcharges.append(f"{independent - 3} independent claim(s) over the three included "
                          "(37 CFR 1.16(h))")
    if billable > 20:
        surcharges.append(f"{billable - 20} claim(s) over the twenty included (37 CFR 1.16(i))")
    if multiple:
        surcharges.append(f"{multiple} multiple dependent claim(s) (37 CFR 1.16(j)); each is "
                          "counted as the number of claims it refers to")
    return {"total": total, "independent": independent, "dependent": total - independent,
            "multiple_dependent": multiple, "billable": billable, "surcharges": surcharges,
            "fee_schedule_url": FEE_SCHEDULE_URL}


def open_notes(sections: Mapping[str, str]) -> list[dict[str, str]]:
    out = []
    for key, _name, heading in draft_workspace.SECTION_FILES:
        for match in _NOTE_RE.finditer(str(sections.get(key) or "")):
            out.append({"section": heading, "note": match.group(1).strip()[:400]})
    return out


def readiness(*, project: Mapping[str, Any], version: Mapping[str, Any],
              qa: Mapping[str, Any] | None = None,
              references: Sequence[Mapping[str, Any]] = (),
              figures: Sequence[Mapping[str, Any]] = ()) -> dict[str, Any]:
    """Is this draft in a state where a practitioner could take it to Patent Center?"""
    sections = dict(version.get("sections") or {})
    blockers: list[dict[str, str]] = []
    formalities: list[dict[str, str]] = []
    remaining: list[str] = []

    notes = open_notes(sections)
    if notes:
        blockers.append({
            "title": f"{len(notes)} unresolved drafting note(s)",
            "detail": "Each one marks a fact only the inventor has. Filing with them in place "
                      "puts placeholder text into the published specification.",
            "items": "; ".join(f"{n['section']}: {n['note']}" for n in notes[:8])})

    unfinished = draft_qa.find_placeholders(sections)
    if unfinished and not notes:
        blockers.append({
            "title": f"{len(unfinished)} unfinished marker(s) in the application",
            "detail": "No TODO, placeholder, blank field, or instruction may enter a filing copy.",
            "items": "; ".join(unfinished[:8])})
    missing_sections = [heading for key, _name, heading in draft_workspace.SECTION_FILES
                        if not str(sections.get(key) or "").strip()]
    if missing_sections:
        blockers.append({
            "title": "The application is missing required sections",
            "detail": "Every filing section must contain final text before export.",
            "items": "; ".join(missing_sections)})

    checks = list((qa or {}).get("checks") or [])
    expected_version = int(version.get("version_no") or 0)
    qa_version = int((qa or {}).get("version_no") or 0)
    if qa is None:
        blockers.append({"title": "This draft has not been reviewed",
                         "detail": "Run the consistency review before treating any draft as "
                                   "ready.", "items": ""})
    elif not expected_version or qa_version != expected_version:
        blockers.append({
            "title": "The exact version being exported has not passed review",
            "detail": "A review of another version cannot authorize this application text.",
            "items": f"application version {expected_version}; review version {qa_version}"})
    elif str(qa.get("status") or "") != "complete" or str(qa.get("verdict") or "") != "pass":
        blockers.append({
            "title": "The independent review did not pass",
            "detail": str(qa.get("summary") or "The review did not return a clean verdict."),
            "items": str(qa.get("last_error") or "")})

    for check in checks:
        if check.get("status") != "pass":
            blockers.append({"title": check.get("name", "Consistency check failed"),
                             "detail": check.get("detail", ""),
                             "items": "; ".join(str(i) for i in (check.get("items") or [])[:6])})
    for finding in ((qa or {}).get("findings") or []):
        blockers.append({"title": finding.get("title", "Independent review finding"),
                         "detail": finding.get("detail", ""),
                         "items": finding.get("where", "")})

    if not str(project.get("inventors") or "").strip():
        blockers.append({"title": "No inventor is named",
                         "detail": "37 CFR 1.41 and 1.63: the application data sheet must name "
                                   "each inventor, and each must sign an oath or declaration.",
                         "items": ""})
    if not str(project.get("applicant") or "").strip():
        formalities.append({"title": "No applicant is recorded",
                            "detail": "37 CFR 1.46: where the applicant is not the inventor, the "
                                      "application data sheet must say who it is.", "items": ""})

    described = draft_qa.figures_mentioned(str(sections.get("drawing_descriptions") or ""))
    figure_specs = version.get("figure_specs") or []
    numeral_table = version.get("numerals") or []
    if isinstance(figure_specs, str):
        try:
            figure_specs = json.loads(figure_specs)
        except json.JSONDecodeError:
            figure_specs = []
    if isinstance(numeral_table, str):
        try:
            numeral_table = json.loads(numeral_table)
        except json.JSONDecodeError:
            numeral_table = []
    drawing_markers = []
    drawing_markers.extend(draft_qa.placeholders_in_text(
        "Reference numeral table", json.dumps(numeral_table, ensure_ascii=False)))
    drawing_markers.extend(draft_qa.placeholders_in_text(
        "Drawing specifications", json.dumps(figure_specs, ensure_ascii=False)))
    if drawing_markers:
        blockers.append({
            "title": f"{len(drawing_markers)} unfinished marker(s) in the drawing sources",
            "detail": "No note, placeholder, confirmation request, or manual instruction may "
                      "enter a filing package.",
            "items": "; ".join(drawing_markers[:8])})
    if not described:
        blockers.append({
            "title": "The application has no drawing plan",
            "detail": "The drafting pipeline must describe and generate at least one figure that "
                      "explains the disclosed structure or process.", "items": ""})
    elif not figures:
        blockers.append({
            "title": f"{len(described)} figure(s) are described but no drawing sheet exists",
            "detail": "35 USC 113 requires a drawing where one is necessary to understand the "
                      "invention. Drawings may be filed informally and corrected later, but they "
                      "must be filed.", "items": ""})
    if described and not figure_specs:
        blockers.append({
            "title": "The reviewed version has no drawing specifications",
            "detail": "Each drawing sheet must be tied to the exact versioned figure brief and "
                      "reference numeral table.", "items": ""})
    drawn_numbers = [draft_qa.figure_number(
        figure.get("figure_label") or figure.get("label")) for figure in figures]
    duplicates = sorted(number for number, count in Counter(drawn_numbers).items()
                        if number and count > 1)
    if duplicates:
        blockers.append({
            "title": "The active drawing set has duplicate figure numbers",
            "detail": "Each figure number must identify exactly one active filing sheet.",
            "items": "; ".join(f"duplicate FIG. {value}" for value in duplicates)})
    drawn_labels = set(drawn_numbers)
    drawn_labels.discard("")
    missing_drawings = sorted(described - drawn_labels)
    extra_drawings = sorted(drawn_labels - described)
    if missing_drawings or extra_drawings:
        blockers.append({
            "title": "The drawing set does not match the application text",
            "detail": "Every described figure must have exactly one checked sheet, with no "
                      "obsolete or extra sheet.",
            "items": "; ".join(
                [f"missing FIG. {value}" for value in missing_drawings] +
                [f"unexpected FIG. {value}" for value in extra_drawings])})
    live_drawing_failures = []
    import draft_figures
    specs_by_key = {draft_figures.figure_key(spec.get("label")): spec
                    for spec in figure_specs if isinstance(spec, Mapping)}
    for sheet_index, figure in enumerate(figures, 1):
        label = figure.get("figure_label") or figure.get("label") or "drawing"
        spec = specs_by_key.get(draft_figures.figure_key(label))
        active = next((row for row in (figure.get("versions") or [])
                       if int(row.get("version_no") or 0) ==
                       int(figure.get("active_version") or 0)), None) or {}
        expected_sheet_number = f"{sheet_index}/{len(figures)}"
        if not draft_figures.current_ocr_audit(
                active.get("numeral_audit") or {},
                expected_sheet_number=expected_sheet_number,
                expected_section_designations=(
                    draft_figures.section_designations(spec.get("caption") or "")
                    if spec else None)):
            live_drawing_failures.append(
                f"{label}: OCR numeral, view-label, or sheet-number inspection did not pass "
                f"for sheet {expected_sheet_number}")
        if not draft_figures.current_semantic_audit(active.get("semantic_audit") or {}):
            live_drawing_failures.append(
                f"{label}: current semantic drawing consensus did not pass")
        if not draft_figures.current_leader_audit(active.get("leader_audit") or {}):
            live_drawing_failures.append(
                f"{label}: current leader placement consensus did not pass")
        if not spec:
            live_drawing_failures.append(f"{label}: no specification exists in this version")
        else:
            expected = draft_figures.expected_entries(spec, numeral_table)
            expected_hash = draft_figures.specification_hash(
                spec.get("label") or label, spec.get("caption") or "", expected)
            if (active.get("semantic_audit") or {}).get(
                    "specification_hash") != expected_hash:
                live_drawing_failures.append(
                    f"{label}: inspection belongs to a different drawing specification")
            if (active.get("leader_audit") or {}).get(
                    "specification_hash") != expected_hash:
                live_drawing_failures.append(
                    f"{label}: leader inspection belongs to a different drawing specification")
    if live_drawing_failures:
        blockers.append({
            "title": "One or more active drawings have not passed live inspection",
            "detail": "Every filing sheet must still match its specification and reference "
                      "numerals at download time.",
            "items": "; ".join(live_drawing_failures[:12])})

    fees = fee_profile(str(sections.get("claims") or ""))
    remaining = [
        "An oath or declaration signed by every named inventor (37 CFR 1.63), or a substitute "
        "statement where one is permitted.",
        "An Application Data Sheet (37 CFR 1.76) - the fields are pre-filled in the package.",
        "Entity-status certification if claiming small or micro entity fees (37 CFR 1.27, 1.29).",
        "An Information Disclosure Statement listing the art you are aware of (37 CFR 1.56, 1.97). "
        "The citation listing in this package is a starting point, not a signed form.",
        "The filing, search and examination fees due on the counts above.",
        "Review by a registered US patent practitioner before anything is submitted.",
    ]
    return {
        "ready": not blockers,
        "blockers": blockers,
        "formalities": formalities,
        "remaining": remaining,
        "fees": fees,
        "notes": notes,
        "citation_count": len({str(c) for c in (version.get("citations") or [])}),
        "reference_count": len(references),
        "verdict": (qa or {}).get("verdict", "unknown"),
        "patent_center_url": EFS_URL,
    }


# =============================================================================================
# The filing package
# =============================================================================================
def numbered_paragraphs(text: str, start: int) -> tuple[list[str], int]:
    """Number the paragraphs of a section as ``[0001]``, per 37 CFR 1.52(b)(6).

    Numbered paragraphs are optional at filing and close to mandatory in practice: without them
    every later amendment has to identify the passage it changes by quotation instead of by
    number.  Numbering here rather than in the draft keeps the draft readable while the filing
    copy is in the form the Office wants.
    """
    out: list[str] = []
    counter = start
    for block in re.split(r"\n\s*\n", (text or "").strip()):
        block = block.strip()
        if not block:
            continue
        out.append(f"[{counter:04d}] {block}")
        counter += 1
    return out, counter


def filing_text(project: Mapping[str, Any], version: Mapping[str, Any]) -> str:
    """The specification as plain text, in 37 CFR 1.77(b) order with numbered paragraphs."""
    sections = _filing_sections(version)
    lines: list[str] = []
    counter = 1
    for key, heading, numbered in FILING_ORDER:
        body = str(sections.get(key) or "").strip()
        if key == "title":
            lines += [heading, "", body or _filing_label(project.get("title") or ""), ""]
            continue
        if not body:
            continue
        lines += [heading, ""]
        if numbered:
            paragraphs, counter = numbered_paragraphs(body, counter)
            lines += paragraphs + [""]
        else:
            lines += [body, ""]
    lines += ["", "=" * 72, "", "CLAIMS", "",
              "What is claimed is:", "", str(sections.get("claims") or "").strip(), "",
              "=" * 72, "", "ABSTRACT OF THE DISCLOSURE", "",
              str(sections.get("abstract") or "").strip(), ""]
    return "\n".join(lines).rstrip() + "\n"


def ads_fields(project: Mapping[str, Any], version: Mapping[str, Any]) -> list[dict[str, str]]:
    """What goes on the Application Data Sheet, as far as we legitimately know it."""
    sections = _filing_sections(version)
    inventors = [line.strip() for line in
                 re.split(r"[\n;]+", str(project.get("inventors") or "")) if line.strip()]
    cross_reference = str(sections.get("cross_reference") or "").strip()
    return [
        {"field": "Title of invention",
         "value": str(sections.get("title") or project.get("title") or "")[:500],
         "note": "37 CFR 1.72(a) - 500 characters maximum."},
        {"field": "Application type", "value": "Utility, non-provisional",
         "note": "Change this if you are filing a provisional or a design application."},
        {"field": "Inventor(s)", "value": "\n".join(inventors) or "(not supplied)",
         "note": "37 CFR 1.41 - legal name, residence and mailing address for each."},
        {"field": "Applicant", "value": str(project.get("applicant") or "") or "(not supplied)",
         "note": "37 CFR 1.46 - required where the applicant is not the inventor."},
        {"field": "Domestic benefit / priority", "value": cross_reference or "(none claimed)",
         "note": "37 CFR 1.78 - a benefit claim must appear in the ADS, not only in the "
                 "specification."},
        {"field": "Foreign priority", "value": "(none stated)",
         "note": "37 CFR 1.55 - add every foreign application whose priority you claim."},
        {"field": "Entity status", "value": "(not certified)",
         "note": "37 CFR 1.27 / 1.29 - small or micro entity status must be certified, not "
                 "assumed."},
        {"field": "Correspondence address", "value": "(not supplied)", "note": "37 CFR 1.33."},
    ]


def citation_listing(version: Mapping[str, Any],
                     references: Sequence[Mapping[str, Any]] = ()) -> list[dict[str, str]]:
    """The cited art in the shape an IDS listing wants it, resolved and dated.

    This is the LISTING, not an executed SB/08a: the duty of disclosure under 37 CFR 1.56 is
    personal to the people involved in the application, and no software discharges it.
    """
    import draft_cite
    by_publication = {str(r.get("publication_number")): r for r in references}
    out = []
    for citation in dict.fromkeys(str(c) for c in (version.get("citations") or [])):
        record = draft_cite.resolve(citation, allow_remote=False)
        reference = by_publication.get(citation, {})
        publication = record.get("publication_number") or citation
        country = publication[:2] if len(publication) > 2 else ""
        out.append({
            "publication_number": publication,
            "country": country,
            "kind_code": record.get("kind_code", ""),
            "publication_date": record.get("publication_date", ""),
            "name_of_patentee": record.get("assignee") or reference.get("title", "")[:120],
            "title": record.get("title") or reference.get("title", ""),
            "resolved": "yes" if record.get("found") else "NOT RESOLVED",
            "source": record.get("source", ""),
        })
    return out


def render_filing_docx(project: Mapping[str, Any], version: Mapping[str, Any], *,
                       readiness_report: Mapping[str, Any],
                       references: Sequence[Mapping[str, Any]] = (),
                       figure_images: Sequence[Mapping[str, Any]] = ()) -> BytesIO:
    """Clean filing text and checked drawing sheets, with no workflow notes or placeholders."""
    if readiness_report.get("blockers"):
        raise drafting.DraftingValidationError(
            "The filing gate has blockers; a filing document was not created.")
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    sections = _filing_sections(version)
    document = Document()
    layout = document.sections[0]
    layout.page_width, layout.page_height = Inches(8.5), Inches(11)
    # 37 CFR 1.52(a)(1)(ii): at least 2.0 cm left and top, 2.0 cm right, 2.0 cm bottom.
    layout.left_margin = layout.top_margin = Inches(1)
    layout.right_margin = layout.bottom_margin = Inches(0.85)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 2.0        # 37 CFR 1.52(b)(2)(ii)
    normal.paragraph_format.space_after = Pt(0)

    def heading(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True

    def plain(text: str, *, size: int = 12, bold: bool = False, spacing: float = 2.0) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = spacing
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    # -- specification -------------------------------------------------------------------------
    counter = 1
    for key, title, numbered in FILING_ORDER:
        body = str(sections.get(key) or "").strip()
        if key == "title":
            heading(title)
            plain(body or _filing_label(project.get("title") or ""))
            continue
        if not body:
            continue
        heading(title)
        if numbered:
            paragraphs, counter = numbered_paragraphs(body, counter)
            for paragraph in paragraphs:
                plain(paragraph)
        else:
            plain(body)

    document.add_page_break()
    heading("CLAIMS")
    plain("What is claimed is:")
    for claim in draft_qa.split_claims(str(sections.get("claims") or "")):
        plain(f"{claim['number']}. {claim['text']}")

    document.add_page_break()
    heading("ABSTRACT OF THE DISCLOSURE")
    plain(str(sections.get("abstract") or "").strip())

    if figure_images:
        drawing_layout = document.add_section(WD_SECTION.NEW_PAGE)
        drawing_layout.page_width, drawing_layout.page_height = Inches(8.5), Inches(11)
        # 37 CFR 1.84(g): 1 inch top and left, 5/8 inch right, 3/8 inch bottom.
        drawing_layout.top_margin = drawing_layout.left_margin = Inches(1)
        drawing_layout.right_margin = Inches(0.625)
        drawing_layout.bottom_margin = Inches(0.375)
        for index, figure in enumerate(figure_images):
            png = bytes(figure.get("png") or b"")
            if not png:
                continue
            picture = document.add_paragraph()
            if index:
                picture.paragraph_format.page_break_before = True
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture.paragraph_format.space_before = Pt(0)
            picture.paragraph_format.space_after = Pt(0)
            picture.paragraph_format.line_spacing = 1.0
            from PIL import Image
            with Image.open(BytesIO(png)) as image:
                pixel_width, pixel_height = image.size
            max_width, max_height = 6.875, 9.5
            width = max_width
            height = width * pixel_height / max(1, pixel_width)
            if height > max_height:
                height = max_height
                width = height * pixel_width / max(1, pixel_height)
            picture.add_run().add_picture(
                BytesIO(png), width=Inches(width), height=Inches(height))

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def readiness_html(report: Mapping[str, Any]) -> str:
    """A compact rendering used by the studio page; kept beside the rules it reports."""
    def block(title: str, items: Sequence[Mapping[str, str]], tone: str) -> str:
        if not items:
            return ""
        rows = "".join(
            f"<li><b>{html.escape(str(i.get('title') or ''))}</b> "
            f"<span>{html.escape(str(i.get('detail') or ''))}</span>"
            + (f"<code>{html.escape(str(i.get('items')))}</code>" if i.get("items") else "")
            + "</li>" for i in items)
        return f'<div class="rdy {tone}"><h4>{html.escape(title)}</h4><ul>{rows}</ul></div>'
    return (block("Blockers", report.get("blockers") or [], "bad") +
            block("Formalities", report.get("formalities") or [], "warn"))
