"""Everything that goes to the Office, built as the Office wants it.

WHAT WAS HERE BEFORE.  One .docx holding the specification, the claims, the abstract and the
drawing images, and a page telling the user that an oath, an ADS, an entity certification and a fee
payment "remain".  A reviewer put it plainly: the application data sheet did not exist, the
declaration did not exist, the fee determination did not exist, there was no guidance on filing
format, and the specification was a manuscript rather than a filing document.  Every one of those
is a paper somebody then had to make by hand, at the point where they are least equipped to make
it, and three of the four were made wrong on the first attempt.

WHAT IS HERE NOW.  A package.  Nine files, each one either complete or explicit about the single
field a human still has to supply:

    00-READ-ME-FIRST.txt          what to upload, in what order, under which document description
    01-Specification.docx         37 CFR 1.52 throughout, and clean through the DOCX validator
    02-Drawings.pdf               37 CFR 1.84 sheets, one view each, numbered 1/N
    03-Application-Data-Sheet.pdf 37 CFR 1.76, every section heading, ready to sign
    03-Web-ADS-values.txt         the same data laid out as Patent Center's web form asks for it
    04-Declaration.pdf            37 CFR 1.63, every required statement, ready to sign
    05-Fee-worksheet.pdf          37 CFR 1.16, the counts and which surcharges they trigger
    06-Information-Disclosure.pdf 37 CFR 1.98 listing of the art this draft cites
    AUDIT.txt                     every check in filing_rules, run against the files above

TWO RULES THIS MODULE FOLLOWS.

  IT NEVER PRINTS A FEE AMOUNT.  The counts that drive the fees are computed exactly and the
  surcharges they trigger are named, with the fee code for each. The dollar figures move by
  rulemaking, Patent Center totals them at submission, and a number baked in here would be quietly
  wrong within a year in the one document a person would trust.

  IT NEVER INVENTS A PARTY DETAIL.  Where a field is not known, the package says which field, on
  the paper that needs it, in the words the ADS uses. An address this module made up would be
  filed.
"""
from __future__ import annotations

import io
import re
import zipfile
from datetime import date
from typing import Any, Mapping, Sequence

import draft_qa
import filing_profile
import filing_rules
import pdf_fonts

SPEC_NAME = "01-Specification.docx"
DRAWINGS_NAME = "02-Drawings.pdf"
ADS_NAME = "03-Application-Data-Sheet.pdf"
WEB_ADS_NAME = "03-Web-ADS-values.txt"
DECLARATION_NAME = "04-Declaration.pdf"
FEES_NAME = "05-Fee-worksheet.pdf"
IDS_NAME = "06-Information-Disclosure-Statement.pdf"
README_NAME = "00-READ-ME-FIRST.txt"
AUDIT_NAME = "AUDIT.txt"

#  37 CFR 1.77(b), with the heading each of our section keys is filed under.
FILING_ORDER = (
    ("title", "TITLE OF THE INVENTION", False),
    ("cross_reference", "CROSS-REFERENCE TO RELATED APPLICATIONS", True),
    ("government_support",
     "STATEMENT REGARDING FEDERALLY SPONSORED RESEARCH OR DEVELOPMENT", True),
    ("field", "FIELD OF THE INVENTION", True),
    ("background", "BACKGROUND OF THE INVENTION", True),
    ("summary", "BRIEF SUMMARY OF THE INVENTION", True),
    ("drawing_descriptions", "BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS", True),
    ("detailed_description", "DETAILED DESCRIPTION OF THE INVENTION", True),
)

#  37 CFR 1.16. No amounts: the code and what drives it, which is the part that does not move.
FEE_CODES = {
    "basic": ("1011 / 2011 / 3011", "Basic filing fee, utility"),
    "search": ("1111 / 2111 / 3111", "Utility search fee"),
    "examination": ("1311 / 2311 / 3311", "Utility examination fee"),
    "independent": ("1201 / 2201 / 3201", "Each independent claim over three, 37 CFR 1.16(h)"),
    "excess": ("1202 / 2202 / 3202", "Each claim over twenty, 37 CFR 1.16(i)"),
    "multiple": ("1203 / 2203 / 3203", "Multiple dependent claim, 37 CFR 1.16(j)"),
    "size": ("1081 / 2081 / 3081", "Application size fee, each 50 sheets over 100, 37 CFR 1.16(s)"),
    "nondocx": ("1082", "Non-DOCX filing surcharge, 37 CFR 1.16(u)"),
}


# =============================================================================================
# Text preparation
# =============================================================================================
def _sections(version: Mapping[str, Any]) -> dict[str, str]:
    """The stored sections, with drafting-only citation keys made filing-readable and the
    typography folded down to what a DOCX validator accepts."""
    import draft_cite
    import drafting
    out = {}
    for key, value in dict(version.get("sections") or {}).items():
        try:
            text = draft_cite.filing_citations(str(value or ""))
        except ValueError as exc:
            raise drafting.DraftingValidationError(str(exc)) from exc
        out[str(key)] = filing_rules.to_filing_text(text)
    return out


def numbered_paragraphs(text: str, start: int) -> tuple[list[str], int]:
    """37 CFR 1.52(b)(6): four numerals in square brackets, at the left margin, then a gap."""
    out: list[str] = []
    counter = start
    for block in re.split(r"\n\s*\n", (text or "").strip()):
        block = " ".join(line.strip() for line in block.strip().splitlines()).strip()
        if not block:
            continue
        out.append(f"[{counter:04d}]|{block}")
        counter += 1
    return out, counter


# =============================================================================================
# 01 - the specification
# =============================================================================================
def specification_docx(project: Mapping[str, Any], version: Mapping[str, Any]) -> bytes:
    """The specification as a filing document rather than a manuscript."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    sections = _sections(version)
    document = Document()
    layout = document.sections[0]
    layout.page_width, layout.page_height = Inches(filing_rules.PAGE_WIDTH_IN), \
        Inches(filing_rules.PAGE_HEIGHT_IN)
    layout.left_margin = layout.right_margin = Inches(filing_rules.SPEC_MARGIN_IN)
    layout.top_margin = layout.bottom_margin = Inches(filing_rules.SPEC_MARGIN_IN)

    normal = document.styles["Normal"]
    normal.font.name = filing_rules.SPEC_FONT
    normal.font.size = Pt(filing_rules.SPEC_FONT_PT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), filing_rules.SPEC_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), filing_rules.SPEC_FONT)
    normal.paragraph_format.line_spacing = filing_rules.SPEC_LINE_SPACING
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    _page_number_footer(document, WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, OxmlElement, qn)

    def heading(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.keep_with_next = True
        #  Plain uppercase, not bold. 37 CFR 1.77 asks for the headings; the Office's own
        #  published examples set them in plain capitals, and a bold heading is one more thing
        #  the DOCX converter has to carry.
        paragraph.add_run(text)

    def numbered(marker: str, body: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.5))
        run = paragraph.add_run(marker)
        run.bold = True
        paragraph.add_run("\t" + body)

    def plain(body: str, *, first_line_indent: float = 0.0) -> None:
        paragraph = document.add_paragraph()
        if first_line_indent:
            paragraph.paragraph_format.first_line_indent = Inches(first_line_indent)
        paragraph.add_run(body)

    # -- title -----------------------------------------------------------------------------------
    title = (sections.get("title") or
             filing_rules.to_filing_text(project.get("title") or "")).strip()
    heading("TITLE OF THE INVENTION")
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(title.upper() if title.isupper() else title)

    counter = 1
    for key, label, number_them in FILING_ORDER:
        if key == "title":
            continue
        body = (sections.get(key) or "").strip()
        if not body:
            continue
        heading(label)
        if number_them:
            paragraphs, counter = numbered_paragraphs(body, counter)
            for entry in paragraphs:
                marker, text = entry.split("|", 1)
                numbered(marker, text)
        else:
            plain(body)

    # -- claims, on their own page: 37 CFR 1.75(h) -----------------------------------------------
    document.add_page_break()
    heading("CLAIMS")
    plain("What is claimed is:")
    for claim in draft_qa.split_claims(sections.get("claims") or ""):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
        paragraph.add_run(f"{claim['number']}. {' '.join(str(claim['text']).split())}")

    # -- abstract, on its own sheet: 37 CFR 1.72(b) ----------------------------------------------
    document.add_page_break()
    heading("ABSTRACT OF THE DISCLOSURE")
    plain(" ".join((sections.get("abstract") or "").split()), first_line_indent=0.5)

    document.core_properties.title = title[:255]
    document.core_properties.subject = "US utility patent application"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = ""

    output = io.BytesIO()
    document.save(output)
    return _strip_unwanted_parts(output.getvalue())


def _page_number_footer(document: Any, align: Any, tab_align: Any, element_factory: Any,
                        qn: Any) -> None:
    """37 CFR 1.52(b)(5): consecutive page numbers, centred below the text."""
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = align.CENTER
    run = paragraph.add_run()
    begin = element_factory("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = element_factory("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = element_factory("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)
    del tab_align


_UNWANTED_PARTS = ("docProps/thumbnail.jpeg", "word/stylesWithEffects.xml")


def _strip_unwanted_parts(blob: bytes) -> bytes:
    """Take out the parts the Patent Center DOCX validator warns about.

    python-docx builds every document from one template, and that template carries a preview
    image, a legacy Word 2010 style part and a customXml item nobody put there. Word leaves an
    empty comments part behind as well. None of them is text and all of them are one more thing
    the validator has to accept on a document that is otherwise perfect.
    """
    source = zipfile.ZipFile(io.BytesIO(blob))
    drop = set()
    for name in source.namelist():
        if name in _UNWANTED_PARTS or name.startswith("customXml/"):
            drop.add(name)
        elif re.fullmatch(r"word/comments\d*\.xml", name):
            if not re.search(r"<w:comment\b", source.read(name).decode("utf-8", "replace")):
                drop.add(name)
    if not drop:
        return blob
    extensions = {name.rsplit(".", 1)[-1].lower() for name in drop if "." in name}
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            if item.filename in drop:
                continue
            data = source.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = _prune_content_types(data, drop, extensions)
            elif item.filename.endswith(".rels"):
                data = _prune_relationships(data, drop, item.filename)
            target.writestr(item, data)
    return output.getvalue()


def _prune_content_types(data: bytes, drop: set[str], extensions: set[str]) -> bytes:
    text = data.decode("utf-8")
    for name in drop:
        text = re.sub(r'<Override[^>]*PartName="/' + re.escape(name) + r'"[^>]*/>', "", text)
    #  A Default entry for an extension nothing uses any more is harmless and is left alone: it
    #  declares a content type, not a part, and removing one that something else still needs would
    #  break the package.
    del extensions
    return text.encode("utf-8")


def _prune_relationships(data: bytes, drop: set[str], rels_name: str) -> bytes:
    """Drop every relationship whose target is a part we removed.

    A target is relative to the directory of the part that owns the .rels file, and it can climb:
    word/_rels/document.xml.rels points at ``../customXml/item1.xml``. Matching on the literal
    string is how the first attempt left a dangling relationship and made the package unreadable,
    so the target is resolved to a package path before it is compared.
    """
    import posixpath
    base = posixpath.dirname(rels_name.replace("/_rels/", "/").lstrip("/")) \
        if "/_rels/" in rels_name else ""
    text = data.decode("utf-8")

    def resolve(target: str) -> str:
        if target.startswith("/"):
            return target.lstrip("/")
        return posixpath.normpath(posixpath.join(base, target)).lstrip("./")

    def keep(match: re.Match[str]) -> str:
        target = re.search(r'Target="([^"]+)"', match.group(0))
        if target and resolve(target.group(1)) in drop:
            return ""
        return match.group(0)

    return re.sub(r"<Relationship\b[^>]*/>", keep, text).encode("utf-8")


# =============================================================================================
# 02 - the drawings
# =============================================================================================
def drawing_sheets(figures: Sequence[Mapping[str, Any]],
                   facts_by_sha: Mapping[str, Mapping[str, Any]] | None = None
                   ) -> list[dict[str, Any]]:
    """One entry per filing sheet, splitting an upload that carries more than one view.

    A composite carrying FIG. 2 and FIG. 3 side by side is two sheets, not one: 37 CFR 1.84(u)
    numbers views independently of sheets, and one view to a sheet is also the only way to get the
    reference characters above the 0.32 cm floor when the source artwork is 1,400 pixels wide.
    Where the inventory pass gave a bounding box for each view, the crop is exact. Where it did
    not, the sheet is left whole and the audit says so.
    """
    import figure_facts
    from PIL import Image
    out: list[dict[str, Any]] = []
    for figure in figures:
        png = bytes(figure.get("png") or b"")
        if not png:
            continue
        label = str(figure.get("label") or "")
        facts = dict((facts_by_sha or {}).get(figure_facts.sheet_key(png)) or {})
        #  An unnumbered picture is a view too, and it travels as its own sheet so the audit can
        #  say it has no number. Dropping it here would take a defect off the page by hiding it.
        views = [dict(view) for view in facts.get("views") or []]
        views += [{"legend": "", "bbox": item.get("bbox") or [0, 0, 1, 1],
                   "character_height": item.get("character_height") or 0.0}
                  for item in facts.get("unlabelled_views") or []]
        if len(views) <= 1:
            single = views[0] if views else {}
            out.append({"label": (single.get("legend") if views else "") or label, "png": png,
                        "facts": facts, "cropped": False,
                        "character_height": single.get("character_height") or 0.0})
            continue
        with Image.open(io.BytesIO(png)) as image:
            width, height = image.size
            ink = _ink_profiles(image)
            boxes = []
            for view in views:
                others = [item.get("bbox") or [0, 0, 1, 1] for item in views if item is not view]
                box = _snap_box(view.get("bbox") or [0, 0, 1, 1], others,
                                ink=ink, width=width, height=height)
                box = _trim_rules(image, box)
                box = _drop_edge_debris(image, box)
                boxes.append(box)
            if not _cuts_are_clean(image, boxes):
                #  REFUSE THE SPLIT rather than file a mutilated view. Where the views were laid
                #  out with no white corridor between them, every cut line runs through ink: the
                #  first build of this cropped the leading 4 off a numeral 48 and put "18" on a
                #  filing sheet, which is a worse defect than the one the split was fixing. The
                #  sheet goes over whole and the audit says what has to be re-supplied.
                out.append({"label": label, "png": png, "facts": facts, "cropped": False,
                            "split_refused": True,
                            "character_height":
                                facts.get("smallest_reference_character_height_fraction") or 0.0})
                continue
            for view, (x0, y0, x1, y1) in zip(views, boxes):
                if x1 - x0 < 80 or y1 - y0 < 80:
                    continue
                buffer = io.BytesIO()
                image.crop((x0, y0, x1, y1)).save(buffer, format="PNG")
                #  A view with no legend keeps no legend. Borrowing the upload's label here
                #  would hide the very defect the audit exists to report.
                out.append({"label": str(view.get("legend") or ""), "png": buffer.getvalue(),
                            "facts": facts, "cropped": True,
                            "character_height": view.get("character_height") or 0.0,
                            "source_fraction": ((x1 - x0) / width, (y1 - y0) / height)})
    return _in_figure_order(out)


def _cuts_are_clean(image: Any, boxes: Sequence[tuple[int, int, int, int]]) -> bool:
    """Does every proposed crop line fall on white?

    A cut through ink takes a lead line, half a numeral or part of the drawing with it. There is
    no way to repair that afterwards and no way to see it in a thumbnail, so the split only
    proceeds when every one of its four borders per view is clear.
    """
    import numpy
    grey = numpy.asarray(image.convert("L")) < 200
    height, width = grey.shape
    for x0, y0, x1, y1 in boxes:
        if x1 - x0 < 80 or y1 - y0 < 80:
            continue
        for column in (x0, x1 - 1):
            if 0 < column < width - 1 and grey[y0:y1, column].any():
                return False
        for row in (y0, y1 - 1):
            if 0 < row < height - 1 and grey[row, x0:x1].any():
                return False
    return True


def _ink_profiles(image: Any) -> tuple[Any, Any]:
    """Which columns and which rows of the artwork carry any ink at all."""
    import numpy
    dark = numpy.asarray(image.convert("L")) < 200
    return dark.any(axis=0), dark.any(axis=1)


def _snap_box(box: Sequence[float], others: Sequence[Sequence[float]], *, ink: tuple[Any, Any],
              width: int, height: int) -> tuple[int, int, int, int]:
    """Grow a view's bounding box out to the white gutter that actually separates the views.

    A box a model draws round a view is tight on the drawing and clips what sits outside it: a
    lead line running off to the left, the numerals for the substrate layers, half of a section
    arrow. Cropping on that box loses reference characters, which is a real defect introduced by
    the tool rather than found by it. So the box is only a starting point. Each edge walks
    outward through ink until it reaches a run of empty rows or columns wide enough to be the
    gutter between two views, and stops there. It never crosses into another view's own box.
    """
    columns, rows = ink
    gutter_x = max(6, int(width * 0.012))
    gutter_y = max(6, int(height * 0.012))

    #  A box may grow only HALF WAY to its neighbour. Growing right up to the neighbour's own
    #  edge is what put the left half of one arrangement, and the word OR between them, onto the
    #  sheet belonging to the other: the model's boxes are tight on the drawings and the space
    #  between two of them belongs to neither. Meeting in the middle cannot take a neighbour's
    #  linework, and the white it keeps is trimmed off again below.
    def limit_low(index: int, axis: int) -> int:
        span = width if axis == 0 else height
        edges = [int(other[axis + 2] * span) for other in others
                 if other[axis + 2] <= box[axis] + 1e-6]
        nearest = max([0] + [edge for edge in edges if edge <= index])
        return (index + nearest) // 2 if nearest > 0 else 0

    def limit_high(index: int, axis: int) -> int:
        span = width if axis == 0 else height
        edges = [int(other[axis] * span) for other in others if other[axis] >= box[axis + 2] - 1e-6]
        nearest = min([span] + [edge for edge in edges if edge >= index])
        return (index + nearest + 1) // 2 if nearest < span else span

    def walk_down(profile: Any, start: int, floor: int, gutter: int) -> int:
        position, empty = start, 0
        while position > floor:
            position -= 1
            empty = 0 if profile[position] else empty + 1
            if empty >= gutter:
                return position + empty
        return floor

    def walk_up(profile: Any, start: int, ceiling: int, gutter: int) -> int:
        position, empty = start, 0
        while position < ceiling - 1:
            position += 1
            empty = 0 if profile[position - 1] else empty + 1
            if empty >= gutter:
                return position - empty
        return ceiling

    floor_x, ceiling_x = limit_low(int(box[0] * width), 0), limit_high(int(box[2] * width), 0)
    floor_y, ceiling_y = limit_low(int(box[1] * height), 1), limit_high(int(box[3] * height), 1)
    x0 = walk_down(columns, max(0, int(box[0] * width)), floor_x, gutter_x)
    x1 = walk_up(columns, min(width, int(box[2] * width)), ceiling_x, gutter_x)
    y0 = walk_down(rows, max(0, int(box[1] * height)), floor_y, gutter_y)
    y1 = walk_up(rows, min(height, int(box[3] * height)), ceiling_y, gutter_y)
    #  Then trim any all-white band left at the edges, so the view is placed on its own sheet at
    #  the largest scale the sight allows.
    while x0 < x1 and not columns[x0]:
        x0 += 1
    while x1 > x0 and not columns[x1 - 1]:
        x1 -= 1
    while y0 < y1 and not rows[y0]:
        y0 += 1
    while y1 > y0 and not rows[y1 - 1]:
        y1 -= 1
    #  The breathing space never crosses a limit either: padding past the half-way line is how a
    #  clamp that was computed correctly ended up six pixels inside the neighbour's drawing.
    pad_x, pad_y = max(4, gutter_x // 3), max(4, gutter_y // 3)
    return (max(floor_x, x0 - pad_x), max(floor_y, y0 - pad_y),
            min(ceiling_x, x1 + pad_x), min(ceiling_y, y1 + pad_y))


def _trim_rules(image: Any, box: tuple[int, int, int, int]) -> tuple[int, int, int, int]:
    """Drop a divider rule the composite drew between two views.

    A line spanning nearly the full width of the crop, at its very top or bottom, is the rule the
    draftsperson used to separate FIG. 2 from FIG. 3. It is not part of either view, and carrying
    it onto a filing sheet puts a stray line on the drawing.
    """
    import numpy
    x0, y0, x1, y1 = box
    dark = numpy.asarray(image.convert("L").crop((x0, y0, x1, y1))) < 200
    if dark.size == 0:
        return box
    coverage = dark.mean(axis=1)
    band = max(2, int((y1 - y0) * 0.03))
    top, bottom = y0, y1
    for offset in range(min(band, len(coverage))):
        if coverage[offset] > 0.9:
            top = y0 + offset + 1
    for offset in range(min(band, len(coverage))):
        if coverage[len(coverage) - 1 - offset] > 0.9:
            bottom = y1 - offset - 1
    if bottom - top < 80:
        return box
    return (x0, top, x1, bottom)


def _drop_edge_debris(image: Any, box: tuple[int, int, int, int]) -> tuple[int, int, int, int]:
    """Cut off the neighbour's leftovers at the edges of a crop.

    Splitting a composite leaves scraps: the word OR that sat between two arrangements, the ". 3"
    tail of a figure caption cut down the middle, the far edge of the enlarged-detail circle. Each
    one is a small island of ink at an edge with a wide white gutter between it and the drawing,
    and each one goes onto a filing sheet as a stray mark that belongs to no view.

    An independent reviewer found all three on the first package this built, which is the reason
    this exists: a tool that introduces a drawing defect while fixing another one has not helped.
    A real view is nearly all of the ink in its own crop, so an edge band holding under a fortieth
    of it and separated by a gutter is not part of it.
    """
    import numpy
    x0, y0, x1, y1 = box
    dark = numpy.asarray(image.convert("L").crop((x0, y0, x1, y1))) < 200
    if dark.size == 0 or not dark.any():
        return box
    total = int(dark.sum())
    columns, rows = dark.sum(axis=0), dark.sum(axis=1)
    gutter_x = max(8, int((x1 - x0) * 0.03))
    gutter_y = max(8, int((y1 - y0) * 0.03))

    def trim(profile: Any, gutter: int) -> tuple[int, int]:
        """How many pixels to cut from the low and the high end."""
        low = high = 0
        length = len(profile)
        run = carried = 0
        for index in range(length):
            if profile[index]:
                if run >= gutter and carried and carried < total * 0.025:
                    low = index - run
                    break
                if run >= gutter:
                    break
                carried += int(profile[index])
                run = 0
            else:
                run += 1
        run = carried = 0
        for index in range(length - 1, -1, -1):
            if profile[index]:
                if run >= gutter and carried and carried < total * 0.025:
                    high = length - 1 - index - run
                    break
                if run >= gutter:
                    break
                carried += int(profile[index])
                run = 0
            else:
                run += 1
        return low, high

    left, right = trim(columns, gutter_x)
    top, bottom = trim(rows, gutter_y)
    trimmed = (x0 + left, y0 + top, x1 - right, y1 - bottom)
    if trimmed[2] - trimmed[0] < 80 or trimmed[3] - trimmed[1] < 80:
        return box
    return trimmed


def _in_figure_order(sheets: Sequence[Mapping[str, Any]]) -> list[dict[str, Any]]:
    def key(sheet: Mapping[str, Any]) -> tuple[int, int, str]:
        match = re.search(r"(\d{1,3})\s*([A-Za-z]?)", str(sheet.get("label") or ""))
        if not match:
            return (1, 0, str(sheet.get("label") or ""))
        return (0, int(match.group(1)), match.group(2).upper())
    return [dict(sheet) for sheet in sorted(sheets, key=key)]


def drawings_pdf(sheets: Sequence[Mapping[str, Any]]) -> tuple[bytes, list[dict[str, Any]]]:
    """37 CFR 1.84 sheets: Letter, the rule's margins, one view each, numbered N/M.

    Returns the PDF and, for each sheet, the height the smallest reference character actually
    renders at, because that is the number 1.84(p)(3) is about and it is only knowable once the
    artwork has been placed on a page.
    """
    import figure_facts
    from PIL import Image
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas as pdfcanvas

    pdf_fonts.ready()
    output = io.BytesIO()
    page = pdfcanvas.Canvas(output, pagesize=letter)
    total = len(sheets)
    measurements: list[dict[str, Any]] = []

    sight_left = filing_rules.SHEET_MARGIN_LEFT_IN * inch
    sight_right = (filing_rules.PAGE_WIDTH_IN - filing_rules.SHEET_MARGIN_RIGHT_IN) * inch
    sight_top = (filing_rules.PAGE_HEIGHT_IN - filing_rules.SHEET_MARGIN_TOP_IN) * inch
    sight_bottom = filing_rules.SHEET_MARGIN_BOTTOM_IN * inch
    sight_width = sight_right - sight_left
    #  The sheet number sits inside the sight, below the top margin: 1.84(t) puts it at the middle
    #  of the top of the sheet and NOT in the margin.
    number_band = 0.32 * inch
    artwork_top = sight_top - number_band
    artwork_height = artwork_top - sight_bottom

    for index, sheet in enumerate(sheets, 1):
        png = bytes(sheet.get("png") or b"")
        with Image.open(io.BytesIO(png)) as image:
            pixel_width, pixel_height = image.size
        scale = min(sight_width / max(1, pixel_width), artwork_height / max(1, pixel_height))
        draw_width = pixel_width * scale
        draw_height = pixel_height * scale
        x = sight_left + (sight_width - draw_width) / 2.0
        y = sight_bottom + (artwork_height - draw_height) / 2.0

        page.drawImage(ImageReader(io.BytesIO(png)), x, y, width=draw_width, height=draw_height,
                       preserveAspectRatio=True, anchor="c", mask=None)
        page.setFont(pdf_fonts.font(pdf_fonts.SANS), filing_rules.SHEET_NUMBER_PT)
        page.drawCentredString(filing_rules.PAGE_WIDTH_IN * inch / 2.0,
                               sight_top - 0.22 * inch, f"{index}/{total}")
        page.showPage()

        #  Measured on the sheet as it will be filed, in pixels, then multiplied by the scale it
        #  is placed at. Nothing is estimated and nothing is converted between crops.
        measure = figure_facts.measure_character_height(png)
        points = float(measure["pixels"]) * scale
        measurements.append({
            "label": str(sheet.get("label") or ""),
            "sheet_number": f"{index}/{total}",
            "character_pixels": round(float(measure["pixels"]), 1),
            "samples": int(measure["samples"]),
            "character_points": points,
            "character_cm": points / 72.0 * 2.54,
            "measured": bool(measure["pixels"]),
        })
    page.save()
    return filing_rules.redistill(output.getvalue()), measurements


# =============================================================================================
# The papers, rendered
# =============================================================================================
def _paper(title: str, blocks: Sequence[Any]) -> bytes:
    """One rendering routine for every form, on embedded faces with no space in a font name."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (KeepTogether, PageBreak, Paragraph, SimpleDocTemplate,
                                    Spacer, Table, TableStyle)
    from reportlab.lib import colors

    pdf_fonts.ready()
    serif = pdf_fonts.font(pdf_fonts.SERIF)
    serif_bold = pdf_fonts.font(pdf_fonts.SERIF_BOLD)
    body = ParagraphStyle("Body", fontName=serif, fontSize=10.5, leading=14, spaceAfter=6)
    small = ParagraphStyle("Small", parent=body, fontSize=9, leading=12)
    head = ParagraphStyle("Head", fontName=serif_bold, fontSize=11.5, leading=15,
                          spaceBefore=12, spaceAfter=6, keepWithNext=True)
    page_title = ParagraphStyle("Title", parent=head, fontSize=14, leading=18, alignment=1,
                                spaceBefore=0, spaceAfter=14)

    story: list[Any] = [Paragraph(_x(title), page_title)]
    for block in blocks:
        kind = block[0]
        if kind == "head":
            story.append(Paragraph(_x(block[1]), head))
        elif kind == "p":
            story.append(Paragraph(_x(block[1]), body))
        elif kind == "small":
            story.append(Paragraph(_x(block[1]), small))
        elif kind == "gap":
            story.append(Spacer(1, float(block[1]) * inch))
        elif kind == "break":
            story.append(PageBreak())
        elif kind == "rows":
            data = [[Paragraph(_x(str(cell)), small) for cell in row] for row in block[1]]
            widths = block[2] if len(block) > 2 else None
            table = Table(data, colWidths=widths, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#888888")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eeeeee")),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
            story.append(table)
            story.append(Spacer(1, 8))
        elif kind == "sign":
            story.append(KeepTogether([
                Spacer(1, 18),
                Paragraph(_x(block[1]), head),
                Spacer(1, 6),
                Table([[Paragraph(_x(label), small), Paragraph("&nbsp;", small)]
                       for label in block[2]],
                      colWidths=[1.7 * inch, 4.6 * inch], hAlign="LEFT",
                      style=TableStyle([
                          ("LINEBELOW", (1, 0), (1, -1), 0.6, colors.black),
                          ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                          ("BOTTOMPADDING", (0, 0), (-1, -1), 10)]))]))

    output = io.BytesIO()
    document = SimpleDocTemplate(
        output, pagesize=letter, leftMargin=inch, rightMargin=inch,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch, title=title,
        subject="US utility patent application")

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont(pdf_fonts.font(pdf_fonts.SANS), 8)
        canvas.drawCentredString(letter[0] / 2.0, 0.5 * inch, f"Page {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return filing_rules.redistill(output.getvalue())


def _x(value: Any) -> str:
    text = filing_rules.to_filing_text(value)
    return (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            .replace("\n", "<br/>"))


def application_data_sheet_pdf(project: Mapping[str, Any], version: Mapping[str, Any],
                               profile: Mapping[str, Any], *, drawing_sheets_count: int) -> bytes:
    """37 CFR 1.76: titled, carrying every section heading, and signable."""
    sections = _sections(version)
    title = (sections.get("title") or str(project.get("title") or "")).strip()
    inventors = list(profile.get("inventors") or [])
    blocks: list[Any] = [
        ("small", "37 CFR 1.76. A benefit or priority claim is only effective if it appears in "
                  "this sheet. Putting it in the specification alone does not make it."),
        ("head", "Inventor Information"),
    ]
    rows = [["#", "Legal name", "Residence", "Mailing address"]]
    for index, row in enumerate(inventors, 1):
        rows.append([str(index),
                     filing_profile.full_name(row) or "(REQUIRED: legal name)",
                     filing_profile.residence(row) or "(REQUIRED: city and country)",
                     filing_profile.mailing_address(row) or "(REQUIRED: mailing address)"])
    if len(rows) == 1:
        rows.append(["1", "(REQUIRED)", "(REQUIRED)", "(REQUIRED)"])
    blocks.append(("rows", rows, [24, 130, 130, 180]))

    correspondence = filing_profile.correspondence_block(profile)
    blocks += [
        ("head", "Correspondence Information"),
        ("p", "\n".join(correspondence) if correspondence else
              "(REQUIRED: the address the Office writes to about this application, 37 CFR "
              "1.33(a))"),
    ]
    if profile.get("customer_number"):
        blocks.append(("p", f"Customer number: {profile['customer_number']}"))

    blocks += [
        ("head", "Application Information"),
        ("rows", [["Field", "Value"],
                  ["Title of the invention", title or "(REQUIRED)"],
                  ["Application type", filing_profile.application_type_label(profile)],
                  ["Subject matter", "Utility"],
                  ["Total number of drawing sheets", str(int(drawing_sheets_count))],
                  ["Suggested figure for publication", "FIG. 1"],
                  ["Attorney docket number", str(profile.get("docket_number") or "(none)")],
                  ["Entity status (37 CFR 1.27, 1.29)",
                   filing_profile.entity_label(profile)]], [180, 284]),
        ("head", "Representative Information"),
        ("p", "None, unless a registered practitioner is appointed by a power of attorney filed "
              "with or after this application."),
        ("head", "Domestic Benefit/National Stage Information"),
        ("p", str(profile.get("domestic_benefit") or "No domestic benefit is claimed under 37 "
                                                     "CFR 1.78.")),
        ("head", "Foreign Priority Information"),
        ("p", str(profile.get("foreign_priority") or "No foreign priority is claimed under 37 "
                                                     "CFR 1.55.")),
        ("head", "Applicant Information"),
        ("p", (f"The applicant is {profile.get('applicant_name')}, which is not the inventor "
               "(37 CFR 1.46)."
               if profile.get("applicant_kind") == "juristic" and profile.get("applicant_name")
               else "The inventor or inventors named above are the applicant.")),
        ("head", "Assignee Information"),
        ("p", ("\n".join(part for part in (str(profile.get("assignee_name") or ""),
                                           str(profile.get("assignee_address") or ""))
                         if part.strip())
               or "No assignee is to be printed on the patent (37 CFR 3.81(a)).")),
        ("sign", "Signature (37 CFR 1.76(e), 1.33(b))",
         ["S-signature  /name/", "Printed name", "Date", "Registration number, if any"]),
        ("small", "An S-signature is the signer's own name typed between forward slashes, for "
                  "example /Jane A. Smith/."),
    ]
    return _paper("APPLICATION DATA SHEET", blocks)


def declaration_pdf(project: Mapping[str, Any], version: Mapping[str, Any],
                    profile: Mapping[str, Any]) -> bytes:
    """37 CFR 1.63, carrying every statement the rule requires and one signature block each.

    The statement that goes missing is 1.63(a)(4), "the application was made or was authorized to
    be made by me". A declaration without it is taken at upload and comes back as a notice months
    later, which is the worst possible place to find out.
    """
    sections = _sections(version)
    title = (sections.get("title") or str(project.get("title") or "")).strip()
    inventors = list(profile.get("inventors") or []) or [{}]
    blocks: list[Any] = [
        ("small", "Declaration under 37 CFR 1.63 for a utility application filed under 35 U.S.C. "
                  "111(a). One declaration is signed by each named inventor."),
        ("head", "This declaration is directed to"),
        ("p", f"The attached application, entitled: {title or '(REQUIRED: title)'}"),
        ("p", "United States application number ____________________, filed on "
              "____________________, where this declaration is filed after the application. "
              "Leave both blank when it is filed with the application."),
        ("head", "Statements"),
        ("p", "1. I believe that I am the original inventor or an original joint inventor of a "
              "claimed invention in the above-identified application. (37 CFR 1.63(a)(3))"),
        ("p", "2. The above-identified application was made or was authorized to be made by me. "
              "(37 CFR 1.63(a)(4))"),
        ("p", "3. I have reviewed and understand the contents of the above-identified "
              "application, including the claims. (37 CFR 1.63(c))"),
        ("p", "4. I acknowledge the duty to disclose to the United States Patent and Trademark "
              "Office all information known to me to be material to patentability as defined in "
              "37 CFR 1.56. (37 CFR 1.63(c))"),
        ("head", "Warning"),
        ("p", "I hereby acknowledge that any willful false statement made in this declaration is "
              "punishable under 18 U.S.C. 1001 by fine or imprisonment of not more than five (5) "
              "years, or both."),
        ("small", "Petitioner or applicant is cautioned to avoid submitting personal information "
                  "in documents filed in a patent application that may contribute to identity "
                  "theft."),
    ]
    for index, row in enumerate(inventors, 1):
        blocks += [
            ("break",) if index > 1 else ("gap", 0.1),
            ("head", f"Legal name of the inventor ({index} of {len(inventors)})"),
            ("rows", [["Field", "Value"],
                      ["Given name", str(row.get("given_name") or "(REQUIRED)")],
                      ["Middle name", str(row.get("middle_name") or "")],
                      ["Family name", str(row.get("family_name") or "(REQUIRED)")],
                      ["Residence", filing_profile.residence(row) or "(REQUIRED)"],
                      ["Mailing address",
                       filing_profile.mailing_address(row) or "(REQUIRED)"]], [140, 324]),
            ("sign", "Inventor's signature",
             ["S-signature  /name/", "Printed name", "Date"]),
        ]
    return _paper("DECLARATION FOR UTILITY OR DESIGN PATENT APPLICATION", blocks)


def fee_profile(claims_text: str, *, specification_sheets: int = 0,
                drawing_sheets_count: int = 0) -> dict[str, Any]:
    """The counts that decide the fees, and which surcharges each one triggers."""
    #  One reading of the claim set, shared with the review and with the page, so a claim the
    #  worksheet bills as independent is the one the page marks independent.
    counted = draft_qa.claim_map(claims_text or "")
    independent = counted["independent"]
    multiple = counted["multiple_dependent"]
    total = counted["total"]
    billable = counted["billable"]
    sheets = int(specification_sheets) + int(drawing_sheets_count)
    size_blocks = 0
    if sheets > filing_rules.SIZE_FEE_FREE_SHEETS:
        over = sheets - filing_rules.SIZE_FEE_FREE_SHEETS
        size_blocks = -(-over // filing_rules.SIZE_FEE_BLOCK_SHEETS)
    triggered = []

    def trigger(key: str, quantity: int) -> None:
        code, what = FEE_CODES[key]
        triggered.append({"key": key, "code": code, "what": what, "quantity": int(quantity)})

    if independent > filing_rules.INDEPENDENT_CLAIMS_INCLUDED:
        trigger("independent", independent - filing_rules.INDEPENDENT_CLAIMS_INCLUDED)
    if billable > filing_rules.CLAIMS_INCLUDED:
        trigger("excess", billable - filing_rules.CLAIMS_INCLUDED)
    if multiple:
        trigger("multiple", multiple)
    if size_blocks:
        trigger("size", size_blocks)
    return {"total": total, "independent": independent, "dependent": total - independent,
            "multiple_dependent": multiple, "billable": billable,
            "specification_sheets": int(specification_sheets),
            "drawing_sheets": int(drawing_sheets_count), "sheets": sheets,
            "size_blocks": size_blocks, "triggered": triggered}


def fee_worksheet_pdf(profile: Mapping[str, Any], fees: Mapping[str, Any]) -> bytes:
    """37 CFR 1.16: what is due and why, with the fee code for each line and no dollar figures.

    NO AMOUNTS ARE PRINTED, on purpose. The schedule moves by rulemaking, Patent Center totals
    the fees at submission from these same counts, and a number baked into a worksheet is quietly
    wrong within a year in the one document somebody would trust.
    """
    rows = [["Fee code (undiscounted / small / micro)", "What it is", "Quantity"],
            [FEE_CODES["basic"][0], FEE_CODES["basic"][1], "1"],
            [FEE_CODES["search"][0], FEE_CODES["search"][1], "1"],
            [FEE_CODES["examination"][0], FEE_CODES["examination"][1], "1"]]
    for item in fees.get("triggered") or []:
        rows.append([item["code"], item["what"], str(item["quantity"])])
    if not (fees.get("triggered") or []):
        rows.append(["-", "No surcharge is triggered by these counts.", "-"])
    counts = [["Count", "Value", "What it buys"],
              ["Total claims", str(fees.get("total")),
               f"{filing_rules.CLAIMS_INCLUDED} are included in the basic fee"],
              ["Independent claims", str(fees.get("independent")),
               f"{filing_rules.INDEPENDENT_CLAIMS_INCLUDED} are included"],
              ["Multiple dependent claims", str(fees.get("multiple_dependent")),
               "Each is counted as the number of claims it refers to (37 CFR 1.75(c))"],
              ["Claims billed", str(fees.get("billable")),
               "Total claims plus the extra claims a multiple dependent claim counts as"],
              ["Specification sheets", str(fees.get("specification_sheets")), ""],
              ["Drawing sheets", str(fees.get("drawing_sheets")), ""],
              ["Sheets in total", str(fees.get("sheets")),
               f"The application size fee starts above {filing_rules.SIZE_FEE_FREE_SHEETS}"]]
    blocks: list[Any] = [
        ("small", "Counts computed from the claim set in this package. Patent Center totals the "
                  "fees at submission from the same counts, and the schedule moves by rulemaking, "
                  "so this worksheet names the fee codes and prints no amounts."),
        ("head", "The counts"),
        ("rows", counts, [150, 60, 254]),
        ("head", "The fees these counts trigger"),
        ("rows", rows, [170, 234, 60]),
        ("head", "Entity status"),
        ("p", f"{filing_profile.entity_label(profile)}. Each fee code above has three columns: "
              "undiscounted, small entity (37 CFR 1.27) and micro entity (37 CFR 1.29). Status is "
              "certified, not assumed, and claiming a discount you are not entitled to is an "
              "improper payment."),
        ("head", "One surcharge that is entirely in your control"),
        ("p", "37 CFR 1.16(u) charges a surcharge for filing the specification, claims and "
              "abstract in any format other than DOCX. This package supplies them as a DOCX for "
              "that reason. Converting it to PDF before uploading pays the surcharge for nothing."),
        ("head", "The current amounts"),
        ("p", filing_rules.FEE_SCHEDULE_URL),
    ]
    return _paper("FEE WORKSHEET", blocks)


def information_disclosure_pdf(citations: Sequence[Mapping[str, Any]]) -> bytes:
    """37 CFR 1.98 listing of what this draft cites. A listing, not an executed statement."""
    rows = [["Publication", "Kind", "Date", "Patentee or applicant", "Title", "Resolved"]]
    for item in citations:
        rows.append([str(item.get("publication_number") or ""),
                     str(item.get("kind_code") or ""),
                     str(item.get("publication_date") or ""),
                     str(item.get("name_of_patentee") or "")[:60],
                     str(item.get("title") or "")[:90],
                     str(item.get("resolved") or "")])
    if len(rows) == 1:
        rows.append(["(none cited)", "", "", "", "", ""])
    blocks: list[Any] = [
        ("small", "THIS IS A LISTING, NOT AN EXECUTED STATEMENT. The duty of disclosure under 37 "
                  "CFR 1.56 is personal to each individual associated with the filing of the "
                  "application, and no software discharges it. Transfer these onto form PTO/SB/08a "
                  "or the Patent Center IDS screen, add anything else you are aware of, and sign "
                  "it yourself."),
        ("head", "US patent documents cited in this application"),
        ("rows", rows, [78, 32, 60, 100, 160, 44]),
        ("head", "Timing"),
        ("p", "An IDS filed within three months of the filing date, or before the first Office "
              "action on the merits, needs no fee and no statement under 37 CFR 1.97(e)."),
        ("head", "A copy is required for some of these"),
        ("p", "37 CFR 1.98(a)(2): a copy of each cited document is required unless it is a US "
              "patent or US patent application publication. A foreign document or a non-patent "
              "publication must be supplied as a copy."),
    ]
    return _paper("INFORMATION DISCLOSURE STATEMENT - CITATION LISTING", blocks)


# =============================================================================================
# 00 - what to do with all of it
# =============================================================================================
def read_me(project: Mapping[str, Any], profile: Mapping[str, Any], *,
            fees: Mapping[str, Any], gaps: Sequence[Mapping[str, str]],
            audit: Sequence[Mapping[str, Any]], sheet_count: int) -> str:
    lines = [
        "FILING PACKAGE",
        f"Built {date.today().isoformat()} for: {str(project.get('title') or '')[:200]}",
        "",
        "This is a complete US utility filing except for the things only a person can do: sign,",
        "certify entity status, and pay. Nothing here is legal advice and nothing here has been",
        "reviewed by a registered practitioner.",
        "",
        "=" * 78,
        "WHAT TO UPLOAD, AND IN WHAT ORDER",
        "=" * 78,
        "",
        f"Patent Center: {filing_rules.PATENT_CENTER_URL}",
        "Start a new application, 'Utility - Nonprovisional application under 35 USC 111(a)'.",
        "",
        f"1. {SPEC_NAME}",
        "   Document description: 'Specification'. Upload it AS A DOCX.",
        "   37 CFR 1.16(u) charges a surcharge for filing the specification, claims and abstract",
        "   in any other format. Converting this file to PDF first pays that surcharge for",
        "   nothing. Patent Center splits the DOCX into specification, claims and abstract by",
        "   itself; check the split it proposes before you continue.",
        "",
        f"2. {DRAWINGS_NAME}",
        "   Document description: 'Drawings only, black and white line drawings'.",
        f"   {sheet_count} sheet(s), numbered 1/{sheet_count} through {sheet_count}/{sheet_count}.",
        "   Upload one image file per view. Where one file carried several views, each view was",
        "   cut onto its own sheet so it could be numbered and scaled above the 0.32 cm",
        "   character minimum; that crop follows a machine reading of where each view ends, so",
        "   check those sheets. Where the views were laid out with no white corridor between",
        "   them the cut was refused and the sheet is here whole, which is not filable.",
        "",
        "3. The Application Data Sheet.",
        "   Use Patent Center's WEB ADS rather than uploading a PDF. The web form validates every",
        f"   field as you type it. {WEB_ADS_NAME} in this package lists every value to enter, in",
        f"   the order the form asks for them. {ADS_NAME} is the same data as a signable paper,",
        "   for the case where you would rather upload one.",
        "",
        f"4. {DECLARATION_NAME}",
        "   Document description: 'Oath or Declaration filed'.",
        "   Print or S-sign it, one signature per named inventor. An S-signature is the signer's",
        "   own name typed between forward slashes: /Jane A. Smith/.",
        "   If you produce the signed PDF yourself, watch the fonts. Patent Center rejects a PDF",
        "   whose font is not embedded, and it also rejects one whose embedded font NAME contains",
        "   spaces, which is what a LibreOffice export produces. Re-distilling through",
        "   'gs -sDEVICE=pdfwrite' fixes it.",
        "",
        f"5. {IDS_NAME}",
        "   Not a filing document as it stands: it is the listing. Transfer it onto PTO/SB/08a or",
        "   the Patent Center IDS screen and sign it. Filing it within three months of the filing",
        "   date needs no fee.",
        "",
        f"6. Fees. {FEES_NAME} has the counts and the fee codes they trigger. Patent Center",
        "   totals the amounts from the same counts at submission.",
        "",
        "=" * 78,
        "WHAT ONLY YOU CAN SUPPLY",
        "=" * 78,
        "",
    ]
    if gaps:
        lines.append("These fields are empty and every one of them is on a paper that gets filed:")
        lines.append("")
        for gap in gaps:
            lines.append(f"  - {gap['field']}  ({gap['rule']})")
        lines.append("")
        lines.append("Fill them in on the Filing tab and rebuild the package. Every paper here")
        lines.append("marks a missing field as (REQUIRED) rather than guessing at it.")
    else:
        lines.append("Nothing. Every field on every paper in this package is filled in.")
    lines += [
        "",
        "Still yours to do in every case:",
        f"  - Certify entity status. This package assumes: {filing_profile.entity_label(profile)}.",
        "  - Sign the declaration, once per named inventor.",
        "  - Sign the ADS, or complete the web ADS, which counts as signing it.",
        "  - Pay the filing, search and examination fees.",
        "  - Have a registered US patent practitioner read it before you submit.",
        "",
        "=" * 78,
        "WHAT THIS PACKAGE CHECKED",
        "=" * 78,
        "",
        f"{len(audit)} finding(s). AUDIT.txt has all of them with the paragraph each comes from.",
        f"Verdict: {filing_rules.verdict(audit)}.",
        "",
        f"Claims: {fees.get('total')} total, {fees.get('independent')} independent, "
        f"{fees.get('multiple_dependent')} multiple dependent, {fees.get('billable')} billed.",
        "",
    ]
    return "\n".join(lines) + "\n"


def audit_text(findings: Sequence[Mapping[str, Any]]) -> str:
    order = {"blocker": 0, "formality": 1, "note": 2}
    lines = ["FILING AUDIT", "",
             "Every check is mechanical and reads the finished file rather than the draft.",
             f"Verdict: {filing_rules.verdict(findings)}", ""]
    for severity, label in (("blocker", "BLOCKERS - the paper is defective or would be refused"),
                            ("formality", "FORMALITIES - the Office objects, the filing date "
                                          "stands"),
                            ("note", "NOTES")):
        group = [item for item in findings if item.get("severity") == severity]
        if not group:
            continue
        lines += ["=" * 78, label, "=" * 78, ""]
        for item in sorted(group, key=lambda row: order.get(row.get("severity"), 9)):
            lines += [f"[{item.get('rule')}]  {item.get('title')}",
                      f"    in: {item.get('where')}"]
            if item.get("detail"):
                lines.append("    " + str(item["detail"]))
            lines.append("")
    if not findings:
        lines.append("No finding. Every check passed.")
    return "\n".join(lines) + "\n"


# =============================================================================================
# Assembly
# =============================================================================================
def build(*, project: Mapping[str, Any], version: Mapping[str, Any],
          profile: Mapping[str, Any], figures: Sequence[Mapping[str, Any]] = (),
          sheet_facts: Mapping[str, Mapping[str, Any]] | None = None,
          citations: Sequence[Mapping[str, Any]] = ()) -> dict[str, Any]:
    """Build every paper, then check every paper, and return both."""
    resolved = filing_profile.resolve(profile, project)
    sections = _sections(version)

    sheets = drawing_sheets(figures, sheet_facts)
    drawings, measurements = drawings_pdf(sheets) if sheets else (b"", [])
    specification = specification_docx(project, version)
    spec_pages = _estimated_pages(specification)
    fees = fee_profile(sections.get("claims") or "", specification_sheets=spec_pages,
                       drawing_sheets_count=len(sheets))

    ads = application_data_sheet_pdf(project, version, resolved,
                                     drawing_sheets_count=len(sheets))
    declaration = declaration_pdf(project, version, resolved)
    worksheet = fee_worksheet_pdf(resolved, fees)
    ids = information_disclosure_pdf(citations)
    web_ads = filing_profile.web_ads_sheet(
        resolved, title=(sections.get("title") or str(project.get("title") or "")).strip(),
        drawing_sheets=len(sheets))

    findings = audit(specification=specification, drawings=drawings, ads=ads,
                     declaration=declaration, worksheet=worksheet, ids=ids,
                     sections=sections, sheets=sheets, measurements=measurements,
                     profile=resolved)
    gaps = filing_profile.gaps(resolved)
    for gap in gaps:
        findings.append(filing_rules.finding(
            gap["rule"], "blocker", ADS_NAME,
            f"{gap['field']} is not supplied", gap["detail"]))

    files = {
        README_NAME: read_me(project, resolved, fees=fees, gaps=gaps, audit=findings,
                             sheet_count=len(sheets)).encode("utf-8"),
        SPEC_NAME: specification,
        ADS_NAME: ads,
        WEB_ADS_NAME: web_ads.encode("utf-8"),
        DECLARATION_NAME: declaration,
        FEES_NAME: worksheet,
        IDS_NAME: ids,
        AUDIT_NAME: audit_text(findings).encode("utf-8"),
    }
    if drawings:
        files[DRAWINGS_NAME] = drawings
    return {"files": files, "findings": findings, "fees": fees, "gaps": gaps,
            "sheets": [{"label": item["label"], "cropped": bool(item.get("cropped")),
                        "split_refused": bool(item.get("split_refused"))}
                       for item in sheets],
            "measurements": measurements,
            "verdict": filing_rules.verdict(findings),
            "ready": not filing_rules.blockers(findings)}


def _estimated_pages(specification: bytes) -> int:
    """How many sheets the specification will be, for the 37 CFR 1.16(s) size fee.

    A .docx has no page count until something lays it out, and nothing on this box does. The
    estimate is characters over a page's worth of 12 point at 1.5 spacing, rounded up, which is
    close enough to answer the only question asked of it: is this application over 100 sheets.
    """
    from docx import Document
    document = Document(io.BytesIO(specification))
    characters = sum(len(paragraph.text) for paragraph in document.paragraphs)
    breaks = sum(1 for paragraph in document.paragraphs
                 if paragraph.paragraph_format.page_break_before or
                 "w:br" in paragraph._p.xml and 'w:type="page"' in paragraph._p.xml)
    return max(1, -(-characters // 2100) + breaks)


def audit(*, specification: bytes, drawings: bytes, ads: bytes, declaration: bytes,
          worksheet: bytes, ids: bytes, sections: Mapping[str, str],
          sheets: Sequence[Mapping[str, Any]], measurements: Sequence[Mapping[str, Any]],
          profile: Mapping[str, Any]) -> list[filing_rules.Finding]:
    """Run every check in filing_rules over the files that were just built."""
    findings: list[filing_rules.Finding] = []
    findings += filing_rules.audit_specification_docx(specification, where=SPEC_NAME)
    for blob, name in ((ads, ADS_NAME), (declaration, DECLARATION_NAME),
                       (worksheet, FEES_NAME), (ids, IDS_NAME)):
        findings += filing_rules.audit_pdf(blob, where=name)
    findings += filing_rules.audit_declaration_text(_pdf_text(declaration), where=DECLARATION_NAME)
    findings += filing_rules.audit_ads_text(_pdf_text(ads), where=ADS_NAME)
    if pdf_fonts.missing():
        findings.append(filing_rules.finding(
            "Patent Center PDF validation", "blocker", "server",
            "A font this server files with is not installed",
            "Missing: " + ", ".join(pdf_fonts.missing()) +
            ". Install fonts-liberation2 and fonts-droid-fallback. Without them reportlab falls "
            "back to the PDF base-14, which is never embedded."))

    if drawings:
        findings += filing_rules.audit_pdf(drawings, where=DRAWINGS_NAME)
        smallest = min((item["character_points"] for item in measurements
                        if item.get("measured")), default=0.0)
        findings += filing_rules.audit_drawing_geometry(
            sheet_count=len(sheets),
            described_figures=_described_figures(sections),
            sheet_labels=[([str(item.get("label"))] if item.get("label") else [])
                          for item in sheets],
            character_points=smallest, where=DRAWINGS_NAME)
        refused = [item for item in sheets if item.get("split_refused")]
        if refused:
            findings.append(filing_rules.finding(
                "37 CFR 1.84(u)", "blocker", DRAWINGS_NAME,
                f"{len(refused)} upload(s) hold several views and could not be separated",
                "The views are laid out with no white corridor between them, so every cut line "
                "would run through a lead line or a numeral. The sheet is in the package whole. "
                "Supply one view per image file: that is also what gets the reference characters "
                "above the 0.32 cm minimum, which a shared sheet rarely clears."))
        split = [item for item in sheets if item.get("cropped")]
        if split:
            findings.append(filing_rules.finding(
                "37 CFR 1.84(u)", "formality", DRAWINGS_NAME,
                f"{len(split)} filing sheet(s) were cut out of a composite upload",
                "Each view was separated onto its own sheet so it could be numbered and scaled. "
                "That is a recovery, not the supported path: the crop follows a machine reading "
                "of where each view ends. Check every one of those sheets before filing, or "
                "re-upload one view per image file."))
        unmeasured = [item["label"] for item in measurements if not item.get("measured")]
        if unmeasured:
            findings.append(filing_rules.finding(
                "37 CFR 1.84(p)(3)", "note", DRAWINGS_NAME,
                "The reference character height was not measured on every sheet",
                "Not measured on: " + ", ".join(unmeasured) + ". The 0.32 cm floor could not be "
                "confirmed for those sheets."))
    else:
        findings += filing_rules.audit_drawing_geometry(
            sheet_count=0, described_figures=_described_figures(sections),
            sheet_labels=[], character_points=0.0, where=DRAWINGS_NAME)
    del profile
    return findings


def _described_figures(sections: Mapping[str, str]) -> list[str]:
    import figure_facts
    return figure_facts.figure_numbers(sections.get("drawing_descriptions") or "")


def _pdf_text(blob: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(blob))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:                                              # noqa: BLE001
        return ""


def zip_bytes(files: Mapping[str, bytes]) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for name in sorted(files):
            archive.writestr(name, files[name])
    return output.getvalue()
