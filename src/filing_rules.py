"""The formal requirements of a US utility filing, as checks that read the finished paper.

WHY THIS IS SEPARATE FROM ``draft_qa``.  ``draft_qa`` reads the DRAFT: is numeral 34 introduced
before it is used, does claim 7 have antecedent basis, is the abstract inside 150 words.  Every one
of those questions is about the text.  This module reads the ARTEFACT: the .docx the Office will
ingest, the drawing sheets it will scan, the declaration somebody has to sign.  A draft can be
flawless and still bounce at upload because the specification was 12 pt single spaced, because the
claims did not commence on their own page, or because a curly quotation mark made the DOCX
validator report a non-Latin script.

Those two failure modes have nothing to do with each other, they are found by completely different
means, and conflating them is how a product ends up telling a user their application is ready when
what it has checked is the half a filing clerk never looks at.

WHAT A FINDING IS.  Every check returns zero or more findings, and every finding names:

    rule        the paragraph it comes from, e.g. "37 CFR 1.52(b)(6)"
    severity    "blocker"   the paper is defective or would be refused as filed
                "formality" the Office will object but the filing date stands
                "note"      worth knowing, not a defect
    where       which file in the package
    title       one line, the defect
    detail      what is wrong and what the rule wants instead

NOTHING HERE IS A LEGAL OPINION and no check says anything about patentability.  These are the
mechanical requirements of 37 CFR 1.52, 1.63, 1.72, 1.75, 1.76, 1.77, 1.83 and 1.84, plus the two
Patent Center ingestion rules that are not in the CFR at all and bounce more filings than any of
them: every font in a PDF must be embedded, and a DOCX must not carry characters the validator
reads as a non-Latin script.
"""
from __future__ import annotations

import io
import re
import shutil
import subprocess
import unicodedata
import zipfile
from typing import Any, Iterable, Mapping, Sequence

# =============================================================================================
# The numbers, in one place, each beside the paragraph it comes from
# =============================================================================================
#  37 CFR 1.52(a)(1)(ii). Letter or A4; we file Letter. The minima are 2.0 cm top, 2.5 cm left,
#  2.0 cm right and 2.0 cm bottom. We use one inch on every side, which clears all four and is
#  what a practitioner expects to see.
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_MIN_TOP_IN = 2.0 / 2.54
MARGIN_MIN_LEFT_IN = 2.5 / 2.54
MARGIN_MIN_RIGHT_IN = 2.0 / 2.54
MARGIN_MIN_BOTTOM_IN = 2.0 / 2.54
SPEC_MARGIN_IN = 1.0

#  37 CFR 1.52(b)(2)(ii): nonscript type, preferably 12 point, and 1 1/2 or double spaced.
SPEC_FONT = "Times New Roman"
SPEC_FONT_PT = 12
SPEC_LINE_SPACING = 1.5
NONSCRIPT_FONTS = frozenset({
    "times new roman", "times", "arial", "helvetica", "courier", "courier new",
    "liberation serif", "liberation sans", "calibri", "cambria", "georgia", "verdana",
})

#  37 CFR 1.72(a) and (b), 1.75(c).
TITLE_CHAR_LIMIT = 500
ABSTRACT_WORD_LIMIT = 150
ABSTRACT_WORD_FLOOR = 50

#  37 CFR 1.84(g): each drawing sheet takes a 2.5 cm top and left margin, a 1.5 cm right margin
#  and a 1.0 cm bottom margin, with a sight of at least 17.0 by 26.2 cm.
SHEET_MARGIN_TOP_IN = 1.0
SHEET_MARGIN_LEFT_IN = 1.0
SHEET_MARGIN_RIGHT_IN = 0.625
SHEET_MARGIN_BOTTOM_IN = 0.375
#  37 CFR 1.84(p)(3): reference characters, sheet numbers and view numbers are at least 0.32 cm
#  (1/8 inch) high. At 72 points to the inch that is just over 9 points.
MIN_CHARACTER_IN = 0.32 / 2.54
MIN_CHARACTER_PT = MIN_CHARACTER_IN * 72.0
#  37 CFR 1.84(t): the sheet number is larger than the reference characters.
SHEET_NUMBER_PT = 14.0
VIEW_LEGEND_PT = 16.0

#  37 CFR 1.16: what the claim counts buy before a surcharge starts.
CLAIMS_INCLUDED = 20
INDEPENDENT_CLAIMS_INCLUDED = 3
#  37 CFR 1.16(s): the application size fee starts above 100 sheets of specification and drawings.
SIZE_FEE_FREE_SHEETS = 100
SIZE_FEE_BLOCK_SHEETS = 50

FEE_SCHEDULE_URL = \
    "https://www.uspto.gov/learning-and-resources/fees-and-payment/uspto-fee-schedule"
PATENT_CENTER_URL = "https://patentcenter.uspto.gov/"

#  37 CFR 1.77(b): the order of the parts of a utility specification.
SECTION_ORDER_1_77 = (
    "TITLE OF THE INVENTION",
    "CROSS-REFERENCE TO RELATED APPLICATIONS",
    "STATEMENT REGARDING FEDERALLY SPONSORED RESEARCH OR DEVELOPMENT",
    "THE NAMES OF THE PARTIES TO A JOINT RESEARCH AGREEMENT",
    "REFERENCE TO A SEQUENCE LISTING",
    "BACKGROUND OF THE INVENTION",
    "BRIEF SUMMARY OF THE INVENTION",
    "BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS",
    "DETAILED DESCRIPTION OF THE INVENTION",
    "CLAIMS",
    "ABSTRACT OF THE DISCLOSURE",
)

#  37 CFR 1.63(a) and (c). Each one is a statement the paper must actually carry; a declaration
#  missing any of them is not a declaration, and Patent Center will take it and the Office will
#  send a notice months later. The patterns are deliberately loose about wording and strict about
#  meaning, because every practitioner phrases these slightly differently.
DECLARATION_STATEMENTS = (
    ("37 CFR 1.63(a)(1)",
     "names the inventor",
     re.compile(r"\b(?:legal\s+name|name\s+of\s+(?:the\s+)?(?:sole\s+|joint\s+)?inventor)\b",
                re.IGNORECASE)),
    ("37 CFR 1.63(a)(2)",
     "identifies the application it is directed to",
     re.compile(r"\bthis\s+declaration\s+is\s+directed\s+to\b|"
                r"\babove[- ]identified\s+application\b", re.IGNORECASE)),
    ("37 CFR 1.63(a)(3)",
     "states the belief in original inventorship",
     re.compile(r"\bbelieve\b[^.]{0,120}\boriginal\b[^.]{0,80}\binventor\b", re.IGNORECASE)),
    ("37 CFR 1.63(a)(4)",
     "states that the application was made or authorized to be made by the signer",
     re.compile(r"\bwas\s+made\s+or\s+was\s+authorized\s+to\s+be\s+made\b", re.IGNORECASE)),
    ("37 CFR 1.63(c)",
     "states that the signer has reviewed and understands the application",
     re.compile(r"\breviewed\s+and\s+understand\b", re.IGNORECASE)),
    ("37 CFR 1.56 / 1.63(c)",
     "acknowledges the duty to disclose",
     re.compile(r"\bduty\s+to\s+disclose\b", re.IGNORECASE)),
    ("18 U.S.C. 1001",
     "carries the willful false statement warning",
     re.compile(r"\bwillful\s+false\s+statement", re.IGNORECASE)),
)

#  37 CFR 1.76(b): the section headings an application data sheet must contain.
ADS_SECTION_HEADINGS = (
    "Inventor Information",
    "Correspondence Information",
    "Application Information",
    "Representative Information",
    "Domestic Benefit/National Stage Information",
    "Foreign Priority Information",
    "Applicant Information",
    "Assignee Information",
)

#  Characters that look ordinary in a word processor and are the single commonest reason a DOCX
#  comes back from the Patent Center validator with "non-Latin script detected". Sixteen of them
#  in four paragraphs did it on a real filing in August 2026.
SMART_CHARACTERS = {
    "‘": "'", "’": "'", "‚": "'", "‛": "'",
    "“": '"', "”": '"', "„": '"', "‟": '"',
    "–": "-", "—": "-", "‒": "-", "―": "-", "−": "-",
    "…": "...", " ": " ", " ": " ", " ": " ", " ": " ",
    "′": "'", "″": '"', "­": "",
    "×": "x", "⁄": "/", "•": "-", "‐": "-", "‑": "-",
}
#  Characters that are legitimately needed in a specification and are plain Latin-1, so the
#  validator is content with them.
ALLOWED_NON_ASCII = frozenset("°µ±½¼¾éèüöä")


class Finding(dict):
    """A defect, as a plain dict so it crosses JSON, a template and a subprocess unchanged."""


def finding(rule: str, severity: str, where: str, title: str, detail: str = "") -> Finding:
    return Finding(rule=rule, severity=severity, where=where, title=title, detail=detail)


def blockers(findings: Iterable[Mapping[str, Any]]) -> list[dict[str, Any]]:
    return [dict(item) for item in findings if item.get("severity") == "blocker"]


def verdict(findings: Sequence[Mapping[str, Any]]) -> str:
    if any(item.get("severity") == "blocker" for item in findings):
        return "not ready"
    if any(item.get("severity") == "formality" for item in findings):
        return "ready, with formalities the Office may object to"
    return "ready"


# =============================================================================================
# Text hygiene
# =============================================================================================
def to_filing_text(value: Any) -> str:
    """Fold a draft's typography down to what a DOCX validator will not complain about.

    A word processor turns a straight quotation mark into a curly one without being asked, and a
    model writes them because its training data is full of them. Neither is wrong prose. Both make
    Patent Center report a script it does not expect, on a document that is otherwise perfect.
    """
    text = str(value or "")
    for bad, good in SMART_CHARACTERS.items():
        text = text.replace(bad, good)
    text = unicodedata.normalize("NFC", text)
    return text.replace("\r\n", "\n").replace("\r", "\n")


def non_latin_characters(text: str) -> list[tuple[str, int]]:
    """Every character a DOCX validator would call out, with how often it appears."""
    counts: dict[str, int] = {}
    for character in str(text or ""):
        if ord(character) < 128 or character in ALLOWED_NON_ASCII:
            continue
        counts[character] = counts.get(character, 0) + 1
    return sorted(counts.items(), key=lambda item: (-item[1], item[0]))


def describe_character(character: str) -> str:
    try:
        name = unicodedata.name(character)
    except ValueError:
        name = "unnamed"
    return f"U+{ord(character):04X} {name}"


# =============================================================================================
# The specification document
# =============================================================================================
_PARAGRAPH_NUMBER_RE = re.compile(r"^\[(\d{4,})\]")
_CLAIM_START_RE = re.compile(r"^\s*1\s*[.)]\s+\S")


def audit_specification_docx(blob: bytes, *, where: str = "specification.docx",
                             expect_paragraph_numbers: bool = True) -> list[Finding]:
    """Read the finished .docx the way the Office's ingestion does, and report what it would."""
    from docx import Document
    from docx.shared import Inches

    out: list[Finding] = []
    try:
        document = Document(io.BytesIO(blob))
    except Exception as exc:                                        # noqa: BLE001
        return [finding("37 CFR 1.52", "blocker", where, "The specification file is unreadable",
                        f"{type(exc).__name__}: {exc}")]

    # -- page and margins ----------------------------------------------------------------------
    for index, section in enumerate(document.sections, 1):
        width = (section.page_width or 0) / Inches(1)
        height = (section.page_height or 0) / Inches(1)
        if abs(width - PAGE_WIDTH_IN) > 0.02 or abs(height - PAGE_HEIGHT_IN) > 0.02:
            out.append(finding(
                "37 CFR 1.52(a)(1)(ii)", "blocker", where,
                "The specification is not on Letter paper",
                f"Section {index} is {width:.2f} by {height:.2f} inches. It must be 8.5 by 11 "
                "inches, or A4."))
        for name, value, minimum in (
                ("top", (section.top_margin or 0) / Inches(1), MARGIN_MIN_TOP_IN),
                ("left", (section.left_margin or 0) / Inches(1), MARGIN_MIN_LEFT_IN),
                ("right", (section.right_margin or 0) / Inches(1), MARGIN_MIN_RIGHT_IN),
                ("bottom", (section.bottom_margin or 0) / Inches(1), MARGIN_MIN_BOTTOM_IN)):
            if value + 0.005 < minimum:
                out.append(finding(
                    "37 CFR 1.52(a)(1)(ii)", "blocker", where,
                    f"The {name} margin is below the minimum",
                    f"Section {index} has {value:.2f} inch; the rule requires at least "
                    f"{minimum:.2f} inch."))
        columns = section._sectPr.xpath("./w:cols/@w:num")
        if columns and int(columns[0]) > 1:
            out.append(finding("37 CFR 1.52(b)(2)", "blocker", where,
                               "The specification is set in more than one column",
                               f"Section {index} declares {columns[0]} columns."))

    # -- type and spacing ----------------------------------------------------------------------
    normal = document.styles["Normal"]
    font_name = (normal.font.name or "").strip()
    if font_name.lower() not in NONSCRIPT_FONTS:
        out.append(finding("37 CFR 1.52(b)(2)(ii)", "formality", where,
                           "The body type is not a recognised nonscript font",
                           f"The Normal style is set in {font_name or 'an unnamed face'}. The rule "
                           "wants a nonscript font such as Times New Roman, Arial or Courier."))
    size = normal.font.size
    if size is None or round(size.pt) != SPEC_FONT_PT:
        out.append(finding("37 CFR 1.52(b)(2)(ii)", "formality", where,
                           "The body type is not 12 point",
                           f"The Normal style is {size.pt if size else 'unset'} point."))
    spacing = normal.paragraph_format.line_spacing
    if spacing is None or float(spacing) + 1e-6 < SPEC_LINE_SPACING:
        out.append(finding("37 CFR 1.52(b)(2)(ii)", "blocker", where,
                           "The specification is not 1.5 or double spaced",
                           f"The Normal style is set to {spacing if spacing else 'single'} "
                           "spacing. The rule requires 1 1/2 or double."))

    paragraphs = [p for p in document.paragraphs]
    texts = [p.text.strip() for p in paragraphs]

    # -- the characters that bounce a DOCX -----------------------------------------------------
    whole = "\n".join(texts)
    strays = non_latin_characters(whole)
    if strays:
        out.append(finding(
            "Patent Center DOCX validation", "blocker", where,
            f"{sum(count for _c, count in strays)} character(s) the validator reads as a "
            "non-Latin script",
            "; ".join(f"{describe_character(char)} x{count}" for char, count in strays[:8])
            + ". Curly quotation marks and dashes are the usual cause. Replace them with the "
              "straight ASCII forms."))

    # -- paragraph numbering, 1.52(b)(6) -------------------------------------------------------
    if expect_paragraph_numbers:
        numbers = []
        for paragraph in paragraphs:
            match = _PARAGRAPH_NUMBER_RE.match(paragraph.text.strip())
            if match:
                numbers.append((int(match.group(1)), paragraph))
        if not numbers:
            out.append(finding(
                "37 CFR 1.52(b)(6)", "formality", where,
                "No paragraph is numbered",
                "Numbered paragraphs are optional at filing and close to mandatory in practice: "
                "without them every later amendment has to identify the passage it changes by "
                "quotation instead of by number."))
        else:
            expected = 1
            for value, _paragraph in numbers:
                if value != expected:
                    out.append(finding(
                        "37 CFR 1.52(b)(6)", "blocker", where,
                        "The paragraph numbers are not consecutive",
                        f"[{expected:04d}] was expected and [{value:04d}] was found. Numbering "
                        "must run consecutively so an amendment can name a paragraph."))
                    break
                expected = value + 1
            first_run = numbers[0][1].runs[0] if numbers[0][1].runs else None
            if first_run is not None and not first_run.bold:
                out.append(finding(
                    "37 CFR 1.52(b)(6)", "formality", where,
                    "The paragraph numbers are not bold",
                    "The rule asks for the bracketed number in bold at the left margin, followed "
                    "by a gap of about four spaces."))

    # -- the two page breaks the rules actually require ----------------------------------------
    breaks = _page_break_indexes(paragraphs)
    claims_at = _index_of(texts, lambda t: t.upper().startswith("CLAIM") and len(t) < 40)
    what_at = _index_of(texts, lambda t: t.lower().startswith("what is claimed"))
    abstract_at = _index_of(texts, lambda t: t.upper().startswith("ABSTRACT") and len(t) < 40)

    if what_at < 0:
        out.append(finding("MPEP 608.01(m)", "formality", where,
                           "The claims are not introduced",
                           'A claim set commences with a phrase such as "What is claimed is:".'))
    if claims_at < 0:
        out.append(finding("37 CFR 1.75", "blocker", where, "The document has no claims heading",
                           "A utility application must contain at least one claim."))
    elif not any(index <= claims_at for index in breaks if index > 0 and
                 claims_at - index <= 2):
        out.append(finding(
            "37 CFR 1.75(h)", "blocker", where,
            "The claims do not commence on a separate page",
            "The rule requires the claim or claims to commence on a separate physical sheet or "
            "electronic page. No page break precedes the claims heading."))
    if abstract_at < 0:
        out.append(finding("37 CFR 1.72(b)", "blocker", where, "The document has no abstract",
                           "A separate abstract on its own sheet is required."))
    elif not any(index <= abstract_at for index in breaks if index > 0 and
                 abstract_at - index <= 2):
        out.append(finding(
            "37 CFR 1.72(b)", "blocker", where,
            "The abstract does not commence on a separate sheet",
            "The rule requires the abstract to commence on a separate sheet, preferably following "
            "the claims."))

    # -- abstract form -------------------------------------------------------------------------
    if abstract_at >= 0:
        body = [t for t in texts[abstract_at + 1:] if t]
        words = len(" ".join(body).split())
        if words > ABSTRACT_WORD_LIMIT:
            out.append(finding("37 CFR 1.72(b)", "formality", where,
                               f"The abstract is {words} words",
                               "The rule sets a range of 50 to 150 words in a single paragraph."))
        elif words and words < ABSTRACT_WORD_FLOOR:
            out.append(finding("37 CFR 1.72(b)", "note", where,
                               f"The abstract is {words} words",
                               "The rule asks for 50 to 150."))
        if len(body) > 1:
            out.append(finding("37 CFR 1.72(b)", "formality", where,
                               "The abstract is more than one paragraph",
                               "The rule asks for a single paragraph."))

    # -- title ---------------------------------------------------------------------------------
    if texts and len(texts[0]) > TITLE_CHAR_LIMIT:
        out.append(finding("37 CFR 1.72(a)", "formality", where,
                           f"The title is {len(texts[0])} characters",
                           f"The Office truncates a title over {TITLE_CHAR_LIMIT} characters."))

    # -- 1.77(b) order -------------------------------------------------------------------------
    out += _heading_order_findings(texts, where)

    # -- page numbers --------------------------------------------------------------------------
    if not _has_page_field(document):
        out.append(finding("37 CFR 1.52(b)(5)", "formality", where,
                           "The pages are not numbered",
                           "The rule asks for consecutive page numbers, centred above or below "
                           "the text."))

    # -- the package itself --------------------------------------------------------------------
    out += audit_docx_package(blob, where=where)
    return out


def _index_of(texts: Sequence[str], predicate) -> int:
    for index, text in enumerate(texts):
        if text and predicate(text):
            return index
    return -1


def _page_break_indexes(paragraphs: Sequence[Any]) -> list[int]:
    """Paragraph indexes at or before which the page changes."""
    from docx.oxml.ns import qn
    out: list[int] = []
    for index, paragraph in enumerate(paragraphs):
        if paragraph.paragraph_format.page_break_before:
            out.append(index)
            continue
        for run in paragraph.runs:
            if run._element.findall(qn("w:br")):
                for element in run._element.findall(qn("w:br")):
                    if element.get(qn("w:type")) == "page":
                        out.append(index)
                        break
    return out


def _has_page_field(document: Any) -> bool:
    for section in document.sections:
        for part in (section.footer, section.header):
            for paragraph in part.paragraphs:
                xml = paragraph._p.xml
                if "PAGE" in xml and ("instrText" in xml or "fldSimple" in xml):
                    return True
    return False


def _heading_order_findings(texts: Sequence[str], where: str) -> list[Finding]:
    positions = []
    for heading in SECTION_ORDER_1_77:
        for index, text in enumerate(texts):
            if text.upper() == heading:
                positions.append((index, heading))
                break
    ordered = [heading for _index, heading in sorted(positions)]
    expected = [heading for heading in SECTION_ORDER_1_77 if heading in ordered]
    if ordered != expected:
        return [finding("37 CFR 1.77(b)", "formality", where,
                        "The sections are not in the order the rule sets out",
                        "Found: " + " / ".join(ordered) + ". Expected: " + " / ".join(expected))]
    return []


def audit_docx_package(blob: bytes, *, where: str = "specification.docx") -> list[Finding]:
    """Parts inside the .docx that make the validator warn on a document nobody can see wrong.

    An empty comments part is the one that catches people: python-docx and Word both leave it in a
    document that has never had a comment, and Patent Center reports it.
    """
    out: list[Finding] = []
    try:
        archive = zipfile.ZipFile(io.BytesIO(blob))
    except Exception as exc:                                        # noqa: BLE001
        return [finding("Patent Center DOCX validation", "blocker", where,
                        "The .docx is not a readable package", str(exc)[:200])]
    names = set(archive.namelist())
    for part in sorted(names):
        if re.fullmatch(r"word/comments\d*\.xml", part):
            body = archive.read(part).decode("utf-8", "replace")
            if not re.search(r"<w:comment\b", body):
                out.append(finding(
                    "Patent Center DOCX validation", "formality", where,
                    "The document carries an empty comments part",
                    f"{part} exists and holds no comment. The validator reports it. Remove the "
                    "part rather than the comments."))
    for part, why in (("docProps/thumbnail.jpeg", "a preview image"),
                      ("word/stylesWithEffects.xml", "a legacy Word 2010 style part")):
        if part in names:
            out.append(finding("Patent Center DOCX validation", "note", where,
                               f"The document carries {why}",
                               f"{part} is not needed in a filing copy and is one more part the "
                               "validator has to accept."))
    if any(name.startswith("customXml/") for name in names):
        out.append(finding("Patent Center DOCX validation", "note", where,
                           "The document carries a customXml part",
                           "Left behind by the template it was built from. Not needed."))
    if any(name.startswith("word/embeddings/") or name.endswith(".bin") for name in names):
        out.append(finding("Patent Center DOCX validation", "blocker", where,
                           "The document carries an embedded object",
                           "Embedded objects and macros are refused. Everything must be text and "
                           "images."))
    settings = ""
    if "word/settings.xml" in names:
        settings = archive.read("word/settings.xml").decode("utf-8", "replace")
    if "w:trackChanges" in settings:
        out.append(finding("Patent Center DOCX validation", "blocker", where,
                           "Track changes is switched on",
                           "A filing copy must be the final text with no revision marks."))
    document = archive.read("word/document.xml").decode("utf-8", "replace") \
        if "word/document.xml" in names else ""
    if "<w:ins " in document or "<w:del " in document:
        out.append(finding("Patent Center DOCX validation", "blocker", where,
                           "The document contains tracked insertions or deletions",
                           "Accept or reject every revision before filing."))
    return out


# =============================================================================================
# PDFs
# =============================================================================================
def audit_pdf(blob: bytes, *, where: str, expect_letter: bool = True) -> list[Finding]:
    """Every font embedded, no space in a font name, and the page the Office expects.

    THE FONT NAME IS NOT A DETAIL. Patent Center rejected a declaration in August 2026 for a
    non-embedded font when the font was embedded: the BaseFont name contained spaces, written as
    ``Liberation#20Sans#20Regular``, and the check that reads the name read the escape as a
    defect. Re-distilling through ghostscript renames it and the same file goes through.
    """
    out: list[Finding] = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(blob))
    except Exception as exc:                                        # noqa: BLE001
        return [finding("Patent Center PDF validation", "blocker", where,
                        "The PDF is unreadable", f"{type(exc).__name__}: {exc}")]
    if expect_letter:
        for index, page in enumerate(reader.pages, 1):
            width = float(page.mediabox.width) / 72.0
            height = float(page.mediabox.height) / 72.0
            letter = abs(width - PAGE_WIDTH_IN) < 0.02 and abs(height - PAGE_HEIGHT_IN) < 0.02
            a4 = abs(width - 8.27) < 0.05 and abs(height - 11.69) < 0.05
            if not (letter or a4):
                out.append(finding(
                    "37 CFR 1.52(a)(1)(ii)", "blocker", where,
                    f"Page {index} is not Letter or A4",
                    f"It measures {width:.2f} by {height:.2f} inches."))
                break
    out += audit_pdf_fonts(blob, where=where)
    return out


def audit_pdf_fonts(blob: bytes, *, where: str) -> list[Finding]:
    """Ask pdffonts, because that is what the other side is doing."""
    if not shutil.which("pdffonts"):
        return [finding("Patent Center PDF validation", "note", where,
                        "Font embedding was not verified",
                        "pdffonts is not installed on this server, so the check that every glyph "
                        "is embedded could not run.")]
    import tempfile
    out: list[Finding] = []
    with tempfile.NamedTemporaryFile(suffix=".pdf") as handle:
        handle.write(blob)
        handle.flush()
        try:
            result = subprocess.run(["pdffonts", handle.name], capture_output=True, text=True,
                                    timeout=60)
        except Exception as exc:                                    # noqa: BLE001
            return [finding("Patent Center PDF validation", "note", where,
                            "Font embedding was not verified", str(exc)[:200])]
    lines = [line for line in (result.stdout or "").splitlines()[2:] if line.strip()]
    for line in lines:
        parts = line.split()
        if len(parts) < 4:
            continue
        name, embedded = parts[0], parts[3]
        if embedded.lower() != "yes":
            out.append(finding(
                "Patent Center PDF validation", "blocker", where,
                f"The font {name} is not embedded",
                "An unembedded font is a listed validation failure. Every face must be embedded, "
                "including the one a drawing legend is set in."))
        base = name.split("+", 1)[-1]
        if "#20" in name or " " in base:
            out.append(finding(
                "Patent Center PDF validation", "blocker", where,
                f"The font name {name} contains a space",
                "A space in a BaseFont name is written as #20 and is read as a defective font "
                "reference even when the face is embedded. Re-distil the file through ghostscript "
                "or register the face under a name with no spaces."))
    return out


def redistill(blob: bytes) -> bytes:
    """Run a PDF back through ghostscript so every font is embedded and subset under a clean name.

    This is the fix for the BaseFont-with-spaces rejection, and it is cheap enough to be
    unconditional on any PDF this product hands over.
    """
    if not shutil.which("gs"):
        return blob
    import tempfile
    from pathlib import Path
    with tempfile.TemporaryDirectory() as directory:
        source = Path(directory) / "in.pdf"
        target = Path(directory) / "out.pdf"
        source.write_bytes(blob)
        try:
            result = subprocess.run(
                ["gs", "-q", "-dNOPAUSE", "-dBATCH", "-dSAFER", "-sDEVICE=pdfwrite",
                 "-dPDFSETTINGS=/prepress", "-dEmbedAllFonts=true", "-dSubsetFonts=true",
                 "-dCompatibilityLevel=1.7", "-dAutoRotatePages=/None",
                 f"-sOutputFile={target}", str(source)],
                capture_output=True, text=True, timeout=180)
        except Exception:                                           # noqa: BLE001
            return blob
        if result.returncode == 0 and target.exists() and target.stat().st_size > 400:
            return target.read_bytes()
    return blob


# =============================================================================================
# The papers that are not the specification
# =============================================================================================
def audit_declaration_text(text: str, *, where: str = "declaration.pdf") -> list[Finding]:
    """Every statement 37 CFR 1.63 requires, present in the paper somebody will sign."""
    body = str(text or "")
    out: list[Finding] = []
    for rule, what, pattern in DECLARATION_STATEMENTS:
        if not pattern.search(body):
            out.append(finding(
                rule, "blocker", where,
                f"The declaration does not state that it {what}",
                "A declaration missing a required statement is accepted at upload and comes back "
                "as a notice to file corrected application papers months later."))
    if not re.search(r"/[^/\n]{2,80}/", body) and "Signature" not in body:
        out.append(finding("37 CFR 1.4(d)(2)", "formality", where,
                           "There is nowhere to sign",
                           "An S-signature is the signer's own name between forward slashes, on a "
                           "line the paper provides."))
    return out


def audit_ads_text(text: str, *, where: str = "application-data-sheet.pdf") -> list[Finding]:
    """37 CFR 1.76(b): titled, and carrying every section heading."""
    body = str(text or "")
    out: list[Finding] = []
    if "application data sheet" not in body.lower():
        out.append(finding("37 CFR 1.76(a)", "blocker", where,
                           'The sheet is not titled "Application Data Sheet"',
                           "The rule requires the title, and the Office treats an untitled sheet "
                           "as not an ADS at all."))
    missing = [heading for heading in ADS_SECTION_HEADINGS
               if heading.split("/")[0].lower() not in body.lower()]
    if missing:
        out.append(finding("37 CFR 1.76(b)", "blocker", where,
                           f"{len(missing)} required section heading(s) are missing",
                           "; ".join(missing)))
    if not re.search(r"/[^/\n]{2,80}/", body) and "Signature" not in body:
        out.append(finding("37 CFR 1.76(e)", "formality", where,
                           "The application data sheet is not signed",
                           "1.76(e) requires a signature in compliance with 1.33(b)."))
    return out


# =============================================================================================
# Drawings
# =============================================================================================
def audit_drawing_geometry(*, sheet_count: int, described_figures: Sequence[str],
                           sheet_labels: Sequence[Sequence[str]],
                           character_points: float,
                           where: str = "drawings.pdf") -> list[Finding]:
    """The parts of 37 CFR 1.84 that are about the sheet rather than about the picture."""
    out: list[Finding] = []
    if not sheet_count:
        out.append(finding("35 USC 113", "blocker", where, "There are no drawing sheets",
                           "A drawing is required where it is necessary to understand the "
                           "invention. Drawings may be informal at filing, but they must be "
                           "filed: a drawing added later is new matter."))
        return out
    for index, labels in enumerate(sheet_labels, 1):
        if not labels:
            out.append(finding("37 CFR 1.84(u)", "blocker", where,
                               f"Sheet {index}/{sheet_count} carries no view number",
                               "Each view must be numbered consecutively and independently of the "
                               'sheet, in the form "FIG. 1".'))
        elif len(labels) > 1:
            out.append(finding(
                "37 CFR 1.84(u)", "blocker", where,
                f"Sheet {index}/{sheet_count} carries {len(labels)} views under one number",
                "Views on one sheet are each numbered separately: " + ", ".join(labels) +
                ". Two arrangements shown as alternatives are two views, not one."))
    drawn = [label for labels in sheet_labels for label in labels]
    described = list(dict.fromkeys(str(value) for value in described_figures))
    missing = [value for value in described if value not in drawn]
    extra = [value for value in drawn if value not in described]
    if missing:
        out.append(finding("37 CFR 1.83(a)", "blocker", where,
                           f"{len(missing)} described view(s) are not drawn",
                           "The Brief Description names " + ", ".join(missing) +
                           " and no sheet carries them."))
    if extra:
        out.append(finding("37 CFR 1.84(u)", "blocker", where,
                           f"{len(extra)} drawn view(s) are not described",
                           "The sheets carry " + ", ".join(extra) +
                           " and the Brief Description does not mention them."))
    #  The height is estimated from a measurement of the artwork, so it carries real error. Two
    #  tiers, because being wrong in either direction costs something: calling a compliant sheet
    #  defective stops a filing nobody could then unblock, and calling a defective sheet
    #  compliant costs a drawing objection, which is correctable without losing the filing date.
    if character_points:
        centimetres = character_points / 72.0 * 2.54
        if character_points < MIN_CHARACTER_PT * 0.9:
            out.append(finding(
                "37 CFR 1.84(p)(3)", "blocker", where,
                "The reference characters are below the minimum height",
                f"They render at {centimetres:.2f} cm and the rule sets a floor of 0.32 cm. "
                "One view to a sheet is what buys the scale: a view that shares a sheet is drawn "
                "at half the width and its numerals at half the height. Artwork below about 1,400 "
                "pixels across cannot clear the floor however it is placed, and needs redrawing "
                "at a higher resolution."))
        elif character_points + 1e-6 < MIN_CHARACTER_PT:
            out.append(finding(
                "37 CFR 1.84(p)(3)", "formality", where,
                "The reference characters are close to the minimum height",
                f"They measure {centimetres:.2f} cm against a 0.32 cm floor. The Office may ask "
                "for replacement drawings; that does not affect the filing date."))
    return out
