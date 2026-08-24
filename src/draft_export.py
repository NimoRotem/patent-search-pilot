"""Portable exports for an immutable US-application draft version."""
from __future__ import annotations

import html
import re
from collections.abc import Mapping, Sequence
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate

import drafting
import draft_cite

WORKING_DRAFT_NOTICE = ""


def application_sections(version: Mapping[str, Any]) -> dict[str, str]:
    """Application sections with drafting-only citation keys made filing-readable."""
    try:
        return {str(key): draft_cite.filing_citations(str(value or ""))
                for key, value in dict(version.get("sections") or {}).items()}
    except ValueError as exc:
        raise drafting.DraftingValidationError(str(exc)) from exc


def _filing_label(value: Any) -> str:
    try:
        return draft_cite.filing_citations(str(value or ""))
    except ValueError as exc:
        raise drafting.DraftingValidationError(str(exc)) from exc


def _clean_filename(value: str, fallback: str = "us-patent-draft") -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "-", (value or "").strip()).strip("-.")
    return (value[:90] or fallback).lower()


def download_name(project: Mapping[str, Any], version_no: int, suffix: str) -> str:
    return f"{_clean_filename(str(project.get('title') or ''))}-v{int(version_no)}.{suffix}"


def render_markdown(project: Mapping[str, Any], version: Mapping[str, Any],
                    references: Sequence[Mapping[str, Any]] = ()) -> str:
    """Render only the clean application text in filing section order."""
    sections = application_sections(version)
    return drafting.render_application_markdown(sections).strip() + "\n"


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _add_text(doc: Document, text: str, *, claims: bool = False) -> None:
    blocks = re.split(r"\n\s*\n", (text or "").strip())
    for block in blocks:
        lines = [line.rstrip() for line in block.splitlines()]
        if claims and len(lines) > 1:
            for line in lines:
                if not line.strip():
                    continue
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.keep_together = False
            continue
        p = doc.add_paragraph("\n".join(lines).strip())
        p.paragraph_format.keep_together = False


def render_docx(project: Mapping[str, Any], version: Mapping[str, Any],
                references: Sequence[Mapping[str, Any]] = ()) -> BytesIO:
    """Build clean editable application text with USPTO-oriented layout defaults."""
    sections = application_sections(version)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.left_margin = Inches(1)
    sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for style_name in ("Title", "Heading 1"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.core_properties.title = _filing_label(
        sections.get("title") or project.get("title") or "")[:255]
    doc.core_properties.subject = "US utility patent application"

    for index, (key, heading) in enumerate(drafting.SECTION_ORDER):
        if key in {"claims", "abstract"}:
            doc.add_page_break()
        if index == 0:
            title = doc.add_paragraph(style="Title")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.add_run(_filing_label(
                sections.get(key) or project.get("title") or "Untitled"))
            continue
        doc.add_heading(heading.upper(), level=1)
        _add_text(doc, str(sections.get(key) or ""), claims=(key == "claims"))

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def render_pdf(project: Mapping[str, Any], version: Mapping[str, Any],
               references: Sequence[Mapping[str, Any]] = ()) -> BytesIO:
    """Render clean, paginated application text."""
    output = BytesIO()
    doc = SimpleDocTemplate(
        output, pagesize=letter, leftMargin=inch, rightMargin=0.75 * inch,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        title=_filing_label(
            (version.get("sections") or {}).get("title") or project.get("title") or "")[:255],
        subject="US utility patent application",
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "PatentBody", parent=styles["BodyText"], fontName="Times-Roman", fontSize=12,
        leading=18, spaceAfter=8, splitLongWords=True)
    heading = ParagraphStyle(
        "PatentHeading", parent=styles["Heading1"], fontName="Times-Bold", fontSize=12,
        leading=16, spaceBefore=12, spaceAfter=8, keepWithNext=True)
    title_style = ParagraphStyle(
        "PatentTitle", parent=heading, alignment=1, fontSize=14, leading=18, spaceAfter=18)
    story = []
    sections = application_sections(version)
    for index, (key, label) in enumerate(drafting.SECTION_ORDER):
        if key in {"claims", "abstract"}:
            story.append(PageBreak())
        content = str(sections.get(key) or "").strip()
        if index == 0:
            story.append(Paragraph(html.escape(
                content or _filing_label(project.get("title") or "Untitled")),
                                   title_style))
            continue
        story.append(Paragraph(html.escape(label.upper()), heading))
        for block in re.split(r"\n\s*\n", content):
            if block.strip():
                story.append(Paragraph(html.escape(block.strip()).replace("\n", "<br/>"), body))
    def footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(7.75 * inch, 0.45 * inch, f"Page {_doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    output.seek(0)
    return output
